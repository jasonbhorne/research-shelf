#!/usr/bin/env python3
"""Generate research report: James Garfield, Chester Arthur, and Death by Lightning."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Paragraph spacing
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Title
title = doc.add_heading('James Garfield, Chester A. Arthur, and Death by Lightning', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('The True Story Behind the Assassination, the Succession, and the Netflix Series')
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Calibri'

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Research Report | March 8, 2026')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = None

doc.add_paragraph()

# ─── Section 1: Definition & Background ───
doc.add_heading('1. Definition and Background', level=1)

doc.add_heading('The Topic', level=2)
doc.add_paragraph(
    '"Death by Lightning" is a four-episode Netflix historical drama miniseries that premiered on '
    'November 6, 2025. Created by Mike Makowsky and based on Candice Millard\'s acclaimed 2011 '
    'nonfiction book "Destiny of the Republic: A Tale of Madness, Medicine and the Murder of a '
    'President," the series chronicles the assassination of President James A. Garfield in 1881 and '
    'its aftermath, including the unlikely transformation of his successor, Chester A. Arthur.'
)
doc.add_paragraph(
    'The title comes from Garfield\'s own words. In a letter to Secretary of the Treasury John Sherman, '
    'responding to concerns about threats to his life, the President-elect wrote: "Assassination can be '
    'no more guarded against than death by lightning." The quote proved tragically prophetic.'
)

doc.add_heading('Key Figures', level=2)

key_figures = [
    ('James A. Garfield (1831-1881)', 'The 20th President of the United States. Born in a log cabin in '
     'Orange Township, Ohio, he was the last of the "log cabin presidents." Fatherless and raised in '
     'poverty, Garfield became a voracious reader, graduated from Williams College in 1856, and rose to '
     'become a college president by age 27, a Union Army general during the Civil War, and an eight-term '
     'congressman. He championed the 13th, 14th, and 15th Amendments and helped found the federal '
     'Department of Education. Inaugurated March 4, 1881, he served only 120 days before being shot and '
     'survived another 80 days before dying on September 19, 1881.'),
    ('Charles J. Guiteau (1841-1882)', 'A failed lawyer and delusional office-seeker from Freeport, Illinois. '
     'Guiteau became convinced he had played a major role in Garfield\'s election victory and deserved a '
     'consulship (Paris or Vienna) as a reward. A self-proclaimed "Stalwart of the Stalwarts," he believed '
     'killing Garfield would save the Republican Party and elevate Vice President Arthur, who would then '
     'reward him with a patronage position. He shot Garfield on July 2, 1881, and was tried, convicted, '
     'and hanged on June 30, 1882.'),
    ('Chester A. Arthur (1829-1886)', 'The 21st President. A protege of New York political boss Roscoe '
     'Conkling and a beneficiary of the spoils system (he had served as Collector of the Port of New York, '
     'the most lucrative patronage position in the country). Placed on Garfield\'s ticket as a compromise to '
     'appease the Stalwart faction, Arthur was widely seen as unfit for the presidency. He shocked the '
     'nation by championing civil service reform and signing the Pendleton Act into law.'),
    ('Dr. Willard Bliss (1825-1889)', 'Garfield\'s primary physician. His aggressive, unsanitary treatment '
     'of the president\'s gunshot wound is widely blamed for turning a survivable injury into a fatal '
     'infection. His first name was literally "Doctor," given to him at birth.'),
    ('Alexander Graham Bell (1847-1922)', 'The inventor of the telephone, who designed an experimental '
     'metal detector (an "induction balance") to locate the bullet in Garfield\'s body. The device likely '
     'would have worked, but Dr. Bliss restricted Bell to searching only the right side of Garfield\'s '
     'torso, where Bliss was convinced the bullet had lodged. The autopsy revealed it was on the left side.'),
    ('Roscoe Conkling (1829-1888)', 'U.S. Senator from New York and leader of the Stalwart faction. Known as '
     '"Lord Roscoe," he was the undisputed king of the patronage system and Arthur\'s political mentor.'),
    ('Julia Sand (1848-1933)', 'A physically disabled New York woman who wrote 23 extraordinary letters to '
     'Chester Arthur between August 1881 and September 1883, urging him to rise above machine politics and '
     'become a statesman. She styled herself as his "little dwarf," referencing court dwarfs who spoke '
     'difficult truths to rulers. Arthur preserved her letters while ordering nearly all other personal '
     'correspondence destroyed before his death.'),
]

for name, desc in key_figures:
    p = doc.add_paragraph()
    run = p.add_run(name + ': ')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = None
    # Bold just the name
    run_parts = name.split('(')
    p.runs[0].bold = False
    # Just make the whole name+desc a single paragraph with bold name
    p.clear()
    run_name = p.add_run(name)
    run_name.bold = True
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(11)
    run_desc = p.add_run(' - ' + desc)
    run_desc.font.name = 'Calibri'
    run_desc.font.size = Pt(11)

doc.add_heading('The Political Context: Stalwarts vs. Half-Breeds', level=2)
doc.add_paragraph(
    'The assassination of James Garfield cannot be understood without the political backdrop of the '
    'Gilded Age Republican Party, which was split into two warring factions:'
)
doc.add_paragraph(
    'The Stalwarts, led by Roscoe Conkling, fiercely defended the spoils system, where federal '
    'positions were distributed to loyal party members rather than based on merit. They championed '
    'machine politics and resisted any civil service reform. Chester Arthur was one of Conkling\'s '
    'most prominent allies.',
    style='List Bullet'
)
doc.add_paragraph(
    'The Half-Breeds, led by Maine senator James G. Blaine, favored moderate civil service reform '
    'and a merit-based system for government appointments. Garfield was aligned with this faction.',
    style='List Bullet'
)
doc.add_paragraph(
    'At the 1880 Republican National Convention, neither faction could secure the nomination. The '
    'Stalwart candidate, former president Ulysses S. Grant, was pitted against the Half-Breed candidate '
    'Blaine. After 36 ballots, both sides compromised on Ohio senator James A. Garfield. To appease '
    'the Stalwarts, Arthur was placed on the ticket as Vice President, a move Conkling opposed.'
)

# ─── Section 2: Current State of Knowledge ───
doc.add_heading('2. Current State of Knowledge', level=1)

doc.add_heading('The Historical Record', level=2)
doc.add_paragraph(
    'The assassination of James Garfield is one of the most well-documented events of the Gilded Age, '
    'though it remained relatively obscure in popular culture until Candice Millard\'s 2011 book brought '
    'it renewed attention. Key primary sources include:'
)
doc.add_paragraph('Congressional records and trial transcripts from United States v. Guiteau (1881-1882)', style='List Bullet')
doc.add_paragraph('The Chester Alan Arthur Papers at the Library of Congress, including Julia Sand\'s 23 letters', style='List Bullet')
doc.add_paragraph('Contemporary newspaper accounts from the New York Times, Washington Post, and others', style='List Bullet')
doc.add_paragraph('Medical records and autopsy reports', style='List Bullet')
doc.add_paragraph('Alexander Graham Bell\'s personal journals documenting his metal detector experiments', style='List Bullet')

doc.add_heading('The Medical Consensus', level=2)
doc.add_paragraph(
    'Modern medical historians broadly agree that Garfield\'s gunshot wound was survivable. The bullet '
    'pierced his back and lodged behind his pancreas but did not damage any vital organs. As Candice '
    'Millard argues, "had the president not been treated at all for the gunshot wounds, he would have '
    'survived the shooting." Modern medicine would have released him within days.'
)
doc.add_paragraph(
    'Instead, a parade of doctors probed the wound with unsterilized hands and instruments. They made a '
    '20-inch incision searching for the bullet. No anesthesia was used for these procedures. Infection '
    'set in, developing into sepsis. Garfield\'s weight dropped from approximately 210 pounds to 130 '
    'pounds over 79 agonizing days. He died on September 19, 1881, at his family\'s cottage in Elberon, '
    'New Jersey. The official cause of death was attributed to heart attack, massive hemorrhaging, and '
    'sepsis, with a possible ruptured gallbladder contributing.'
)
doc.add_paragraph(
    'Guiteau himself made this argument at trial: "I did not kill the president. The doctors did that. '
    'I merely shot him." The jury was unmoved and convicted him nonetheless.'
)

doc.add_heading('The Psychological Assessment of Guiteau', level=2)
doc.add_paragraph(
    'Guiteau\'s mental state has been debated since his trial. He pleaded insanity, and his behavior in '
    'the courtroom was erratic. He objected to his court-appointed defense attorney and attempted to act '
    'as his own counsel. When denied the right to deliver an opening statement, he declared the judge had '
    '"no right to muzzle" him. Modern experts have suggested he may have suffered from schizophrenia or '
    'other severe mental illness, though no definitive diagnosis is possible from historical records alone.'
)

# ─── Section 3: The True Story vs. the Netflix Series ───
doc.add_heading('3. What the Series Got Right vs. What It Changed', level=1)

doc.add_heading('What Death by Lightning Got Right', level=2)

accurate_items = [
    ('Garfield\'s rags-to-riches biography', 'The series accurately depicts Garfield\'s rise from '
     'poverty in a log cabin to the presidency, including his roles as a college president, Civil War '
     'general, and congressman.'),
    ('The 1880 convention deadlock', 'The 36-ballot deadlock between Grant and Blaine, and Garfield\'s '
     'emergence as a compromise candidate, is historically accurate.'),
    ('Guiteau\'s delusion and stalking', 'Guiteau\'s obsessive belief that he deserved a consulship, his '
     'repeated visits to the White House seeking an appointment, and his eventual decision to kill Garfield '
     'are all grounded in the historical record.'),
    ('The assassination itself', 'The shooting at the Baltimore and Potomac Railroad Station on July 2, 1881, '
     'with Garfield walking with his sons and Secretary of State Blaine, matches the historical account.'),
    ('Guiteau\'s declaration', 'His cry of "I am a Stalwart of the Stalwarts! Arthur is President now!" '
     'upon surrendering is documented in contemporary sources.'),
    ('The botched medical treatment', 'The series portrays the disastrous medical care that turned a '
     'survivable wound into a death sentence, including the unsterilized probing and Dr. Bliss\'s stubborn '
     'incompetence.'),
    ('Alexander Graham Bell\'s intervention', 'Bell\'s attempt to locate the bullet with an experimental '
     'metal detector, and Dr. Bliss\'s refusal to let him search the correct side of Garfield\'s body, '
     'are historically documented.'),
    ('Arthur\'s transformation', 'The broad arc of Arthur going from a spoils-system politician to a '
     'champion of civil service reform is accurate, though the mechanism of his transformation is '
     'fictionalized (see below).'),
    ('Conkling\'s opposition to Arthur\'s VP nomination', 'Arthur accepted the vice-presidential '
     'nomination without Conkling\'s approval and refused to rescind it when Conkling demanded he do so.'),
    ('The Pendleton Act', 'Arthur\'s signing of the Pendleton Civil Service Reform Act on January 16, 1883, '
     'is accurately depicted as the culmination of reform efforts catalyzed by Garfield\'s death.'),
]

for title_text, desc in accurate_items:
    p = doc.add_paragraph()
    run_t = p.add_run(title_text + ': ')
    run_t.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.name = 'Calibri'
    run_d.font.size = Pt(11)

doc.add_heading('What Death by Lightning Got Wrong or Fictionalized', level=2)

fiction_items = [
    ('Lucretia Garfield\'s inspirational speech to Arthur',
     'The series depicts Lucretia (Crete) Garfield slapping Arthur and delivering a speech that inspires '
     'his reform. This never happened. The despondent Arthur remained in New York while Garfield was on his '
     'deathbed, specifically to avoid the impression that he was awaiting the presidency. He never had any '
     'meeting with Lucretia Garfield. The words attributed to her in the series actually came from Julia '
     'Sand, the New York woman who wrote 23 letters to Arthur over two years. Sand is not depicted in the '
     'series at all, which historians consider a significant omission.'),
    ('Arthur portrayed as a drunk and unfaithful husband',
     'The series shows Arthur making drunken spectacles of himself and being unfaithful to his wife. '
     'Historians note there is no evidence of either. Arthur was, by most accounts, a dignified and '
     'reserved man who was genuinely devastated by the death of his wife Ellen in January 1880, more '
     'than a year before Garfield\'s assassination.'),
    ('Arthur and Guiteau as friends',
     'The series depicts frequent encounters between Arthur and Guiteau. In reality, they met only once '
     'during the 1880 campaign. There was no friendship or ongoing relationship.'),
    ('James Blaine\'s role in VP selection',
     'The series suggests Blaine played a role in selecting Arthur as VP. Historically, Blaine had no '
     'role in this decision and actually opposed it. Blaine later resigned as Secretary of State under '
     'Arthur to demonstrate his distrust.'),
    ('Arthur\'s Port Collector timeline',
     'The series implies Arthur was still serving as Collector of the Port of New York during the campaign. '
     'He had actually been removed from this position two years earlier by President Rutherford B. Hayes.'),
    ('Cabinet nominations kidnapped',
     'A plotline involving the kidnapping of cabinet nominations is entirely fictional. No such event '
     'occurred.'),
    ('Senate resignations and sex scandals',
     'The series attributes the failed reappointment of Conkling and Thomas Platt to sexual misconduct '
     'scandals. This is largely fabricated. Their resignations and failed reappointment were about patronage '
     'disputes, not personal scandals.'),
    ('Use of the term "progressive"',
     'The series uses the word "progressive" to describe the reform faction. This is anachronistic; the '
     'term was not used in this political context until approximately 15 years after these events. The '
     'correct contemporary term was "Half-Breeds."'),
]

for title_text, desc in fiction_items:
    p = doc.add_paragraph()
    run_t = p.add_run(title_text + ': ')
    run_t.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.name = 'Calibri'
    run_d.font.size = Pt(11)

# ─── Section 4: The Julia Sand Story ───
doc.add_heading('4. The Julia Sand Story: The Series\' Biggest Omission', level=1)

doc.add_paragraph(
    'Perhaps the most significant departure from history in Death by Lightning is the complete absence '
    'of Julia Sand. Her story is one of the most remarkable in American presidential history.'
)
doc.add_paragraph(
    'Julia Isabella Sand was born in New York in April 1848 to a German immigrant father and an '
    'Irish-descended mother. She lived at 46 East 74th Street in Manhattan with her mother and siblings. '
    'Despite significant physical disabilities (she described herself as "an invalid" with deafness, '
    'lameness, and spinal troubles), she cultivated an expansive intellectual life. She was educated, '
    'politically engaged, and published literary criticism under a pseudonym in magazines like Century '
    'and Harper\'s.'
)
doc.add_paragraph(
    'Sand wrote her first letter to Arthur on August 27, 1881, as Garfield lay dying. She addressed him '
    'directly: "The hours of Garfields life are numbered. Before this meets your eye, you may be '
    'President." Over the next two years, she wrote 23 letters covering:'
)
doc.add_paragraph('Moral transformation: She believed the presidency could change Arthur from a patronage politician into a statesman', style='List Bullet')
doc.add_paragraph('Machine politics: She repeatedly warned against associations with New York\'s corrupt political establishment', style='List Bullet')
doc.add_paragraph('Cabinet appointments: She praised some decisions and critiqued others', style='List Bullet')
doc.add_paragraph('Chinese Exclusion: She vehemently opposed exclusionary legislation, calling it "mean & cowardly, more than that, it is a step back into barbarism"', style='List Bullet')
doc.add_paragraph('Civil service reform: She celebrated his eventual signing of the Pendleton Act', style='List Bullet')

doc.add_paragraph(
    'Arthur never wrote back to Sand, despite her occasional pleas for replies. However, on August 20, '
    '1882, he made a single in-person visit to her home. During this awkward encounter (attended by most '
    'of her family), Arthur acknowledged reading her letters and remarked that she had offered perspectives '
    'he found valuable despite disagreeing with some.'
)
doc.add_paragraph(
    'Arthur\'s biographer Thomas C. Reeves characterized Sand as Arthur\'s "self-appointed conscience." '
    'When Arthur became seriously ill in November 1886, he ordered nearly all his personal correspondence '
    'burned, but specifically excluded Julia Sand\'s letters. They survive today in the Chester Alan Arthur '
    'Papers at the Library of Congress.'
)

# ─── Section 5: The Source Material ───
doc.add_heading('5. The Source Material: Destiny of the Republic', level=1)

doc.add_paragraph(
    'Death by Lightning is adapted from "Destiny of the Republic: A Tale of Madness, Medicine and the '
    'Murder of a President" by Candice Millard, published in 2011. The book won the Edgar Award for '
    'Best Fact Crime in 2012 and received positive reviews from the New York Times, Washington Post, '
    'Washington Times, and Seattle Times.'
)
doc.add_paragraph(
    'Millard\'s book weaves together four parallel narratives: the life of James Garfield, the delusions '
    'of Charles Guiteau, the medical disasters of Dr. Willard Bliss, and Alexander Graham Bell\'s '
    'desperate attempt to save the president with technology. The Washington Post praised Millard for '
    'having "crafted a fresh narrative that plumbs some of the most dramatic days in U.S. presidential '
    'history."'
)
doc.add_paragraph(
    'Series creator Mike Makowsky described reading the book "in one sitting, because it was one of the '
    'most insane true stories I had ever heard." He found "a very deeply ingrained situational absurdity '
    'to roughly all of the proceedings." The adaptation, however, takes significant dramatic liberties, '
    'particularly in compressing timelines, inventing interpersonal scenes, and omitting Julia Sand '
    'entirely in favor of giving her role to Lucretia Garfield.'
)

# ─── Section 6: The Netflix Series ───
doc.add_heading('6. The Netflix Series: Production and Reception', level=1)

doc.add_paragraph(
    'Death by Lightning premiered on Netflix on November 6, 2025, as a four-episode limited series. '
    'The cast features:'
)

cast_list = [
    'Michael Shannon as President James A. Garfield',
    'Matthew Macfadyen as Charles J. Guiteau',
    'Nick Offerman as Chester A. Arthur',
    'Betty Gilpin as Lucretia "Crete" Garfield',
    'Bradley Whitford as James G. Blaine',
    'Shea Whigham as Roscoe Conkling',
]
for item in cast_list:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'The series received widespread critical acclaim, with particular praise for the performances of '
    'Shannon and Macfadyen. It was named one of the best television programs of the year by the '
    'American Film Institute. It carries a 7.6 rating on IMDb.'
)
doc.add_paragraph(
    'Despite its acclaim, historians and critics have noted the show\'s tendency to prioritize dramatic '
    'entertainment over strict historical accuracy. Daniel Drezner, writing on Substack, titled his '
    'review "Why Death by Lightning Was So Frustrating," and the Library of Congress blog "Unfolding '
    'History" published a detailed analysis sourcing the show\'s dialogue against historical documents.'
)

# ─── Section 7: Aftermath and Legacy ───
doc.add_heading('7. Aftermath and Legacy', level=1)

doc.add_heading('Chester Arthur\'s Presidency', level=2)
doc.add_paragraph(
    'Arthur served as president from September 19, 1881, to March 4, 1885. To the shock of his former '
    'Stalwart allies, the onetime Collector of the Port of New York became a champion of civil service '
    'reform. His signature achievement was the Pendleton Civil Service Reform Act (January 16, 1883), '
    'which opened ten percent of federal positions to a merit system and established the Civil Service '
    'Commission. The act was a direct response to the outrage over Garfield\'s assassination by a '
    'patronage-seeker.'
)
doc.add_paragraph(
    'Arthur\'s administration also saw the largest expansion of the U.S. Navy and the implementation '
    'of harsher immigration restrictions. He did not seek re-election in 1884 (he was suffering from '
    'Bright\'s disease, which he kept secret) and died on November 18, 1886, at age 57.'
)

doc.add_heading('The End of the Spoils System', level=2)
doc.add_paragraph(
    'Garfield\'s assassination served as the catalyst that finally broke the spoils system. The public '
    'horror at a president being killed by a disgruntled office-seeker created overwhelming political '
    'pressure for reform. The Pendleton Act, while initially modest, established the principle that '
    'federal employment should be based on merit rather than political loyalty, a principle that '
    'expanded steadily over subsequent decades.'
)

doc.add_heading('Garfield\'s "What If"', level=2)
doc.add_paragraph(
    'Series creator Makowsky described Garfield\'s story as "one of the great what-ifs in American '
    'history," suggesting he might have become "one of our great presidents" had he survived. Garfield '
    'was a brilliant, self-made man with strong views on civil rights and government reform. His death '
    'at 49, after only 200 days in office (120 before the shooting, 80 after), left his potential '
    'forever unrealized.'
)

# ─── Section 8: Bottom Line ───
doc.add_heading('8. Bottom Line', level=1)

doc.add_paragraph(
    'The assassination of James Garfield in 1881 is one of those historical events that sounds too '
    'dramatic to be true: a brilliant, self-made president cut down by a delusional office-seeker, '
    'killed not by the bullet but by his own doctors, while Alexander Graham Bell races to save him '
    'with a prototype metal detector that would have worked if the lead physician had not been too '
    'stubborn to let him use it properly. Meanwhile, the vice president everyone dismissed as a '
    'machine-politics hack transforms into a genuine reformer, partly inspired by letters from a '
    'disabled woman he had never met. It is an extraordinary American story.'
)
doc.add_paragraph(
    'Netflix\'s Death by Lightning, based on Candice Millard\'s excellent "Destiny of the Republic," '
    'brings this largely forgotten chapter of American history to a mass audience with strong '
    'performances and compelling drama. The broad strokes are accurate: Garfield\'s rise from poverty, '
    'the political warfare between Stalwarts and Half-Breeds, Guiteau\'s delusion, the medical '
    'catastrophe, Bell\'s intervention, and Arthur\'s transformation. These core elements of the story '
    'are well-documented in the historical record and faithfully rendered.'
)
doc.add_paragraph(
    'However, the series takes significant dramatic liberties. The most consequential is the erasure of '
    'Julia Sand, whose 23 letters to Arthur represent one of the most remarkable citizen-to-president '
    'correspondences in American history. The show gives her role to Lucretia Garfield, inventing a '
    'confrontation that never happened. It also portrays Arthur as a drunk and philanderer without '
    'historical evidence, invents a friendship between Arthur and Guiteau, and includes entirely '
    'fictional plotlines like the kidnapping of cabinet nominations. Viewers should enjoy the series '
    'as drama but read Millard\'s book for the actual history.'
)
doc.add_paragraph(
    'The lasting significance of this story is the Pendleton Civil Service Reform Act of 1883, born '
    'directly from the public outrage over Garfield\'s death. A president\'s murder by a patronage-seeker '
    'accomplished what years of reform advocacy could not: it broke the spoils system and established '
    'merit-based federal employment. Chester Arthur, the man everyone expected to protect the old system, '
    'became the one who signed its death warrant. That irony remains the truest and most remarkable part '
    'of the whole story.'
)

# ─── Sources ───
doc.add_heading('Sources', level=1)

sources = [
    'Millard, Candice. "Destiny of the Republic: A Tale of Madness, Medicine and the Murder of a President." Doubleday, 2011.',
    '"Death by Lightning Ending and True Story Explained." Netflix Tudum, November 2025. https://www.netflix.com/tudum/articles/death-by-lightning-ending-explained-true-story',
    '"The Real Story Behind Netflix\'s Death by Lightning." Smithsonian Magazine, November 2025. https://www.smithsonianmag.com/history/the-real-story-behind-netflixs-death-by-lightning-and-the-shocking-assassination-of-president-james-a-garfield-180987598/',
    '"Death by Lightning: Fact vs. Fiction." RealClearHistory, November 21, 2025. https://www.realclearhistory.com/articles/2025/11/21/death_by_lightning_fact_vs_fiction_1148258.html',
    '"The Tragic True Story Behind Death by Lightning." TIME, November 2025. https://time.com/7331326/death-by-lightning-netflix/',
    '"Death by Lightning: Fact vs. Fiction." Competitive Enterprise Institute, 2025. https://cei.org/opeds_articles/death-by-lightning-fact-vs-fiction/',
    '"Netflix\'s Death by Lightning True Story: Fact vs. Fiction." Slate, November 2025. https://slate.com/culture/2025/11/death-by-lightning-netflix-show-president-james-garfield-true-story.html',
    '"Chester A. Arthur\'s Little Dwarf: The Correspondence of Julia I. Sand." Library of Congress, Chester Alan Arthur Papers. https://www.loc.gov/collections/chester-alan-arthur-papers/articles-and-essays/correspondence-of-julia-i-sand/',
    '"How Did James A. Garfield Die?" Britannica. https://www.britannica.com/topic/How-Did-James-A-Garfield-Die',
    '"Famous Inventor Tried to Help Save President\'s Life." National Park Service. https://www.nps.gov/articles/000/famous-inventor-tried-to-help-save-president-s-life.htm',
    '"Stalwarts, Half Breeds, and Political Assassination." National Park Service. https://www.nps.gov/articles/000/stalwarts-half-breeds-and-political-assassination.htm',
    '"The Trial of Charles Guiteau: An Account." Famous Trials. https://famous-trials.com/guiteau/2197-home',
    '"The Execution of Charles Guiteau." National Park Service. https://www.nps.gov/articles/000/the-execution-of-charles-guiteau.htm',
    '"Did Lucretia Garfield Really Say That? Sourcing the Dialogue in Death by Lightning." Library of Congress, Unfolding History Blog. https://blogs.loc.gov/manuscripts/?p=8205',
    '"Death by Lightning." Wikipedia. https://en.wikipedia.org/wiki/Death_by_Lightning',
    '"Assassination of James A. Garfield." Wikipedia. https://en.wikipedia.org/wiki/Assassination_of_James_A._Garfield',
    '"Charles J. Guiteau." Wikipedia. https://en.wikipedia.org/wiki/Charles_J._Guiteau',
    '"Chester A. Arthur." Wikipedia. https://en.wikipedia.org/wiki/Chester_A._Arthur',
    '"Pendleton Civil Service Reform Act." Wikipedia. https://en.wikipedia.org/wiki/Pendleton_Civil_Service_Reform_Act',
    'Reeves, Thomas C. "Gentleman Boss: The Life of Chester Alan Arthur." Alfred A. Knopf, 1975.',
    'Peskin, Allan. "Garfield." Kent State University Press, 1978.',
    'Goodyear, C.W. "President Garfield: From Radical to Unifier." Simon & Schuster, 2023.',
]

for source in sources:
    p = doc.add_paragraph(source, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

# Save
output_path = os.path.expanduser(
    '~/Documents/Research/garfield-assassination-death-by-lightning/'
    '2026-03-08 Garfield and Death by Lightning Research Report.docx'
)
doc.save(output_path)
print(f'Report saved to: {output_path}')
