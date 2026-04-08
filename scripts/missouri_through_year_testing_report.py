#!/usr/bin/env python3
"""Generate research report: Missouri's Through-Year Testing Program (SRSA)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

OUTPUT_DIR = os.path.expanduser("~/Documents/Research/missouri-year-round-testing")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "2026-03-30 Missouri Through-Year Testing Research Report.docx")

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    hs.font.name = 'Calibri'


def add_para(text, bold=False, italic=False, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


# =====================================================================
# TITLE PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Missouri's Through-Year Testing Program")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("The Success-Ready Student Assessment (SRSA)\nand Implications for Tennessee School Districts")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("Research Report\nMarch 30, 2026")
run.font.size = Pt(13)

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Prepared for Greeneville City Schools")
run.font.size = Pt(12)
run.italic = True

doc.add_page_break()

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
doc.add_heading('Executive Summary', level=1)

add_para(
    'In July 2025, the U.S. Department of Education approved Missouri as the only state to participate in '
    'Round 3 of the Innovative Assessment Demonstration Authority (IADA), launching the Success-Ready Student '
    'Assessment (SRSA). This through-year modular testing system replaces the traditional end-of-year Missouri '
    'Assessment Program (MAP) with assessments distributed across fall, winter, and spring testing windows in '
    'grades 3-8 for ELA and mathematics. The program began piloting in 2025-26 with approximately 5 districts '
    'and plans to scale statewide by 2029-30.'
)

add_para(
    'The SRSA represents the most ambitious current effort to fundamentally redesign state-level student '
    'assessment in the United States. Its key innovations include score banking (students who demonstrate '
    'proficiency on a module retain that score), AI-powered scoring with 24-48 hour result turnaround, and '
    'explicit support for competency-based education alongside traditional instructional models. Missouri is '
    'building on work by the Success-Ready Student Network (SRSN), a coalition of 108+ districts that have '
    'been collaborating on assessment reform since 2022.'
)

add_para(
    'The evidence base for through-year assessment as an accountability tool is thin. No state has successfully '
    'scaled a through-year system statewide with proven comparability to end-of-year tests. The most rigorous '
    'meta-analyses find that formative and continuous assessment produces a small but positive effect on student '
    'achievement (d = 0.20-0.29), substantially smaller than the d = 0.40-0.70 commonly cited from older, less '
    'rigorous reviews (Black & Wiliam, 1998). The strongest causal evidence, from a cluster randomized trial by '
    'Konstantopoulos et al. (2016), shows interim assessments help most in math and for lower-achieving students, '
    'with effects of 0.11-0.39 SD depending on grade and subject.'
)

add_para(
    'Several states have attempted through-year assessment with mixed results. Nebraska is the only state to '
    'fully implement a through-year model statewide (NSCAS Growth with NWEA). New Hampshire\'s PACE program, '
    'the longest-running pilot, showed small positive effects after three years but no significant difference '
    'from traditional testing after five years. Georgia withdrew from IADA after its dual-consortium approach '
    'fell behind schedule. Texas paused its pilot until 2029. Florida\'s progress monitoring model did not '
    'reduce overall testing burden and was layered on top of existing district assessments.'
)

add_para(
    'Tennessee does not currently have a through-year assessment model and has not applied for IADA. TDOE\'s '
    'assistant commissioner explicitly warned against spread-out testing in September 2025, stating that states '
    'which attempted it abandoned the practice "because it essentially enforces a statewide version of a pacing '
    'calendar, which is wildly unpopular." Tennessee\'s Joint Advisory Committee on Innovations in K-12 '
    'Education is studying assessment reform, with recommendations expected by end of 2025. However, Tennessee '
    'already has elements of through-year assessment at the local level through required K-3 universal reading '
    'screeners administered three times per year under the Literacy Success Act.'
)

add_para(
    'For a district like Greeneville City Schools, which consistently outperforms state averages on TCAP and '
    'has seen multi-year proficiency gains of 18-26 percentage points across content areas, the immediate '
    'practical implications are limited. Missouri\'s SRSA is a pilot that will take years to yield outcome data. '
    'The stronger near-term action is to monitor Missouri\'s results, watch for Tennessee legislative action on '
    'assessment reform, and ensure that local interim assessment practices (MAP, iReady, or similar) are being '
    'used consistently throughout the year, as research shows this is where the instructional benefit lies, '
    'regardless of what happens at the state accountability level.',
    space_after=Pt(12)
)

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (manual)
# =====================================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Executive Summary',
    '1. Definition and Background',
    '2. Evidence of Effectiveness',
    '3. Federal and State Policy Context',
    '4. What Works: Implementation Models',
    '5. Cost-Benefit Analysis',
    '6. Criticisms, Limitations, and Risks',
    '7. Measurable Outcomes',
    '8. Bottom Line',
    '9. References',
    'Appendix A: Tennessee Data',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# =====================================================================
# SECTION 1: Definition & Background
# =====================================================================
doc.add_heading('1. Definition and Background', level=1)

doc.add_heading('What Is Through-Year Assessment?', level=2)

add_para(
    'Through-year assessment (also called "through-course" or "year-round" assessment) is a model that '
    'distributes state-mandated student testing across multiple windows throughout the school year rather '
    'than concentrating it in a single end-of-year summative administration. Instead of one high-stakes '
    'test in April or May, students take modular assessments in fall, winter, and spring (or more frequently), '
    'with results combined into an annual summative determination of proficiency.'
)

add_para(
    'The concept sits at the intersection of several assessment traditions. Formative assessment, which '
    'Black and Wiliam (1998) famously described as "inside the black box," provides feedback during '
    'learning to guide instruction. Interim or benchmark assessments (MAP Growth, iReady, STAR) are '
    'administered periodically to track progress and predict performance. Summative assessments measure '
    'achievement at the end of a defined period for accountability purposes. Through-year models attempt '
    'to serve all three functions simultaneously, which, as we will see, creates fundamental design tensions.'
)

doc.add_heading('Missouri\'s Success-Ready Student Assessment (SRSA)', level=2)

add_para(
    'Missouri\'s SRSA, approved under IADA in July 2025, is the most recent and arguably most ambitious '
    'through-year testing initiative in the country. It replaces the traditional MAP end-of-year test with '
    'modular assessments in grades 3-8 for ELA and mathematics.'
)

add_para(
    'The program emerged from the Success-Ready Students Work Group (SRSWG), a coalition of 106 stakeholders '
    'recruited in 2022, representing 68 unique districts out of Missouri\'s 554. By late 2024, the Success-Ready '
    'Student Network (SRSN) had grown to 108+ districts, with 37 authorized by the Missouri State Board of '
    'Education for innovation waivers. Missouri received a $1 million federal Competitive Grant for State '
    'Assessments in fall 2024 to fund initial development.'
)

add_para('The ELA assessment design (piloting in Grade 4) uses three equivalent parallel test forms administered '
    'in fall, winter, and spring windows. Each form covers the full range of Missouri Learning Standards across '
    'five modules: Reading, Language, Listening, Research, and Writing. AI scoring provides feedback on reading, '
    'listening, and writing modules, with results returned within 24-48 hours. A hierarchical Item Response '
    'Theory (IRT) model produces overall scores.'
)

add_para(
    'The mathematics assessment (piloting in Grade 5) takes a different approach: 11 cluster assessments of 10 '
    'items each, aligned to Missouri Learning Standards clusters, administered as computer-adaptive tests. '
    'Students are classified as mastery or not-mastery on each cluster, with pacing requirements of 3 clusters '
    'by fall, 7 by winter, and all 11 by spring. An end-of-grade performance task assesses conceptual '
    'understanding and mathematical reasoning.'
)

add_para(
    'Both designs feature score banking: once a student demonstrates proficiency on a module or cluster, that '
    'score is retained and the student is not retested on that content. This is the single most distinctive '
    'feature of the SRSA, designed to reduce redundant testing and support competency-based progression.'
)

doc.add_heading('Historical Context', level=2)

add_para(
    'Through-year assessment is not a new idea. ESSA, signed in 2015, explicitly authorized states to use '
    '"multiple assessments (e.g., curriculum-embedded, interim, or through-course tests) administered throughout '
    'the year" as part of their assessment systems. The IADA provision (Section 1204) created a structured pilot '
    'pathway for states to develop innovative assessment systems. Six states have been approved: New Hampshire '
    '(2018), Louisiana (2018), Georgia (2019, withdrawn), North Carolina (2019), Massachusetts (2020), and '
    'Missouri (2025).'
)

add_para(
    'The movement gained momentum from frustration with traditional end-of-year testing. Results from spring '
    'administrations typically arrive months later, too late for instructional adjustments. The testing window '
    'itself can disrupt two to four weeks of instruction. And a single-sitting, high-stakes test creates anxiety '
    'that may not reflect students\' true capabilities, particularly for lower-achieving students, English '
    'learners, and students with disabilities.'
)

doc.add_page_break()

# =====================================================================
# SECTION 2: Evidence of Effectiveness
# =====================================================================
doc.add_heading('2. Evidence of Effectiveness', level=1)

doc.add_heading('Meta-Analyses: What the Research Actually Shows', level=2)

add_para(
    'The evidence on formative and continuous assessment has been widely cited but frequently misrepresented. '
    'Understanding the actual effect sizes matters for setting realistic expectations about what through-year '
    'testing can deliver.'
)

add_table(
    ['Study', 'N Studies', 'Effect Size', 'Key Finding'],
    [
        ['Black & Wiliam (1998)', '~250 articles', 'd = 0.40-0.70', 'Broad review; low-attaining students benefited most'],
        ['Hattie (2009)', '800+ meta-analyses', 'd = 0.90', 'Formative evaluation ranked 3rd of 138 interventions'],
        ['Kingston & Nash (2011)', '13 studies', 'd = 0.20', 'Only 13 of 300+ candidates met inclusion criteria'],
        ['Lee et al. (2020)', '33 studies', 'd = 0.29', 'US K-12 specific; medium-cycle assessments most effective'],
        ['Yan et al. (2024)', '118 studies, 258 ESs', "Hedges' g = 0.25", 'Most comprehensive recent meta-analysis'],
    ]
)

add_para(
    'The pattern is clear: earlier, less rigorous reviews (Black & Wiliam, Hattie) reported large effects '
    '(d = 0.40-0.90), while more recent, methodologically rigorous meta-analyses consistently find smaller '
    'effects (d = 0.20-0.29). Kingston and Nash (2011) is particularly telling: of 300+ candidate studies, '
    'only 13 met basic inclusion criteria for rigorous research design. A 2024 systematic review of 13 '
    'meta-analyses found that the robustness of evidence was "very low" in 9 of 13 reviews. The best available '
    'evidence suggests formative and continuous assessment produces a small but positive effect on student '
    'achievement, not the transformative impact sometimes claimed.'
)

doc.add_heading('Strongest Causal Evidence', level=2)

add_para(
    'Konstantopoulos, Li, Miller, and van der Ploeg (2016) conducted a cluster randomized controlled trial '
    'in 57 Indiana schools with approximately 25,000 students, testing mCLASS (K-2) and Acuity (3-8) interim '
    'assessments. This is one of the strongest causal studies in the field. In math (grades 3-8), they found '
    'consistent positive effects across the achievement distribution, with the strongest effects at the 10th '
    'percentile (~0.20 SD) and Grade 5 effects of 0.25-0.39 SD. In reading, significant effects appeared only '
    'at the 10th percentile (~0.11 SD) with minimal effects elsewhere. K-2 findings were minimal and '
    'non-significant.'
)

add_para(
    'McClure (2025) analyzed Utah\'s statewide benchmark module data using fixed effects regressions across '
    'grades 3-8 from 2020-2023. The key finding: most teachers only used benchmark modules right before the '
    'summative exam, producing negligible effects. However, when modules were assigned regularly throughout '
    'the year, students scored higher on end-of-year state exams. This reinforces that implementation fidelity '
    'is the critical variable, not the assessment design itself.'
)

doc.add_heading('Testing Frequency: Diminishing Returns', level=2)

add_para(
    'Bangert-Drowns, Kulik, and Kulik (1991) meta-analyzed 35 studies on classroom testing frequency and '
    'found an overall effect of d = 0.23. Critically, they found that achievement gains decreased as the '
    'number of tests increased, particularly for longer tests. When tests were used diagnostically with '
    'remedial support, benefits were greater. This suggests an optimal testing frequency exists, and more '
    'testing is not automatically better. The finding has direct implications for through-year models: if '
    'multiple assessment windows simply mean more testing without changed instructional practice, the benefit '
    'may be negligible or even negative.'
)

doc.add_heading('Through-Year Assessment as Accountability: No Proven Models', level=2)

add_para(
    'No head-to-head randomized controlled trial has compared through-year assessment to end-of-year testing '
    'for state accountability purposes. The IES evaluation of IADA states (NCEE 2023-004) found that after '
    '2-3 years, IADA systems had made "limited progress" and may not be on track to meet the 5-year statewide '
    'scale-up goal. The evaluation reported no student achievement outcome data. New Hampshire\'s PACE, the '
    'longest-running pilot, showed small positive effects after three years but no significant difference from '
    'traditional testing after five years (Dadey & Gong, 2023). Georgia withdrew from IADA. No through-year '
    'system has been implemented statewide with results accepted as comparable to end-of-year summative scores.'
)

doc.add_page_break()

# =====================================================================
# SECTION 3: Federal & State Policy Context
# =====================================================================
doc.add_heading('3. Federal and State Policy Context', level=1)

doc.add_heading('Federal Framework: ESSA Creates the Opening', level=2)

add_para(
    'ESSA (2015) requires annual statewide assessments in reading/language arts and mathematics in grades 3-8 '
    'and once in high school, with at least 95% participation for all students and subgroups. Critically, ESSA '
    'does not mandate that annual assessments be a single end-of-year test. Section 1111(b)(2)(H) explicitly '
    'permits "multiple assessments (e.g., curriculum-embedded, interim, or through-course tests) administered '
    'throughout the year," provided they collectively produce a valid, reliable annual proficiency determination.'
)

add_para(
    'The Innovative Assessment Demonstration Authority (IADA), created by ESSA Section 1204, provides the '
    'primary pathway for states to pilot through-year systems. IADA allows up to 7 states a demonstration '
    'period of up to 5 years (with possible 2-year extension) to develop innovative assessment systems with '
    'temporary flexibility from certain requirements. States must demonstrate validity, reliability, '
    'comparability across demographic groups, and maintain disaggregated reporting. Six states have been '
    'approved to date.'
)

add_table(
    ['State', 'Year', 'System', 'Design Type', 'Status'],
    [
        ['New Hampshire', '2018', 'PACE', 'Competency-based performance tasks', 'Operational (~20 districts)'],
        ['Louisiana', '2018', 'LEAP Through-Year', 'Curriculum-embedded, 3 windows', 'Active development'],
        ['Georgia', '2019', 'MAP/Navvy', 'Through-year (dual consortia)', 'Withdrawn from IADA'],
        ['North Carolina', '2019', 'NCPAT', 'Through-year, personalized', 'Active'],
        ['Massachusetts', '2020', 'Science performance tasks', 'End-of-year, multiple components', 'Active'],
        ['Missouri', '2025', 'SRSA', 'Through-year modular', 'Pilot (2025-26)'],
    ]
)

add_para(
    'The Trump administration has signaled strong support for state assessment flexibility. A July 29, 2025 '
    'Dear Colleague letter encouraged states to request ESEA waivers under Section 8401, and Oklahoma has '
    'requested the most aggressive waiver: permission to switch to through-year testing AND allow districts '
    'to choose their own assessments without federal peer review. The current policy environment is the most '
    'favorable for through-year assessment innovation since ESSA\'s passage.'
)

doc.add_heading('IDEA and Equity Guardrails', level=2)

add_para(
    'Both ESSA and IDEA require all students with disabilities to be included in state assessments with '
    'appropriate accommodations. Through-year models raise specific considerations: IEP teams must determine '
    'accommodations for each testing occasion, not just one. The 1% statewide cap on alternate assessments '
    'applies regardless of format. Assessment systems must demonstrate that multiple testing windows do not '
    'create disparate impact on protected groups, and accommodations must be consistently available across '
    'all windows. OCR\'s civil rights protections, including Title VI, Title IX, and Section 504, cannot be '
    'waived under any circumstances.'
)

doc.add_heading('Tennessee\'s Assessment Landscape', level=2)

add_para(
    'Tennessee does not currently have a through-year assessment model and has not applied for IADA. The '
    'state\'s assessment system centers on TCAP, a traditional end-of-year summative administered in a '
    'concentrated April-May window. TDOE has been explicit in its skepticism: Assistant Commissioner David '
    'Laird told the Joint Advisory Committee on Innovations in K-12 Education in September 2025 that states '
    'which attempted spread-out testing abandoned it "because it essentially enforces a statewide version of '
    'a pacing calendar, which is wildly unpopular in a lot of places."'
)

add_para(
    'The 2025 legislative session made assessment policy a flashpoint. The House and Senate disagreed '
    'significantly, with the House pushing to reduce testing requirements and the Senate defending the current '
    'system. The compromise was HB0675, creating a Joint Advisory Committee to study assessment and teacher '
    'evaluation systems, with recommendations due by end of 2025. Teachers testified about scheduling upheaval '
    'and student stress during the two-week TCAP window, with at least one educator (Dr. Kevin Schaaf, MNPS) '
    'proposing splitting TCAP into three parts throughout the year.'
)

add_para(
    'Tennessee already has through-year assessment elements at the local level: K-3 universal reading '
    'screeners are required three times per year under the Literacy Success Act, and districts widely use '
    'interim assessments (MAP, iReady, STAR) for screening and progress monitoring. TDOE provides optional '
    'start-of-year checkpoints and a formative assessment platform (Schoolnet). But these are used locally '
    'for instructional decisions, not state accountability.'
)

add_para(
    'Changing Tennessee\'s assessment model would have cascading effects. TISA\'s outcomes funding is tied to '
    'TCAP results. TVAAS requires consistent year-to-year summative data to calculate teacher growth scores, '
    'with 35% of teacher evaluations based on student growth measures. The Advisory Committee is studying both '
    'assessments and evaluations together, acknowledging they are inextricably linked.'
)

doc.add_page_break()

# =====================================================================
# SECTION 4: What Works - Implementation Models
# =====================================================================
doc.add_heading('4. What Works: Implementation Models', level=1)

doc.add_heading('Missouri SRSA: The Most Detailed Blueprint', level=2)

add_para(
    'Missouri\'s IADA application provides the most comprehensive implementation blueprint available for '
    'through-year assessment. The five-year rollout plan moves from small-scale pilot (Year 1, ~5 districts, '
    '~100 students per grade) to full statewide implementation (Year 5, all 554 districts, ~65,000 students '
    'per grade). Districts must meet readiness criteria across five domains: instructional readiness, data use '
    'capacity, assessment infrastructure, staff commitment, and instructional leadership.'
)

add_table(
    ['Year', 'School Year', 'Scale', 'Students/Grade', 'Key Activities'],
    [
        ['1', '2025-26', '~5 LEAs', '~100', 'Small-scale pilot Gr 4 ELA, Gr 5 Math; vendor RFP'],
        ['2', '2026-27', '10-15 LEAs', '~3,000', 'Limited pilot; item development grades 3-8'],
        ['3', '2027-28', '~50 LEAs', '10,000-15,000', 'Operational field test Gr 4-5; pilot other grades'],
        ['4', '2028-29', 'Statewide field test', '~65,000', 'All grades 3-8; accountability alignment'],
        ['5', '2029-30', 'Full statewide', '~65,000', 'Full operational administration'],
    ]
)

add_para(
    'The SRSA supports three instructional models simultaneously: traditional (seat-time based), '
    'competency-based (mastery-focused, students advance when proficient), and hybrid (traditional scope with '
    'readiness-based acceleration). This flexibility is enabled by Missouri statute Section 161.380, which '
    'authorizes competency-based learning frameworks. Professional development is delivered through 9 Regional '
    'Professional Development Centers (RPDCs), a model that specifically addresses rural and dispersed district '
    'access.'
)

doc.add_heading('Nebraska NSCAS Growth: The Only Statewide Implementation', level=2)

add_para(
    'Nebraska is the only state to fully implement a through-year model statewide. NSCAS Growth, built on '
    'NWEA\'s through-year assessment platform, replaced both the interim MAP Growth and the summative NSCAS '
    'General test beginning in 2022-23. The system administers adaptive assessments in fall, winter, and spring '
    'for grades 3-8 in ELA, mathematics, and science. It measures both growth and grade-level proficiency, '
    'eliminating the need for separate interim and summative tests. The statewide contract with NWEA is valued '
    'at approximately $29 million.'
)

doc.add_heading('Montana MAST: Short Testlets, Rapid Scaling', level=2)

add_para(
    'Montana\'s MAST (Montana Aligned to Standards Through-Year) uses a distinctive design: 5 testlets of '
    'approximately 15 minutes each throughout the year, developed by New Meridian. The model scaled rapidly '
    'from ~4,400 students in 67 schools during the 2022-23 pilot to ~20,700 students in 129 schools for '
    'the 2023-24 field test. Montana received a rare federal field-testing flexibility waiver in August 2023. '
    'The short-testlet design is less disruptive to instruction than longer testing sessions and has shown '
    'the ability to work across Montana\'s many small and rural districts.'
)

doc.add_heading('What Doesn\'t Work', level=2)

add_para(
    'Curriculum-anchored approaches are hard to scale. Louisiana\'s model ties assessments to specific '
    'curricula (initially ELA Guidebooks 2.0), but implementation varies in pacing and unit selection even '
    'among districts using the same materials. Georgia\'s dual-consortium approach (MAP Partnership plus '
    'Navvy) added complexity and split focus; both fell behind schedule and Georgia withdrew from IADA. '
    'Requiring fundamental instructional changes alongside assessment changes, as Georgia\'s Navvy system '
    'did, proved too much for districts to absorb at once.'
)

add_para(
    'Dual testing during pilots is consistently problematic. Texas\'s TTAP does not eliminate the obligation '
    'to administer STAAR, so pilot districts run both systems simultaneously, increasing rather than decreasing '
    'the testing burden. The Bellwether four-state study (Delaware, Florida, Nebraska, Texas) found that '
    'districts continued administering interim tests on top of through-year assessments in all four states. '
    'The promised testing reduction does not materialize during pilot phases.'
)

doc.add_heading('Cross-State Comparison', level=2)

add_table(
    ['State', 'Program', 'Windows/Year', 'Status', 'Scale'],
    [
        ['Missouri', 'SRSA', '3', 'Pilot (2025-26)', '5 LEAs, scaling to statewide by 2029-30'],
        ['Nebraska', 'NSCAS Growth', '3', 'Full statewide', 'All districts'],
        ['Montana', 'MAST', '5 testlets', 'Transitioning statewide', '54+ districts'],
        ['New Hampshire', 'PACE', 'Multiple', 'Operational', '~20 districts'],
        ['Louisiana', 'LEAP TY', '3', 'Active development', 'Limited districts'],
        ['Texas', 'TTAP', '3', 'Paused until 2029', '~90 districts'],
        ['Georgia', 'MAP/Navvy', 'Varies', 'Withdrawn', 'N/A'],
    ]
)

doc.add_page_break()

# =====================================================================
# SECTION 5: Cost-Benefit Analysis
# =====================================================================
doc.add_heading('5. Cost-Benefit Analysis', level=1)

doc.add_heading('Missouri SRSA Costs', level=2)

add_para(
    'Missouri\'s Year 1 SRSA budget is $1,140,000, including $120,000 for psychometric and data services, '
    'funded primarily through a $1 million federal Competitive Grant for State Assessments. Year 1 covers '
    'planning, vendor procurement, initial item development, stakeholder engagement, and limited pilot design. '
    'Detailed per-pupil cost data is not yet publicly available since the program just began piloting.'
)

doc.add_heading('Assessment Platform Costs', level=2)

add_table(
    ['Platform', 'Component', 'Est. Cost/Student/Year', 'Notes'],
    [
        ['iReady', 'Diagnostic only', '~$6/subject', 'Minimum 150 licenses'],
        ['iReady', 'Diagnostic + Instruction', '~$30-34/subject', 'Volume discounts available'],
        ['NWEA MAP Growth', 'Full suite', 'Custom (not public)', 'Includes scoring, reporting, PD'],
        ['NWEA (Nebraska)', 'Statewide contract', '~$29M total', 'Statewide for all tested grades'],
        ['SBAC Basic', 'Summative only', '$22.50/student', '$6.20 consortium + $16.30 state'],
        ['SBAC Complete', 'Summative + interim', '$27.30/student', '$9.55 consortium + $17.75 state'],
    ]
)

add_para(
    'The national average for state assessment primary contracts is approximately $27 per student (Brookings '
    'estimate), while full assessment spending across all tests averages approximately $65 per student. '
    'Through-year models have the potential to consolidate costs if they eliminate separate interim/diagnostic '
    'purchases ($6-34/student/year for platforms like iReady or MAP Growth) and separate summative test '
    'purchases ($20-30/student/year). However, this consolidation has not been demonstrated at scale.'
)

doc.add_heading('Hidden Costs', level=2)

add_para(
    'The AFT estimates that when factoring in teacher salary time, test prep, and instructional disruption, '
    'true testing costs can reach $848-$1,792 per student (Nelson, 2013). Teachers spend approximately twice '
    'the actual student testing time on testing-related activities (e.g., 27.7 hours of teacher time vs. 13.7 '
    'hours of student testing time in 3rd grade). Through-year models may reduce this if they eliminate '
    'separate test-prep cycles, but may increase it if each testing window triggers its own prep cycle.'
)

add_para(
    'Montana\'s experience illustrates the instructional time concern: the through-year model used approximately '
    '10 testing days across the year versus 4 days for end-of-year testing (Ravalli Republic), and Montana '
    'educator feedback indicated significant difficulty aligning testing windows with planned curriculum. '
    'Technology maintenance costs are also '
    'higher because sustained device readiness is required across 3+ testing windows rather than one.'
)

doc.add_heading('Potential Savings', level=2)

add_para(
    'If a through-year assessment eliminates the need for separate diagnostic/interim benchmarks, districts '
    'could save $6-34 per student per year on those platforms. If it also replaces the state summative test '
    'purchase, an additional $20-30 per student could be saved. Faster results (24-48 hours vs. months) reduce '
    'the need for separate local data collection. However, ESSER funds that many districts used to purchase '
    'assessment platforms have expired, creating a funding cliff that makes any transition more costly in the '
    'near term.'
)

doc.add_page_break()

# =====================================================================
# SECTION 6: Criticisms, Limitations & Risks
# =====================================================================
doc.add_heading('6. Criticisms, Limitations, and Risks', level=1)

doc.add_heading('The Fundamental Design Tension', level=2)

add_para(
    'The Center for Assessment (NCIEA) identifies a logical contradiction at the heart of through-year '
    'assessment systems. For summative/accountability use, the system assumes the student does not change '
    'significantly between assessment and year-end. For instructional use, the system explicitly expects '
    'students to improve before year-end. These purposes are "logically opposed to each other." Assessment '
    'scholars including Dylan Wiliam have argued that formative and summative assessment are so different that '
    'the same system cannot fulfill both functions, with formative assessment being "driven out" by summative '
    'demands when they coexist.'
)

doc.add_heading('Missouri-Specific Criticism', level=2)

add_para(
    'Missouri\'s SRSA faces bipartisan Senate opposition to the related A-F school grading plan, which relies '
    'heavily on test scores. State senators from both parties argue the system "focuses only on teaching to '
    'the test" and "doesn\'t give the full picture of what schools are really dealing with." Both the Missouri '
    'NEA and Missouri State Teachers Association oppose systems that rely heavily on student test scores. A '
    'parent opt-out movement in St. Louis predates the through-year proposal, and 20 districts had previously '
    'sought exemption from the Missouri Assessment Program, indicating deep dissatisfaction with the existing '
    'testing regime.'
)

doc.add_heading('Testing Burden: More, Not Less', level=2)

add_para(
    'FutureEd (2024) found it is "not clear that through-year assessments will lead to less testing time '
    'overall." The Bellwether four-state study found that districts in Delaware, Florida, Nebraska, and Texas '
    'continued administering interim tests on top of through-year assessments, making the testing burden '
    'additive rather than a replacement. Florida\'s FAST model is the most documented example: students in '
    'grades 3-10 now take three statewide assessments each in ELA and math (six total), layered over '
    'pre-existing district progress monitoring exams. FairTest states directly: "There is no actual evidence '
    'that through-year assessments provide educators with more usable, timely information."'
)

doc.add_heading('Student Impact', level=2)

add_para(
    'Research consistently finds that 10-40% of all students experience some level of test anxiety (Edutopia; '
    'ERIC). Harvard GSE research (2019) found students had 15% more cortisol the homeroom period before a '
    'standardized test than on days with no high-stakes testing. Through-year models multiply these stress '
    'points from one per year to three or more. Young students experience "anxiety, panic, irritability, '
    'frustration, boredom, crying, headaches, and loss of sleep" during high-stakes tests. If through-year '
    'testing transforms test prep from a seasonal activity to a year-round one, an EdWeek survey (2023) found '
    'that 36% of educators report test prep already consumes "pretty much the entire school year" under '
    'current systems.'
)

doc.add_heading('Equity Concerns', level=2)

add_para(
    'Students from low-income backgrounds experience significantly more learning loss on MAP assessments, '
    'particularly in ELA, meaning through-year testing may repeatedly measure and highlight disadvantage '
    'without providing resources to address it. Georgia\'s experience showed participation rates dropped to '
    'as low as 15% in the highest-poverty districts. Schools in high-poverty neighborhoods lack the human '
    'and material resources to translate assessment data into instructional changes, turning data into burden '
    'without benefit. ELL students face compounded disadvantage when tested more frequently in a language '
    'they are still acquiring. When regression-based methods handle missing data, there is evidence of '
    'differential performance based on student race, ethnicity, and school poverty (IES research).'
)

doc.add_heading('Adaptive Testing for Accountability: A Contested Choice', level=2)

add_para(
    'Using adaptive tests like MAP Growth for high-stakes accountability raises concerns. A Virginia '
    'validation study found MAP Growth assessments "do not assess the full range of SOL content, nor do they '
    'assess at the highest levels of cognitive demand." Adaptive tests give different students different items, '
    'which conflicts with the expectation that a "fair" test gives everyone the same items. Disability rights '
    'advocates warned that adaptive testing could repeat pre-NCLB patterns of giving students below-level '
    'tests and calling them proficient. Jon Cohen of the American Institutes for Research criticized "a big '
    'push from test vendors selling adaptive tests that are not tied to a strong testing blueprint" as "a '
    'step backwards."'
)

doc.add_heading('Conflicting Findings Flag', level=2)

add_para(
    'The research base contains significant internal contradictions. Black and Wiliam (1998) reported effect '
    'sizes of 0.40-0.70 for formative assessment, while Kingston and Nash (2011) found only 0.20 using '
    'stricter inclusion criteria. Hattie (2009) ranked formative evaluation at d = 0.90, but Yan et al. '
    '(2024) found g = 0.25 across 258 effect sizes. Phelps (2019) reported d = 0.84 for frequency + stakes + '
    'feedback combined, but acknowledged lower effects for K-12 versus postsecondary settings and for '
    'randomized studies. These discrepancies reflect differences in study quality, inclusion criteria, and '
    'how "formative assessment" is defined. Decision-makers should rely on the more rigorous recent estimates '
    '(d = 0.20-0.29) rather than the headline-grabbing larger figures.',
    space_after=Pt(12)
)

doc.add_page_break()

# =====================================================================
# SECTION 7: Measurable Outcomes
# =====================================================================
doc.add_heading('7. Measurable Outcomes', level=1)

add_para(
    'Through-year assessment proponents claim several measurable benefits. Here is each claim alongside the '
    'strength of the supporting evidence.'
)

add_table(
    ['Claimed Outcome', 'Evidence Strength', 'Key Finding'],
    [
        ['Improved student achievement', 'Moderate', 'd = 0.20-0.29 in rigorous meta-analyses; positive but modest'],
        ['Benefits for low-achieving students', 'Moderate-Strong', 'Konstantopoulos RCT: effects strongest at 10th percentile'],
        ['Faster instructional feedback', 'Strong (face validity)', '24-48 hour results vs. months for traditional; no outcome study'],
        ['Reduced testing burden', 'Weak', 'Bellwether: districts added tests; Florida FAST increased testing'],
        ['Improved teacher data use', 'Weak', 'No controlled studies; implementation quality is key mediator'],
        ['Support for CBE', 'Theoretical', 'Score banking enables mastery-based progression; no outcome data'],
        ['Lower test anxiety', 'Unproven', 'Multiplies test occasions; may increase rather than decrease stress'],
        ['Improved equity outcomes', 'Weak-Mixed', 'Potential benefits for low achievers; participation gaps in high-poverty'],
        ['Cost savings', 'Weak', 'Consolidation potential exists but not demonstrated at scale'],
    ]
)

add_para(
    'The most consistently supported finding is that formative and interim assessment benefits '
    'lower-achieving students the most. This finding appears across multiple meta-analyses and in the '
    'Konstantopoulos RCT. However, this benefit depends on what teachers do with the data, not on the '
    'assessment structure itself. McClure (2025) found that benchmark modules only improved outcomes when '
    'used consistently throughout the year, not when crammed before testing. Implementation fidelity '
    'is the critical variable.'
)

add_para(
    'New Hampshire\'s PACE provides the longest-running outcome data for a through-year pilot. After three '
    'years, small positive effects were observed across grades 8 and 11 in ELA and math, with lower-achieving '
    'students showing small positive differential effects. After five years, however, there was no significant '
    'difference between PACE and non-PACE schools. This trajectory, initial positive effects that fade over '
    'time, should temper expectations about long-term impact.'
)

add_para(
    'A competency-based education study (IES) found that math competency scores accurately predicted state '
    'test proficiency for only 40% of students, and literacy competency scores for 59%. Competency scores '
    'had statistically significant but weak relationships with state test scores (R-squared = 0.03-0.04). '
    'However, students below grade level completed their level faster: 43-47% finished in 3 or fewer quarters '
    'versus 17-22% in traditional systems.'
)

doc.add_page_break()

# =====================================================================
# SECTION 8: Bottom Line
# =====================================================================
doc.add_heading('8. Bottom Line', level=1)

add_para(
    'Missouri\'s SRSA is an ambitious and well-designed pilot that addresses real frustrations with '
    'traditional end-of-year testing: delayed results, single-sitting high-stakes pressure, and disconnection '
    'from daily instruction. The score banking mechanism, AI-powered rapid feedback, and support for '
    'competency-based education represent genuine innovations in state assessment design. Missouri has also '
    'done the stakeholder engagement work, with 108+ districts participating in the SRSN and a detailed '
    'five-year implementation plan. If any through-year system can succeed, Missouri has put itself in a '
    'strong position.'
)

add_para(
    'But the evidence demands caution. No state has scaled a through-year assessment statewide with proven '
    'comparability to traditional tests. The best research shows formative assessment produces modest '
    'achievement effects (d = 0.20-0.29), not the transformative impact sometimes claimed. New Hampshire\'s '
    'positive results faded after five years. Georgia withdrew from IADA. Florida\'s model increased rather '
    'than decreased testing. The fundamental tension between formative and summative purposes remains '
    'unresolved in every implementation. And the hidden costs, from teacher time to technology infrastructure '
    'to dual-testing during pilots, are real and underappreciated.'
)

add_para(
    'For Tennessee, and specifically for Greeneville City Schools, the immediate implications are '
    'observational, not operational. TDOE has explicitly pushed back on spread-out testing. The Joint Advisory '
    'Committee is studying assessment reform but has not recommended a through-year model. TISA\'s outcomes '
    'funding and TVAAS teacher evaluations are built on the current TCAP framework, and changing either would '
    'require careful recalibration. GCS already outperforms state averages by significant margins (9.7-19.7 '
    'percentage points across subjects), so the case for disrupting a working system is not urgent.'
)

add_para(
    'The recommended next steps for a district like GCS are practical. First, monitor Missouri\'s SRSA pilot '
    'results as they emerge over the next 2-3 years, looking specifically for achievement outcome data, not '
    'just implementation milestones. Second, ensure that local interim assessments (MAP, iReady, or similar) '
    'are being used consistently throughout the year with clear protocols for data-to-instruction cycles, as '
    'this is where the research shows the real benefit lies. Third, track Tennessee\'s Advisory Committee '
    'recommendations and any legislative action on assessment reform in the 2026 session. Fourth, if through-year '
    'assessment gains traction nationally, investigate Nebraska\'s NSCAS Growth as the only proven statewide '
    'model and evaluate whether the NWEA through-year platform could integrate with existing GCS assessment '
    'practices.'
)

add_para(
    'The questions that remain unanswered are significant. Can score banking produce valid, comparable '
    'accountability scores? Does through-year testing actually reduce the overall testing burden, or does it '
    'just redistribute and potentially increase it? What happens to teacher evaluation models built on '
    'end-of-year growth scores when the data inputs change? And most importantly: does any of this improve '
    'student learning more than simply using existing interim assessments well? Missouri will begin to answer '
    'these questions. Tennessee should watch carefully, but there is no reason to act before the evidence arrives.'
)

doc.add_page_break()

# =====================================================================
# SECTION 9: References
# =====================================================================
doc.add_heading('9. References', level=1)

references = [
    'Aligned. (2025). Missouri to Pilot a New Way to Test Students. https://www.wearealigned.org/blogs/missouri-pilot-new-way-test-students',
    'American Enterprise Institute. Progress Monitoring in Florida: A New Solution to an Old Education Problem. https://www.aei.org/research-products/report/progress-monitoring-in-florida-a-new-solution-to-an-old-education-problem/',
    'Aurora Institute. (2019). The U.S. Department of Education Is Inviting New Applications for Round 3 of the IADA. https://aurora-institute.org/blog/the-u-s-department-of-education-is-inviting-new-applications-for-round-3-of-the-innovative-assessment-demonstration-authority-iada/',
    'Aurora Institute. (2020). How Systems of Assessments Aligned with Competency-Based Education Can Support Equity. http://www.aurora-institute.org/wp-content/uploads/how-systems-of-assessment-aligned-with-competency-based-education-can-support-equity-jan-2020-web.pdf',
    'Bangert-Drowns, R. L., Kulik, J. A., & Kulik, C. C. (1991). Effects of frequent classroom testing. The Journal of Educational Research, 85(2), 89-99.',
    'Bellwether. (2023). Demystifying Statewide Standardized Assessments. https://bellwether.org/wp-content/uploads/2023/03/DemystifyingStandardizedAssessments_Brief-1_Bellwether_April2023.pdf',
    'Bellwether. (2024). Testing the Waters: Insights Into Parent Perspectives on Through-Year Assessment Implementation. https://bellwether.org/publications/testing-the-waters/',
    'Bipartisan Policy Center. Federal Education Policy 101: State Assessments. https://bipartisanpolicy.org/explainer/federal-education-policy-101-state-assessments/',
    'Black, P., & Wiliam, D. (1998). Assessment and classroom learning. Assessment in Education: Principles, Policy & Practice, 5(1), 7-74.',
    'Black, P., & Wiliam, D. (1998). Inside the black box: Raising standards through classroom assessment. Phi Delta Kappan, 80(2), 139-148.',
    'Brookings Institution. (2012). Strength in Numbers: State Spending on K-12 Assessment Systems. https://www.brookings.edu/articles/strength-in-numbers-state-spending-on-k-12-assessment-systems/',
    'Brookings Institution. (2015). Testing Costs a Drop in the Bucket. https://www.brookings.edu/articles/testing-costs-a-drop-in-the-bucket/',
    'Center for American Progress. (2014). Testing Overload in America\'s Schools. https://cdn.americanprogress.org/wp-content/uploads/2014/10/LazarinOvertestingReport.pdf',
    'Center for American Progress. (2024). Fact Sheet: 3 Trends in K-12 Assessments Across the Country. https://www.americanprogress.org/article/fact-sheet-3-trends-in-k-12-assessments-across-the-country/',
    'Center for Assessment (NCIEA). Why Has It Been So Difficult to Develop a Viable Through-Year Assessment? https://www.nciea.org/blog/why-has-it-been-so-difficult-to-develop-a-viable-through-year-assessment/',
    'Chalkbeat Tennessee. (2025). Should Tennessee Dial Back State Testing Requirements? Some Teachers Think So. https://www.chalkbeat.org/tennessee/2025/09/26/testing-requirements-too-burdensome-some-teachers-say/',
    'Congressional Research Service. Secretarial Waiver Authority Under ESEA, Section 8401. https://www.congress.gov/crs-product/IF13152',
    'Consortium for Citizens with Disabilities. Concerns About Adaptive Testing. Referenced in Education Week (2013).',
    'Dadey, N. & Gong, B. (2023). Keeping Up the PACE: Evaluating Grade 8 Student Achievement Outcomes. Applied Measurement in Education, 36(2).',
    'Data Quality Campaign. What Are Through-Year Assessments? https://dataqualitycampaign.org/what-are-through-year-assessments/',
    'Education Trust. (2023). Future of Assessments: Centering Equity and the Lived Experiences of Students, Families, and Educators. https://edtrust.org/wp-content/uploads/2014/09/Future_Assessments_FINAL.pdf',
    'Education Trust. 5 Things Every Equity Advocate Should Know About Through-Year Assessments. https://edtrust.org/rti/5-things-every-equity-advocate-should-know-about-through-year-assessments/',
    'Education Week. (2013). Adaptive Testing Gains Momentum, Prompts Concerns. https://www.edweek.org/teaching-learning/adaptive-testing-gains-momentum-prompts-concerns/2013/07',
    'Education Week. (2023). States Eye Assessment Throughout the Year as Frustration With Standardized Testing Mounts. https://www.edweek.org/teaching-learning/states-eye-assessment-throughout-the-year-as-frustration-with-standardized-testing-mounts/2023/09',
    'Education Week. (2025). The Future of Annual State Testing Is in the Trump Admin.\'s Hands. https://www.edweek.org/policy-politics/the-future-of-annual-state-testing-is-in-the-trump-admin-s-hands/2025/09',
    'ExcelinEd. (2024). 7 Education Policy Trends for State Lawmakers in 2025. https://excelined.org/2024/12/11/7-education-policy-trends-for-state-lawmakers-in-2025/',
    'ExcelinEd. (2025). Tennessee 2025 Legislative Session Delivers Bold Education Policies. https://excelinedinaction.org/2025/06/02/tennessee-2025-legislative-session-delivers-bold-education-policies/',
    'FairTest. (2024). Interim, Formative and Through-Year Assessments: Keeping Assessments Authentic and Useful. https://fairtest.org/wp-content/uploads/2024/07/Interim-Formative-and-Through-Year-Assessments-Keeping-Assessments-Authentic-and-Useful.pdf',
    'Frontiers in Psychology. (2022). The effectiveness of formative assessment for enhancing reading achievement in K-12 classrooms: A meta-analysis. https://www.frontiersin.org/journals/psychology/articles/10.3389/fpsyg.2022.990196/full',
    'FutureEd. (2024). The Limitations of Through-Year Assessments. https://www.future-ed.org/wp-content/uploads/2024/05/The-Limitations-of-Through-Year-Assessments.pdf',
    'Getting Smart. (2023). Why 20 Missouri School Districts Are Seeking New Innovation Waivers. https://www.gettingsmart.com/2023/08/29/why-20-missouri-school-districts-are-seeking-new-innovation-waivers-to-rethink-the-way-they-test-students/',
    'Greeneville City Schools. (2025). Greeneville Students Achieve Record TCAP Gains. https://www4.gcschools.net/article/2303995',
    'Harvard Graduate School of Education. (2019). Tests and Stress Bias. https://www.gse.harvard.edu/ideas/usable-knowledge/19/02/tests-and-stress-bias',
    'Hattie, J. (2009). Visible Learning: A Synthesis of Over 800 Meta-Analyses Relating to Achievement. Routledge.',
    'Institute of Education Sciences (IES). (2023). Evaluating the Federal Innovative Assessment Demonstration Authority (NCEE 2023-004). https://ies.ed.gov/use-work/resource-library/report/evaluation-report/evaluating-federal-innovative-assessment-demonstration-authority-early-implementation-and-progress',
    'Kingston, N., & Nash, B. (2011). Formative assessment: A meta-analysis and a call for research. Educational Measurement: Issues and Practice, 30(4), 28-37.',
    'Konstantopoulos, S., Li, W., Miller, S. R., & van der Ploeg, A. (2016). Effects of interim assessments across the achievement distribution. Educational and Psychological Measurement, 76(4), 587-608.',
    'KSMU. (2024). School districts in Missouri collaborate to rethink assessment. https://www.ksmu.org/news/2024-12-16/school-districts-in-missouri-collaborate-to-rethink-assessment',
    'Lane, S. (2020). Test-Based Accountability Systems: The Importance of Paying Attention to Consequences. ETS Research Report Series.',
    'Lee, H., Chung, H. Q., Zhang, Y., Abedi, J., & Warschauer, M. (2020). The effectiveness and features of formative assessment in US K-12 education. Applied Measurement in Education, 33(2), 124-140.',
    'Louisiana DOE. (2024). Louisiana Awarded $3 Million Federal Grant for Innovative Assessments. https://doe.louisiana.gov/about/newsroom/news-releases/release/2024/07/20/',
    'McClure, K. N. (2025). Benchmark modules: A better interim assessment? Evidence from statewide use in Utah. AERA Open.',
    'Missouri DESE. (2025). Missouri Assessment Program IADA Application. https://www.ed.gov/media/document/missouri-application-2025-110444.pdf',
    'Missouri DESE. (2025). Missouri Chosen to Pilot Innovative Statewide Assessment Program. https://dese.mo.gov/communications/missouri-chosen-pilot-innovative-statewide-assessment-program',
    'Missouri Independent. (2026). Missouri pilots new standardized test model for public schools with faster results. https://missouriindependent.com/2026/03/27/missouri-pilots-new-standardized-test-model-for-public-schools-with-faster-results/',
    'Montana OPI. Montana Aligned to Standards Through-Year Assessment. https://opi.mt.gov/Leadership/Assessment-Accountability/Montana-Aligned-to-Standards-Through-Year',
    'Nebraska DOE. (2019). Evolution of NSCAS: The NSCAS Adaptive Through-Year Model FAQ. https://www.education.ne.gov/wp-content/uploads/2019/11/NSCAS-Through-Year-Assessment-Internal-FAQ-November-2019-003.pdf',
    'Nelson, H. (2013). Testing More, Teaching Less. American Federation of Teachers. https://www.aft.org/sites/default/files/media/2014/testingmore2013.pdf',
    'New Hampshire DOE. Performance Assessment of Competency Education (PACE). https://www.education.nh.gov/who-we-are/division-of-learner-support/bureau-of-instructional-support/performance-assessment-competency-education',
    'New Meridian Corp. Through-Year Assessment Defined: How This Model Really Works. https://newmeridiancorp.org/through-year-assessment-defined-how-this-model-really-works/',
    'NWEA. MAP Growth Technical Report (2019). https://www.nwea.org/uploads/2021/11/MAP-Growth-Technical-Report-2019_NWEA.pdf',
    'NWEA. Through-Year Assessment. https://www.nwea.org/through-year-assessment/',
    'Phelps, R. P. (2019). Test frequency, stakes, and feedback in student achievement: A meta-analysis. Evaluation Review, 43(3-4), 111-151.',
    'Show-Me Institute. 20 Missouri Districts Seek Exemption from the Missouri Assessment Program. https://showmeinstitute.org/blog/performance/20-missouri-districts-seek-exemption-from-the-missouri-assessment-program/',
    'Sustainability (MDPI). (2024). A systematic review of meta-analyses on the impact of formative assessment on K-12 students\' learning. Sustainability, 16(17), 7826.',
    'Tennessee Department of Education. 2025-26 TCAP Testing Calendar. https://www.tn.gov/content/dam/tn/education/events/2025-26_TCAP_Testing_Calendar.pdf',
    'Tennessee Department of Education. (2025). TDOE Announces Year-Over-Year Improvement Across All Subject Areas. https://www.tn.gov/education/news/2025/7/8/',
    'Tennessee Department of Education. Best for All Strategic Plan. https://www.tn.gov/education/best-for-all.html',
    'Tennessee Department of Education. Overview of Testing in Tennessee. https://www.tn.gov/education/districts/lea-operations/assessment/testing-overview.html',
    'Tennessee Department of Education. TISA Formula. https://www.tn.gov/education/best-for-all/tnedufunding.html',
    'Tennessee Firefly. (2025). Lawmakers to Consider Reducing Student Testing and Teacher Evaluation Requirements. https://www.tnfirefly.com/news/lawmakers-to-consider-reducing-student-testing-and-teacher-evaluation-requirements',
    'Tennessee SCORE. The Future of State and Local Assessments in Tennessee. https://tnscore.org/resources/the-future-of-state-and-local-assessments-in-tennessee',
    'Texas Education Agency. Texas Through-Year Assessment Pilot. https://tea.texas.gov/student-assessment/assessment-initiatives/texas-through-year-assessment-pilot',
    'U.S. Department of Education. (2025). Dear Colleague Letter: ESEA Flexibility and Waivers (July 29, 2025). https://www.ed.gov/media/document/dear-colleague-letter-esea-flexibility-and-waivers-july-29-2025-110440.pdf',
    'U.S. Department of Education. (2025). Innovative Assessment Demonstration Authority (IADA). https://www.ed.gov/grants-and-programs/formula-grants/school-improvement/iada',
    'U.S. Department of Education. (2025). U.S. Department of Education Approves Missouri to Pilot Innovative Statewide Assessment Program. https://www.ed.gov/about/news/press-release/us-department-of-education-approves-missouri-pilot-innovative-statewide-assessment-program',
    'WGEM. (2026). Missouri governor\'s A-F school grading plan faces bipartisan Senate opposition. https://www.wgem.com/2026/03/27/missouri-governors-a-f-school-grading-plan-faces-bipartisan-senate-opposition/',
    'Yan, Z., et al. (2024). The impact of formative assessment on K-12 learning: A meta-analysis. Educational Research and Evaluation, 29(7-8).',
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_page_break()

# =====================================================================
# APPENDIX A: Tennessee Data
# =====================================================================
doc.add_heading('Appendix A: Tennessee Data', level=1)

doc.add_heading('Greeneville City Schools 2024-25 TCAP Results', level=2)

add_table(
    ['Subject', 'GCS Proficiency', 'State Average', 'Difference'],
    [
        ['ELA (Grades 3-8)', '48.6%', '38.9%', '+9.7 pp'],
        ['Math (Grades 3-8)', '49.0%', '42.2%', '+6.8 pp'],
        ['Social Studies (Grades 6-8)', '64.9%', '49.0%', '+15.9 pp'],
        ['English II EOC', '74.0%', '54.3%', '+19.7 pp'],
        ['English I EOC', '43.7%', '39.2%', '+4.5 pp'],
    ]
)

add_para(
    'Over four years, GCS proficiency increased approximately 19 percentage points in ELA, 20 in math, '
    '18 in science, and 26 in social studies. GCS outperforms state averages in every content area.'
)

doc.add_heading('Tennessee vs. Missouri: Assessment Model Comparison', level=2)

add_table(
    ['Feature', 'Missouri (SRSA Pilot)', 'Tennessee (TCAP)'],
    [
        ['Model', 'Through-year modular (IADA pilot)', 'Traditional end-of-year summative'],
        ['Administration', 'Minimum 3 occasions per year', 'Single testing window (April-May)'],
        ['Teacher Flexibility', 'Choose when to administer modules', 'Fixed state testing windows'],
        ['Score Reporting', '24-48 hours', 'Months later (summer release)'],
        ['Federal Authority', 'IADA approved for 2025-26', 'No IADA application'],
        ['Score Banking', 'Students can bank proficiency scores', 'No equivalent'],
        ['CBE Support', 'Supports competency-based frameworks', 'Traditional standards-based'],
        ['Status', 'Pilot (select districts, 2025-26)', 'Fully operational statewide'],
    ]
)

doc.add_heading('Tennessee Assessment Policy Timeline (2025-26)', level=2)

add_table(
    ['Date', 'Event', 'Details'],
    [
        ['May 2025', 'HB0675 passed', 'Created Joint Advisory Committee on Innovations in K-12 Education'],
        ['Sept 2025', 'Advisory Committee hearing', 'TDOE warns against spread-out testing; teachers testify on TCAP burden'],
        ['Late 2025', 'Committee recommendations due', 'Studying both assessments and teacher evaluations'],
        ['Spring 2026', 'SB1585 advances', 'Would allow voucher schools to use national tests instead of TCAP'],
        ['Spring 2027', 'Elementary TCAP online', 'Elementary TCAP migration to online administration begins'],
    ]
)

doc.add_heading('Tennessee Local Assessment Landscape', level=2)

add_para(
    'While Tennessee does not have a through-year state assessment, districts already administer interim '
    'assessments throughout the year. K-3 universal reading screeners are required three times per year '
    '(fall, winter, spring) under the Literacy Success Act. Approved screeners include MAP, iReady, STAR, '
    'and others. Students scoring at or below the 15th percentile receive Home Literacy Reports. TDOE '
    'provides optional start-of-year checkpoints and the Schoolnet formative assessment platform for teacher-'
    'created assessments.'
)

doc.add_heading('TISA and Assessment Accountability Connection', level=2)

add_para(
    'Tennessee\'s Investment in Student Achievement (TISA) formula includes outcomes funding tied to TCAP '
    'results. Districts must submit annual accountability reports including plans to achieve 70% or more of '
    '3rd graders scoring "met expectations" or "exceeded expectations" on TCAP ELA. Changing the assessment '
    'model would directly impact TISA accountability metrics, outcomes funding calculations, and TVAAS growth '
    'scores that comprise 35% of teacher evaluations.'
)

# =====================================================================
# Save
# =====================================================================
doc.save(OUTPUT_FILE)
print(f"Report saved to: {OUTPUT_FILE}")
