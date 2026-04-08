#!/usr/bin/env python3
"""
Generate research report: The Impact of School Operations on Student Achievement
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

# Title
title = doc.add_heading('The Impact of School Operations on Student Achievement', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Maintenance, Facilities, Nursing, Transportation, HR, Finance, and Nutrition')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(89, 89, 89)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('A Comprehensive Research Synthesis')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(89, 89, 89)
run.italic = True

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Helper functions
def add_body(text):
    doc.add_paragraph(text)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)

add_body(
    'School operations, the infrastructure and support systems that surround classroom instruction, '
    'exert a measurable and often underappreciated influence on student achievement. This report synthesizes '
    'peer-reviewed research, federal and state policy frameworks, implementation case studies, cost-benefit '
    'analyses, and critical counterarguments to provide a comprehensive picture of how maintenance, facilities, '
    'nursing, transportation, human resources, finance, and nutrition services affect student outcomes.'
)

add_body(
    'The evidence is clear on several points. First, teacher quality, an HR function at its core, is the '
    'single most important school-based determinant of student achievement, with replacing a bottom-5% teacher '
    'with an average one increasing the present value of students\' lifetime income by approximately $250,000 '
    'per classroom (Chetty, Friedman, & Rockoff, 2014). Second, indoor environmental quality, particularly '
    'air quality, ventilation, and temperature, has surprisingly strong causal evidence, with ventilation '
    'improvements yielding 0.07-0.15 standard deviation gains in math and reading (Stafford, 2015). Third, '
    'per-pupil spending increases produce meaningful long-term outcomes: a 10% increase across all 12 years '
    'of schooling leads to 7.25% higher adult wages and a 3.67 percentage-point reduction in poverty '
    '(Jackson, Johnson, & Persico, 2016).'
)

add_body(
    'Most operational areas affect achievement indirectly through attendance. School nursing, transportation, '
    'nutrition, and custodial services improve student presence in school, which then drives academic gains. '
    'Effects are nearly always larger for disadvantaged students, making operational investments inherently '
    'equity investments. However, the evidence also shows that operations investments face diminishing returns '
    'above a threshold of basic adequacy, and that facility construction alone, without attention to the '
    'instructional and social environment within the building, produces disappointingly small academic effects.'
)

add_body(
    'For a district like Greeneville City Schools, already performing above state averages with Exemplary '
    'District designation, the evidence points toward targeted operational investments in indoor air quality, '
    'preventive maintenance, full-time school nursing, and teacher retention as the highest-return strategies. '
    'Tennessee\'s $9.8 billion statewide infrastructure gap and 1:3,000 nurse-to-student ratio (versus the '
    'recommended 1:750) represent significant policy contexts that shape local decision-making.'
)

# ============================================================
# SECTION 1: DEFINITION & BACKGROUND
# ============================================================
doc.add_heading('1. Definition and Background', level=1)

add_body(
    'School operations encompass the non-instructional systems and services that create the conditions '
    'for teaching and learning to occur. While classroom instruction is the primary mechanism through which '
    'schools produce academic outcomes, instruction does not happen in a vacuum. Students must be transported '
    'safely to school, housed in buildings with adequate heating, cooling, lighting, and air quality, fed '
    'nutritious meals, provided health services when ill or injured, and supported by staff who are recruited, '
    'retained, and compensated through human resource systems. All of these functions are funded through '
    'financial management systems that determine how limited resources are allocated.'
)

add_body(
    'The major domains of school operations include:'
)

domains = [
    ('Facilities and Maintenance', 'Building conditions, HVAC systems, lighting, air quality, custodial services, preventive and reactive maintenance, capital planning, and construction.'),
    ('Transportation', 'Bus routing, ride times, vehicle maintenance, driver staffing, access for students with disabilities and experiencing homelessness.'),
    ('School Nursing and Health Services', 'School nurse staffing, chronic condition management, immunization compliance, mental health screening, and school-based health centers.'),
    ('Human Resources', 'Teacher and staff recruitment, retention, compensation, professional development, evaluation, and working conditions.'),
    ('Finance', 'Per-pupil spending, budget allocation, funding formulas, capital funding, grant management, and financial efficiency.'),
    ('Nutrition Services', 'School breakfast and lunch programs, Community Eligibility Provision, food quality, and meal participation rates.'),
]

for domain, desc in domains:
    p = doc.add_paragraph()
    run = p.add_run(f'{domain}: ')
    run.bold = True
    p.add_run(desc)

add_body(
    'Historically, these operational functions have been treated as support services, distinct from and '
    'subordinate to instruction. The assumption was that operations were necessary but not sufficient for '
    'student learning, a cost center rather than an investment in outcomes. This report examines whether '
    'that assumption holds up against the evidence.'
)

add_body(
    'The conceptual framework connecting operations to achievement operates through multiple pathways. '
    'Some are direct: poor air quality impairs cognitive function in real time. Others are indirect: '
    'unreliable transportation leads to chronic absenteeism, which leads to missed instruction, which '
    'leads to lower achievement. Still others are systemic: inadequate compensation and poor working '
    'conditions drive teacher turnover, which destabilizes instruction across an entire school. Understanding '
    'these pathways is critical for district leaders making resource allocation decisions.'
)

# ============================================================
# SECTION 2: EVIDENCE OF EFFECTIVENESS
# ============================================================
doc.add_heading('2. Evidence of Effectiveness', level=1)

doc.add_heading('2.1 Indoor Air Quality and Ventilation', level=2)

add_body(
    'Indoor air quality has some of the strongest causal evidence linking operational conditions to '
    'student achievement. Stafford (2015) found that ventilation improvement projects improved math scores '
    'by 0.07 standard deviations and reading scores by 0.11 standard deviations, and increased the '
    'probability of passing standardized tests by 2-3%. Mold remediation projects improved math scores by '
    '0.15 SD and reading scores by 0.14 SD. Roth\'s research on indoor air particulate exposure demonstrated '
    'detrimental impacts on math performance up to 0.5 SD.'
)

add_body(
    'The EPA reviewed 11 studies and found that 8 showed statistically significant improvements in student '
    'performance with increased ventilation rates. A study of 100 U.S. schools found that doubling ventilation '
    'rates from approximately 7.5 to 15 CFM per person was associated with an 8% improvement in academic '
    'performance. Lawrence Berkeley National Laboratory documented 13-14% improvements when ventilation rates '
    'met ASHRAE standards. Allen et al. (2016) found cognitive function scores doubled when CO2 levels decreased '
    'from 1,400 to 550 ppm. These findings suggest that air quality interventions may be a more cost-effective '
    'way to improve standardized test scores than class size reductions.'
)

doc.add_heading('2.2 Temperature', level=2)

add_body(
    'Park et al. (2020) found that for every one-degree Fahrenheit increase in average temperature over the school '
    'year, students demonstrated a 1% decrease in learning. Controlled experimental settings showed thermal '
    'discomfort led to declines in short-term memory of 12% and verbal ability of 24% (Leyten & Kurvers, '
    '2021). These findings have been replicated across quasi-experimental designs and are '
    'particularly relevant for schools in regions with extreme temperatures or aging HVAC systems.'
)

doc.add_heading('2.3 Lighting and Daylighting', level=2)

add_body(
    'The Heschong Mahone Group\'s seminal study (1999) found students in classrooms with the most daylight '
    'progressed 20% faster in math and 26% faster in reading over one year compared to those in classrooms '
    'with the least daylight. While the evidence is correlational rather than experimental, it has been '
    'replicated across multiple settings and is consistent with neuroscience research on circadian rhythms '
    'and cognitive performance.'
)

doc.add_heading('2.4 Overall Building Condition', level=2)

add_body(
    'Stewart\'s (2014) meta-analytical synthesis reviewed 42 studies and found consistent evidence that '
    'students in well-maintained, modern facilities outperform those in deteriorating buildings by 3-17 '
    'percentage points on standardized tests. Schools without major maintenance backlogs have higher average '
    'daily attendance by 4-5 students per 1,000 and annual dropout rates lower by 10-13 students per 1,000. '
    'In low-wealth districts, facilities upgrades led to test score increases equivalent to 10% of the gap '
    'between high- and low-income districts\' academic outcomes.'
)

doc.add_heading('2.5 Teacher Quality and HR Operations', level=2)

add_body(
    'Teacher quality is the most robustly documented operational factor affecting student achievement. '
    'Chetty, Friedman, and Rockoff (2014) analyzed 2.5 million children linked to tax records and found '
    'that a one standard deviation increase in teacher value-added in a single grade results in increased '
    'earnings at age 28 of $182 (0.9% of mean earnings), higher college attendance rates, and higher adult '
    'salaries. Replacing a bottom-5% teacher with an average teacher increases the present value of students\' '
    'lifetime income by approximately $250,000 per classroom.'
)

add_body(
    'Hattie\'s Visible Learning meta-meta-analysis (2009, updated to 1,400+ meta-analyses) ranks '
    'teacher-student relationships at an effect size of 0.72, well above the 0.40 "hinge point" of one '
    'year\'s academic growth. Teacher estimates of achievement have the highest effect size at 1.62. '
    'Borman and Dowling (2008) found that high teacher turnover negatively affects student achievement '
    'even for students whose own teacher did not leave, and that school organizational characteristics '
    '(administrative support, student discipline, collaboration, professional development) are key '
    'moderators of teacher retention. These findings make HR operations, specifically the systems that '
    'recruit, retain, and support effective teachers, arguably the highest-leverage operational investment '
    'a district can make.'
)

doc.add_heading('2.6 Per-Pupil Spending', level=2)

add_body(
    'The research consensus on school spending has shifted dramatically. Jackson, Johnson, and Persico '
    '(2016) used school finance reform timing as an exogenous spending shifter and found that a 10% increase '
    'in per-pupil spending for all 12 years of public school leads to 0.27 more completed years of education, '
    '7.25% higher wages, and a 3.67 percentage-point reduction in adult poverty. Effects are much more '
    'pronounced for children from low-income families: 0.43 additional years of education, 9.5% higher '
    'earnings, and a 6.8 percentage-point reduction in adult poverty.'
)

add_body(
    'Baron (2022) found that increasing operational spending specifically by $1,000 per pupil increased '
    'test scores by approximately 0.15 SD and graduation rates by approximately 9 percentage points. '
    'Lafortune, Rothstein, and Schanzenbach (2018) found that a 20% increase in per-pupil spending for '
    'low-income children over 12 years increased educational attainment by a full year. Jackson\'s (2018) '
    'comprehensive review for NBER concluded that "the question of whether money matters is essentially '
    'settled," with a majority of rigorous causal studies finding positive effects.'
)

doc.add_heading('2.7 School Nursing and Health Services', level=2)

add_body(
    'Best et al. (2017) found that comprehensive school nursing services were associated with increased '
    'academic achievement grades and decreased absenteeism and academic procrastination behaviors. However, '
    'Maughan (2018) noted that the presence of a school nurse is associated with reduced absenteeism and '
    'missed class time but not consistently with direct academic achievement gains. The primary pathway '
    'from nursing to achievement runs through attendance: school nurses reduce absenteeism through chronic '
    'condition management, immunization compliance, early illness intervention, and mental health identification. '
    'The research base is described as "weak" methodologically, with most studies being descriptive rather '
    'than experimental (Bohnenkamp et al., 2019).'
)

doc.add_heading('2.8 Transportation', level=2)

add_body(
    'Cordes, Rick, and Schwartz (2022) found that long bus rides have deleterious effects on attendance '
    'and chronic absenteeism. Hemelt et al. (2024) found that school bus eligibility increases attendance '
    'rates for economically disadvantaged students by 0.63 percentage points (approximately 1 additional '
    'day per 180-day school year) but does not have statistically significant effects on student achievement '
    'directly. Every one-minute increase in school commute time is associated with a 1.3-minute reduction '
    'in sleep (Lutz et al., 2024). Transportation primarily affects achievement indirectly through attendance '
    'and school access, with effects concentrated among disadvantaged students.'
)

doc.add_heading('2.9 School Nutrition', level=2)

add_body(
    'Cohen et al. (2024) published a systematic review in JAMA Network Open finding that universal free '
    'school meals were associated with increased lunch participation and modestly improved attendance, though '
    'no included studies directly assessed academic achievement. Community Eligibility Provision (CEP) '
    'adoption led to approximately 0.06 SD increase in math test scores for elementary students in South '
    'Carolina (Gordanier et al., 2020). The School Breakfast Program has been associated with 17.5% higher '
    'math scores and 1.5 more school days per year among participants, with NAEP data showing +2.2 points '
    'in math and +2.0 points in reading (Frisvold, 2015). Long-term effects include 4 additional months '
    'of educational attainment for women and 1 year for men (NSLP participation).'
)

# Summary table
doc.add_heading('2.10 Summary: Effect Sizes and Evidence Quality', level=2)

add_table(
    ['Operational Area', 'Key Effect Size(s)', 'Evidence Quality', 'Primary Pathway'],
    [
        ['Air Quality/Ventilation', '0.07-0.15 SD; up to 0.5 SD', 'Strong (quasi-experimental)', 'Direct cognitive impact'],
        ['Temperature', '1% learning decrease per 1 F', 'Moderate-Strong', 'Cognitive performance'],
        ['Lighting/Daylight', '20-26% faster progress', 'Moderate (correlational)', 'Cognitive performance'],
        ['Building Condition', '3-17 percentile point gap', 'Moderate (correlational)', 'Health, attendance, cognition'],
        ['Teacher Quality (HR)', 'ES 0.72; $250K lifetime/class', 'Very Strong (causal)', 'Direct instruction'],
        ['Per-Pupil Spending', '0.15 SD per $1,000', 'Strong (causal)', 'Staffing, resources'],
        ['School Nursing', 'Reduced absenteeism', 'Weak-Moderate', 'Attendance'],
        ['Transportation', '0.63 pp attendance gain', 'Moderate (quasi-experimental)', 'Attendance, access'],
        ['Nutrition/Meals', '0.06 SD math (CEP)', 'Moderate (mixed methods)', 'Attendance, health'],
        ['Custodial/Maintenance', 'Up to 40% illness reduction', 'Weak-Moderate', 'Attendance, health'],
    ]
)

# ============================================================
# SECTION 3: FEDERAL & STATE POLICY CONTEXT
# ============================================================
doc.add_heading('3. Federal and State Policy Context', level=1)

doc.add_heading('3.1 Federal Framework', level=2)

add_body(
    'Federal policy creates a multi-layered framework that simultaneously mandates operational standards '
    'and constrains how districts allocate resources. The Every Student Succeeds Act (ESSA), signed in 2015, '
    'requires all states to report per-pupil expenditures, creating transparency around operational spending. '
    'Title IV, Part A (Student Support and Academic Enrichment) explicitly funds safe and healthy school '
    'conditions, recognizing that operational conditions affect student outcomes. LEAs with allocations over '
    '$30,000 must spend at least 20% on safe and healthy students.'
)

add_body(
    'The Individuals with Disabilities Education Act (IDEA) defines both transportation and health services '
    'as "related services" that districts must provide as part of a student\'s IEP. Section 504 and the ADA '
    'require all school facilities to meet specific accessibility standards. Title VI and Title IX have been '
    'used by the Office for Civil Rights to address inequitable facility conditions and athletics facilities. '
    'The McKinney-Vento Act requires districts to provide transportation to and from the school of origin '
    'for students experiencing homelessness.'
)

add_body(
    'Federal nutrition programs serve roughly 60% of American schoolchildren daily. The Healthy, Hunger-Free '
    'Kids Act of 2010 updated nutrition standards with documented results: students consumed 16% more vegetables '
    'and 23% more fruit, and the mean Healthy Eating Index score among low-income participants rose from 42.7 '
    'to 54.6. The E-Rate program provides $3.9 billion annually to help schools obtain affordable broadband. '
    'ESSER/ARP pandemic relief provided $122 billion (U.S. Department of Education, 2021), with allowable uses explicitly including HVAC improvements '
    'to improve indoor air quality.'
)

add_body(
    'The critical gap in federal policy is facilities construction. The federal government provides only '
    'about 0.2% of capital costs for school construction, with states providing 18% and local governments '
    'bearing the remainder (Center on Budget and Policy Priorities, 2022). National spending on K-12 buildings '
    'falls short by an estimated $90 billion annually (21st Century School Fund, 2025). Proposed legislation '
    'like the Rebuild America\'s Schools Act ($130 billion over five years) has not been enacted.'
)

doc.add_heading('3.2 Tennessee Policy Framework', level=2)

add_body(
    'Tennessee\'s education policy landscape is shaped by several major frameworks that connect to school '
    'operations. The Tennessee Investment in Student Achievement (TISA) Act, signed May 2, 2022, replaced '
    'the Basic Education Program (BEP) beginning in 2023-2024. TISA shifted from a resource-based formula '
    'to a student-based formula with base funding of $6,860 per student (2023-24), rising to $7,075 in '
    '2024-25, with additional weights for economically disadvantaged students, students with disabilities, '
    'rural districts, and English learners. In the first year, the General Assembly invested an additional '
    '$1.16 billion in K-12 education, a 21.6% budget increase. All locally administered districts received '
    'more state money under TISA than under the BEP.'
)

add_body(
    'However, TISA funds instructional and weighted needs while the bulk of capital spending on school '
    'facilities and related debt payments is paid from local revenues, including bonds, adequate facilities '
    'taxes, and dedicated property taxes. This creates inequity between property-wealthy and property-poor '
    'districts. The Tennessee Comptroller\'s November 2024 report revealed staggering statewide infrastructure '
    'needs.'
)

add_table(
    ['Metric', 'Value'],
    [
        ['Total statewide infrastructure needs (5-year)', '$9.8 billion'],
        ['New construction needs', '$4.1 billion'],
        ['Renovation needs', '$5.7 billion'],
        ['Buildings in good or excellent condition', '57%'],
        ['Counties with "monetarily significant" needs', '57%'],
        ['Average cost of new school (2022)', '$51 million'],
        ['Average cost of new school (2012)', '$15 million'],
        ['Buildings in poor condition', '18 (including 11 in Nashville)'],
    ]
)

add_body(
    'Tennessee has one of the least stringent school nurse requirements in the nation, with a '
    'nurse-to-student ratio of 1:3,000 compared to the National Association of School Nurses recommended '
    'ratio of 1:750. Districts determine whether to employ additional school nurses beyond the state minimum '
    'and set their own salary amounts.'
)

add_body(
    'Tennessee\'s chronic absenteeism rate reached 20.3% in 2021-22, an increase of 7 percentage points '
    'since 2019. The Comptroller\'s 2023 report identified transportation instability, housing instability, '
    'and health needs (lack of nursing/health services) as key operational barriers to attendance.'
)

add_body(
    'Teacher workforce challenges are significant. Elementary vacancy rates doubled from 0.6% to 1.4% '
    'between 2023 and 2024. Special education (3.6%), world languages (3.3%), and ESL (2.9%) vacancy rates '
    'all exceed the 2.5% critical shortage threshold. One in five teachers leave within their first three '
    'years, and 74% of school leaders reported dissatisfaction with applicant pools in 2024, up from 56% in '
    '2021. Students at "F" letter grade schools are 7 times more likely to be taught by teachers on emergency '
    'credentials than students at "A" schools.'
)

add_body(
    'The Tennessee Educator Acceleration Model (TEAM) evaluation system connects teacher effectiveness to '
    'student outcomes, with 50% of evaluations based on student achievement data. Operational support '
    '(adequate facilities, materials, staffing) directly affects teacher working conditions and, by extension, '
    'observation performance and retention. The RTI2 framework requires adequate staffing for intervention '
    'specialists, appropriate physical spaces, materials, technology, and data systems, all operational '
    'infrastructure that falls under school operations.'
)

# ============================================================
# SECTION 4: WHAT WORKS - IMPLEMENTATION MODELS
# ============================================================
doc.add_heading('4. What Works: Implementation Models', level=1)

doc.add_heading('4.1 Facility Renovations and New Construction', level=2)

add_body(
    'Neilson and Zimmerman (2014) studied New Haven, Connecticut\'s district-wide school construction program '
    'and found that by six years after building occupancy, school construction increased reading scores by '
    '0.15 standard deviations. School construction also raised home prices by roughly 10% and increased public '
    'school enrollment. Maxwell (1999) examined 21 renovated Syracuse schools and found math score improvements '
    'for 3rd and 6th graders post-renovation. In Los Angeles, improvements raising a school\'s environmental '
    'compliance rating from "worst" to "best" correlated with a 36-point increase in the Academic Performance '
    'Index, a nearly 6% increase.'
)

add_body(
    'A key lesson from the case studies is that facility investment effects take time to materialize. '
    'Reading gains from construction and renovation take 4-6 years to fully appear. However, these investments '
    'also produce community-level benefits (enrollment increases, property value gains) that sustain political '
    'support for ongoing operational investment.'
)

doc.add_heading('4.2 Indoor Air Quality Interventions', level=2)

add_body(
    'One mid-size K-12 district (28 buildings) installed CO2 sensors in every classroom after a parent '
    'complaint. Most fixes were simple: stuck dampers and clogged filters. Within 60 days, nurse visits '
    'for headaches dropped 38%, and teachers reported noticeably better afternoon engagement. This illustrates '
    'that low-cost monitoring can identify high-impact maintenance failures. Rough calculations suggest IAQ '
    'renovations may be a more cost-effective way to improve standardized test scores than class size reductions.'
)

doc.add_heading('4.3 Transportation Optimization', level=2)

add_body(
    'Boston Public Schools implemented the BiRD (Bi-objective Routing Decomposition) algorithm developed by '
    'MIT researchers. Results: 50 buses removed from the fleet (7%), $5 million in annual savings, and no '
    'increase in average student walking or riding times. The reform also led the School Committee to '
    'unanimously approve the first comprehensive start time reform in 30 years (Bertsimas, Delarue, & Martin, '
    '2019). Oakland, California reduced routes lasting over an hour from 70% to 10% within one year of '
    'implementing optimization software. Recent optimization approaches have reduced rural bus ride times '
    'by 37-39%.'
)

doc.add_heading('4.4 Nutrition Program Expansion', level=2)

add_body(
    'Community Eligibility Provision (CEP) adoption shows consistent improvements. In South Carolina, CEP '
    'led to approximately 0.06 SD increase in elementary math scores with increased participation, improved '
    'attendance, and reduced middle school disciplinary referrals. Wisconsin\'s School Breakfast Program '
    'implementation was associated with a 3.5 percentage-point reduction in low attendance and a 0.08 SD '
    'increase in reading scores. Evidence is strongest when universal free meal provisions include lunch, '
    'not just breakfast, and effects are largest for elementary and high-poverty populations.'
)

doc.add_heading('4.5 Staffing and HR Reforms', level=2)

add_body(
    'Kreisman and Steinberg (2019) found that an additional $1,000 per year in base funding (roughly 10% '
    'expenditure increase) improved reading scores by 0.1 SD and math scores by more than 0.07 SD in Texas '
    'small districts. High school dropout rates decreased by nearly 2 percentage points. The funding was '
    'largely directed toward staffing improvements. California\'s LCFF directed additional resources to '
    'high-need districts, with roughly 84% of the variation in spending effectiveness explained by class size '
    'reductions, teacher salary increases, and reductions in teacher turnover.'
)

add_body(
    'Compensation plays a particularly significant role in attrition rates for small school districts, '
    'as they struggle to match salaries of larger nearby districts. However, financial incentives alone '
    'are insufficient for retention. Working conditions, professional support, and management practices '
    'are critical complements.'
)

doc.add_heading('4.6 Preventive Maintenance Programs', level=2)

add_body(
    'Districts that shifted from reactive to preventive maintenance cut total maintenance costs 18-30% '
    'without new budget lines (NCES, 2003). Every dollar moved from reactive to preventive work generates '
    '$2.50-$4.00 in avoided cost. Without a maintenance plan, 40% of public schools risk unplanned '
    'disruptions and escalating costs. Shifting even 20% of maintenance work from reactive to planned '
    'reduces total spend by 12-18% with zero additional budget.'
)

# Implementation summary table
doc.add_heading('4.7 Implementation Summary', level=2)

add_table(
    ['District/Study', 'Intervention', 'Key Results', 'Timeline'],
    [
        ['New Haven, CT', 'School construction/renovation', '+0.15 SD reading, +10% home prices', '6 years'],
        ['28-building K-12 district', 'CO2 sensors + HVAC fixes', '-38% headache nurse visits', '60 days'],
        ['South Carolina (CEP)', 'Universal free meals', '+0.06 SD elementary math', 'Implementation period'],
        ['Wisconsin (SBP)', 'Universal free breakfast', '-3.5 pp low attendance, +0.08 SD reading', 'Implementation period'],
        ['Boston, MA', 'Bus routing optimization', '$5M savings, 50 fewer buses', 'Fall 2017'],
        ['Oakland, CA', 'Route optimization software', '70% to 10% routes over 1 hour', '1 year'],
        ['Texas small districts', '$1,000/student increase', '+0.1 SD reading, +0.07 SD math', 'Cumulative'],
        ['California (LCFF)', 'Weighted student funding', 'Improved achievement, reduced gaps', 'Multi-year'],
    ]
)

# ============================================================
# SECTION 5: COST-BENEFIT ANALYSIS
# ============================================================
doc.add_heading('5. Cost-Benefit Analysis', level=1)

doc.add_heading('5.1 Facilities and Maintenance', level=2)

add_body(
    'The U.S. faces approximately $300 billion in deferred maintenance to bring schools to "good" condition '
    '(Center on Budget and Policy Priorities, 2022). '
    'Renovation costs range from $50 to $200 per square foot, with new K-12 construction averaging $221-$239 '
    'per square foot nationally. While Cellini, Ferreira, and Rothstein (2010) found that locally financed '
    'capital campaigns produce modest achievement effects, the ROI is highest when addressing the worst '
    'conditions first, particularly in under-resourced districts.'
)

add_table(
    ['Maintenance Approach', 'Cost per Sq Ft', 'Equipment Lifespan (HVAC)', 'Long-term Savings'],
    [
        ['Preventive', '$1-$2', '15+ years', '30-40% lower total costs'],
        ['Reactive', '$3-$5', '~8 years', 'Baseline (highest cost)'],
        ['Predictive', '$1.50-$2.50', '18+ years', '12-18% beyond preventive'],
    ]
)

doc.add_heading('5.2 School Nursing', level=2)

add_body(
    'Wang et al. (2014) in JAMA Pediatrics established the cost-benefit framework for school nursing. '
    'A full-time school nurse generates $1.59 in benefits for every $1 invested, while a part-time nurse '
    'generates $1.29. Benefits are measured as savings in medical procedure costs, teachers\' productivity '
    'loss (addressing health issues instead of teaching), and parents\' productivity loss (early dismissals, '
    'medication administration). Multi-year full-time nursing produces increasing returns of $1.50-$1.67 '
    'annually as cumulative gains in attendance and self-management compound. One important caveat: a study '
    'on LPNs found incremental costs of $68,228 per school with negligible attendance improvements, suggesting '
    'that nursing credential level matters for outcomes.'
)

doc.add_heading('5.3 Transportation', level=2)

add_body(
    'National per-pupil transportation expenditure is $1,152 per student transported. Rural districts spend '
    '40% more per student than urban districts. Route optimization offers significant savings: Boston Public '
    'Schools saved $5 million annually through the BiRD algorithm while maintaining service quality. '
    'Optimization can reduce driving distance by 23%, driving time by 30%, and operating costs by 26%. '
    'Each eliminated route saves approximately $37,000 per year.'
)

doc.add_heading('5.4 Teacher Turnover', level=2)

add_body(
    'The Learning Policy Institute\'s 2024 update provides current per-teacher replacement costs: $11,860 '
    'for small districts (under 10,000 students), $16,450 for medium districts, and $24,930 for large '
    'districts. Training costs account for approximately 87% of total cost per leaver. Nationally, teacher '
    'turnover costs an estimated $2.2 billion annually. A district with 1,000 teachers and 15% annual '
    'turnover faces $3.74 million in annual turnover costs. For a small district like GCS, retaining even '
    'a handful of additional teachers each year could save $50,000-$100,000 while simultaneously improving '
    'instructional continuity.'
)

doc.add_heading('5.5 HVAC and Indoor Air Quality', level=2)

add_body(
    'HVAC improvements represent one of the highest-ROI operational investments with relatively low costs '
    'and measurable academic gains. Proper ventilation versus poor conditions yields 3-8% improvement in '
    'standardized test scores. Cooling system replacement produces +3% of a standard deviation in math; '
    'heating system replacement produces +4%. Payback periods for advanced filtration are typically 12-24 '
    'months. However, 41% of school districts need to update or replace HVAC systems in at least half '
    'their schools (U.S. Joint Economic Committee, 2024), with a national estimated cost of approximately '
    '$72 billion (Center for American Progress, 2022).'
)

doc.add_heading('5.6 Nutrition Programs', level=2)

add_body(
    'School breakfast participation is associated with 17.5% higher math scores and 1.5 more school days '
    'per year (Frisvold, 2015). Each 1% reduction in absenteeism can recover significant per-student funding. Universal '
    'free meal programs reduce food insecurity and behavioral incidents, with the greatest decreases among '
    'food-insecure households. The long-term returns are substantial: NSLP participation is associated '
    'with 4 additional months of educational attainment for women and 1 year for men.'
)

# ROI summary table
doc.add_heading('5.7 ROI Summary by Operational Area', level=2)

add_table(
    ['Operational Area', 'Estimated ROI', 'Evidence Strength', 'Implementation Complexity'],
    [
        ['Preventive maintenance shift', '$2.50-$4.00 per $1', 'Strong', 'Low-Medium'],
        ['School nursing (full-time RN)', '$1.59 per $1', 'Strong', 'Low'],
        ['Route optimization', '23-30% cost reduction', 'Strong (case studies)', 'Medium'],
        ['HVAC/IAQ improvements', '12-24 month payback', 'Strong', 'Medium-High'],
        ['School breakfast/nutrition', 'High academic gains', 'Strong', 'Low'],
        ['Teacher retention programs', '$11,860-$24,930 saved per teacher', 'Strong', 'Medium'],
        ['Per-pupil spending increases', '0.15 SD per $1,000', 'Strong (causal)', 'High (policy-dependent)'],
    ]
)

# ============================================================
# SECTION 6: CRITICISMS, LIMITATIONS & RISKS
# ============================================================
doc.add_heading('6. Criticisms, Limitations, and Risks', level=1)

doc.add_heading('6.1 The Hanushek Critique: Money Does Not Guarantee Results', level=2)

add_body(
    'The most enduring counterargument comes from economist Eric Hanushek (Stanford/Hoover Institution), '
    'whose decades of research have argued that there is no strong or consistent relationship between student '
    'performance and school resources after controlling for family inputs. He analyzed close to 400 studies '
    'and concluded that input-based schooling policies fail to reliably produce achievement gains. His central '
    'claim is not that resources are irrelevant, but that "just putting more money into schools is unlikely '
    'to give us very good results."'
)

add_body(
    'Critics have challenged Hanushek\'s methodology. Hedges, Laine, and Greenwald (1994) re-analyzed his study '
    'sample and found that expenditures do impact achievement. More recent quasi-experimental research '
    '(Jackson, 2018; Lafortune et al., 2018) has generally found positive spending effects. Hanushek himself '
    'acknowledges that a majority of rigorous studies now find positive effects, but maintains that the '
    'relationship is contingent on how money is spent, not simply how much is spent.'
)

doc.add_heading('6.2 Facility Investments: Modest and Inconsistent Returns', level=2)

add_body(
    'Multiple studies have found that school construction produces disappointingly small academic effects. '
    'A rigorous regression discontinuity analysis found "very precise zero estimates of achievement effects," '
    'with only 0.016 and 0.030 SD increases for reading and math. A $12 billion school construction initiative '
    'in New York City yielded only about a 1 percentage-point increase in attendance after three years and no '
    'significant change in test scores. These findings suggest that facility improvements alone, without '
    'attention to the instructional and social environment within the building, are insufficient.'
)

doc.add_heading('6.3 The Threshold/Adequacy Argument', level=2)

add_body(
    'A compelling middle-ground counterargument holds that operational conditions matter primarily up to a '
    'threshold of basic adequacy, after which additional investments yield diminishing returns. Research from '
    'seven states found that operational spending effects are driven by districts below the median in spending '
    '(Baron, 2022). Jackson and Mackevicius (2021) found that spending returns are higher among areas with low previous '
    'investment. The implication: for districts already meeting basic standards, additional operational spending '
    'may produce negligible academic gains. Operational investments should be targeted rather than universal.'
)

doc.add_heading('6.4 Opportunity Cost: Instruction vs. Operations', level=2)

add_body(
    'About 60-61% of current expenditures nationally go to instruction and 35-36% to support services. '
    'Teacher quality is the most important school-related factor influencing achievement, creating a '
    'fundamental opportunity cost argument: every dollar spent on non-instructional operations is a dollar '
    'not spent on teacher recruitment, retention, professional development, or reduced class sizes. Critics '
    'argue that tracking money "into the school building to the classroom" is typically missing from financial '
    'analysis, and that the ratio should shift further toward instruction.'
)

doc.add_heading('6.5 Equity Concerns', level=2)

add_body(
    'School facility funding is heavily dependent on local property wealth, creating systematic inequities. '
    'Districts in the top 20% of assessed property wealth have facility revenue almost twice as high as '
    'districts in the bottom 20%. Districts in the 75th percentile receive more than 3.5 times the GO bond '
    'revenue per pupil as districts in the 25th percentile. Without addressing this funding structure, '
    'increased operational investment may disproportionately benefit affluent districts. High-poverty '
    'districts already spend $300 (30%) less per student on capital construction.'
)

doc.add_heading('6.6 Methodological Weaknesses', level=2)

add_body(
    'The facilities-achievement research suffers from significant limitations: selection bias (students in '
    'better facilities often come from wealthier communities), confounding variables (the inclusion or '
    'exclusion of SES as a control "drastically alters results"), aggregation problems (varying levels of '
    'analysis produce different results), inconsistent definitions across studies, and potential publication '
    'bias. These limitations mean that many observed facility-achievement correlations may overstate the '
    'true causal relationship.'
)

doc.add_heading('6.7 Conflicting Findings', level=2)

add_body(
    'Several areas show genuinely conflicting evidence. Facility construction studies range from near-zero '
    'effects (Cellini et al., 2010) to meaningful gains (Neilson & Zimmerman, 2014), likely reflecting '
    'differences in baseline conditions and study contexts. School breakfast programs show positive effects '
    'in some districts but a rigorous Philadelphia trial found no improvement and minor adverse math effects. '
    'School nursing research finds consistent attendance benefits but inconsistent achievement effects. '
    'These conflicts suggest that context, implementation quality, and baseline conditions are critical '
    'moderators that determine whether operational investments produce academic returns.'
)

# ============================================================
# SECTION 7: MEASURABLE OUTCOMES
# ============================================================
doc.add_heading('7. Measurable Outcomes', level=1)

add_body(
    'The following table synthesizes specific measurable outcomes linked to school operations investments, '
    'organized by outcome type and strength of evidence.'
)

add_table(
    ['Outcome', 'Operational Driver', 'Measured Effect', 'Evidence Strength'],
    [
        ['Standardized test scores', 'Air quality/ventilation', '0.07-0.15 SD improvement', 'Strong'],
        ['Standardized test scores', 'Per-pupil spending (+$1,000)', '0.15 SD improvement', 'Strong'],
        ['Standardized test scores', 'Facility renovation', '0.015-0.15 SD (varies)', 'Mixed'],
        ['Standardized test scores', 'CEP/universal meals', '0.02-0.06 SD math', 'Moderate'],
        ['Attendance', 'School bus eligibility', '+0.63 pp (disadvantaged)', 'Moderate'],
        ['Attendance', 'School nursing', 'Reduced chronic absenteeism', 'Moderate'],
        ['Attendance', 'Universal breakfast', '-3.5 pp low attendance', 'Moderate'],
        ['Attendance', 'Building condition', '+4-5 students per 1,000 ADA', 'Moderate'],
        ['Graduation rates', 'Operational spending (+$1,000)', '+9 percentage points', 'Strong'],
        ['Graduation rates', 'Building condition', '-10-13 dropouts per 1,000', 'Moderate'],
        ['Adult earnings', 'Per-pupil spending (+10%)', '+7.25% wages', 'Strong'],
        ['Adult earnings', 'Teacher quality (+1 SD VA)', '+$182 at age 28', 'Very Strong'],
        ['Adult poverty', 'Per-pupil spending (+10%)', '-3.67 pp', 'Strong'],
        ['College attendance', 'Teacher quality', 'Increased rates', 'Strong'],
        ['Behavioral incidents', 'Universal free meals', 'Reduced suspensions', 'Moderate'],
        ['Teacher retention', 'Working conditions', 'Key moderator of turnover', 'Strong'],
        ['Home prices', 'School construction', '+10% in affected neighborhoods', 'Moderate'],
    ]
)

add_body(
    'Two patterns emerge. First, the strongest and most direct links to academic achievement run through '
    'teacher quality (an HR function) and per-pupil spending (a finance function). Second, most other '
    'operational areas, including nursing, transportation, nutrition, and facilities, affect achievement '
    'primarily through attendance. This makes chronic absenteeism a key intermediate metric for evaluating '
    'operational investments. Districts should track attendance data alongside achievement data when assessing '
    'the return on operational investments.'
)

# ============================================================
# SECTION 8: BOTTOM LINE
# ============================================================
doc.add_heading('8. Bottom Line', level=1)

add_body(
    'The evidence is clear: school operations matter for student achievement, but not equally and not '
    'unconditionally. The strongest causal links run through teacher quality and per-pupil spending, both '
    'of which are fundamentally operational functions (HR and finance). Indoor environmental quality, '
    'particularly air quality and temperature, has surprisingly robust evidence for direct cognitive impacts. '
    'Most other operational areas, including nursing, transportation, nutrition, and custodial services, '
    'affect achievement primarily through attendance, making them equity investments that disproportionately '
    'benefit disadvantaged students.'
)

add_body(
    'The counterarguments deserve serious weight. Facility construction alone produces modest academic effects '
    'unless conditions were genuinely substandard to begin with. Operational spending faces diminishing returns '
    'above a threshold of basic adequacy. And every operational dollar carries an opportunity cost against '
    'direct instructional investment. The evidence does not support blanket increases in operational spending; '
    'it supports targeted investments in specific areas with demonstrated returns, particularly for districts '
    'below the adequacy threshold.'
)

add_body(
    'For a district like Greeneville City Schools, already performing above state averages with per-pupil '
    'spending approximately $600 above the state median and Exemplary District designation, the evidence '
    'points toward several high-priority strategies. First, invest in indoor air quality monitoring and '
    'HVAC maintenance, the highest-ROI, lowest-complexity intervention available. Second, shift maintenance '
    'practices from reactive to preventive, generating $2.50-$4.00 in avoided cost per dollar shifted. '
    'Third, advocate for improved school nurse staffing (Tennessee\'s 1:3,000 ratio is four times worse '
    'than the recommended 1:750). Fourth, prioritize teacher retention through working conditions, '
    'professional support, and competitive compensation, recognizing that each retained teacher saves '
    '$11,860 and preserves instructional continuity. Fifth, ensure full participation in federal nutrition '
    'programs, including CEP where eligible, as a low-cost, evidence-based attendance and equity intervention.'
)

add_body(
    'Questions that remain unanswered include the precise threshold at which operational adequacy is '
    '"good enough" for achievement purposes, the optimal allocation ratio between instructional and '
    'operational spending for different district profiles, and the long-term academic effects of the '
    'post-ESSER decline in facilities funding. Tennessee\'s $9.8 billion infrastructure gap and the '
    'expiration of ESSER funds in January 2025 make these questions increasingly urgent for district '
    'leaders across the state.'
)

# ============================================================
# SECTION 9: REFERENCES
# ============================================================
doc.add_heading('9. References', level=1)

references = [
    'Adolphus, K. et al. (2022). Breakfast and School-Related Outcomes in Children and Adolescents in the US: A Literature Review. PMC.',
    'Allen, J.G. et al. (2016). Associations of Cognitive Function Scores with Carbon Dioxide, Ventilation, and Volatile Organic Compound Exposures in Office Workers. Environmental Health Perspectives, 124(6), 805-812.',
    'Baron, E.J. (2022). School District Operational Spending and Student Outcomes: Evidence from Tax Elections in Seven States. Journal of Public Economics.',
    'Bertsimas, D., Delarue, A., & Martin, S. (2019). Optimizing Schools\' Start Time and Bus Routes. PNAS, 116(13).',
    'Best, N.C., Oppewal, S., & Travers, D. (2017). The Impact of Comprehensive School Nursing Services on Students\' Academic Performance. Journal of School Nursing, PMC5348663.',
    'Bohnenkamp, J.H. et al. (2019). School Nurses and Student Academic Outcomes: An Integrative Review. Journal of School Nursing.',
    'Borman, G.D. & Dowling, N.M. (2008). Teacher Attrition and Retention: A Meta-Analytic and Narrative Review. Review of Educational Research.',
    'Cellini, S.R., Ferreira, F., & Rothstein, J. (2010). The Value of School Facility Investments: Evidence from a Dynamic RDD. Journal of Public Economics.',
    'Center for American Progress. (2022). School Air Filtration and Ventilation Strategies to Improve Health, Education Equity, and Environmental Outcomes.',
    'Center on Budget and Policy Priorities. (2022). America\'s School Infrastructure Needs a Major Investment of Federal Funds.',
    'Chetty, R., Friedman, J.N., & Rockoff, J.E. (2014). Measuring the Impacts of Teachers II: Teacher Value-Added and Student Outcomes in Adulthood. American Economic Review, 104(9), 2633-2679.',
    'Cohen, J.F.W. et al. (2021). Universal School Meals and Associations with Student Participation, Attendance, Academic Performance, Diet Quality, Food Security, and Body Mass Index. Nutrients, PMC8000006.',
    'Cohen, J.F.W. et al. (2024). Universal Free School Meals and School and Student Outcomes: A Systematic Review. JAMA Network Open.',
    'Cordes, S.A., Rick, C., & Schwartz, A.E. (2022). Do Long Bus Rides Drive Down Academic Outcomes? Educational Evaluation and Policy Analysis.',
    'EdTrust Tennessee. (2024). The State of Tennessee\'s Teacher Workforce in 2023-2024.',
    'Frisvold, D. (2015). Nutrition and Cognitive Achievement: An Evaluation of the School Breakfast Program. Journal of Public Economics.',
    'Gordanier, J. et al. (2020). Impact of the Community Eligibility Provision on Student Nutrition, Behavior, and Academic Outcomes. American Journal of Public Health.',
    'Hanushek, E.A. (2003). The Failure of Input-Based Schooling Policies. Economic Journal.',
    'Hattie, J. (2009). Visible Learning: A Synthesis of Over 800 Meta-Analyses Relating to Achievement. Routledge.',
    'Hedges, L.V., Laine, R.D., & Greenwald, R. (1994). Does Money Matter? A Meta-Analysis of Studies of the Effects of Differential School Inputs on Student Outcomes. Educational Researcher, 23(3), 5-14.',
    'Hemelt, S.W. et al. (2024). Another One Rides the Bus: The Impact of School Transportation on Student Outcomes in Michigan. Journal of Education Finance and Policy, 19(1).',
    'Heschong Mahone Group. (1999). Daylighting in Schools: An Investigation into the Relationship Between Daylighting and Human Performance.',
    'Jackson, C.K. (2018). Does School Spending Matter? The New Literature on an Old Question. NBER Working Paper 25368.',
    'Jackson, C.K., Johnson, R.C., & Persico, C. (2016). The Effects of School Spending on Educational and Economic Outcomes. Quarterly Journal of Economics, 131(1), 157-218.',
    'Jackson, C.K. & Mackevicius, C. (2021). The Distribution of School Spending Impacts. NBER Working Paper 28517.',
    'Kreisman, D. & Steinberg, M. (2019). The Effect of Increased Funding on Student Achievement: Evidence from Texas. Journal of Public Economics.',
    'Lafortune, J., Rothstein, J., & Schanzenbach, D. (2018). School Finance Reform and the Distribution of Student Achievement. American Economic Journal.',
    'Learning Policy Institute. (2024). What\'s the Cost of Teacher Turnover? 2024 Update.',
    'Leyten, J.L. & Kurvers, S.R. (2021). Classrooms\' Indoor Environmental Conditions Affecting Academic Achievement: A Systematic Literature Review. PMC.',
    'Lutz, B. et al. (2024). Examining the Impacts of School Bus Travel on Students\' Academic Performance. Canadian Geographer.',
    'Maughan, E.D. (2018). School Nurses: An Investment in Student Achievement. Phi Delta Kappan.',
    'Maxwell, L. (1999). School Building Renovation and Student Performance: One District\'s Experience. ERIC.',
    'NCES. (2003). Planning Guide for Maintaining School Facilities. U.S. Department of Education.',
    'Neilson, C. & Zimmerman, S. (2014). The Effect of School Construction on Test Scores, School Enrollment, and Home Prices. Journal of Public Economics.',
    'Park, R.J. et al. (2020). Heat and Learning. American Economic Journal: Economic Policy, 12(2), 306-339.',
    'Roth, S. (2018). The Effect of Indoor Air Pollution on Cognitive Performance. LSE Working Paper.',
    'Stafford, T.M. (2015). Indoor Air Quality and Academic Performance. Journal of Environmental Economics and Management.',
    'Stewart, J. (2014). School Building Condition and Student Academic Achievement. eScholarship, UC Berkeley.',
    'Tennessee Comptroller, OREA. (2024). An Overview of K-12 Capital Infrastructure and Investment in Tennessee and Other States.',
    'Tennessee Comptroller, OREA. (2025). Tennessee Investment in Student Achievement: First-Year Report.',
    'Tennessee Department of Education. (2024). TISA Formula Overview.',
    'Tennessee SCORE. (2024). Strategic School Staffing: Tennessee\'s Opportunity.',
    '21st Century School Fund. (2025). State of Our Schools Report.',
    'U.S. Department of Education. (2021). Elementary and Secondary School Emergency Relief Fund.',
    'U.S. EPA. (n.d.). Indoor Air Quality Tools for Schools. https://www.epa.gov/iaq-schools.',
    'U.S. Joint Economic Committee. (2024). Improving School Infrastructure Benefits Students, the Economy, and the Environment.',
    'Wang, L.Y. et al. (2014). Cost-Benefit Study of School Nursing Services. JAMA Pediatrics.',
]

for ref in references:
    p = doc.add_paragraph(ref, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)

# ============================================================
# APPENDIX A: TENNESSEE DATA
# ============================================================
doc.add_heading('Appendix A: Tennessee Data', level=1)

doc.add_heading('A.1 Greeneville City Schools Profile', level=2)

add_table(
    ['Metric', 'Greeneville City Schools', 'State Average/Comparison'],
    [
        ['Enrollment', '2,996 students', '-'],
        ['Per-pupil spending', '$12,085', '~$11,478 (state median)'],
        ['Math proficiency', '40%', '34%'],
        ['Reading proficiency', '44%', '37%'],
        ['Number of schools', '8', '-'],
        ['2024-25 designation', 'Exemplary District (1 of 14)', '-'],
        ['Letter grades', 'All A (EV, TV, GMS, GHS)', '-'],
    ]
)

doc.add_heading('A.2 Comparable District: Greene County', level=2)

add_table(
    ['Metric', 'Greene County', 'Greeneville City'],
    [
        ['Enrollment', '6,121', '2,996'],
        ['Per-pupil spending', '$10,281', '$12,085'],
        ['Reading proficiency', '31%', '44%'],
        ['Math proficiency', '40%', '40%'],
        ['Per-pupil difference', '-$1,804 vs. GCS', 'Baseline'],
    ]
)

doc.add_heading('A.3 Tennessee Infrastructure Needs', level=2)

add_table(
    ['Metric', 'Value'],
    [
        ['Total 5-year infrastructure needs', '$9.8 billion'],
        ['New construction needs', '$4.1 billion'],
        ['Renovation needs', '$5.7 billion'],
        ['Buildings in good/excellent condition', '57%'],
        ['Counties with significant needs', '57%'],
        ['Buildings in poor condition', '18'],
        ['Average new school cost (2022)', '$51 million'],
        ['Average new school cost (2012)', '$15 million'],
    ]
)

doc.add_heading('A.4 Tennessee Teacher Workforce', level=2)

add_table(
    ['Metric', 'Value'],
    [
        ['Elementary vacancy rate (2024)', '1.4% (doubled from 0.6% in 2023)'],
        ['Pre-K vacancy rate', '5.0%'],
        ['Special education vacancy rate', '3.6%'],
        ['World languages vacancy rate', '3.3%'],
        ['ESL vacancy rate', '2.9%'],
        ['Critical shortage threshold', '2.5%'],
        ['Teachers lost in first 3 years', '20% (1 in 5)'],
        ['Leaders dissatisfied with applicant pools (2024)', '74%'],
        ['Leaders dissatisfied with applicant pools (2021)', '56%'],
        ['Emergency credential disparity (F vs. A schools)', '7x more likely'],
    ]
)

doc.add_heading('A.5 Tennessee School Nursing', level=2)

add_table(
    ['Metric', 'Value'],
    [
        ['Tennessee nurse-to-student ratio', '1:3,000'],
        ['NASN recommended ratio', '1:750'],
        ['States with better ratios (e.g., AL, VT)', '1:500'],
        ['Minimum salary set by state', 'No (unlike teachers)'],
    ]
)

doc.add_heading('A.6 Tennessee TISA Funding', level=2)

add_table(
    ['Metric', 'Value'],
    [
        ['Base per-student amount (2023-24)', '$6,860'],
        ['Base per-student amount (2024-25)', '$7,075'],
        ['State share of base + weighted', '70%'],
        ['Local share', '30%'],
        ['Additional investment in first year', '$1.16 billion (+21.6%)'],
        ['Capital spending source', 'Primarily local revenue'],
    ]
)

doc.add_heading('A.7 Tennessee Chronic Absenteeism', level=2)

add_table(
    ['Year', 'Chronic Absenteeism Rate'],
    [
        ['2018-19 (pre-pandemic)', '~13%'],
        ['2021-22', '20.3%'],
        ['Change', '+7 percentage points'],
    ]
)

# Save
output_path = '/Users/hornej/Documents/Research/school-operations-student-achievement/2026-03-13 School Operations Student Achievement Research Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
