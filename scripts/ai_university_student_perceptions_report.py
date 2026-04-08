#!/usr/bin/env python3
"""Generate research report: AI in University Settings - Student Preferences & Perceptions"""

import sys
import os
from datetime import date

# Use Anaconda python-docx
sys.path.insert(0, '/opt/anaconda3/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AI in a University Setting:\nStudent Preferences, Perceptions, and the Path Forward')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('A Research Synthesis')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'{date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x47, 0x7a)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x3a, 0x7a, 0xa5)
    return p

def para(doc, text):
    p = doc.add_paragraph(text)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_key_metrics_table(doc):
    h2(doc, 'Key Metrics at a Glance')
    rows = [
        ('Metric', 'Finding', 'Source(s)'),
        ('Overall student AI usage', '86-95% depending on country/survey', 'DEC 2024; HEPI 2025, 2026; Copyleaks 2025'),
        ('Weekly AI use', '54-65%', 'DEC 2024; Tyton Partners 2025'),
        ('Daily AI use', '24-29%', 'DEC 2024; Copyleaks 2025'),
        ('ChatGPT market share among students', '66-74%', 'DEC 2024; Copyleaks 2025'),
        ('Students using AI for assessments (UK)', '88-94%', 'HEPI 2025, 2026'),
        ('Institutions with formal AI policy', '20-28%', 'EDUCAUSE 2025; Tyton Partners 2025; Thesify 2025'),
        ('Students violating AI policies (not seeing it as wrong)', '48%', 'Copyleaks 2025'),
        ('Faculty fearing student AI overreliance', '95%', 'AAC&U/Elon 2026'),
        ('Meta-analytic effect on learning achievement', "SMD = 0.45 to g = 0.86 across reviews (medium to large)", 'Ma 2025; Educational Research Review 2025'),
        ('AI negative impact on critical thinking', 'Coefficient = -0.37 (p < .001)', 'Gerlich 2025'),
        ('AI literacy (self-assessed)', '~3.0/5.0 (moderate); 58% feel underprepared', 'Mansoor et al. 2024; DEC 2024'),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = 'Light Grid Accent 1'
    for i, (c1, c2, c3) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = c1
        cells[1].text = c2
        cells[2].text = c3
        if i == 0:
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph('')

def build_report(output_path):
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc)

    # =========================================================================
    # 1. DEFINITION & BACKGROUND
    # =========================================================================
    h1(doc, '1. Definition and Background')

    para(doc, 'The integration of artificial intelligence into higher education has moved from a niche interest to an institutional imperative in under three years. Since the public release of ChatGPT in November 2022, universities worldwide have grappled with how to respond to a technology that students adopted faster than any previous educational tool. By early 2026, the question is no longer whether students use AI, but how, why, and with what consequences.')

    para(doc, 'This report synthesizes findings from more than 50 peer-reviewed studies, institutional surveys, meta-analyses, and policy reports published between 2024 and early 2026. The focus is on student preferences and perceptions, meaning what students think about AI in their education, how they use it, what they want from their institutions, and where the evidence points toward genuine benefits or legitimate concerns.')

    h2(doc, 'Key Terms')
    bullet(doc, 'Generative AI (GenAI): AI systems that produce text, code, images, or other content in response to prompts. In the university context, this primarily means large language models like ChatGPT, Google Gemini, Microsoft Copilot, Claude, and Grammarly.')
    bullet(doc, 'AI literacy: A student\'s ability to understand what AI is, how it works, when to use it appropriately, and how to critically evaluate its outputs. Encompasses technical understanding, practical application, and ethical awareness.')
    bullet(doc, 'Cognitive offloading: The tendency to delegate thinking tasks to external tools (in this case, AI), potentially reducing the mental effort invested in learning.')
    bullet(doc, 'Academic integrity in the AI context: The evolving norms around what constitutes acceptable vs. unacceptable use of AI in coursework, which students and faculty are still negotiating.')

    h2(doc, 'Historical Context')
    para(doc, 'The timeline of AI adoption in higher education has been remarkably compressed. In late 2022, most universities had no AI policy. By 2023, roughly half of students had tried ChatGPT. By 2025, usage had climbed to 86-95% across multiple international surveys. This pace of adoption has outstripped institutional policy development: as of spring 2025, only 20-28% of institutions had formal AI policies in place (EDUCAUSE, 2025; Tyton Partners, 2025). The result is a landscape where student behavior has raced ahead of institutional guidance, creating tensions around academic integrity, equity, and pedagogical effectiveness.')

    # =========================================================================
    # 2. CURRENT STATE OF KNOWLEDGE
    # =========================================================================
    h1(doc, '2. Current State of Knowledge')

    h2(doc, 'Adoption Rates and Usage Patterns')
    para(doc, 'Student AI adoption is now near-universal in most developed countries and climbing rapidly elsewhere. The HEPI/Kortext surveys in the UK tracked usage from 66% (2024) to 92% (2025) to 95% (2026). The Digital Education Council\'s 16-country survey found 86% usage in mid-2024, with 54% using AI weekly and roughly one-quarter daily. In the US, Copyleaks (2025) found approximately 90% of university students have used AI for academic purposes, with 29% using it daily. The Chegg Global Student Survey (2025) of 11,706 undergraduates across 15 countries found 80% have used GenAI to support their studies.')

    para(doc, 'ChatGPT dominates the market. Copyleaks (2025) found 74% of US students use ChatGPT, followed by Google Gemini (43%), Grammarly (38%), Microsoft Copilot (29%), Claude (25%), and Perplexity (16%). The Digital Education Council found ChatGPT at 66% globally, with students using an average of 2.1 AI tools. OpenAI reported over 700,000 campus licenses sold to approximately 35 US public universities by 2025, with 14 million interactions across 20 campuses in September 2025 alone.')

    h2(doc, 'What Students Use AI For')
    para(doc, 'Students primarily use AI as a learning companion rather than a content generator, though the line blurs. The Inside Higher Ed/Kaplan survey (2025) of 1,047 students found the top uses were brainstorming ideas (55%), asking questions like a tutor (50%), and studying for exams (46%). Fewer used it for completing assignments (25%) or writing full essays (19%). The Chegg survey found 56% primarily use AI to understand concepts, with 29% turning to GenAI first when stuck (up from 10% in 2023).')

    para(doc, 'The Jisc report (2025) documented that students use AI for writing, research, notetaking, revision, presentations, life organization, and job preparation. Disabled and neurodiverse students find AI particularly helpful for explaining concepts, building flexibility, and rehearsing answers, an accessibility benefit worth noting.')

    h2(doc, 'Student Attitudes and Emotions')
    para(doc, 'The largest study on student attitudes, Raman et al. (2025) surveying 23,218 students from 109 countries, found that curiosity and calmness were the most common emotions. Students viewed ChatGPT as useful for simplifying complex information and improving study efficiency, but less reliable for providing accurate information or supporting classroom learning. They believed AI could enhance access to knowledge and improve grades, while recognizing its limitations for critical thinking and decision-making.')

    para(doc, 'A striking finding across multiple studies is what researchers at Taylor and Francis termed "widely used but barely trusted." A 2025 study of 132 UK university students found that students explicitly disagree that AI is a reliable knowledge source, yet continue using it for convenience, speed, and help getting started. At Oregon State University, the most common student emotions were "curious" and "concerned," while open-ended responses most frequently listed "angry" and "disappointed." Students are pragmatic adopters, not enthusiastic believers.')

    h2(doc, 'Meta-Analyses: Effect Sizes on Learning')
    para(doc, 'Two major meta-analyses published in 2025 provide the best available evidence on AI\'s impact on learning outcomes:')
    bullet(doc, 'Ma et al. (2025), published in the Journal of Computer Assisted Learning, analyzed 49 articles and found mean effect sizes of g = 0.857 on learning achievement and g = 0.803 on motivation (both large). Cognitive outcomes showed g = 0.604, behavioral outcomes g = 0.698, and affective outcomes g = 0.478. Critically, effects on higher-order cognitive skills (creating, evaluating) were minimal, and instructed use produced stronger effects than unguided use.')
    bullet(doc, 'A systematic review in Educational Research Review analyzed 68 experimental and quasi-experimental studies, finding a moderate overall positive effect (SMD = 0.45, 95% CI [0.43, 0.47]). Chatbot-based interventions showed the largest impact (effect size = 1.02). Like Ma et al., this review found that GenAI significantly enhances lower-order cognitive outcomes but minimally impacts higher-order skills.')

    para(doc, 'These are encouraging findings for AI\'s potential as a learning tool, but the consistent finding that higher-order thinking benefits are minimal is a significant caveat. AI appears to help students learn facts and procedures more effectively while doing less for the analytical and creative skills that higher education is supposed to develop.')

    add_key_metrics_table(doc)

    # =========================================================================
    # 3. WHAT WORKS VS. WHAT DOESN'T
    # =========================================================================
    h1(doc, '3. What Works vs. What Doesn\'t')

    h2(doc, 'What Works')
    bullet(doc, 'AI as tutor/study companion: Students who use AI to explain concepts, brainstorm, and prepare for exams report the most positive experiences. The Chegg survey found 50% reported improved understanding of complex concepts when using AI this way.')
    bullet(doc, 'Structured, instructed use: Meta-analyses consistently show stronger effects when AI use is guided by instructors rather than left to students\' own devices (Ma et al., 2025). The "how" matters more than the "what."')
    bullet(doc, 'Accessibility support: The Jisc report highlighted that disabled and neurodiverse students find particular value in AI for explaining concepts in different ways, providing flexibility, and practicing responses.')
    bullet(doc, 'Reducing anxiety: A quasi-experimental study on emotions and creative problem-solving (Computers and Education: Artificial Intelligence, 2025) found GenAI reduced anxiety and shame while enhancing correctness and novelty in creative tasks, with self-efficacy as a mediator (beta = 0.256, p < 0.001).')
    bullet(doc, 'Course-specific integration: Students respond best when AI is woven into the curriculum with clear guidelines, not treated as a separate topic or banned without explanation (Jisc, 2025; DEC, 2024).')

    h2(doc, 'What Doesn\'t Work')
    bullet(doc, 'Blanket bans: Despite some institutions banning AI (Sciences Po Paris, NYC Education Department), student usage continues regardless. A Thesify (2025) analysis of top university policies estimated that the vast majority of student AI use goes undetected by instructors. Bans create an underground economy of use without guidance.')
    bullet(doc, 'Unguided use for writing: Students who use AI to write full essays report more negative impacts on critical thinking (12%) compared to those who use it for studying (6%), per the Inside Higher Ed survey (2025).')
    bullet(doc, 'Policy without education: A key finding from the Journal of Academic Ethics (2025) is that students\' ethical beliefs, not institutional policies, are the strongest predictors of both perceived misconduct and actual AI use. Policy awareness had no significant effect on ethical judgments or behavior. Rules alone do not change behavior.')
    bullet(doc, 'General-purpose tools without context: The Chegg survey found 50% of students want education-specific AI tools, suggesting that general-purpose chatbots fall short of what students need for academic work.')
    bullet(doc, 'Ignoring the equity gap: When institutions fail to provide AI tools, students with more resources gain compounding advantages. The EDUCAUSE 2025 study found cost is the number one barrier, and 83% of CTOs are concerned about widening the digital equity divide.')

    # =========================================================================
    # 4. CRITICISMS & LIMITATIONS
    # =========================================================================
    h1(doc, '4. Criticisms and Limitations')

    h2(doc, 'Critical Thinking Erosion')
    para(doc, 'The most substantial criticism of AI in education, supported by multiple studies, is its potential to erode critical thinking. Gerlich (2025), studying 666 participants, found a significant negative correlation between frequent AI use and critical thinking scores (coefficient = -0.37, p < .001), with younger participants (ages 17-25) showing higher AI dependence and lower thinking scores. A systematic review in Education and Information Technologies (2025) confirmed these findings, describing a "generate first, think later" behavioral pattern among frequent AI users.')

    para(doc, 'A study published in Current Psychology (2025) identified the mechanism: AI tool usage positively predicts both epistemic laziness (reduced motivation to seek knowledge independently) and metacognitive weakness (reduced ability to monitor one\'s own thinking), both of which mediate the negative relationship between AI use and critical thinking.')

    para(doc, 'Research covered by the Harvard Gazette (November 2025) found that participants who used ChatGPT for essay writing showed reduced neural connectivity in networks associated with memory and creativity, with memory retention dropping compared to groups using Google Search or no tools at all. This represents secondary reporting of an unpublished study, so the finding should be treated with appropriate caution.')

    h2(doc, 'The Cognitive Paradox')
    para(doc, 'Jose et al. (2025), writing in Frontiers in Psychology, described AI as having a "paradoxical character as both a cognitive amplifier and inhibitor." AI enhances accessibility and reduces cognitive overload, but excessive dependence compromises autonomy. The EDUCAUSE Review (December 2025) captured this succinctly: "Better results, worse thinking." AI produces higher-quality outputs for students but may reduce the cognitive effort that drives actual learning.')

    para(doc, 'This is not a settled debate. A systematic review in Smart Learning Environments (2024) noted that conclusions remain conflicting: some studies find positive effects on learning performance while others document erosion of deeper cognitive skills. The effect likely depends heavily on how AI is used, with passive consumption of AI outputs being more harmful than active, critical engagement with them.')

    h2(doc, 'Methodological Concerns')
    bullet(doc, 'Most studies rely on self-reported data, which may not accurately reflect actual behavior or learning outcomes.')
    bullet(doc, 'Many studies are cross-sectional rather than longitudinal, making it difficult to establish causation.')
    bullet(doc, 'Publication bias may favor studies showing significant effects (positive or negative) over null results.')
    bullet(doc, 'The rapid pace of AI development means findings from even 2024 may not reflect current tool capabilities.')
    bullet(doc, 'Sample sizes vary enormously, from 132 to 23,218, making direct comparisons challenging.')

    # =========================================================================
    # 5. MEASURABLE OUTCOMES
    # =========================================================================
    h1(doc, '5. Measurable Outcomes')

    h2(doc, 'Learning Achievement')
    para(doc, 'Meta-analyses show medium to large positive effects on learning achievement. Ma et al. (2025) reported Hedges\' g = 0.857 for overall learning achievement and g = 0.533 for academic achievement specifically; a separate systematic review in Educational Research Review found a moderate effect (SMD = 0.45). The strongest effects were for lower-order cognitive tasks and chatbot-based interventions (effect size = 1.02). Higher-order skills (analysis, evaluation, creation) show minimal improvement.')

    h2(doc, 'Motivation and Engagement')
    para(doc, 'A systematic review in Computers and Education: Artificial Intelligence (2025) found significant positive impacts on motivation and engagement, with medium effect sizes for behavioral engagement and large effect sizes for cognitive engagement. Subject area, learning strategy, and context all moderate these effects.')

    h2(doc, 'Anxiety and Self-Efficacy')
    para(doc, 'AI use is associated with reduced academic anxiety (beta = 0.093, p < 0.05) and enhanced self-efficacy (beta = 0.256, p < 0.001). Students with lower self-esteem and higher academic anxiety are more likely to use AI, suggesting it serves a compensatory function (Pavone, 2025).')

    h2(doc, 'Critical Thinking (Negative)')
    para(doc, 'Frequent AI use is associated with lower critical thinking scores (coefficient = -0.37, p < .001), with the effect mediated by epistemic laziness and metacognitive weakness. Younger students (17-25) are most affected. This is the most robust negative finding in the literature.')

    h2(doc, 'AI Literacy')
    para(doc, 'Multinational assessments find moderate AI literacy levels (approximately 3.0 out of 5.0), with technical understanding particularly low even when practical application is average. A study across Germany, the UK, and the US (2025) found German students showed higher literacy levels. The Digital Education Council found 58% of students feel they lack sufficient AI knowledge and 48% feel inadequately prepared for an AI-enabled workforce.')

    # =========================================================================
    # 6. PRACTICAL APPLICATIONS
    # =========================================================================
    h1(doc, '6. Practical Applications')

    h2(doc, 'Effective Policy Approaches')
    para(doc, 'The research converges on several principles for effective AI policy in universities:')
    bullet(doc, 'Course-specific rather than institution-wide rules: Students respond better to guidelines tailored to their discipline and course (Jisc, 2025; Thesify, 2025). Faculty authority over course-level rules is becoming the standard approach.')
    bullet(doc, 'Explain the rationale: Duke\'s Center for Teaching and Learning (2025) emphasizes heading off issues through assessment design rather than punishment, and explaining why AI boundaries exist.')
    bullet(doc, 'Teach before you test: Students need practical AI literacy training before being expected to use AI appropriately. The Digital Education Council found 72% desire more AI literacy courses.')
    bullet(doc, 'Disclosure and documentation: Some institutions now require students to submit chat transcripts as appendixes, creating transparency without banning use.')

    h2(doc, 'Institutional Provision of AI Tools')
    para(doc, 'There is a growing gap between what students want and what institutions provide. The HEPI 2025 survey found 53% of students believe institutions should provide AI tools (up from 30% the prior year), but only 24% of institutions actually do so (up from 9%). Half of colleges do not grant students institutional access to GenAI tools, with cost as the primary barrier (EDUCAUSE, 2025). OpenAI has moved into this space aggressively, with campus-wide licenses at $2-3 per user per month.')

    h2(doc, 'Assessment Redesign')
    para(doc, 'The most effective institutional responses focus on redesigning assessments rather than policing AI use. This includes process-based assessment (evaluating how students arrive at answers, not just the final product), oral examinations and presentations, portfolio-based assessment with documented reflection, and authentic assessments tied to real-world problems that require original thinking beyond what AI can provide.')

    h2(doc, 'Demographic-Sensitive Approaches')
    para(doc, 'Research consistently shows demographic differences that policy should account for:')
    bullet(doc, 'Gender: Males use AI more frequently and hold more positive attitudes. Females are more concerned about risks and academic integrity (HEPI 2025; Cachero et al., 2025). Gender-inclusive educational strategies are needed.')
    bullet(doc, 'Discipline: STEM and business students adopt AI at higher rates and hold more positive views than humanities students (Farinosi & Melchior, 2025; EU-JER systematic review, 2025). Humanities students may need different onboarding.')
    bullet(doc, 'Socioeconomic status: Students from wealthier backgrounds use AI more and have access to premium tools. Institutional provision of AI tools is an equity intervention.')

    # =========================================================================
    # 7. BOTTOM LINE
    # =========================================================================
    h1(doc, '7. Bottom Line')

    para(doc, 'The evidence on AI in university settings is now substantial enough to draw firm conclusions on several points, while acknowledging genuine uncertainty on others.')

    para(doc, 'First, student AI adoption is essentially universal in developed countries and accelerating globally. With 86-95% of students using AI tools and the trajectory still climbing, any institutional strategy premised on students not using AI is disconnected from reality. The 48% of students who violate AI policies without seeing it as wrong (Copyleaks, 2025) are not bad actors; they are operating in a world where the rules have not caught up with the technology. The most effective institutions are moving from prohibition to guided integration, building AI literacy into the curriculum and redesigning assessments to work with AI rather than against it.')

    para(doc, 'Second, the learning outcomes picture is genuinely mixed, and the nuance matters. Meta-analyses show medium to large positive effects on achievement and motivation when AI is used with instructor guidance. These are real benefits. But the consistent finding that higher-order cognitive skills, the analytical and creative thinking that defines a university education, show minimal improvement from AI use is a serious concern. The critical thinking erosion documented by Gerlich (2025) and others is not a hypothetical risk but a measurable effect, particularly among younger students. The challenge for educators is harnessing the genuine benefits of AI for learning while structuring its use to preserve the cognitive effort that develops deeper skills.')

    para(doc, 'Third, the equity dimension deserves more attention than it is getting. When half of institutions do not provide students with AI tools and cost is the primary barrier, we are creating a two-tier system where students at well-resourced institutions accumulate AI skills and productivity gains while others are left behind. The Brookings Institution (2024) put it well: wealthier schools may use AI as a supplement to high-quality education, while poorer schools may rely on it as a replacement. Gender and disciplinary divides compound these access issues. Institutional provision of AI tools is not a luxury; it is an equity intervention.')

    para(doc, 'Fourth, student trust in AI is low even as usage is high. Students are pragmatic, not naive. They know AI makes mistakes, they worry about bias and privacy, and they use it anyway because it is fast, convenient, and helpful for getting started. This pragmatic skepticism is actually a healthy foundation for AI literacy, suggesting that students are more ready for critical engagement with AI than they are often given credit for. The finding that ethical beliefs, not institutional policies, predict AI behavior (Journal of Academic Ethics, 2025) suggests that building ethical reasoning is more effective than building better rules.')

    para(doc, 'Several questions remain unanswered. We do not yet have robust longitudinal data showing how sustained AI use over a full degree program affects learning outcomes and career readiness. The mechanism behind critical thinking erosion, whether it is cognitive offloading, reduced practice, or something else, needs more precise investigation. The optimal balance between AI-assisted and unassisted learning for different subjects and skill levels is unknown. And the long-term labor market effects that students worry about remain speculative. These are the research frontiers that will shape the next wave of evidence.')

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_page_break()
    h1(doc, 'References')

    refs = [
        'AAC&U & Elon University. (2026). National survey: 95% of college faculty fear student overreliance on AI. https://www.aacu.org/newsroom/national-survey-95-of-college-faculty-fear-student-overreliance-on-ai-and-diminished-critical-thinking-among-learners-who-use-generative-ai-tools',
        'BestColleges. (2025). College student attitudes on AI. https://www.bestcolleges.com/research/college-student-attitudes-on-ai/',
        'Brookings Institution. (2024). AI and the next digital divide in education. https://www.brookings.edu/articles/ai-and-the-next-digital-divide-in-education/',
        'Cachero, C., Tomas, D., & Pujol, M. (2025). Gender bias in AI self-perception among undergraduates. ACM Transactions on Computing Education. https://dl.acm.org/doi/10.1145/3721295',
        'Chegg. (2025). Chegg global student survey 2025: 80% of undergraduates worldwide have used GenAI. https://www.chegg.com/about/newsroom/press-release/chegg-global-student-survey-2025',
        'Copyleaks. (2025). AI in action: 2025 student AI usage report. https://copyleaks.com/blog/ai-in-action-2025-student-ai-usage-report',
        'Copyleaks. (2025). Study: 48% of students admit using AI in violation of school policies. https://www.globenewswire.com/news-release/2025/11/05/3182011/',
        'Digital Education Council. (2024). Global AI student survey 2024. https://www.digitaleducationcouncil.com/post/digital-education-council-global-ai-student-survey-2024',
        'Duke Center for Teaching and Learning. (2025). Artificial intelligence policies in syllabi. https://ctl.duke.edu/ai-and-teaching-at-duke-2/',
        'EDUCAUSE. (2025). 2025 EDUCAUSE AI landscape study. https://www.educause.edu/content/2025/2025-educause-ai-landscape-study/',
        'EDUCAUSE Review. (2025, December). The paradox of AI assistance: Better results, worse thinking. https://er.educause.edu/articles/2025/12/the-paradox-of-ai-assistance-better-results-worse-thinking',
        'Farinosi, M., & Melchior, C. (2025). To adopt or to ban? Student perceptions and use of generative AI in higher education. Humanities and Social Sciences Communications (Nature). https://www.nature.com/articles/s41599-025-05982-7',
        'Gerlich, M. (2025). AI and cognitive offloading. Societies, 15(1), 6. https://www.mdpi.com/2075-4698/15/1/6',
        'HEPI/Kortext. (2025). Student generative AI survey 2025. https://www.hepi.ac.uk/reports/student-generative-ai-survey-2025/',
        'HEPI. (2026). Student generative AI survey 2026. https://www.hepi.ac.uk/reports/student-generative-ai-survey-2026/',
        'Inside Higher Ed. (2025). Survey: College students\' views on AI. https://www.insidehighered.com/news/students/academics/2025/08/29/survey-college-students-views-ai',
        'Jisc. (2025). Student perceptions of AI 2025. https://www.jisc.ac.uk/reports/student-perceptions-of-ai-2025',
        'Jose, S., et al. (2025). The cognitive paradox of AI in education. Frontiers in Psychology, 16. https://www.frontiersin.org/journals/psychology/articles/10.3389/fpsyg.2025.1550621/full',
        'Journal of Academic Ethics. (2025). AI and academic integrity: Exploring student perceptions. https://link.springer.com/article/10.1007/s10805-025-09613-3',
        'Ma, S., et al. (2025). Meta-analysis of GenAI on learning outcomes. Journal of Computer Assisted Learning. https://onlinelibrary.wiley.com/doi/10.1111/jcal.70117',
        'Mansoor, A., et al. (2024). Artificial intelligence literacy among university students: A comparative transnational survey. Frontiers in Communication. https://www.frontiersin.org/journals/communication/articles/10.3389/fcomm.2024.1478476/full',
        'Li, L., et al. (2024). Effects of over-reliance on AI dialogue systems on students\' cognitive abilities: A systematic review. Smart Learning Environments. https://link.springer.com/article/10.1186/s40561-024-00316-7',
        'O\'Dea, R., et al. (2025). Ethical uses of generative AI in assessment. Evaluation Review (SAGE). https://journals.sagepub.com/doi/10.1177/0193841X251399712',
        'OECD & Fondazione Agnelli. (2025). AI adoption in the education system. https://www.oecd.org/en/publications/ai-adoption-in-the-education-system_69bd0a4a-en.html',
        'OpenAI. (2025). College students and ChatGPT. https://openai.com/global-affairs/college-students-and-chatgpt/',
        'Oregon State University Ecampus. (2024). Student AI survey. https://ecampus.oregonstate.edu/research/study/ai-survey/',
        'Pavone, A. (2025). AI, self-esteem, and academic anxiety. Journal of Marketing Education (SAGE). https://journals.sagepub.com/doi/10.1177/02734753251346857',
        'Pearson, H. (2025). AI in higher education. Nature, 646, 788-791. https://www.nature.com/articles/d41586-025-03340-w',
        'Perkins, M., et al. (2024). Students\' perceptions of AI-giarism. Education and Information Technologies. https://link.springer.com/article/10.1007/s10639-024-13151-7',
        'Pew Research Center. (2026, February). How teens use and view AI. https://www.pewresearch.org/internet/2026/02/24/how-teens-use-and-view-ai/',
        'Raman, R., et al. (2025). Global ChatGPT perceptions study. PLOS ONE. https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0315011',
        'Roe, J., & Perkins, M. (2024). Student perspectives on GenAI. International Journal for Educational Integrity. https://link.springer.com/article/10.1007/s40979-024-00149-4',
        'Siddiqi, A. (2025). AI and community colleges. New Directions for Community Colleges. https://onlinelibrary.wiley.com/doi/abs/10.1002/cc.70051',
        'Sullivan, M., et al. (2024). Rapid rise of GenAI and academic integrity. Computers and Education: Artificial Intelligence. https://www.sciencedirect.com/science/article/pii/S2666920X24000766',
        'Systematic review on AI and critical thinking. (2025). Education and Information Technologies. https://files.eric.ed.gov/fulltext/EJ1459623.pdf',
        'Thesify. (2025). Generative AI policies at the world\'s top universities. https://www.thesify.ai/blog/gen-ai-policies-update-2025',
        'Tyton Partners / D2L. (2025). Time for class 2025. https://www.insidehighered.com/news/student-success/academic-life/2025/06/11/65-percent-students-use-gen-ai-chat-bot-weekly',
        'UNESCO. (2025). Survey of higher education institutions on AI guidance. https://www.unesco.org/en/articles/unesco-survey-two-thirds-higher-education-institutions-have-or-are-developing-guidance-ai-use',
        'Wiley. (2024). AI has hurt academic integrity in college courses. https://newsroom.wiley.com/press-releases/press-release-details/2024/AI-Has-Hurt-Academic-Integrity-in-College-Courses/',
        'Current Psychology. (2025). Epistemic laziness and metacognitive weakness: Mediating effects on AI use and critical thinking. https://link.springer.com/article/10.1007/s12144-025-08800-0',
        'Cogent Education. (2025). The role of over-reliance on AI in the negative consequences of student learning. https://www.tandfonline.com/doi/full/10.1080/2331186X.2025.2591503',
        'Harvard Gazette. (2025, November). Is AI dulling our minds? https://news.harvard.edu/gazette/story/2025/11/is-ai-dulling-our-minds/',
        'Computers and Education: Artificial Intelligence. (2025). Systematic review of GenAI on motivation and engagement. https://www.sciencedirect.com/science/article/pii/S2666920X25000955',
        'Educational Research Review. (2025). Systematic review and meta-analysis of GenAI on learning outcomes. https://www.sciencedirect.com/science/article/abs/pii/S1747938X25000740',
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
        pf.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(output_path)
    print(f'Report saved to: {output_path}')

if __name__ == '__main__':
    output = sys.argv[1] if len(sys.argv) > 1 else '/Users/hornej/Documents/Research/ai-university-student-perceptions/2026-03-24 AI University Student Perceptions Research Report.docx'
    build_report(output)
