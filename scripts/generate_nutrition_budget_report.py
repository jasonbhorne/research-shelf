#!/usr/bin/env python3
"""
Generate School Nutrition Budget Constraints Research Report
Output: ~/Documents/Research/school-nutrition-budget-cafeteria-options/2026-03-13 School Nutrition Budget Constraints Research Report.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

doc = Document()

# ----- Styles -----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Title
title = doc.add_heading('School Nutrition Budget Constraints on Cafeteria Options', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('A Research Report for Greeneville City Schools')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
date_line = doc.add_paragraph(f'March 13, 2026')
date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_line.runs[0].font.size = Pt(11)
date_line.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ===== EXECUTIVE SUMMARY =====
doc.add_heading('Executive Summary', level=1)

doc.add_paragraph(
    'This report examines why school nutrition budgets constrain cafeteria menu options, '
    'with specific analysis of Greeneville City Schools (GCS) financial data and the broader '
    'regulatory, operational, and economic factors that shape what districts can serve. The '
    'findings draw on peer-reviewed research, federal and state policy analysis, implementation '
    'case studies from comparable districts, cost-benefit modeling, and a frank assessment of '
    'the risks involved in expanding menu offerings.'
)

doc.add_paragraph(
    'The short answer to the question is multi-layered: GCS\'s nutrition program ran a $515,934 '
    'deficit in FY2024-25, with expenses ($2.58M) outpacing revenue ($2.06M) by a widening '
    'margin. The fund balance has declined from $1.88M to $1.23M over two years. Every dollar '
    'matters in this environment, and adding menu options means adding costs for food, labor, '
    'training, and compliance, all while navigating a dense web of federal regulations that '
    'dictate what can be served, how it must be priced, and what nutritional standards it must meet.'
)

doc.add_paragraph(
    'However, the research also reveals a counterintuitive finding: strategic menu improvements, '
    'particularly through scratch or modified-scratch cooking, can actually improve financial '
    'performance by increasing student participation. More meals served means more federal '
    'reimbursement revenue. Studies show scratch cooking shifts costs from food to labor without '
    'increasing total per-meal cost, and districts that have made this transition report average '
    'participation increases of 9%. For GCS, a 9% participation increase could generate roughly '
    '$112,000 in additional annual reimbursement revenue.'
)

doc.add_paragraph(
    'The report recommends a phased approach: (1) optimize existing menu items using scratch '
    'techniques and the district\'s recently purchased equipment, funded through the required '
    'excess balance spend-down; (2) add targeted new options based on student preference data; '
    'and (3) pursue external funding through the Chef Ann Foundation\'s Get Schools Cooking '
    'program and USDA Farm to School grants. The biggest risk is not acting: the current trajectory '
    'of declining fund balances and rising costs is unsustainable regardless of menu strategy.'
)

# ===== SECTION 1: DEFINITION & BACKGROUND =====
doc.add_heading('1. Definition and Background', level=1)

doc.add_paragraph(
    'School nutrition programs in the United States operate within a unique financial ecosystem. '
    'Unlike most public services funded primarily through local tax revenue, school meal programs '
    'are largely self-funding operations. Revenue comes from three sources: federal reimbursements '
    'from USDA (the dominant source), student and adult meal payments, and state matching funds. '
    'These revenues must cover all program costs, including food, labor, equipment, supplies, and overhead.'
)

doc.add_paragraph(
    'The term "cafeteria options" encompasses two distinct categories of food service, each governed '
    'by different rules:'
)

p = doc.add_paragraph()
p.add_run('Reimbursable meals').bold = True
p.add_run(
    ' are breakfast and lunch combinations that meet USDA meal pattern requirements. These are the '
    'core program, and they generate the bulk of federal reimbursement revenue. Menu flexibility here '
    'is constrained by detailed component requirements (specific quantities of grains, protein, '
    'vegetables by subgroup, fruits, and milk), calorie ranges by grade level, and limits on sodium, '
    'saturated fat, and added sugars.'
)

p = doc.add_paragraph()
p.add_run('Nonprogram foods').bold = True
p.add_run(
    ' include a la carte items, extra entrees, snacks, and beverages sold outside the reimbursable '
    'meal. These must meet Smart Snacks nutrition standards and generate enough revenue to cover their '
    'proportional share of total food costs (the Nonprogram Revenue requirement). At GCS, these items '
    'include ice cream, chips, cookies, drinks, and extra entrees.'
)

doc.add_paragraph(
    'When people ask "why can\'t we have more options in the cafeteria?", they are typically asking '
    'about both categories: why the standard meals seem repetitive, and why there aren\'t more a la '
    'carte choices. The answer involves an interconnected set of financial, regulatory, operational, '
    'and labor constraints that this report examines in detail.'
)

doc.add_heading('GCS Nutrition Program: Financial Snapshot', level=2)

# GCS Financial Table
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Metric', 'FY2023-24', 'FY2024-25']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['Starting Balance (July 1)', '$1,884,244', '$1,742,624'],
    ['Program Revenue', '$2,208,901', '$2,060,150'],
    ['Total Expenses', '$2,350,521', '$2,576,085'],
    ['Surplus/Deficit', '-$141,620', '-$515,934'],
    ['Ending Balance (June 30)', '$1,742,624', '$1,226,690'],
    ['Avg. Monthly Expenses', '$235,052', '$257,608'],
    ['3-Month Expense Threshold', '$705,156', '$772,825'],
    ['Excess Balance', '$1,037,468', '$453,864'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')  # spacer

doc.add_paragraph(
    'The trajectory is clear: revenue is declining while expenses are rising. The $374,314 swing '
    'in the deficit (from -$142K to -$516K) in a single year signals structural financial pressure, '
    'not a one-time event. This is the fundamental context for any discussion of menu expansion.'
)

# ===== SECTION 2: EVIDENCE OF EFFECTIVENESS =====
doc.add_heading('2. Evidence of Effectiveness', level=1)

doc.add_heading('The Reimbursement-Cost Gap', level=2)
doc.add_paragraph(
    'The most well-documented mechanism linking budgets to menu limitations is the persistent gap '
    'between federal reimbursement rates and actual meal production costs. The USDA\'s School '
    'Nutrition and Meal Cost Study (SNMCS), the first nationally representative assessment since '
    'updated nutrition standards took effect, found that the average reported cost to produce a '
    'school lunch exceeded the free lunch subsidy by $0.49, and breakfast exceeded by $0.84 '
    '(USDA FNS, 2019). For context, the mean reported cost to produce a reimbursable breakfast '
    'was $2.72, while the federal subsidy was $1.88.'
)

doc.add_paragraph(
    'This gap has widened since that 2014-15 study. The School Nutrition Association\'s (SNA) '
    'annual Trends Reports show a steady increase in the percentage of programs reporting '
    'insufficient reimbursement: by SY 2024-25, 70% of meal program directors reported that '
    'federal reimbursement rates are insufficient to cover lunch production costs, up from 64% '
    'the prior year. Furthermore, 91.6% of programs report serious or moderate concern about '
    'financial sustainability three years out (SNA, 2025).'
)

doc.add_heading('Food Costs vs. Menu Diversity', level=2)
doc.add_paragraph(
    'Budget-constrained programs face a direct trade-off between food cost per meal and menu '
    'diversity. A multi-objective optimization study in Public Health Nutrition (Eustachio Colombo '
    'et al., 2023) demonstrated that least-cost optimized menus could reduce per-student costs by '
    '$2.60 per week, but achieving these savings required constraining variety. The study specifically '
    'included variety constraints because "variety is important for nutrition and student acceptance." '
    'In other words, cheaper menus are less diverse menus.'
)

doc.add_heading('Scratch Cooking: Cost-Neutral but Operationally Different', level=2)
doc.add_paragraph(
    'A seminal study in the Journal of the Academy of Nutrition and Dietetics (Bodas et al., 2014) '
    'found that entrees with the highest scratch-cooking scores had significantly lower food costs '
    'but significantly higher labor costs, with no significant difference in total costs. A 2020 '
    'California study confirmed this: nutrition departments with high scratch cooking spent the same '
    'total percentage of their budgets on food and labor (87%) as those doing little scratch cooking. '
    'The cost simply shifts between line items.'
)

doc.add_paragraph(
    'A 2025 study in the Journal of School Health (Zuercher et al.) surveyed 430 California school '
    'food authorities and found that 82% used scratch-cooked foods and 80% used locally grown foods. '
    'More scratch cooking was associated with fewer perceptions of student nonparticipation due to '
    'concerns about meal taste and freshness. This provides some of the strongest evidence that '
    'scratch cooking improves both perceived quality and participation.'
)

doc.add_heading('Plate Waste: A Hidden Budget Drain', level=2)
doc.add_paragraph(
    'The USDA\'s SNMCS found that roughly 25-30% of calories and nutrients served in schools are '
    'wasted. Vegetables had the highest waste rate at 29%, while entrees had the lowest at 12%. '
    'A 2024 narrative review in the Journal of Child Nutrition and Management found that schools '
    'permitting students to choose their own milk portion reported 76% less milk waste (SNA, 2024). '
    'Waste represents a direct financial loss that compounds the reimbursement-cost gap: GCS is '
    'paying for food that students throw away.'
)

doc.add_heading('Healthy Hunger-Free Kids Act Impact', level=2)
doc.add_paragraph(
    'The HHFKA of 2010 expanded menu requirements significantly. A systematic review (Schwartz et al., '
    '2020) documented that Healthy Eating Index scores increased from 58% to 82% of the maximum '
    'possible score post-implementation. Critically, Cohen et al. (2014) found no significant effect '
    'on school lunch participation rates, contrary to fears that healthier standards would drive '
    'students away. Harvard research also showed no increase in food waste. However, implementation '
    'costs were real, and some districts experienced initial participation dips before recovery.'
)

# Key Findings Table
doc.add_heading('Key Research Findings', level=2)
table = doc.add_table(rows=11, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Finding', 'Value', 'Source']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
findings_data = [
    ['Avg. lunch cost exceeds free subsidy by', '$0.49', 'USDA SNMCS, 2019'],
    ['Directors reporting insufficient reimbursement', '70%', 'SNA Trends, 2025'],
    ['Programs with financial sustainability concerns', '91.6%', 'SNA Trends, 2025'],
    ['Top challenge: food costs', '98% of directors', 'SNA Trends, 2025'],
    ['Top challenge: labor costs', '95% of directors', 'SNA Trends, 2025'],
    ['Plate waste (% of calories served)', '25-30%', 'USDA SNMCS'],
    ['Vegetable waste rate', '29%', 'USDA SNMCS'],
    ['Scratch cooking total cost vs. non-scratch', 'No significant difference', 'Bodas et al., 2014'],
    ['CA schools using scratch cooking', '82%', 'Zuercher et al., 2025'],
    ['HEI score improvement post-HHFKA', '58% to 82%', 'Schwartz et al., 2020'],
]
for r, row_data in enumerate(findings_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

# ===== SECTION 3: FEDERAL & STATE POLICY CONTEXT =====
doc.add_heading('3. Federal and State Policy Context', level=1)

doc.add_heading('Federal Meal Pattern Requirements', level=2)
doc.add_paragraph(
    'The USDA meal pattern requirements (7 CFR 210, 220) impose detailed constraints on every '
    'reimbursable meal. Schools must offer specific component quantities across five food groups: '
    'grains, meat/meat alternates, vegetables (with weekly subgroup requirements for dark green, '
    'red/orange, beans/peas, starchy, and other), fruits, and fluid milk. At least 80% of weekly '
    'grains must be whole grain-rich. Calorie ranges are set by grade group (e.g., K-5 lunch: '
    '550-650 kcal; 9-12 lunch: 750-850 kcal). Saturated fat is limited to less than 10% of '
    'total calories.'
)

doc.add_paragraph(
    'The April 2024 final rule phases in additional requirements: added sugar limits on breakfast '
    'cereals, yogurt, and flavored milk starting SY 2025-26, and sodium reductions of 15% for '
    'lunch and 10% for breakfast by SY 2027-28. These requirements limit which products schools '
    'can purchase, as manufacturers must reformulate products to meet the standards. For GCS, this '
    'means some current menu items may need to be replaced, further constraining options and '
    'potentially increasing procurement costs (USDA FNS, 2024).'
)

doc.add_heading('Smart Snacks and Competitive Foods', level=2)
doc.add_paragraph(
    'All foods sold outside of reimbursable meals during the school day must meet Smart Snacks '
    'nutrition standards. Snack items are limited to 200 calories, 200mg sodium, and no more than '
    '35% of calories from fat or 35% of weight from sugar. Entree items sold a la carte are capped '
    'at 350 calories and 480mg sodium. These standards directly limit which a la carte items GCS can '
    'sell. Conventional versions of ice cream, cookies, and chips are largely excluded, requiring '
    'compliant reformulations that often cost more and may be less appealing to students (USDA FNS).'
)

doc.add_heading('Nonprogram Revenue (NPR) Requirements', level=2)
doc.add_paragraph(
    'Under 7 CFR 210.14(f), the proportion of revenue from nonprogram food sales must equal or '
    'exceed the proportion of total food costs attributable to nonprogram foods. In practice, if '
    'a la carte food costs represent 11.4% of total food costs (as GCS\'s reference week data '
    'suggests: $2,424 nonprogram food cost / $21,347 total food cost), then at least 11.4% of '
    'total food service revenue must come from nonprogram food sales. GCS\'s actual nonprogram '
    'revenue for the reference week ($6,754) exceeded the minimum required ($4,113), indicating '
    'compliance. However, expanding a la carte offerings changes this ratio and requires ongoing '
    'monitoring to maintain compliance.'
)

doc.add_heading('Community Eligibility Provision (CEP)', level=2)
doc.add_paragraph(
    'GCS participates in CEP at its four elementary schools (Hal Henard, Highland, EastView, '
    'Tusculum View) for 2025-2026, meaning all students at these schools receive free meals with '
    'no application required. GMS and GHS use traditional free/reduced-price applications. CEP '
    'fundamentally changes the revenue model: it eliminates family payment revenue but increases '
    'federal reimbursement. The claiming percentage (ISP x 1.6) determines what share of meals '
    'are reimbursed at the free rate vs. the lower paid rate. Schools using CEP were 3.6 times '
    'more likely to report breaking even compared to non-CEP schools (No Kid Hungry).'
)

doc.add_heading('Tennessee-Specific Policy', level=2)
doc.add_paragraph(
    'Tennessee\'s school nutrition program is administered through TDOE and the TMAC electronic '
    'claiming system. Key Tennessee-specific factors affecting GCS:'
)

p = doc.add_paragraph()
p.add_run('Excess Balance Rules: ').bold = True
p.add_run(
    'Federal regulation caps the nonprofit food service fund balance at three months\' average '
    'operating expenditures. GCS\'s FY25 ending balance of $1.23M against a three-month threshold '
    'of approximately $773K creates an excess of $454K that must be spent down under a TDOE-approved '
    'plan. GCS has appropriately used excess funds for staff stipends ($25K/year) and capital '
    'equipment. This forced spend-down represents an opportunity to invest in menu improvement '
    'infrastructure.'
)

p = doc.add_paragraph()
p.add_run('State Matching Funds: ').bold = True
p.add_run(
    'Tennessee distributes state matching funds based on prior-year lunches served. At least 10% '
    'must be used for certified food service supervision, manager accreditation, or professional '
    'training. GCS received $14,903 in state matching in March 2025.'
)

p = doc.add_paragraph()
p.add_run('Farm-to-School Losses: ').bold = True
p.add_run(
    'In March 2025, USDA cancelled $1 billion nationally in local food purchasing programs, with '
    'Tennessee losing $12.6M from Local Food for Schools and $7.6M from the Local Food Purchase '
    'Assistance program. This significantly reduces the opportunity for GCS to access locally grown '
    'foods at competitive prices. The FY2026 Patrick Leahy Farm to School Grants ($18M nationally) '
    'remain available but are competitive.'
)

p = doc.add_paragraph()
p.add_run('Pending Legislation: ').bold = True
p.add_run(
    'Several Tennessee bills (SB 0740, HB 1458, HB 12) would establish universal free school meals '
    'statewide, which would significantly affect nutrition program finances if passed. A separate bill '
    'would add Medicaid-enrolled students to TISA\'s economically disadvantaged definition, potentially '
    'increasing weighted funding for districts like GCS.'
)

doc.add_heading('USDA Reimbursement Rates SY 2025-26', level=2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Meal Type', 'Free Rate', 'Reduced Rate', 'Paid Rate']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
rate_data = [
    ['Lunch (standard)', '~$4.60', '~$4.20', '~$0.44'],
    ['Lunch (60%+ F/RP)', '~$4.62', '~$4.22', '~$0.46'],
    ['Performance bonus', '+$0.09/lunch', '', ''],
    ['Breakfast (standard)', '~$2.46', '~$2.16', '~$0.39'],
    ['Severe Need Breakfast', '~$2.73', '~$2.43', '~$0.38'],
    ['USDA Foods per lunch', '$0.45', '', ''],
]
for r, row_data in enumerate(rate_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'Rates increased 3.85% for SY 2025-26. However, 96.8% of directors say increased funding '
    'is needed to comply with forthcoming sodium and sugar limits. USDA estimates the new standards '
    'will cost schools an average of $206 million annually nationwide (SNA, 2025).'
)

# ===== SECTION 4: WHAT WORKS =====
doc.add_heading('4. What Works: Implementation Models', level=1)

doc.add_heading('Osborn School District #8, Phoenix, AZ (~3,500 students)', level=2)
doc.add_paragraph(
    'The strongest parallel to GCS. This small, Title I district transitioned from heat-and-serve '
    'to roughly 60% scratch-cooked meals. They enrolled in CEP, driving participation to 90-95% '
    'on surveyed days. Higher participation boosted federal reimbursement revenue, creating a '
    'virtuous financial cycle. Kitchens now scratch-bake breads, prepare picadillo with locally '
    'sourced beef, and make signature dishes like chicken pot pie. A local nonprofit (Blue Watermelon) '
    'provided chef-led professional development, reducing the training cost burden on the district '
    '(Ed Week, 2026).'
)

doc.add_heading('Wisconsin Rapids Public Schools, WI (~5,000 students)', level=2)
doc.add_paragraph(
    'Participated in Chef Ann Foundation\'s Get Schools Cooking program (3-year grant, $35,000 plus '
    'training). Reduced heat-and-serve meals by half, replacing them with speed-scratch recipes. '
    'Implemented salad bars as the "gateway" to scratch cooking. Used grant funds for equipment '
    'purchases. The program includes technical assistance, training, and annual evaluations at no '
    'additional cost (Chef Ann Foundation).'
)

doc.add_heading('Manhattan-Ogden USD 383, KS (~6,753 students)', level=2)
doc.add_paragraph(
    'Operates a central production kitchen serving 15 buildings. More than doubled scratch-made '
    'menu items by hiring one additional cook (bringing the central kitchen team to 3). Menu '
    'includes pulled pork, barbacoa beef, local Wagyu burgers, and teriyaki chicken. The central '
    'kitchen model allows batch production and distribution to satellite finishing kitchens, a '
    'model GCS could explore with its 6 cafeterias using a hub-and-spoke approach '
    '(Chef Ann Foundation).'
)

doc.add_heading('Key Success Strategies', level=2)

doc.add_paragraph(
    'Across all case studies, several strategies consistently drive success:'
)

strategies = [
    'Start with speed-scratch (hybrid of pre-made components and on-site preparation) rather than full scratch cooking',
    'Use student taste-testing to guide menu decisions; require 75% approval before adding items (LAUSD model)',
    'Leverage CEP enrollment to maximize participation and reimbursement revenue',
    'Apply for the Chef Ann Foundation Get Schools Cooking program (free, 3-year intensive assessment and strategic planning)',
    'Implement salad bars as a low-risk entry point for fresh food offerings',
    'Use a hub-and-spoke kitchen model where one site does primary preparation and others function as finishing kitchens',
    'Incorporate culturally relevant dishes that students recognize and want to eat',
    'Phase the transition over 3+ years to avoid staff burnout and quality inconsistency',
]
for s in strategies:
    doc.add_paragraph(s, style='List Bullet')

# ===== SECTION 5: COST-BENEFIT ANALYSIS =====
doc.add_heading('5. Cost-Benefit Analysis', level=1)

doc.add_heading('The Economics of a School Meal', level=2)
doc.add_paragraph(
    'The average cost to produce a school lunch is approximately $3.81, with food accounting for '
    '44.7% and labor accounting for 44.5% of total meal costs. The remaining ~11% covers overhead. '
    'This near-equal split between food and labor is important because scratch cooking shifts the '
    'ratio (less food cost, more labor cost) without changing the total.'
)

doc.add_heading('The Participation-Revenue Cycle', level=2)
doc.add_paragraph(
    'This is the most critical financial lever for school nutrition programs. At the SY 2025-26 '
    'free lunch reimbursement rate of approximately $4.60, every additional student participating '
    'daily generates roughly $828/year in reimbursement revenue (180 school days). Research shows '
    'that menu improvements drive average daily participation increases of 3-16% (average 9%).'
)

doc.add_heading('Revenue Impact Model for GCS', level=2)
doc.add_paragraph(
    'Using GCS\'s actual data and research-based assumptions:'
)

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
model_data = [
    ['Metric', 'Value'],
    ['Estimated current daily lunch participation', '~1,500 meals'],
    ['Projected participation increase (9%)', '+135 meals/day'],
    ['Free lunch reimbursement rate', '$4.60/meal'],
    ['Additional annual reimbursement', '~$111,780'],
    ['Variable cost of additional meals ($3.81 x 135 x 180)', '~$92,583'],
    ['Net annual revenue gain from participation', '~$19,197'],
]
for r, row_data in enumerate(model_data):
    table.rows[r].cells[0].text = row_data[0]
    table.rows[r].cells[1].text = row_data[1]

doc.add_paragraph('')
doc.add_paragraph(
    'This model is conservative. It does not account for breakfast participation gains, snack '
    'program revenue, or the compounding effect of higher participation on commodity allocations '
    '(USDA Foods are distributed based on meals served). It also does not account for reduced '
    'plate waste, which research suggests can be achieved alongside menu improvements.'
)

doc.add_heading('GCS Nonprogram Food Revenue Analysis', level=2)
doc.add_paragraph(
    'GCS\'s a la carte items show strong margins for several products:'
)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Item', 'Raw Cost', 'Selling Price', 'Margin']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
npr_data = [
    ['Cookies', '$0.15', '$1.00', '567%'],
    ['Extra entrees', '$0.59', '$2.00', '239%'],
    ['Ice cream', '$0.46', '$1.00', '117%'],
    ['Chips', '$0.41', '$0.75', '83%'],
    ['Gatorade', '$0.73', '$1.50', '105%'],
    ['Tea', '$0.28', '$1.25', '346%'],
]
for r, row_data in enumerate(npr_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'Expanding high-margin a la carte items (particularly fresh-baked goods if transitioning to '
    'scratch cooking) could generate additional nonprogram revenue while improving food quality. '
    'However, all items must meet Smart Snacks standards and the overall nonprogram revenue must '
    'maintain proportional compliance.'
)

doc.add_heading('Leveraging the Excess Balance', level=2)
doc.add_paragraph(
    'GCS\'s $454K excess balance creates a strategic opportunity. The funds must be spent down '
    'regardless, so directing them toward menu improvement infrastructure is sound strategy. '
    'Allowable uses include: equipment purchases for scratch cooking (food processors, mixers, '
    'prep tables), staff training programs, one-time menu development costs (recipe testing, '
    'nutritional analysis), small wares and supplies, and technology (menu planning software, '
    'inventory management systems). Rather than viewing the spend-down as a burden, it can fund '
    'the upfront costs of a menu improvement initiative that generates ongoing participation and '
    'revenue gains.'
)

doc.add_heading('Grant Opportunities', level=2)
grants = [
    ('Chef Ann Foundation Get Schools Cooking', 'Free 3-year program including $35,000 grant, technical assistance, training, and evaluation. GCS is an ideal candidate.'),
    ('USDA Equipment Assistance Grants', '$10-20M nationally, competitive subgrants through state agencies for equipment >$1,000 that supports healthier meals.'),
    ('Patrick Leahy Farm to School Grants (FY 2027)', '$18M nationally, minimum $100K request, requires 25% match. GCS should begin establishing agricultural partnerships now.'),
    ('Team Nutrition Training Grants', 'Fund hands-on culinary training for cafeteria staff. Multi-year programs coordinated through Tennessee DOE.'),
]
for name, desc in grants:
    p = doc.add_paragraph()
    p.add_run(name + ': ').bold = True
    p.add_run(desc)

# ===== SECTION 6: CRITICISMS, LIMITATIONS & RISKS =====
doc.add_heading('6. Criticisms, Limitations, and Risks', level=1)

doc.add_heading('Plate Waste May Increase', level=2)
doc.add_paragraph(
    'Research consistently shows that expanding menu options can paradoxically increase food waste. '
    'A USDA study found that middle-school students left nearly 50% of fresh fruit and 37% of '
    'canned fruit unconsumed. Students reported throwing away at least one-quarter of their lunch '
    'an average of 2.1 days per week, primarily because they did not like the taste. While variety '
    'in entrees correlates with less waste, increased fruit and vegetable variety was associated '
    'with increased waste through cycle menus (Cohen et al., 2023).'
)

doc.add_heading('Choice Overload', level=2)
doc.add_paragraph(
    'The "paradox of choice" literature suggests that exceeding two to four options can slow '
    'service lines, reduce eating time, and lead students to select items they ultimately do not '
    'consume. This is most acute for younger students with limited lunch periods of 20-25 minutes. '
    'Any menu expansion at GCS should be calibrated to offer meaningful choice without overwhelming '
    'students or slowing throughput.'
)

doc.add_heading('Food Safety Risks of Scratch Cooking', level=2)
doc.add_paragraph(
    'Scratch cooking introduces elevated food safety concerns. Foods that are cooked, cooled, and '
    'reheated are the highest-risk items, requiring strict HACCP monitoring. USDA research notes '
    'that 100% scratch cooking is difficult to execute safely because it is hard to keep within '
    'safety limits when not everything can be done the same day. Most school kitchens were '
    'originally built for reheating pre-packaged meals, not cooking from scratch. GCS\'s recent '
    'equipment investments (dish machines, serving lines, ovens) improve this baseline, but '
    'additional ventilation, prep space, and refrigeration capacity should be assessed '
    '(USDA FNS; Action for Healthy Kids).'
)

doc.add_heading('Labor Market Constraints', level=2)
doc.add_paragraph(
    'Ninety percent of school nutrition directors reported staffing shortages heading into 2023-24. '
    'The Southeast FNS region (including Tennessee) reported statistically significant higher '
    'challenges with food costs, labor costs, and equipment costs compared to the national average. '
    'Rural districts like GCS face compounded challenges: most food service positions are low-paying, '
    'part-time, and often without benefits. Menu expansion requiring scratch cooking demands '
    'culinary-trained staff that rural labor markets may not supply (SNA, 2025; CRS).'
)

doc.add_heading('Menu Variety Alone Does Not Improve Participation', level=2)
doc.add_paragraph(
    'Research shows that the primary barriers to school meal participation are not menu variety. '
    'During pandemic-era universal free meals, 81.5% of directors cited students\' preference to '
    'eat elsewhere as the top barrier, followed by negative perceptions of food taste (67%). '
    'Studies found limited evidence that taste tests or modified menu options alone significantly '
    'increased participation. In contrast, structural changes like alternative breakfast models '
    '(breakfast in the classroom, grab-and-go) showed consistent increases. This suggests menu '
    'expansion addresses a secondary driver at best (Healthy Eating Research, 2022; Chriqui et al., 2025).'
)

doc.add_heading('Farm-to-School Sustainability Concerns', level=2)
doc.add_paragraph(
    'Districts most likely to cease farm-to-school programs are smaller, more reliant on federal '
    'assistance, and serve higher proportions of students on federal benefits, a profile that '
    'matches GCS. Recent federal funding cuts have disproportionately impacted small districts, '
    'with well-resourced districts able to continue from their own funds while under-resourced '
    'districts lose access (ScienceDirect, 2021; Capital & Main, 2025).'
)

doc.add_heading('Equity Considerations', level=2)
doc.add_paragraph(
    'Menu expansion without adequate resourcing risks benefiting some schools more than others '
    'within the same district. Research documents that districts serving majority students of '
    'color use less scratch cooking compared to majority-White districts. If GCS expands menu '
    'options, ensuring equitable implementation across all six cafeterias is essential. A '
    'hub-and-spoke model where one kitchen supplies all schools could help address this '
    '(Prescott et al., 2025).'
)

# ===== SECTION 7: MEASURABLE OUTCOMES =====
doc.add_heading('7. Measurable Outcomes', level=1)

doc.add_paragraph(
    'Research links school meal program improvements to several measurable outcomes, though the '
    'strength of evidence varies:'
)

table = doc.add_table(rows=9, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Outcome', 'Expected Impact', 'Evidence Strength']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
outcomes_data = [
    ['Meal participation rates', '+3-16% (avg. 9%) with scratch cooking', 'Moderate-Strong'],
    ['Federal reimbursement revenue', 'Proportional to participation increase', 'Strong (direct)'],
    ['Plate waste reduction', 'Up to 76% for milk with self-serve', 'Moderate'],
    ['Diet quality (HEI scores)', '+24 points post-HHFKA', 'Strong'],
    ['Student satisfaction', 'Higher with scratch vs. heat-and-serve', 'Moderate'],
    ['Food costs per meal', 'Lower with scratch cooking', 'Strong'],
    ['Labor costs per meal', 'Higher with scratch cooking', 'Strong'],
    ['Total cost per meal', 'No significant difference', 'Strong'],
]
for r, row_data in enumerate(outcomes_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

# ===== SECTION 8: BOTTOM LINE =====
doc.add_heading('8. Bottom Line', level=1)

doc.add_paragraph(
    'GCS\'s nutrition budget constrains cafeteria options through five interconnected mechanisms: '
    '(1) federal meal pattern requirements that dictate what can be served in reimbursable meals; '
    '(2) Smart Snacks standards that limit a la carte offerings; (3) nonprogram revenue rules that '
    'constrain a la carte pricing; (4) a structural deficit that is eroding the fund balance; and '
    '(5) rising costs for food, labor, and compliance that outpace reimbursement rate increases.'
)

doc.add_paragraph(
    'However, the research reveals that the relationship between budgets and menu options is not '
    'purely a story of constraint. Strategic menu improvements, particularly through scratch or '
    'modified-scratch cooking, can actually improve financial performance by increasing student '
    'participation. The evidence is clear that scratch cooking does not increase total per-meal '
    'costs; it redistributes them from food to labor. And the participation-reimbursement cycle '
    'means that more appealing meals generate more revenue, creating the potential for a virtuous '
    'cycle rather than a vicious one.'
)

doc.add_paragraph(
    'GCS is actually better positioned than many small districts for this transition. Recent '
    'equipment investments ($385K+ in serving lines, dish machines, ovens, and a freezer) have '
    'modernized kitchen infrastructure. The excess fund balance ($454K) provides forced investment '
    'capital that can be directed toward training, small equipment, and menu development. CEP '
    'participation at all four elementary schools simplifies administration and boosts participation. '
    'And perfect health inspection scores across all cafeterias demonstrate operational competence.'
)

doc.add_paragraph(
    'The recommended path forward for GCS involves three phases:'
)

phases = [
    'Phase 1 (Year 1): Optimize existing menu items using speed-scratch techniques. Use excess balance funds for staff training (culinary skills workshops, potentially through Chef Ann Foundation or Tennessee School Nutrition Association), small equipment (food processors, mixers), and recipe development. Implement student taste-testing at each school. Target: 2-3 scratch-made items per week.',
    'Phase 2 (Year 2): Add 1-2 new daily options based on Phase 1 taste-testing data. Explore a hub-and-spoke model where GHS or GMS (largest kitchens) does primary preparation for distribution to elementary sites. Expand high-margin a la carte offerings with scratch-baked goods. Apply for Chef Ann Foundation Get Schools Cooking and FY2027 Farm to School grants. Target: 50% scratch or speed-scratch meals.',
    'Phase 3 (Year 3): Scale successful items district-wide. Evaluate participation data and adjust. Explore local sourcing partnerships with East Tennessee farms. Target: Participation increase of 5-10% with stable or improved financial performance.',
]
for i, phase in enumerate(phases):
    doc.add_paragraph(phase, style='List Number')

doc.add_paragraph(
    'The remaining open questions are primarily operational: Can GCS recruit or train staff with '
    'the culinary skills needed for scratch cooking in the East Tennessee labor market? Will the '
    'hub-and-spoke model work across GCS\'s six geographically dispersed schools? And will student '
    'taste preferences align with the healthier, scratch-made items? These questions are best '
    'answered through the phased pilot approach recommended above rather than through additional '
    'research.'
)

doc.add_paragraph(
    'The biggest risk is inaction. The current trajectory of declining fund balances ($1.88M to '
    '$1.23M in two years) and widening deficits (-$142K to -$516K) is unsustainable. Whether or '
    'not GCS expands menu options, the program needs a strategic intervention to bend the cost curve '
    'or boost revenue. Menu improvements that drive participation represent the most promising path '
    'to doing both simultaneously.'
)

# ===== SECTION 9: REFERENCES =====
doc.add_heading('9. References', level=1)

references = [
    'Action for Healthy Kids. Cafeteria Chronicles: Scratch Cooking Best Practices for School Kitchens. https://healthymealsincentives.org/cafeteria-chronicles-scratch-cooking-best-practices-for-school-kitchens/',
    'Asperin, A. et al. (2018). Operating School Meal Programs in Rural Districts: Challenges and Solutions. Journal of Child Nutrition & Management. https://schoolnutrition.org/wp-content/uploads/2022/06/Operating-School-Meals-in-Rural-Districts-Challenges-and-Solutions-Spring2018.pdf',
    'Bala et al. (2024). Future research directions in choice overload and its moderators. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC11111947/',
    'Bodas, A.R. et al. (2014). Is Scratch-Cooking a Cost-Effective Way to Prepare Healthy School Meals with US Department of Agriculture Foods? Journal of the Academy of Nutrition and Dietetics. https://pubmed.ncbi.nlm.nih.gov/25043447/',
    'Byker-Shanks, C. et al. (2017). Food Waste in the National School Lunch Program 1978-2015: A Systematic Review. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC5660654/',
    'Center for American Progress. Challenges and Opportunities of Providing Free School Meals for All. https://www.americanprogress.org/article/challenges-and-opportunities-of-providing-free-school-meals-for-all/',
    'Chef Ann Foundation. Get Schools Cooking. https://www.chefannfoundation.org/what-we-do/get-schools-cooking/',
    'Chriqui et al. (2025). Foodservice Directors\' Perceived Barriers to Student Participation in School Meals. PubMed. https://pubmed.ncbi.nlm.nih.gov/40524480/',
    'Cohen, J. et al. (2023). Food choice, plate waste and nutrient intake of elementary- and middle-school students. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC10282278/',
    'Congressional Research Service. School Meals and Other Child Nutrition Programs: Background and Funding (R46234). https://www.congress.gov/crs-product/R46234',
    'Congressional Research Service. The School Foodservice Workforce: Characteristics and Labor Market Outcomes (R47199). https://crsreports.congress.gov/product/pdf/R/R47199',
    'Congressional Research Service. USDA\'s Latest Update to Nutrition Standards for School Meals (R47522). https://www.congress.gov/crs-product/R47522',
    'Ed Week (2026). How One Arizona District Turned School Cafeterias Into Scratch Kitchens. https://www.edweek.org/leaders/2026/how-one-arizona-district-turned-school-cafeterias-into-scratch-kitchens',
    'Eustachio Colombo, P. et al. (2023). Improving school lunch menus with multi-objective optimisation. Public Health Nutrition. https://pmc.ncbi.nlm.nih.gov/articles/PMC10410403/',
    'Food Research & Action Center (2025). Reach of School Breakfast and Lunch 2023. https://frac.org/reach-report-2025',
    'Gerber, N. et al. (2023). A qualitative investigation of food waste in a universal free School Breakfast Program. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC10271748/',
    'Greeneville Sun. All City School Cafeterias Earn Perfect Health Scores. https://www.greenevillesun.com/news/local_news/all-city-school-cafeterias-earn-perfect-health-scores/',
    'Healthy Eating Research (2022). Promising Strategies to Increase Student Participation in School Meals. https://healthyeatingresearch.org/wp-content/uploads/2022/11/HER-Meal-Participation-Brief_final.pdf',
    'Kinsey, E.W. et al. (2021). Universal School Meals and Associations with Student Participation. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC8000006/',
    'No Kid Hungry (2023). The Dakota 10: School Nutrition Success with Shared Services. https://bestpractices.nokidhungry.org/sites/default/files/2023-05/THE%20DAKOTA%2010_Fv5.pdf',
    'Prescott, M.P. et al. (2025). Considerations for diverse, equitable, and inclusive school food programs. Health Promotion International. https://pmc.ncbi.nlm.nih.gov/articles/PMC11986204/',
    'School Nutrition Association (2025). SY 2024-25 School Nutrition Trends Report. https://schoolnutrition.org/wp-content/uploads/2025/01/2024-25-School-Nutrition-Trends-Report.pdf',
    'School Nutrition Association (2025). Position Paper: Increase Reimbursements. https://schoolnutrition.org/resource/position-paper-2025-increase-reimbursements/',
    'School Nutrition Association (2024). Strategies to Address Food Waste in K-12 Schools: A Narrative Review. https://schoolnutrition.org/journal/spring-2024-strategies-to-address-food-waste-in-k-12-schools-a-narrative-review/',
    'Schwartz, M.B. et al. (2020). Documented Success and Future Potential of the Healthy, Hunger-Free Kids Act. PMC. https://pmc.ncbi.nlm.nih.gov/articles/PMC7216560/',
    'Tennessee Department of Education. School Nutrition Programs. https://www.tn.gov/education/districts/snp-resources/snp-programs.html',
    'Tennessee Department of Education. Community Eligibility Provision. https://www.tn.gov/education/districts/health-and-safety/school-nutrition/community-eligibility-provision-cep.html',
    'Tennessee Lookout (2025). Tennessee students, farmers bear impact of USDA cuts. https://tennesseelookout.com/2025/03/17/tennessee-students-farmers-bear-impact-of-usda-cuts-to-local-food-programs-for-schools-food-banks/',
    'U.S. GAO (2023). School Meals: USDA Should Address Challenges (GAO-23-105697). https://www.gao.gov/products/gao-23-105697',
    'USDA Economic Research Service. Balancing Nutrition, Participation, and Cost in the NSLP. https://www.ers.usda.gov/amber-waves/2008/september/balancing-nutrition-participation-and-cost-in-the-national-school-lunch-program',
    'USDA Food and Nutrition Service (2019). School Nutrition and Meal Cost Study: Summary of Findings. https://fns-prod.azureedge.us/sites/default/files/resource-files/SNMCS_Summary-Findings.pdf',
    'USDA Food and Nutrition Service (2024). Final Rule: Child Nutrition Programs Meal Patterns. https://www.fns.usda.gov/cn/fr-042524',
    'USDA Food and Nutrition Service (2025). Reimbursement Rates SY 2025-26. https://www.fns.usda.gov/schoolmeals/fr-072425',
    'USDA Food and Nutrition Service. Community Eligibility Provision. https://www.fns.usda.gov/cn/cep',
    'USDA Food and Nutrition Service. NSLP Equipment Assistance Grants. https://www.fns.usda.gov/grant/nslp-equipment-assistance',
    'USDA Food and Nutrition Service. Nonprogram Food Revenue Requirements. https://www.fns.usda.gov/cn/nonprofit-school-food-service-account-nonprogram-food-revenue-requirements',
    'USDA Food and Nutrition Service. Smart Snacks in Schools. https://www.fns.usda.gov/school-meals/nutrition-standards/smartsnacks',
    'USDA Food and Nutrition Service. USDA Foods in Schools. https://www.fns.usda.gov/usda-fis',
    'WBIR (2025). Tennessee to lose $12.6M in funding to bring local food to students. https://www.wbir.com/article/news/local/tennessee-lose-12m-funding-bring-local-food-students-across-the-state/',
    'Zuercher, M.D. et al. (2025). Factors and Outcomes Associated With Using Scratch-Cooked Foods in School Meals in California. Journal of School Health. https://pmc.ncbi.nlm.nih.gov/articles/PMC11860737/',
]

for ref in references:
    doc.add_paragraph(ref, style='List Number')

# ===== APPENDIX A: TENNESSEE DATA =====
doc.add_heading('Appendix A: Tennessee Data', level=1)

doc.add_heading('GCS Fund Balance Trajectory', level=2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Metric', 'FY2023-24', 'FY2024-25']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
gcs_data = [
    ['Starting Balance', '$1,884,244', '$1,742,624'],
    ['Total Reimbursement', '$1,328,345', '$1,581,514'],
    ['Program Revenue', '$2,208,901', '$2,060,150'],
    ['Total Expenses', '$2,350,521', '$2,576,085'],
    ['Surplus/Deficit', '-$141,620', '-$515,934'],
    ['Ending Balance', '$1,742,624', '$1,226,690'],
    ['Excess Balance Determination', '$1,037,468', '$453,864'],
    ['3-Month Average Expenses', '$705,156', '$772,825'],
]
for r, row_data in enumerate(gcs_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('GCS State Claims by Month (FY2022-23 and FY2023-24)', level=2)
table = doc.add_table(rows=13, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Month', 'FY2022-23', 'FY2023-24']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
claims_data = [
    ['August', '$128,136', '$128,749'],
    ['September', '$154,377', '$141,467'],
    ['October', '$120,882', '$121,009'],
    ['November', '$133,079', '$150,806'],
    ['December', '$80,071', '$72,153'],
    ['January', '$132,768', '$118,705'],
    ['February', '$119,318', '$148,169'],
    ['March', '$126,476', '$115,731'],
    ['April', '$144,020', '$181,830'],
    ['May', '$109,240', '$105,558'],
    ['June', '$40,695', '$44,168'],
    ['Total', '$1,289,063', '$1,328,345'],
]
for r, row_data in enumerate(claims_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('GCS Meal Comparison: FY2024-25 vs FY2025-26 (August)', level=2)
table = doc.add_table(rows=8, cols=5)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['School', 'Breakfast 24-25', 'Lunch 24-25', 'Breakfast 25-26', 'Lunch 25-26']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
meal_data = [
    ['Hal Henard', '4,014', '5,287', '4,071', '4,966'],
    ['EastView', '2,900', '4,426', '2,938', '4,387'],
    ['GHS', '2,427', '10,251', '2,351', '9,163'],
    ['GMS', '2,385', '7,304', '2,748', '7,643'],
    ['Highland', '1,971', '2,026', '2,011', '2,372'],
    ['Tusculum View', '3,631', '5,133', '3,969', '5,219'],
    ['Total', '17,328', '34,427', '18,088', '33,750'],
]
for r, row_data in enumerate(meal_data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('Tennessee Lost Federal Local Food Funding (2025)', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
loss_data = [
    ['Program', 'Amount Lost'],
    ['Local Food for Schools', '$12.6 million'],
    ['Local Food Purchase Assistance', '$7.6 million'],
    ['Total', '$20.2 million'],
]
for r, row_data in enumerate(loss_data):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('GCS Excess Balance Spending Plan (FY2024-25)', level=2)
doc.add_paragraph(
    'Total available at 10%: $25,000. Distribution plan included staff stipends across '
    'managers ($5,625), full-time associates ($6,675), part-time associates ($4,160), and '
    'coordinator balance ($8,600). Planned capital expenditures from excess balance included: '
    'water softener pellets ($1,000), waterless hot well at Tusculum View ($9,000), off-site '
    'freezer ($80,000), SN stipend ($25,000), and small wares ($1,500).'
)

doc.add_paragraph(
    'FY2023-24 excess balance spending was more extensive, including: dish machines at TV and EV '
    '($90,000 each), pass-thru warmers at HH ($11,552), GHS ($30,000), and EV ($20,000), pass-thru '
    'fridge at EV ($15,109), Duke waterless serving lines x2 ($175,000), stockroom shelving at all '
    'schools ($31,000 total), Blodgett ovens at GMS and EV ($20,000 each), and salary adjustments '
    'for staff ($52,000).'
)

# Save
output_path = os.path.expanduser(
    "~/Documents/Research/school-nutrition-budget-cafeteria-options/"
    "2026-03-13 School Nutrition Budget Constraints Research Report.docx"
)
doc.save(output_path)
print(f"Report saved to: {output_path}")
