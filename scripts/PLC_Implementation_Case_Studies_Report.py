#!/usr/bin/env python3
"""Generate PLC Implementation Case Studies Research Report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title
title = doc.add_heading('Professional Learning Communities: Implementation Case Studies and Impact on Student Achievement', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Comprehensive Research Summary for District Leadership')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('Prepared February 16, 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ============================================================
# SECTION 1: INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction and Purpose', level=1)
doc.add_paragraph(
    'Professional Learning Communities (PLCs) have become one of the most widely adopted '
    'frameworks for school improvement in American public education over the past three decades. '
    'Rooted in the foundational work of Richard DuFour, Robert Eaker, and Rebecca DuFour at '
    'Adlai E. Stevenson High School in Lincolnshire, Illinois, and further developed through '
    'Shirley Hord\'s research at the Southwest Educational Development Laboratory (SEDL), the '
    'PLC model centers on a deceptively simple premise: when educators work collaboratively in '
    'recurring cycles of collective inquiry and action research, student outcomes improve.'
)
doc.add_paragraph(
    'This report synthesizes implementation case studies from districts across the United States, '
    'with particular attention to small-to-mid-size districts (1,000 to 10,000 students) and '
    'rural or small-city contexts. The goal is to provide district leaders with actionable '
    'intelligence: what specific structures, schedules, staffing models, protocols, and timelines '
    'have produced measurable gains in student achievement, and what pitfalls have derailed '
    'otherwise promising implementations.'
)

# ============================================================
# SECTION 2: THEORETICAL FRAMEWORKS
# ============================================================
doc.add_heading('2. Theoretical Frameworks: DuFour vs. Hord', level=1)

doc.add_heading('2.1 The DuFour/PLC at Work Model', level=2)
doc.add_paragraph(
    'The DuFour model, now branded as "PLC at Work" through Solution Tree, is organized around '
    'four critical questions that drive collaborative team work:'
)
doc.add_paragraph('1. What do we want our students to learn? (Essential standards)', style='List Number')
doc.add_paragraph('2. How will we know our students are learning? (Common formative assessments)', style='List Number')
doc.add_paragraph('3. How will we respond when some students do not learn? (Systematic interventions)', style='List Number')
doc.add_paragraph('4. How will we extend learning for students who are already proficient? (Enrichment)', style='List Number')
doc.add_paragraph(
    'DuFour and Eaker (1998) identified key characteristics of effective PLCs: shared mission, '
    'vision, and values; collective inquiry; collaborative teams; action orientation and '
    'experimentation; continuous improvement; and a results orientation. The model emphasizes '
    'that PLCs are not meetings -- they are an ongoing process embedded in the culture of a school.'
)

doc.add_heading('2.2 The Hord/SEDL Model', level=2)
doc.add_paragraph(
    'Shirley Hord (1997), working through SEDL, identified five attributes of professional '
    'learning communities: (1) supportive and shared leadership, (2) collective creativity, '
    '(3) shared values and vision, (4) supportive conditions (both structural and relational), '
    'and (5) shared personal practice. Hord\'s model places particular emphasis on reflective '
    'dialogue as a vehicle for collective learning and on peer observation as a mechanism for '
    'shared personal practice. While DuFour\'s model leans heavily toward data-driven cycles of '
    'assessment and intervention, Hord\'s framework gives more weight to the relational '
    'conditions and cultural shifts that must precede effective collaboration.'
)
doc.add_paragraph(
    'Both models converge on the critical role of the principal in establishing supportive '
    'conditions, developing a shared vision, and distributing leadership. In practice, most '
    'districts implementing PLCs today draw from both frameworks, though the DuFour/PLC at Work '
    'model has become the dominant commercial framework through Solution Tree\'s professional '
    'development services.'
)

# ============================================================
# SECTION 3: CASE STUDIES
# ============================================================
doc.add_heading('3. District Implementation Case Studies', level=1)

# --- Case Study 1: Sanger ---
doc.add_heading('3.1 Sanger Unified School District, California', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 11,500 students. Rural, predominantly Latino community in the '
          'Central Valley of California.')
doc.add_paragraph(
    'Sanger Unified is one of the most cited turnaround stories in PLC literature. In 2004, '
    'the district was designated one of the lowest-performing districts in California. By 2012, '
    'the district posted a 94% graduation rate for Latino students and a 97% graduation rate '
    'district-wide. Sanger High School\'s Academic Performance Index (API) score rose from 576 '
    'in 1999 to 794 in 2013, a gain of 218 points. The district outperformed average state gains '
    'every year from 2005 forward.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'Schools facilitated weekly PLC meetings as designated time for teachers to plan as a team, '
    'collaborate, and support each other. The district invested heavily in relationship building '
    'both within schools and with the community. Decisions were made based on evidence, with '
    'collaborative teams analyzing student data regularly. A robust leadership development '
    'pipeline was established to build capacity at every level of the system. PLCs were supported '
    'at all levels -- from classroom teams to district leadership.'
)
p = doc.add_paragraph()
run = p.add_run('Key Takeaway: ')
run.bold = True
p.add_run(
    'Sanger demonstrates that a high-poverty, rural district can achieve dramatic gains through '
    'sustained PLC implementation over 8-10 years. The combination of weekly collaboration time, '
    'evidence-based decision making, and leadership pipeline development was critical.'
)

# --- Case Study 2: Whittier Union ---
doc.add_heading('3.2 Whittier Union High School District, California', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 13,000 students across five comprehensive high schools in Los Angeles County.')
doc.add_paragraph(
    'Whittier Union\'s API grew from 635 in 2004 to 765 in 2010, with all five comprehensive '
    'high schools improving. Individual school gains ranged from 106 points (Pioneer High) to '
    '150 points (Whittier High). The achievement gap between ethnic minorities and '
    'socioeconomically disadvantaged students and their peers narrowed significantly, with '
    'traditionally underserved populations showing the greatest overall gains.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'The district created new staffing roles to support PLC work: Intervention Coordinators, '
    'Site Assessment Technicians, Course Leads, and Link Crew Coordinators. Department Chair '
    'job descriptions were revised to emphasize collaborative leadership. The district invested '
    'heavily in building teacher capacity to work collaboratively and take collective '
    'responsibility for student achievement. Common assessments were developed, administered, '
    'and analyzed through structured PLC protocols at every school.'
)
p = doc.add_paragraph()
run = p.add_run('Key Takeaway: ')
run.bold = True
p.add_run(
    'Whittier Union demonstrates the importance of creating new roles and restructuring existing '
    'positions to sustain PLC work. The district recognized that effective teacher leadership is '
    'the most critical component of improving school culture.'
)

# --- Case Study 3: Fillmore ---
doc.add_heading('3.3 Fillmore Unified School District, California', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 3,400 students. Small-city district in Ventura County, California. '
          'Recognized as a Model PLC at Work district.')
doc.add_paragraph(
    'Fillmore\'s mathematics growth score improved from 36 in 2021 to 60 in 2023, and the '
    'English language arts/reading score went from 70 in 2021 to 80 in 2023. One student cohort '
    'improved ELA proficiency by 28% over three years. San Cayetano Elementary, a school within '
    'the district, saw its California API increase from 702 to 822 between 2005 and 2012.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'The district employed Teachers on Special Assignment (TOSAs) as instructional coaches who '
    'provided targeted demonstration lessons, PLC coaching, and data analysis training for '
    'teachers and site administrators. San Cayetano Elementary moved away from teachers working '
    'in silos and compiling data without direction, instead using common assessment results to '
    'improve individual and collective practice and to meet extension and intervention needs of '
    'students. The district partnered with UCLA for learning labs and math professional development.'
)
p = doc.add_paragraph()
run = p.add_run('Key Takeaway: ')
run.bold = True
p.add_run(
    'As a district of only 3,400 students, Fillmore demonstrates that small districts can '
    'successfully implement the PLC at Work model with meaningful results. The use of TOSAs as '
    'embedded coaches is a scalable staffing model for small districts.'
)

# --- Case Study 4: White River ---
doc.add_heading('3.4 White River School District, Washington', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 3,800 students. Small-city/rural district in Buckley, Washington, '
          'recognized as a Model PLC at Work district.')
doc.add_paragraph(
    'White River\'s PLC journey spans over 25 years and represents one of the longest sustained '
    'implementations in the country. In 1998, Mountain Meadow Elementary was the only school '
    'operating as a PLC at Work, and it outperformed other elementary schools in the district, '
    'the region, and most elementary schools in Washington State. From 2021 to 2023, Mountain '
    'Meadow was the highest-performing elementary school in Pierce County. The school was named '
    'an Ambassador Model PLC at Work school in 2024, and White River High School has been a '
    'Model PLC at Work school since 2016.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'District leadership recognized early that inconsistent implementation across schools would '
    'limit results. After seeing Mountain Meadow\'s success, leadership mandated consistent '
    'implementation of the PLC process in every school. The district has sustained its work for '
    '17+ years by staying the course, using the four critical questions of learning to drive all '
    'work, and intentionally limiting competing district initiatives that could dilute focus.'
)
p = doc.add_paragraph()
run = p.add_run('Key Takeaway: ')
run.bold = True
p.add_run(
    'White River demonstrates the power of starting with a proof-of-concept school, then '
    'scaling district-wide. Their strategy of limiting competing initiatives to protect PLC '
    'focus is a critical lesson for small districts with limited bandwidth.'
)

# --- Case Study 5: Mason Crest ---
doc.add_heading('3.5 Mason Crest Elementary, Fairfax County, Virginia', level=2)
p = doc.add_paragraph()
run = p.add_run('School Profile: ')
run.bold = True
p.add_run('Title I elementary school in Annandale, Virginia, within Fairfax County Public Schools.')
doc.add_paragraph(
    'In year one of PLC implementation, Mason Crest\'s test scores exceeded the Commonwealth of '
    'Virginia\'s accreditation benchmarks in all four subject categories: English, math, history, '
    'and science. Within two years, English and science passing percentages had each increased by '
    'at least 7 percent, and Mason Crest was beating the commonwealth\'s passing rates by 5 to '
    '16 percent depending on the subject. By year four, Mason Crest was recognized as a National '
    'Title I Distinguished School and became the first recipient of the DuFour Award. The school '
    'has hosted more than 300 site visits from educators, including visitors from Taiwan and Australia.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'Mason Crest used results from common assessments to improve both individual and collective '
    'practice and to build a systematic intervention/extension system. The school implemented a '
    'multi-year Social Emotional Learning (SEL) integration using monthly grade-level PLC meetings '
    'focused on SEL needs. They piloted lessons using a co-teaching model so successfully that '
    'they scaled it to all grades the following year.'
)

# --- Case Study 6: Kildeer ---
doc.add_heading('3.6 Kildeer Countryside School District 96, Illinois', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 3,500 students. Suburban K-8 district in Lake County, Illinois.')
doc.add_paragraph(
    'Kildeer Countryside began its PLC journey over 22 years ago and represents one of the '
    'most mature implementations in the country. All seven schools in the district have been '
    'recognized as Model PLC at Work schools. Woodlawn Middle School received the 2018 DuFour '
    'Award, one of the highest honors in PLC implementation.'
)
p = doc.add_paragraph()
run = p.add_run('Key Takeaway: ')
run.bold = True
p.add_run(
    'As a K-8 district of approximately 3,500 students, Kildeer demonstrates that small '
    'districts can achieve the highest levels of PLC recognition when the work is sustained '
    'over many years with consistent leadership commitment.'
)

# --- Case Study 7: Arkansas State Initiative ---
doc.add_heading('3.7 Arkansas Statewide PLC at Work Project', level=2)
p = doc.add_paragraph()
run = p.add_run('Scope: ')
run.bold = True
p.add_run('Statewide initiative launched in 2017-2018 involving multiple cohorts of schools, '
          'including many small and rural districts.')
doc.add_paragraph(
    'An independent evaluation by Education Northwest found that within two years, the PLC at '
    'Work model had positive impacts on math achievement test scores. Arkansas\' math gains '
    'exceeded those attributed to other professional learning programs. The evaluation was '
    'designed to meet What Works Clearinghouse (WWC) standards with reservations and establish '
    'ESSA Tier II evidence. Educators reported positive changes in instructional practice, '
    'culture of collaboration, and collective responsibility for student learning.'
)
doc.add_paragraph(
    'Specific school results from the Arkansas initiative include: Marked Tree School District '
    '(approximately 530 students) saw tenth-grade reading scores increase by 23% and '
    'seventh-grade English scores by 15%. Camden Fairview High School achieved its highest-ever '
    'graduation rate of 91%, with reading readiness scores increasing from 12% to 17% and '
    'behavior issues decreasing by 14%. As of September 2024, Arkansas has thirty-nine Model '
    'PLC schools or districts and eight Promising Practices schools.'
)
p = doc.add_paragraph()
run = p.add_run('Caveats: ')
run.bold = True
p.add_run(
    'A University of Arkansas study found no statistically significant improvements in overall '
    'student achievement or growth, with concerning trends for economically disadvantaged '
    'students. Additionally, one study examining the DuFour model found no significant '
    'relationship between implementation phase and third-grade MCA math proficiency. These '
    'mixed findings underscore that implementation fidelity and context matter enormously.'
)

# --- Case Study 8: Academy District 20 ---
doc.add_heading('3.8 Academy District 20, Colorado Springs, Colorado', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 27,000 students. Suburban district in Colorado Springs.')
doc.add_paragraph(
    'Academy District 20 has been Accredited with Distinction for 15 straight years -- one of '
    'only four districts in Colorado to achieve distinction every year since ratings launched in '
    '2008. Five D20 schools ranked in Colorado\'s top 10 schools of 2023. The district has '
    'shown steadily improving results in Math and ELA on standardized assessments, including '
    'both achievement and growth measures.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'The district uses a designated late-start PLC time structure, providing concentrated time '
    'for teacher teams. In addition, teacher collaboration occurs at various times throughout '
    'a typical week. The late starts provide time for teachers to meet with special educators, '
    'special service providers, interventionists, teachers from other grade levels, and even '
    'teachers from other schools.'
)

# --- Case Study 9: Williamson County ---
doc.add_heading('3.9 Williamson County Schools, Tennessee', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 42,000 students across 52 schools. The first district in Tennessee '
          'to be named a Model PLC at Work District.')
doc.add_paragraph(
    'Williamson County Schools posted a state-best ACT composite of 25.3 for the Class of 2025 '
    '(versus the state average of 19.3). The district was designated an Advancing District, with '
    '35 schools designated as Reward Schools. Elementary proficiency rates stand at 67% for '
    'reading and 76% for math. While Williamson County is a larger, more affluent district than '
    'the primary focus of this report, its Tennessee context and Model PLC designation make it '
    'a relevant reference point for Tennessee districts of any size.'
)

# --- Case Study 10: Mount Vernon ---
doc.add_heading('3.10 Mount Vernon City School District, New York', level=2)
p = doc.add_paragraph()
run = p.add_run('District Profile: ')
run.bold = True
p.add_run('Approximately 8,000 students in Westchester County, New York.')
doc.add_paragraph(
    'Mount Vernon designed and implemented a systemic PLC structure for ongoing, job-embedded '
    'professional development for both teachers and principals across grade, school, and district '
    'levels. The district modified master schedules to provide educators the time and space '
    'needed for PLC meetings. They developed a comprehensive PLC Handbook with district-specific '
    'protocols and norms, and used Office 365 to facilitate collaboration and resource sharing.'
)
p = doc.add_paragraph()
run = p.add_run('What They Did: ')
run.bold = True
p.add_run(
    'Mount Vernon organized PLCs into cross-school clusters, pairing schools with robust PLCs '
    'with those just beginning. They paired the highest-performing schools with the lowest-performing '
    'schools and ensured that each cluster contained representation from elementary, middle, and '
    'high schools. PLCs established norms as the first order of business, with protocols for '
    'looking at student work that included criteria for authentic student performance, assessment, '
    'and instruction.'
)

# --- Case Study 11: Stevenson ---
doc.add_heading('3.11 Adlai E. Stevenson High School, Lincolnshire, Illinois (The Origin Story)', level=2)
p = doc.add_paragraph()
run = p.add_run('School Profile: ')
run.bold = True
p.add_run('Large suburban high school (approximately 4,500 students). The birthplace of PLC at Work.')
doc.add_paragraph(
    'Under Richard DuFour\'s leadership as principal (1983-1991) and superintendent (1991-2002), '
    'Stevenson became the first public high school in Illinois to receive four Blue Ribbon Awards '
    'for Excellence from the U.S. Department of Education (1987, 1991, 1998, 2002). The U.S. '
    'Department of Education called it "the most recognized and celebrated high school in the '
    'United States in the 1990s." Teachers met regularly in course-specific and interdisciplinary '
    'teams to improve teaching strategies, create common assessments, and refine lesson plans. '
    'Stevenson remains a foundational case study that demonstrates the full arc of PLC maturation.'
)

# ============================================================
# SECTION 4: ACHIEVEMENT DATA TABLES
# ============================================================
doc.add_heading('4. Before-and-After Achievement Data', level=1)
doc.add_paragraph(
    'The following tables consolidate publicly available achievement data from districts and '
    'schools that implemented PLCs. Data should be interpreted cautiously, as improvements may '
    'reflect multiple concurrent initiatives.'
)

# Table 1: API/Achievement Scores
doc.add_heading('Table 1: Academic Performance Index / Achievement Score Changes', level=3)
table1 = doc.add_table(rows=7, cols=5)
table1.style = 'Light Grid Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['District/School', 'Enrollment', 'Baseline Score (Year)', 'Post-PLC Score (Year)', 'Change']
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

data1 = [
    ['Sanger High School, CA', '~2,800', 'API: 576 (1999)', 'API: 794 (2013)', '+218 pts'],
    ['Whittier Union HSD, CA (district)', '~13,000', 'API: 635 (2004)', 'API: 765 (2010)', '+130 pts'],
    ['San Cayetano Elem., CA', '~485', 'API: 702 (2005)', 'API: 822 (2012)', '+120 pts'],
    ['Fillmore USD, CA (Math Growth)', '~3,400', '36 (2021)', '60 (2023)', '+24 pts'],
    ['Fillmore USD, CA (ELA)', '~3,400', '70 (2021)', '80 (2023)', '+10 pts'],
    ['Camden Fairview HS, AR (Grad Rate)', '~1,200', 'Below 91%', '91% (highest ever)', 'Record high'],
]
for row_idx, row_data in enumerate(data1, 1):
    for col_idx, val in enumerate(row_data):
        cell = table1.rows[row_idx].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Table 2: Arkansas PLC at Work Results
doc.add_heading('Table 2: Arkansas PLC at Work Project -- Selected School Results', level=3)
table2 = doc.add_table(rows=5, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['School/Metric', 'Before PLC', 'After PLC', 'Change']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

data2 = [
    ['Marked Tree -- 10th Grade Reading', 'Baseline', '+23% increase', '+23%'],
    ['Marked Tree -- 7th Grade English', 'Baseline', '+15% increase', '+15%'],
    ['Camden Fairview HS -- Reading Readiness', '12%', '17%', '+5 pct pts'],
    ['Camden Fairview HS -- Behavior Issues', 'Baseline', '-14% decrease', '-14%'],
]
for row_idx, row_data in enumerate(data2, 1):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

# Table 3: Mason Crest
doc.add_heading('Table 3: Mason Crest Elementary, VA -- State Accreditation Pass Rates', level=3)
table3 = doc.add_table(rows=3, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Subject', 'Year 1 Status', 'Year 2 Change', 'Comparison to State']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

data3 = [
    ['English', 'Exceeded state benchmark', '+7% or more by Year 2', 'Beat state by 5-16%'],
    ['Science', 'Exceeded state benchmark', '+7% or more by Year 2', 'Beat state by 5-16%'],
]
for row_idx, row_data in enumerate(data3, 1):
    for col_idx, val in enumerate(row_data):
        cell = table3.rows[row_idx].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# ============================================================
# SECTION 5: SCHEDULING AND PROTECTED TIME
# ============================================================
doc.add_heading('5. How Districts Protect PLC Time in the Master Schedule', level=1)
doc.add_paragraph(
    'Time is the most frequently cited barrier to effective PLC implementation. Districts that '
    'have achieved sustained results have employed several scheduling strategies:'
)

doc.add_heading('5.1 Late-Start or Early-Release Days', level=2)
doc.add_paragraph(
    'Academy District 20 in Colorado uses a designated late-start PLC time, providing '
    'concentrated blocks for teacher collaboration. This approach allows teams to work with '
    'special educators, interventionists, and teachers from other grade levels. Some districts '
    'implement weekly late starts (e.g., 90 minutes before students arrive) while others use '
    'bi-weekly or monthly structures.'
)

doc.add_heading('5.2 Embedded Intervention/Enrichment Blocks', level=2)
doc.add_paragraph(
    'Several districts have restructured their bell schedules to embed 20-30 minutes of '
    'intervention and enrichment time into each instructional day. This serves a dual purpose: '
    'it provides systematic response time for students who need additional support (Question 3 '
    'of the PLC framework) and it creates common time for teachers to meet when specialists '
    'cover the intervention block.'
)

doc.add_heading('5.3 Common Planning Periods by Grade Level or Content Area', level=2)
doc.add_paragraph(
    'Elementary schools increasingly construct master schedules with common planning time for '
    'grade-level teams. Schools identify and color-code times in the schedule that are essential '
    '(protected instructional time) versus non-essential, ensuring students are not pulled for '
    'services during core learning and that PLC meeting times are genuinely available. Some PLC '
    'structures provide more than three hours of weekly common planning time, plus an additional '
    'full day per quarter for extended collaboration.'
)

doc.add_heading('5.4 The Jigsaw Approach to Elementary Scheduling', level=2)
doc.add_paragraph(
    'As described by AllThingsPLC, a jigsaw approach to the elementary master schedule rotates '
    'specialist classes (art, music, PE, library) to free grade-level teams for simultaneous '
    'planning blocks. This requires careful coordination but yields consistent, weekly '
    'collaboration time without reducing instructional minutes.'
)

doc.add_heading('5.5 Adjusted School Day for Secondary PLCs', level=2)
doc.add_paragraph(
    'The Littlestown Area School District in Pennsylvania adjusted the school day to allow '
    'dedicated time for teachers to collaborate. Five-plus years into implementation, results '
    'showed a re-cultured school system that included systematic, job-embedded, and site-specific '
    'collaboration. For secondary schools, creating common planning for content-area teams '
    'often requires the most creative scheduling, particularly when dealing with singleton '
    'teachers (the only teacher of a particular subject) -- a common challenge in small districts.'
)

# ============================================================
# SECTION 6: STAFFING MODELS
# ============================================================
doc.add_heading('6. Staffing Models: Who Leads PLCs and How Facilitators Are Trained', level=1)

doc.add_heading('6.1 Instructional Coaches / Teachers on Special Assignment (TOSAs)', level=2)
doc.add_paragraph(
    'Fillmore USD and many Colorado districts assign one instructional coach to each PLC or '
    'cluster of PLCs. Coaches provide feedback, support, and serve as the point of contact '
    'for emergent needs. They are trained specifically on data analysis, reflection practices, '
    'and facilitation skills for supporting healthy team dynamics. In one high-implementation '
    'district, the coach introduced an instructional practice or standard in the first week of '
    'a three-week PLC cycle, with teachers implementing in week two and reflecting on results '
    'in week three.'
)

doc.add_heading('6.2 Teacher Leaders and Course Leads', level=2)
doc.add_paragraph(
    'Whittier Union created Course Lead positions and revised Department Chair job descriptions '
    'to emphasize collaborative leadership. ESC Region 13 in Texas offers specialized training '
    'for campus and teacher leaders on tools and strategies for leading PLC teams. The '
    'Supporting Teacher Effectiveness Project (STEP) uses an asset-based approach, helping '
    'educators discover and replicate what is already working in their schools rather than '
    'importing external solutions.'
)

doc.add_heading('6.3 Principal as PLC Leader', level=2)
doc.add_paragraph(
    'SEDL\'s multi-state case study research found that the role of campus-based leadership is '
    'the most critical factor in PLC progression. Principals played the main role in starting '
    'positive change, sharing leadership, and creating commitment to common goals. Both the '
    'DuFour and Hord models place strong emphasis on the principal\'s role in establishing '
    'supportive conditions and developing a shared vision.'
)

doc.add_heading('6.4 Building Regional Capacity', level=2)
doc.add_paragraph(
    'Colorado\'s Department of Education trains multiple teachers as coaches and establishes '
    'weekly PLCs using district-wide protocols, building regional capacity through educational '
    'cooperatives. Kentucky\'s Department of Education uses PLCs as a key strategy within its '
    'achievement gap closing framework, with engagement through the Centers of Regional '
    'Excellence model that Tennessee also employs through its eight CORE offices.'
)

# ============================================================
# SECTION 7: PROTOCOLS AND TOOLS
# ============================================================
doc.add_heading('7. Meeting Protocols, Data Analysis Tools, and Resources', level=1)

doc.add_heading('7.1 The Four-Question Protocol', level=2)
doc.add_paragraph(
    'The foundational meeting structure in most PLC implementations revolves around the four '
    'critical questions. PLC meeting agendas typically include: team norms (reviewed at the '
    'start), SMART goal(s), the specific critical question being addressed, topics of '
    'discussion, estimated time per topic, and action items with owners and deadlines. '
    'Massachusetts\' Department of Secondary and Elementary Education provides a detailed '
    'PLC Guide (Module 3) with specific facilitation protocols.'
)

doc.add_heading('7.2 Common Formative Assessment Data Protocol', level=2)
doc.add_paragraph(
    'Solution Tree\'s "Analyzing Common Formative Assessment Data Protocol" is widely used. '
    'The process works as follows: during planning (ideally before the unit begins), teachers '
    'chunk assessment items by essential learning standards, determine rigor of items, make '
    'common scoring agreements, and determine proficiency thresholds. After assessment, '
    'teachers determine the percentage of students proficient by target and compile data onto '
    'one document for the collaborative team to view. The team then identifies the teacher '
    'with the strongest results per target and analyzes what instructional strategies produced '
    'those results.'
)

doc.add_heading('7.3 Student Work Analysis Protocols', level=2)
doc.add_paragraph(
    'When meeting to share student work, PLC teams follow structured protocols for professional '
    'dialogue. These protocols include criteria for evaluating authentic student performance, '
    'assessment quality, and instructional alignment. The goal is to determine whether students '
    'are showing progress toward meeting priority standards and to identify inconsistencies in '
    'academic language, rigor expectations, or instructional approaches across classrooms.'
)

doc.add_heading('7.4 Norms Development', level=2)
doc.add_paragraph(
    'Mount Vernon\'s PLC Handbook emphasizes that the first task for any PLC team is to '
    'establish norms and a method to monitor them. Norms represent protocols and commitments '
    'to guide members in working together. Effective practice includes recording norms on each '
    'meeting\'s agenda and having members rate their adherence to norms at the end of each '
    'meeting. Common norms address: starting and ending on time, staying focused on agenda, '
    'listening respectfully, using data rather than opinions, committing to action items, and '
    'maintaining confidentiality about individual student/teacher data.'
)

doc.add_heading('7.5 Technology Tools', level=2)
doc.add_paragraph(
    'Assessment platforms like Formative provide data-collection tools designed specifically for '
    'results-oriented PLCs, facilitating collaboration and helping teams become action-oriented '
    'with lesson delivery, assessment, and data analysis in one platform. PLC leaders can be '
    'given admin rights to share and analyze team data, and activities can be published to an '
    'organization-wide library. Mount Vernon used Office 365 sites for cross-school collaboration '
    'and resource sharing.'
)

# ============================================================
# SECTION 8: TIMELINES
# ============================================================
doc.add_heading('8. Implementation Timelines: How Long Before Results Appear', level=1)
doc.add_paragraph(
    'Research and case studies consistently show that meaningful PLC implementation is a '
    'multi-year endeavor. The following timeline patterns emerge from the evidence:'
)

doc.add_paragraph('Years 1-2: Foundation building. Establishing norms, developing common '
    'assessments, building trust, restructuring schedules. Some districts see early gains in '
    'targeted areas (Mason Crest exceeded state benchmarks in Year 1). The Arkansas evaluation '
    'found positive math impacts within two years.', style='List Bullet')
doc.add_paragraph('Years 3-5: Deepening practice. Collaborative teams become more skilled at '
    'data analysis, intervention design, and honest reflection. Sanger, Whittier, and Fillmore '
    'all showed significant gains in this window. The Littlestown Area School District showed a '
    '"re-cultured" system after five years.', style='List Bullet')
doc.add_paragraph('Years 5-10: Systemic transformation. White River expanded from one school to '
    'district-wide implementation over this period. Sanger went from one of California\'s '
    'lowest-performing districts to a recognized turnaround success.', style='List Bullet')
doc.add_paragraph('Years 10+: Mature and sustained. Kildeer Countryside (22+ years), White River '
    '(25+ years), and Stevenson High School represent mature PLCs where the process is deeply '
    'embedded in district culture. These districts show that sustained commitment yields '
    'sustained results, but also that the work never truly "finishes."', style='List Bullet')

doc.add_paragraph(
    'Minnesota\'s PLC Roadmap notes that the time required depends on where a district enters '
    'the process and emphasizes that the process moves more efficiently when led by committed '
    'district and school leaders, Leadership Implementation Teams, and adults who share a '
    'commitment toward achieving high-functioning PLCs.'
)

# ============================================================
# SECTION 9: WHAT DIDN'T WORK
# ============================================================
doc.add_heading('9. Lessons Learned: What Did Not Work', level=1)

doc.add_paragraph(
    'The research literature and practitioner accounts identify several common failure modes in '
    'PLC implementation:'
)

doc.add_heading('9.1 PLCs as Meetings Rather Than a Process', level=2)
doc.add_paragraph(
    'The most frequently cited failure is treating PLCs as scheduled meetings rather than as a '
    'fundamental shift in how educators work together. When PLCs function more like '
    'administrative meetings than meaningful collaborative professional learning, they fail to '
    'improve instruction. In one district, teams would bring different data to analyze every '
    'week and discuss strategies, but once they left the room, those strategies did not translate '
    'into classroom action because teams were always "moving on to the next thing" without '
    'follow-through.'
)

doc.add_heading('9.2 The "One Expert" Problem', level=2)
doc.add_paragraph(
    'When the lead or veteran teacher with the best results dominates PLC discussions, the '
    'PLC happens TO people instead of WITH them. Teachers mimic collaboration only to return '
    'to their classrooms and do what they think will work individually. True PLCs require '
    'collective inquiry, not one-directional knowledge transfer.'
)

doc.add_heading('9.3 Insufficient Time and Competing Priorities', level=2)
doc.add_paragraph(
    'Time is the most cited barrier. Adding collaborative planning without carving out dedicated, '
    'protected time leads to burnout and low engagement. Districts that layer PLC expectations '
    'on top of existing meeting loads without removing other obligations consistently report '
    'poor results. White River\'s strategy of intentionally limiting competing district '
    'initiatives is the counter-example.'
)

doc.add_heading('9.4 Singleton and Multi-Prep Teacher Challenges', level=2)
doc.add_paragraph(
    'Secondary schools, particularly small ones, struggle with teachers who teach multiple preps '
    'or are the only teacher of a particular subject. Administrative teams often do not know what '
    'to do with singletons, and when paired with secondary teachers\' content-driven nature, this '
    'creates an environment that is not conducive to cross-disciplinary collaboration.'
)

doc.add_heading('9.5 Leadership Turnover', level=2)
doc.add_paragraph(
    'Staff turnover, especially at the leadership level, is a significant obstacle. When a '
    'principal or superintendent who championed PLCs departs, the work can stall or revert if '
    'the culture has not been deeply embedded. Distributed leadership and documented systems '
    '(like Mount Vernon\'s PLC Handbook) help mitigate this risk.'
)

doc.add_heading('9.6 Incoherence and Insularity', level=2)
doc.add_paragraph(
    'Learning Forward identifies five common challenges: incoherence (PLC work disconnected from '
    'school improvement goals), insularity (teams working in isolation from other teams), unequal '
    'participation, congeniality (being "nice" instead of honest), and privacy (unwillingness to '
    'share practice). Each of these can hollow out PLCs from the inside while maintaining the '
    'appearance of implementation.'
)

doc.add_heading('9.7 Mixed Research Results as a Caution', level=2)
doc.add_paragraph(
    'The University of Arkansas study found no statistically significant improvements in overall '
    'student achievement from the Arkansas PLC at Work project, with concerning trends for '
    'economically disadvantaged students. A separate study found no significant relationship '
    'between implementation phase of the DuFour model and third-grade math proficiency. These '
    'findings do not mean PLCs are ineffective, but they do mean that implementation quality '
    'and fidelity matter enormously. Simply adopting the label without transforming practice '
    'does not produce results.'
)

# ============================================================
# SECTION 10: RESEARCH BASE AND INDEPENDENT EVALUATIONS
# ============================================================
doc.add_heading('10. Research Base and Independent Evaluations', level=1)
doc.add_paragraph(
    'The PLC at Work model has been evaluated through multiple lenses. The Evidence for ESSA '
    'database lists PLC at Work as a recognized program. Education Northwest\'s independent '
    'evaluation of the Arkansas project established ESSA Tier II evidence for the math '
    'achievement impact. A large-scale study involving more than 200 schools in four countries '
    'with more than 750,000 students found a strong link between gains in student achievement '
    'and long-term PLC implementation. Vescio, Ross, and Adams conducted a review of research on '
    'the impact of PLCs on teaching practice and student learning, finding that well-developed '
    'PLCs have positive impacts on both. However, the research base also shows that poorly '
    'implemented PLCs produce negligible or no results, reinforcing that the process, not just '
    'the structure, is what matters.'
)

# ============================================================
# SECTION 11: IMPLICATIONS FOR SMALL/MID-SIZE DISTRICTS
# ============================================================
doc.add_heading('11. Implications for Small-to-Mid-Size and Rural Districts', level=1)
doc.add_paragraph(
    'Several patterns emerge from the case studies that are particularly relevant to small '
    'districts (1,000-10,000 students) and rural contexts:'
)

doc.add_paragraph('Start small and scale. White River started with one school and expanded '
    'district-wide once proof of concept was established. Small districts can pilot PLC '
    'practices in one or two schools before a full rollout.', style='List Bullet')
doc.add_paragraph('Use TOSAs or part-time coaching models. Fillmore and Colorado districts '
    'show that even small districts can embed coaching by reassigning existing staff as '
    'Teachers on Special Assignment rather than hiring new positions.', style='List Bullet')
doc.add_paragraph('Protect time fiercely. The schedule is the budget of a school. Without '
    'dedicated, protected, and consistent PLC time in the master schedule, implementation will '
    'fail regardless of other investments.', style='List Bullet')
doc.add_paragraph('Address the singleton problem directly. Small secondary schools must be '
    'creative about forming collaborative teams -- vertical teams, interdisciplinary teams, '
    'or cross-district virtual teams are all viable options.', style='List Bullet')
doc.add_paragraph('Leverage close relationships. Research on rural PLCs notes that the close '
    'interpersonal relationships common in small districts can be an asset for building trust '
    'and collaboration, though these same relationships can also enable avoidance of difficult '
    'conversations about practice.', style='List Bullet')
doc.add_paragraph('Limit competing initiatives. White River\'s strategy of staying the course '
    'and limiting district initiatives to protect PLC focus is especially important for small '
    'districts with limited administrative bandwidth.', style='List Bullet')
doc.add_paragraph('Plan for 3-5 years minimum. Districts should set realistic expectations that '
    'significant, measurable gains typically appear in the 3-5 year window, with deeper cultural '
    'transformation requiring 5-10 years.', style='List Bullet')
doc.add_paragraph('Document everything. Mount Vernon\'s PLC Handbook approach protects against '
    'leadership turnover and ensures institutional knowledge persists. Small districts are '
    'especially vulnerable to the loss of a single champion.', style='List Bullet')

# ============================================================
# SECTION 12: SOURCES
# ============================================================
doc.add_heading('12. Sources', level=1)

sources = [
    '1. AllThingsPLC (Solution Tree). "Model PLC at Work." https://allthingsplc.info/model-plc-at-work/',
    '2. AllThingsPLC (Solution Tree). "PLC Master Schedules Provide Time for Collaboration and Interventions During the Day." https://allthingsplc.info/plc-master-schedules-provide-time-for-collaboration-and-interventions-during-the-day/',
    '3. AllThingsPLC (Solution Tree). "From Theory to Practice: A Jigsaw Approach to an Elementary Master Schedule." https://allthingsplc.info/from-theory-to-practice-a-jigsaw-approach-to-an-elementary-master-schedule/',
    '4. AllThingsPLC (Solution Tree). "10 Steps to Creating a PLC Culture." https://allthingsplc.info/blog/view/155/10-steps-to-creating-a-plc-culture',
    '5. Arkansas Division of Elementary and Secondary Education. "Professional Learning Communities for Arkansas." https://dese.ade.arkansas.gov/Offices/special-projects/professional-learning-communities-for-arkansas',
    '6. California Department of Education. "Alignment and Coherence: Sanger Unified." https://www.cde.ca.gov/ci/pl/plstoriesacsusd.asp',
    '7. Colorado Department of Education. "Professional Learning Communities (PLCs) Strategy Guide v2.1." https://www.cde.state.co.us/uip/strategyguide-plcv2',
    '8. Colorado Department of Education. "Coaching Strategy Guide." https://www.cde.state.co.us/uip/strategyguide-coaching',
    '9. Corwin Connect (2023). "Overcoming Challenges in PLC." https://corwin-connect.com/2023/03/overcoming-challenges-in-plc/',
    '10. DuFour, R. (2004). "Professional Learning Communities: A Bandwagon, an Idea Worth Considering, or Our Best Hope for High Levels of Learning?" ResearchGate. https://www.researchgate.net/publication/234611729',
    '11. Education Northwest. "Independent Evaluation Validates Success of PLC at Work Project in Arkansas." https://educationnorthwest.org/insights/independent-evaluation-validates-success-plc-work-project-arkansas',
    '12. Education Northwest. "The PLC at Work in Arkansas Evaluation: Implementation, Impact, and Methodology." https://educationnorthwest.org/insights/plc-work-arkansas-evaluation-implementation-impact-and-methodology',
    '13. Education World. "Why Don\'t Professional Learning Communities Work?" https://www.educationworld.com/a_admin/professional-learning-community-pitfalls-best-practices.shtml',
    '14. Edutopia. "Designing a Master Schedule to Support All Teachers." https://www.edutopia.org/article/designing-master-schedule-schools/',
    '15. ERIC (2018). "The Effective Implementation of Professional Learning Communities." EJ1194725. https://files.eric.ed.gov/fulltext/EJ1194725.pdf',
    '16. Evidence for ESSA. "Professional Learning Communities at Work (PLC at Work)." https://www.evidenceforessa.org/program/professional-learning-communities-at-work-plc-at-work/',
    '17. Fisher, S. "PLCs: A Recipe for Success?" St. John Fisher University ETD. https://fisherpub.sjf.edu/cgi/viewcontent.cgi?article=1603&context=education_etd',
    '18. Hord, S. (1997). "Professional Learning Communities: Communities of Continuous Inquiry and Improvement." SEDL. https://sedl.org/pubs/change34/plc-cha34.pdf',
    '19. Innovate Public Schools. "Turnaround Spotlight: Sanger Unified School District." https://innovateschools.org/effective-education-policies/turnaround-spotlight-sanger-unified-school-district/',
    '20. Insight Education Group. "PLC Model: Supporting Teacher Effectiveness Project." https://www.insighteducationgroup.com/plc-supporting-teacher-effectiveness-project',
    '21. Kentucky Department of Education. "Professional Learning Communities (PLCs)." https://www.education.ky.gov/school/stratclsgap/Pages/plc.aspx',
    '22. Learning Forward. "Overcome 5 PLC Challenges." https://learningforward.org/journal/learning-better-by-learning-together/overcome-5-plc-challenges/',
    '23. Loyola Marymount University. "The Positive Impacts of a Professional Learning Community Model on Student Achievement in Small Schools." https://digitalcommons.lmu.edu/ce/v26/iss2/2/',
    '24. Martinez, B. "Professional Learning Community: Perspectives of Rural School Teachers." Abilene Christian University. https://digitalcommons.acu.edu/etd/294/',
    '25. Massachusetts Department of Elementary and Secondary Education. "DSAC/ESE PLC Guide, Module 3: Building Effective PLC Teams." https://www.mass.gov/doc/download-module-3-0/download',
    '26. Minnesota Department of Education. "Professional Learning Community (PLC) Roadmap." https://education.mn.gov/MDE/dse/edev/prac/plc/',
    '27. New Teacher Center. "Beyond Meetings: Designing PLCs That Drive Instructional Growth." https://newteachercenter.org/resources/beyond-meetings-designing-plcs-that-drive-instructional-growth-community-school-transformation/',
    '28. New York State Education Department. "Mount Vernon City School District PLC Handbook." https://www.nysed.gov/sites/default/files/STLE%20PLC%20Handbook%20MVCSD.pdf',
    '29. Nova Southeastern University. "Stakeholder Participation and Perceptions in Professional Learning Communities: A Case Study in a Small, Rural School District." https://nsuworks.nova.edu/fse_etd/596/',
    '30. Rivet Education. "What\'s the Right Way to Structure PLCs?" https://riveteducation.org/whats-the-right-way-to-structure-plcs/',
    '31. Rosebrock, C. "The Complete Guide for Implementing Professional Learning Communities." https://www.carrierosebrock.com/post/the-complete-guide-for-implementing-professional-learning-communities',
    '32. Schoolytics (2023). "A Data-Driven Guide to PLCs." https://www.schoolytics.com/blog/2023/04/data-driven-plc',
    '33. SEDL (2014). "Implementing Effective Professional Learning Communities." Insights Vol. 2, No. 3. https://sedl.org/insights/2-3/',
    '34. Solution Tree. "Evidence of Excellence for PLC at Work." https://www.solutiontree.com/plc-at-work/evidence-of-excellence',
    '35. Solution Tree. "PLC at Work Project Leads to Higher Student Achievement in Arkansas." https://www.solutiontree.com/press-release/plc-at-work-higher-student-achievement-arkansas',
    '36. Solution Tree. "Evidence of Excellence: Mason Crest Elementary." https://www.solutiontree.com/plc-at-work/evidence-of-excellence/mason-crest',
    '37. Solution Tree. "Analyzing Common Formative Assessment Data Protocol." https://cloudfront-s3.solutiontree.com/pdfs/Reproducibles_CTPBPLC/analyzingcommonformativeassessmentdataprotocol.pdf',
    '38. Southern Regional Education Board. "Tennessee - Professional Learning." https://www.sreb.org/post/tennessee-professional-learning',
    '39. Spark Repository, Bethel University. "The Relationship between the Implementation Phase of the DuFour Model of Professional Learning Communities and Students\' Achievement." https://spark.bethel.edu/etd/742/',
    '40. Texas Tech University. "A Comparative Case Study of 21st Century Focused Professional Development Embedded Within Biology PLCs." https://ttu-ir.tdl.org/items/03f51390-f9a8-409d-af3a-1f27d4d095f9',
    '41. The Education Magazine. "Adlai E. Stevenson High School: How a Suburban Behemoth Became a Blueprint for Public Education." https://www.theeducationmagazine.com/adlai-e-stevenson-high-school/',
    '42. University of Arkansas Office for Education Policy. "Professional Learning Communities and Student Outcomes." https://scholarworks.uark.edu/oepreport/101/',
    '43. University of North Texas. "Professional Learning Communities and School Improvement: Implications for District Leadership." https://digital.library.unt.edu/ark:/67531/metadc1157636/',
    '44. Vescio, V., Ross, D., & Adams, A. "A Review of Research on the Impact of Professional Learning Communities on Teaching Practice and Student Learning." https://www.psycholosphere.com/A%20review%20on%20research%20on%20the%20impact%20of%20PLCs%20on%20teaching%20practice%20&%20student%20learning%20by%20Vescio,%20Ross%20&%20Adams.pdf',
    '45. Whittier Daily News (2009). "Whittier Union Emulated by School Districts Near and Far." https://www.whittierdailynews.com/2009/12/25/whittier-union-emulated-by-school-districts-near-and-far/',
    '46. Williamson Scene (2024). "WCS First in State to Be Named Model PLC at Work District." https://www.williamsonscene.com/schools/wcs-first-in-state-to-be-named-model-plc-at-work-district/',
    '47. Academy District 20. "Professional Learning Communities (PLC)." https://www.asd20.org/professional-learning-communities/',
    '48. School Scheduling Associates. "Adding an Extended Planning Block into the Master Schedule for PLC Time." https://www.schoolschedulingassociates.com/adding-an-extended-planning-block-into-the-master-schedule-for-plc-time/',
    '49. EdWeek (2015). "Professional Learning Communities Still Work (If Done Right)." https://www.edweek.org/leadership/opinion-professional-learning-communities-still-work-if-done-right/2015/10',
    '50. Frontiers in Education (2021). "Practices of Professional Learning Communities." https://www.frontiersin.org/journals/education/articles/10.3389/feduc.2021.617613/full',
]

for s in sources:
    p = doc.add_paragraph(s, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

# Save
output_dir = os.path.expanduser('~/Documents/Research Reports')
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'PLC Implementation Case Studies - Research Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
