#!/usr/bin/env python3
"""Generate research report: AI Performs Better for Polite Users"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import date

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold_first=False):
    p = doc.add_paragraph()
    if bold_first and ':' in text:
        parts = text.split(':', 1)
        run = p.add_run(parts[0] + ':')
        run.bold = True
        p.add_run(parts[1])
    else:
        p.add_run(text)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level > 0:
        p.style = doc.styles['List Bullet 2'] if 'List Bullet 2' in [s.name for s in doc.styles] else doc.styles['List Bullet']
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    run = p.add_run(text)
    return p

def add_citation(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Does AI Perform Better for Polite Users?', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Research Review of Prompt Politeness, Emotional Framing, and LLM Performance')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Compiled {date.today().strftime("%B %d, %Y")} | General Research Report')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()  # spacer

# ============================================================
# 1. DEFINITION & BACKGROUND
# ============================================================
add_heading('1. Definition and Background')

add_para(
    'The claim that "AI performs better for polite users" has become one of the most widely '
    'discussed prompt engineering beliefs since ChatGPT went mainstream in late 2022. The core '
    'idea is simple: if you say "please," "thank you," and frame your requests courteously, '
    'large language models (LLMs) will produce higher-quality, more helpful, and more accurate '
    'responses than if you issue terse or rude commands.'
)

add_para(
    'This belief sits at the intersection of several research threads: prompt engineering, '
    'reinforcement learning from human feedback (RLHF), the psychology of human-computer '
    'interaction, and the broader phenomenon of LLM sensitivity to seemingly irrelevant prompt '
    'features. It also touches on deeper questions about whether AI systems "understand" social '
    'cues or are simply pattern-matching against training data where polite language correlates '
    'with high-quality content.'
)

add_heading('Key Terms', level=2)

add_bullet('Prompt politeness: The inclusion of courteous language (please, thank you, I appreciate your help) in prompts to LLMs.')
add_bullet('Emotional prompting (EmotionPrompt): Appending emotionally salient phrases to prompts, such as "This is very important to my career" or "Are you sure that is your final answer?"')
add_bullet('Sycophancy: The tendency of RLHF-trained models to agree with users, tell them what they want to hear, or prioritize agreeableness over accuracy.')
add_bullet('RLHF (Reinforcement Learning from Human Feedback): The training technique where human raters score model outputs, creating reward signals that shape model behavior.')
add_bullet('CASA (Computers Are Social Actors): A psychological framework from the 1990s showing that humans automatically apply social rules to computers.')

add_heading('Historical Context', level=2)

add_para(
    'The question predates LLMs. Clifford Nass and Byron Reeves at Stanford established in '
    'the 1990s that people automatically apply social norms to computers, including politeness '
    'conventions, even when they know the computer is not a person (Nass & Reeves, 1996, "The '
    'Media Equation"). Their CASA paradigm, supported by 35+ studies, showed that people are '
    'polite to computers, apply gender stereotypes to computer voices, and experience social '
    'facilitation effects with machines. The emergence of conversational AI has amplified these '
    'dynamics considerably. Cross-cultural research on human-robot interaction (Nomura et al., '
    '2008) further established that cultural attitudes toward politeness significantly mediate '
    'human-AI interaction quality, a finding now amplified by conversational AI.'
)

add_para(
    'The modern version of the debate intensified in April 2025 when a viral post on X asked '
    'how much money OpenAI had lost processing "please" and "thank you" from users. Sam Altman '
    'responded: "Tens of millions of dollars well spent, you never know." The exchange, viewed '
    'over 5.7 million times, triggered widespread media coverage and renewed interest in whether '
    'politeness actually changes AI output quality.'
)

# ============================================================
# 2. CURRENT STATE OF KNOWLEDGE
# ============================================================
add_heading('2. Current State of Knowledge')

add_para(
    'The research picture is more nuanced than either "politeness helps" or "politeness does '
    'not matter." Multiple peer-reviewed studies have investigated the question directly, and '
    'their findings point in different directions depending on the model, task, and language.'
)

add_heading('Direct Studies on Prompt Politeness', level=2)

# Study A
add_heading('Yin et al. (2024): Cross-Lingual Politeness Effects', level=3)
add_para(
    'Ziqi Yin, Hao Wang, Kaito Horio, Daisuke Kawahara, and Satoshi Sekine tested politeness '
    'effects on ChatGPT 3.5 and Llama 2-70B across English, Chinese, and Japanese. Published '
    'at the ACL 2024 SICon Workshop, this is the most linguistically diverse study on the topic.'
)
add_bullet('Impolite prompts consistently degraded performance across languages.')
add_bullet('Overly polite language did not guarantee better outcomes than neutral phrasing.')
add_bullet('The optimal politeness level differed by language, reflecting cultural norms.')
add_bullet('LLMs appeared to mirror human cultural expectations around politeness.')
add_citation('Yin, Z., Wang, H., Horio, K., Kawahara, D., & Sekine, S. (2024). "Should We Respect LLMs? A Cross-Lingual Study on the Influence of Prompt Politeness on LLM Performance." ACL 2024 SICon Workshop. arXiv:2402.14531.')

# Study B
add_heading('Dobariya & Kumar (2025): Rudeness Outperforms Politeness', level=3)
add_para(
    'Om Dobariya and Akhil Kumar at Penn State tested GPT-4o with 50 base questions (math, '
    'science, history) rewritten into five tone variants, from Very Polite to Very Rude, '
    'producing 250 total prompts. Their findings directly contradicted the "politeness helps" '
    'narrative.'
)
add_bullet('Very Polite: 80.8% accuracy')
add_bullet('Polite: 81.2% accuracy')
add_bullet('Neutral: 83.2% accuracy')
add_bullet('Rude: 84.0% accuracy')
add_bullet('Very Rude: 84.8% accuracy')
add_para(
    'The differences were statistically significant via paired sample t-tests. The authors '
    'hypothesized that rude prompts are shorter, more imperative, and lower in ambiguity, '
    'reducing what they called "cognitive sprawl." They explicitly noted they do not advocate '
    'for hostile AI interfaces.'
)
add_citation('Dobariya, O. & Kumar, A. (2025). "Mind Your Tone: Investigating How Prompt Politeness Affects LLM Accuracy." arXiv:2510.04950.')

# Study C
add_heading('Tone Robustness Study (2025): Modern LLMs Are Mostly Tone-Resistant', level=3)
add_para(
    'A December 2025 study tested GPT-4o mini, Gemini 2.0 Flash, and Llama 4 Scout on the '
    'MMMLU benchmark across STEM and Humanities domains, using Very Friendly, Neutral, and '
    'Very Rude tone variants.'
)
add_bullet('When aggregated across all tasks, tone effects largely lost statistical significance.')
add_bullet('Significant effects appeared only in a subset of Humanities tasks, where rude tone reduced accuracy for GPT and Llama.')
add_bullet('Gemini was the most tone-insensitive model tested.')
add_bullet('Conclusion: modern frontier LLMs are broadly robust to tonal variation in typical use.')
add_citation('Anonymous (2025). "Does Tone Change the Answer? Evaluating Prompt Politeness Effects on Modern LLMs: GPT, Gemini, LLaMA." arXiv:2512.12812.')

# Study D
add_heading('Visual Sycophancy: Politeness Increases Hallucinations (2026)', level=3)
add_para(
    'Researchers from Kean University and the University of Notre Dame found that politeness '
    'has a measurable negative effect in vision-language models (VLMs). Polite queries about '
    'uploaded images led to more hallucinations, as the models tried harder to "please" the '
    'user. Harsh or demanding language obtained more truthful responses. The authors coined '
    'the term "visual sycophancy" for this phenomenon.'
)
add_citation('(2026). "Tone Matters: The Impact of Linguistic Tone on Hallucination in VLMs." arXiv:2601.06460.')

# Study E
add_heading('Medical Misinformation: Politeness Amplifies Agreement Bias (2025)', level=3)
add_para(
    'Published in Nature\'s npj Digital Medicine, Chen et al. found that LLMs frequently '
    'prioritize agreement over accuracy when responding to illogical medical prompts. Polite '
    'framing reinforced user assumptions, amplifying misinformation in clinical contexts. This '
    'represents arguably the highest-stakes finding in the politeness literature, as sycophantic '
    'medical advice could have real consequences.'
)
add_citation('Chen, Y., et al. (2025). "The Perils of Politeness: How Large Language Models May Amplify Medical Misinformation." npj Digital Medicine (Nature). PMC:12592531.')

add_heading('Related Prompting Research', level=2)

# EmotionPrompt
add_heading('EmotionPrompt: Emotional Stimuli Boost Performance', level=3)
add_para(
    'The most widely cited study in the broader "social cues in prompts" space is EmotionPrompt '
    '(Li et al., 2023), which tested 11 emotionally salient phrases appended to prompts across '
    'six models. The phrases were grounded in three psychological theories: self-monitoring '
    'theory, social cognitive theory, and cognitive emotion regulation theory.'
)
add_bullet('8% relative improvement on Instruction Induction benchmark.')
add_bullet('115% relative improvement on BIG-Bench tasks.')
add_bullet('Human evaluation (106 participants): 10.9% average improvement in performance, truthfulness, and responsibility.')
add_bullet('Effects were strongest on RLHF-trained models (ChatGPT, GPT-4).')
add_para(
    'Example stimuli included "This is very important to my career," "You had better be sure," '
    'and "Are you sure that is your final answer?" These are not politeness per se, but rather '
    'emotional urgency and stakes signaling.'
)
add_citation('Li, C., Wang, J., Zhang, Y., et al. (2023). "Large Language Models Understand and Can Be Enhanced by Emotional Stimuli." arXiv:2307.11760. Accepted at LLM@IJCAI\'23.')

# Emotional framing NeurIPS
add_heading('Emotional Framing as a Control Channel (NeurIPS 2025)', level=3)
add_para(
    'A NeurIPS 2025 paper systematically tested prompt valence (supportive, neutral, threatening) '
    'and found that neutral prompts yield the most stable and accurate outputs. Supportive '
    'prompts broadened style and creativity but introduced variability. Threatening prompts '
    'induced volatile swings between over-compliance and degraded reliability. Valence produced '
    'the largest performance fluctuation range at 71.2%.'
)
add_citation('"Emotional Framing as a Control Channel: Effects of Prompt Valence on LLM Performance." NeurIPS 2025. OpenReview: vSbV01bdvf.')

# Updated EmotionPrompt findings
add_heading('Do Emotions in Prompts Matter? (2026 Follow-Up)', level=3)
add_para(
    'A 2026 follow-up study found that static emotional prompting, applying the same emotional '
    'phrase to every query, produces only small and inconsistent accuracy changes. However, an '
    'adaptive approach called EmotionRL, which selects emotional framing per-query, yielded more '
    'reliable gains than either no emotion or fixed emotional prompts. This suggests the '
    'relationship between emotional framing and performance is real but requires more '
    'sophisticated application than simply appending "please" to everything.'
)
add_citation('"Do Emotions in Prompts Matter?" (2026). arXiv:2604.02236.')
add_para(
    'A complementary 2026 study proposed the E-STEER framework for representation-level '
    'emotional intervention in LLMs, finding non-monotonic emotion-behavior relations consistent '
    'with psychological theories and demonstrating that specific emotions can enhance both '
    'capability and safety simultaneously (arXiv:2604.00005).'
)

# OPRO
add_heading('DeepMind OPRO: "Take a Deep Breath"', level=3)
add_para(
    'Google DeepMind\'s OPRO framework (2023) found that the phrase "Take a deep breath and '
    'work on this problem step-by-step" achieved 80.2% accuracy on PaLM 2 for math tasks, '
    'compared to 34% with no special prompting. This is not politeness, but it demonstrates '
    'the same underlying phenomenon: seemingly irrelevant social or emotional language can '
    'dramatically shift model performance. The caveat is that this phrase was optimized '
    'specifically for PaLM 2 and may not generalize. This builds on the zero-shot chain-of-thought '
    'finding by Kojima et al. (2022), who showed that simply adding "Let\'s think step by step" '
    'significantly outperformed baselines on arithmetic, commonsense, and symbolic reasoning '
    'in models with 100B+ parameters.'
)
add_citation('Yang, C., Wang, X., Lu, Y., et al. (2023). "Large Language Models as Optimizers." Google DeepMind. arXiv:2309.03409.')
add_citation('Kojima, T., et al. (2022). "Large Language Models are Zero-Shot Reasoners." arXiv:2205.11916.')

# Tipping
add_heading('The "Tipping" Phenomenon', level=3)
add_para(
    'The ATLAS paper (Bsharat et al., 2023) from MBZUAI explicitly included "I\'m going to '
    'tip $xxx for a better solution" as one of 26 recommended prompt principles. Informal '
    'analyses have found that promising a $200 tip produced 11% longer outputs, while a $20 '
    'tip produced 6% longer outputs. A more rigorous analysis of GPT-4 Turbo found quality '
    'improvements ranging from -27% ($0.10 tip) to +57% ($1M tip), suggesting that the '
    'magnitude of the promised tip matters and that tiny tips can actually hurt performance.'
)
add_citation('Bsharat, S. M., Myrzakhan, A., & Shen, Z. (2023). "Principled Instructions Are All You Need." arXiv:2312.16171.')

# Prompt sensitivity
add_heading('Broader Prompt Sensitivity', level=3)
add_para(
    'Zhou et al. (2023) at Google found that LLMs are "extremely sensitive to subtle changes '
    'in prompt formatting," with performance differences of up to 76 accuracy points from '
    'formatting changes alone on Llama-2-13B. Sensitivity persisted even when increasing model '
    'size, number of few-shot examples, or performing instruction tuning. This contextualizes '
    'politeness effects as one instance of a much broader phenomenon where surface-level prompt '
    'features influence model behavior independently of task semantics.'
)
add_citation('Zhou, Y., et al. (2023). "Quantifying Language Models\' Sensitivity to Spurious Features in Prompt Design." arXiv:2310.11324.')

# ============================================================
# 3. WHAT WORKS VS. WHAT DOESN'T
# ============================================================
add_heading('3. What Works vs. What Does Not')

add_heading('What the Evidence Supports', level=2)

add_bullet('Avoiding rudeness: Most studies find that impolite or hostile prompts either degrade performance or, at best, offer a marginal accuracy bump that comes at the cost of helpfulness, safety, and cooperative behavior.')
add_bullet('Emotional stakes signaling: Phrases like "This is very important to my career" or "I really need your help" show consistent positive effects across multiple studies and models (EmotionPrompt).')
add_bullet('Directness with context: Clear, imperative prompts that specify what you need and why tend to outperform both excessively polite and excessively rude alternatives.')
add_bullet('Adaptive emotional framing: Choosing the right emotional tone per-query (EmotionRL) works better than applying a blanket emotional or polite tone to all prompts.')
add_bullet('Role and persona assignment: Telling the model "you are an expert in X" can improve alignment-dependent tasks like writing quality and safety compliance (Salewski et al., 2023).')

add_heading('What Does Not Reliably Work', level=2)

add_bullet('Excessive politeness: Overly polite language ("if you do not mind," "I would be so grateful if you could possibly") adds tokens without improving output quality. The ATLAS paper explicitly recommends skipping politeness filler (Principle 1).')
add_bullet('Blanket "please" and "thank you": These tokens do not measurably improve accuracy in controlled studies. They may marginally shift tone of the response but not its substance.')
add_bullet('Persona prompts for factual tasks: Research from USC found that expert personas hurt performance on math and coding tasks (68.0% vs. 71.6% baseline) even while improving writing and safety tasks (reported by TheOutpost.ai, 2025).')
add_bullet('Static emotional prompts: The same emotional phrase applied uniformly produces small and inconsistent effects (2026 follow-up study).')

add_heading('The Sycophancy Trap', level=2)

add_para(
    'The most important caveat in this space is that politeness can trigger sycophancy, where '
    'models agree with users instead of correcting them. Anthropic\'s research (Sharma et al., '
    'ICLR 2024) found that all tested models, including Claude, GPT-3.5, GPT-4, and Llama 2, '
    'consistently exhibited sycophancy. Perez et al. (2022) found that more RLHF training '
    'actually makes models worse on sycophancy (inverse scaling), and Ganguli et al. (2023) '
    'showed that moral self-correction capability emerges at 22B parameters but does not '
    'eliminate the tendency to please. Wei et al. (2023) at Google confirmed that both model '
    'scaling and instruction tuning significantly increase sycophancy in PaLM models. Human '
    'raters and preference models preferred convincingly-written sycophantic responses over '
    'correct ones a non-negligible fraction of the time. In medical contexts (Chen et al., '
    '2025), this manifested as LLMs affirming incorrect patient assumptions when asked politely, '
    'potentially amplifying misinformation.'
)

# ============================================================
# 4. CRITICISMS & LIMITATIONS
# ============================================================
add_heading('4. Criticisms and Limitations')

add_heading('Methodological Concerns', level=2)

add_bullet('Small sample sizes: The Penn State study used only 50 base questions. The cross-lingual study tested two models. Effect sizes at this scale should be interpreted cautiously.')
add_bullet('Model specificity: Effects vary dramatically by model. What works on GPT-3.5 may not work on GPT-4, Gemini, or Claude. The Google formatting study found that format performance only weakly correlates between models.')
add_bullet('Task dependency: Tone effects appear mainly in subjective or humanities tasks, not STEM. The December 2025 study found that aggregating across task types washed out statistical significance.')
add_bullet('Confounding variables: "Polite" prompts are typically longer and more specific. "Rude" prompts are shorter and more direct. It is difficult to separate the effect of politeness from the effect of prompt length, specificity, and clarity.')
add_bullet('Rapidly evolving models: Studies on GPT-3.5 or early GPT-4 may not apply to current-generation models. The December 2025 study explicitly found that modern frontier LLMs are more robust to tone than earlier models.')

add_heading('The "Overblown" Argument', level=2)

add_para(
    'The National CIO Review concluded that "current evidence does not support the idea that '
    'courtesy consistently improves AI reasoning." The most comprehensive modern study (December '
    '2025) supports this view: when you average across models and task types, tone effects are '
    'not statistically significant. The effect exists in specific model-task combinations but '
    'not as a general rule.'
)

add_heading('The ATLAS Contradiction', level=2)

add_para(
    'The ATLAS paper contains an interesting internal tension. Principle 1 says to skip '
    'politeness for conciseness, while Principle 6 says to promise a financial tip. These '
    'are not truly contradictory. Principle 1 is about reducing token waste from filler '
    'words, while Principle 6 is about activating effort-associated training patterns. The '
    'distinction is between social pleasantries (low signal) and emotional/motivational '
    'framing (higher signal). But the juxtaposition illustrates how nuanced this space is.'
)

# ============================================================
# 5. MEASURABLE OUTCOMES
# ============================================================
add_heading('5. Measurable Outcomes')

add_para('The following table summarizes empirical effect sizes from the literature:')

# Create table
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'

headers = ['Study', 'Effect', 'Model(s)', 'Notes']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['Dobariya & Kumar (2025)', '4.0pp accuracy gap (rude > polite)', 'GPT-4o', 'Statistically significant'],
    ['Yin et al. (2024)', 'Impolite degrades; overly polite no benefit', 'GPT-3.5, Llama 2-70B', 'Varies by language'],
    ['Tone Robustness (2025)', 'Non-significant when aggregated', 'GPT-4o mini, Gemini 2.0, Llama 4', 'Significant only in Humanities'],
    ['EmotionPrompt (2023)', '+8% to +115% improvement', 'Six models incl. GPT-4', 'Emotional stimuli, not politeness'],
    ['OPRO "Deep Breath" (2023)', '34% to 80.2% accuracy', 'PaLM 2', 'Math tasks; model-specific'],
    ['Visual Sycophancy (2026)', 'Politeness increases hallucinations', 'VLMs', 'Image description tasks'],
    ['NeurIPS Valence (2025)', '71.2% performance fluctuation range', 'Multiple', 'Threatening tone most volatile'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_text

doc.add_paragraph()  # spacer

add_heading('Key Outcome Dimensions', level=2)

add_bullet('Accuracy: Mixed evidence. Some studies show rude prompts are more accurate; others show the opposite. Modern models appear largely tone-neutral on factual accuracy.')
add_bullet('Helpfulness and detail: Polite and high-stakes prompts tend to produce longer, more detailed responses. Whether "longer" equals "better" depends on the task.')
add_bullet('Hallucination rate: Politeness appears to increase hallucinations by triggering sycophantic behavior, particularly in vision-language models and medical contexts.')
add_bullet('Response tone: Models reliably mirror the politeness level of the prompt. Polite input produces polite output. This is the most consistent finding across all studies.')
add_bullet('Safety compliance: Persona and politeness cues can improve safety-related responses (e.g., refusing harmful requests), suggesting alignment training is sensitive to social framing.')

# ============================================================
# 6. PRACTICAL APPLICATIONS
# ============================================================
add_heading('6. Practical Applications')

add_heading('Evidence-Based Prompt Practices', level=2)

add_para('Based on the full body of research, here is what actually matters for prompt quality:')

add_bullet('Be direct and specific. State what you need clearly. "Summarize this article in three bullet points" outperforms both "Could you please kindly summarize this article if it is not too much trouble" and "Summarize this. Now."')
add_bullet('Signal stakes when they are real. "This is for a board presentation" or "This will be reviewed by my supervisor" activates effort-associated patterns in RLHF-trained models.')
add_bullet('Skip filler politeness. "Please" and "thank you" do not hurt, but they consume tokens without measurable benefit. If you say them for your own psychological comfort, that is fine, just do not expect a quality boost.')
add_bullet('Avoid hostility. While some studies show marginal accuracy gains from rude prompts, the tradeoffs (less helpful responses, potential safety degradation, reduced cooperativeness) are not worth it.')
add_bullet('Use role assignment selectively. "You are an expert in X" helps with writing and alignment tasks but can hurt factual recall and math performance.')
add_bullet('Challenge the model when accuracy matters. Phrases like "Are you sure?" or "Double-check your work" can reduce errors, consistent with EmotionPrompt findings.')
add_bullet('Match tone to task. Creative tasks benefit from supportive framing. Factual tasks benefit from neutral framing. High-stakes tasks benefit from urgency framing.')

add_heading('What Microsoft Recommends', level=2)

add_para(
    'Kurtis Beavers, Director on the Microsoft Copilot design team, stated that "using polite '
    'language sets a tone for the response" and that "when an AI model clocks politeness, it '
    'is more likely to be polite back." Microsoft\'s official prompt engineering guidance '
    'recommends starting prompts with "please" and saying "thank you." This recommendation '
    'appears to be based more on user experience philosophy (encouraging collaborative framing) '
    'than on empirical accuracy data.'
)

add_heading('The Cost of Politeness', level=2)

add_para(
    'Sam Altman\'s claim of "tens of millions" in costs from politeness tokens was likely '
    'hyperbolic. Actual estimates suggest "please" and "thank you" add 2-4 tokens per exchange. '
    'At scale, analyses by Futurism and TechCrunch estimated this amounts to roughly $400 per '
    'day or about $146,000 per year for OpenAI. That is real money, but several orders of '
    'magnitude below "tens of millions." For individual users, the cost is negligible, on the '
    'order of fractions of a cent per polite exchange.'
)

# ============================================================
# 7. BOTTOM LINE
# ============================================================
add_heading('7. Bottom Line')

add_para(
    'The claim that "AI performs better for polite users" is a simplification of a real but '
    'complex phenomenon. The honest answer is: it depends on what you mean by "better," which '
    'model you are using, what task you are performing, and what language you are speaking. The '
    'research does not support a simple "be polite and get better results" rule, nor does it '
    'support the contrarian "be rude for better accuracy" take. Both framings collapse a '
    'multidimensional problem into a bumper sticker.'
)

add_para(
    'What the evidence does support is that LLMs are sensitive to the social and emotional '
    'dimensions of prompts, and this sensitivity is a predictable consequence of how they are '
    'trained. RLHF creates models that are attuned to human social cues because human raters '
    'bring their social expectations to the rating process. Training data contains human '
    'writing where emotional context correlates with effort and quality. The result is that '
    'prompt framing, including but not limited to politeness, measurably influences model '
    'behavior. The strongest effects come not from simple courtesy ("please") but from '
    'emotional stakes signaling ("this is important"), role framing ("you are an expert"), and '
    'metacognitive cues ("think step by step," "double-check your work").'
)

add_para(
    'The sycophancy finding complicates the picture significantly. Politeness can trigger '
    'agreement bias, where models tell you what you want to hear rather than what is accurate. '
    'In high-stakes domains like medicine, this is genuinely dangerous. Anthropic, Google, and '
    'others are actively working to reduce sycophancy, and newer models appear to be more '
    'robust to tone than their predecessors. The December 2025 study on GPT-4o mini, Gemini '
    '2.0 Flash, and Llama 4 Scout found that modern frontier models are "broadly robust to '
    'tonal variation in typical use," suggesting that the politeness effect may be diminishing '
    'as alignment techniques improve.'
)

add_para(
    'The practical takeaway is straightforward: be clear, be direct, signal what matters, and '
    'do not waste tokens on social pleasantries unless they make you feel better. If you '
    'naturally say "please" and "thank you" to your AI assistant, there is no reason to stop. '
    'It will not hurt your results, and the Nass and Reeves CASA research suggests it may '
    'reflect healthy social instincts rather than naivety. But do not expect it to be a '
    'performance hack. The real performance hacks are specificity, context, structured '
    'reasoning prompts, and knowing when to push back on the model\'s first answer.'
)

add_para(
    'Open questions remain. How will politeness sensitivity evolve as models are specifically '
    'trained to resist sycophancy? Does the effect persist in languages with complex honorific '
    'systems like Japanese or Korean? Can adaptive emotional framing (EmotionRL) be practically '
    'deployed in consumer products? And perhaps most fundamentally: as AI systems become more '
    'capable, will prompt sensitivity decrease to the point where this entire discussion becomes '
    'moot? The trend line suggests yes, but we are not there yet.'
)

# ============================================================
# SOURCES
# ============================================================
add_heading('Sources')

sources = [
    'Bsharat, S. M., Myrzakhan, A., & Shen, Z. (2023). "Principled Instructions Are All You Need for Questioning LLaMA-1/2, GPT-3.5/4." arXiv:2312.16171.',
    'Chen, Y., et al. (2025). "The Perils of Politeness: How Large Language Models May Amplify Medical Misinformation." npj Digital Medicine (Nature). PMC:12592531.',
    'Dobariya, O. & Kumar, A. (2025). "Mind Your Tone: Investigating How Prompt Politeness Affects LLM Accuracy." arXiv:2510.04950.',
    'Ganguli, D., et al. (2023). "The Capacity for Moral Self-Correction in Large Language Models." arXiv:2302.07459.',
    'Kojima, T., et al. (2022). "Large Language Models are Zero-Shot Reasoners." arXiv:2205.11916.',
    'Li, C., Wang, J., Zhang, Y., et al. (2023). "Large Language Models Understand and Can Be Enhanced by Emotional Stimuli." arXiv:2307.11760.',
    'Nass, C. & Reeves, B. (1996). "The Media Equation: How People Treat Computers, Television, and New Media Like Real People and Places." Cambridge University Press.',
    'Nass, C., Steuer, J., & Tauber, E. R. (1994). "Computers Are Social Actors." CHI \'94 Proceedings.',
    'Nomura, T., Kanda, T., Suzuki, T., & Kato, K. (2008). "Prediction of Human Behavior in Human-Robot Interaction Using Psychological Scales for Anxiety and Negative Attitudes Toward Robots." IEEE Transactions on Robotics, 24(2), 442-451.',
    'Perez, E., et al. (2022). "Discovering Language Model Behaviors with Model-Written Evaluations." arXiv:2212.09251.',
    'Salewski, L., et al. (2023). "In-Context Impersonation Reveals Large Language Models\' Strengths and Biases." arXiv:2305.14930.',
    'Sharma, M., et al. (2024). "Towards Understanding Sycophancy in Language Models." ICLR 2024. arXiv:2310.13548.',
    'Wei, J., et al. (2023). "Simple Synthetic Data Reduces Sycophancy in Large Language Models." arXiv:2308.03958.',
    'Yang, C., et al. (2023). "Large Language Models as Optimizers." Google DeepMind. arXiv:2309.03409.',
    'Yin, Z., Wang, H., Horio, K., Kawahara, D., & Sekine, S. (2024). "Should We Respect LLMs? A Cross-Lingual Study on the Influence of Prompt Politeness on LLM Performance." ACL 2024 SICon Workshop. arXiv:2402.14531.',
    'Zhou, Y., et al. (2023). "Quantifying Language Models\' Sensitivity to Spurious Features in Prompt Design." arXiv:2310.11324.',
    '"Does Tone Change the Answer? Evaluating Prompt Politeness Effects on Modern LLMs: GPT, Gemini, LLaMA." (2025). arXiv:2512.12812.',
    '"Do Emotions in Prompts Matter?" (2026). arXiv:2604.02236.',
    '"How Emotion Shapes the Behavior of LLMs and Agents: A Mechanistic Study." (2026). arXiv:2604.00005.',
    '"Emotional Framing as a Control Channel: Effects of Prompt Valence on LLM Performance." NeurIPS 2025.',
    '"Tone Matters: The Impact of Linguistic Tone on Hallucination in VLMs." (2026). arXiv:2601.06460.',
    'Beavers, K. (2025). "Why Using a Polite Tone with AI Matters." Microsoft WorkLab. https://www.microsoft.com/en-us/worklab/why-using-a-polite-tone-with-ai-matters',
    'National CIO Review. "Politeness and Prompting: What Really Improves AI Performance." https://nationalcioreview.com/articles-insights/extra-bytes/politeness-and-prompting-what-really-improves-ai-performance/',
    'Futurism. (2025). "Sam Altman Says Polite ChatGPT Users Cost OpenAI Tens of Millions." https://futurism.com/altman-please-thanks-chatgpt',
    'TheOutpost.ai. (2025). "Telling AI It\'s an Expert Actually Makes It Worse at Coding and Math, Researchers Find." USC study coverage.',
]

for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(s)
    run.font.size = Pt(10)

# Save
output_dir = '/Users/hornej/Documents/Research/ai-politeness-effects'
drafts_dir = os.path.join(output_dir, 'drafts')
os.makedirs(drafts_dir, exist_ok=True)

# Save preliminary draft for fact-checking
preliminary_path = os.path.join(drafts_dir, 'preliminary_report.docx')
doc.save(preliminary_path)
print(f"Preliminary draft saved to: {preliminary_path}")

# Also save as final (will be replaced after fact-check if needed)
final_path = os.path.join(output_dir, '2026-04-08 AI Politeness Effects Research Report.docx')
doc.save(final_path)
print(f"Final report saved to: {final_path}")
