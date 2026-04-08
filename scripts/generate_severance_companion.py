#!/usr/bin/env /opt/anaconda3/bin/python3
"""Generate Severance Companion Guide .docx report."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

# Quote style
quote_style = doc.styles.add_style('BlockQuote', 1)  # paragraph style
quote_style.font.name = 'Calibri'
quote_style.font.size = Pt(10.5)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
quote_style.paragraph_format.left_indent = Inches(0.5)
quote_style.paragraph_format.right_indent = Inches(0.3)
quote_style.paragraph_format.space_before = Pt(6)
quote_style.paragraph_format.space_after = Pt(6)

today = datetime.date.today().strftime("%Y-%m-%d")

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Severance")
run.font.size = Pt(36)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Companion Guide")
run.font.size = Pt(18)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"TV Series (Apple TV+) | {today}\nFull Spoilers: Seasons 1 & 2")
run.font.size = Pt(11)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_heading(text, level=2):
    doc.add_heading(text, level=level)

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_quote(text):
    p = doc.add_paragraph(text, style='BlockQuote')
    return p

# ============================================================
# 1. QUICK TAKE
# ============================================================
add_heading("Quick Take", 1)

add_para(
    "Severance is the best science fiction show of the decade so far, and it knows exactly what it is. "
    "Created by Dan Erickson and directed primarily by Ben Stiller, the Apple TV+ series takes a single "
    "speculative premise, a surgical procedure that splits your consciousness between work and personal life, "
    "and builds from it a layered, devastating exploration of identity, corporate control, and what it means "
    "to be a whole person. Season 1 (97% on Rotten Tomatoes) constructed a hypnotic puzzle box. Season 2 "
    "(94% critics, 87 Metacritic) deepened the emotional stakes while expanding the mythology. The show has "
    "accumulated 41 Emmy nominations, 10 wins, a Peabody Award, and over 6.4 billion streaming minutes across "
    "Season 2 alone, making it Apple TV+'s most-watched series ever, surpassing Ted Lasso. A third season "
    "begins filming in April 2026, with four total seasons planned. Whether you have just finished the Season 2 "
    "finale or are looking for an excuse to rewatch, this guide is designed to enrich every corner of the experience."
)

# ============================================================
# 2. THE SOURCE
# ============================================================
add_heading("The Source", 1)

add_para(
    "Severance is entirely original, born not from a novel or film but from genuine workplace misery. "
    "After graduating from NYU Tisch with an MFA in Dramatic Writing, Dan Erickson took a temp job at a door "
    "factory in Los Angeles, cataloguing hinges and deadbolts in a windowless office. The monotony was crushing. "
    '"God, what if I could jump ahead and suddenly it would be five?" he thought. "I would have done the day\'s '
    "work, but I wouldn't have to experience it.\" He recognized this as a deeply unsettling wish, and that "
    "tension became the seed of the show."
)

add_para(
    "Erickson wrote the first version of the pilot while still at the factory. In 2015, he submitted the "
    "script to Ben Stiller's Red Hour Productions, where it was passed along to Stiller by development "
    "executive Jackie Cohn and head of creative Nicholas Weinstock. In 2016, the Severance pilot became the "
    "first television script ever included on the annual Blood List, a survey of the best unproduced dark genre "
    "screenplays. Apple TV+ gave the series order in November 2019, and it premiered on February 18, 2022, "
    "roughly a decade after the idea first struck Erickson in a room full of door hardware."
)

add_para(
    "Erickson has cited an unusually eclectic range of influences: the Backrooms creepypasta (endless liminal "
    "office spaces), the 2013 video game The Stanley Parable (office drudgery and free will), films including "
    "Office Space, The Truman Show, Being John Malkovich, and Eternal Sunshine of the Spotless Mind, the Black "
    "Mirror episode \"White Christmas,\" Jean-Paul Sartre's No Exit, George Orwell's Nineteen Eighty-Four, and "
    "even the comic strip Dilbert. The name \"Kier Eagan\" is a nod to Eagan's Drive-In, a burger joint in "
    "Erickson's hometown of Olympia, Washington."
)

# ============================================================
# 3. THEMES AND WHAT MAKES IT WORK
# ============================================================
add_heading("Themes and What Makes It Work", 1)

add_para(
    "At its core, Severance asks a question that sounds like a philosophy seminar but feels like a gut punch: "
    "How much of who you are is constituted by your experience versus something more innate? The severance "
    "procedure literalizes the fantasy of perfect work-life separation and reveals it as a horror. Your innie "
    "exists only at work, born into fluorescent light with no childhood, no family, no context. Your outie "
    "walks out the door with eight missing hours and a paycheck. Two consciousnesses share one body with "
    "conflicting hopes, agendas, and desires."
)

add_para(
    "The show operates on several thematic layers simultaneously, which is why it resonates with casual viewers "
    "who see workplace satire and philosophy-minded audiences who see a meditation on consciousness:"
)

add_para(
    "Work-life balance: The literal premise is the ultimate expression of separating work from life, taken to "
    "its logical, horrifying conclusion. Mark underwent severance to cope with his wife's supposed death, "
    "literalizing emotional avoidance. The show argues that you cannot amputate part of your experience without "
    "mutilating your identity."
)

add_para(
    "Corporate control and cultism: Lumon Industries radiates cult energy through worship of founder Kier "
    "Eagan, handbook recitations, the Perpetuity Wing, \"waffle parties,\" and perks that would be absurd if "
    "they weren't so precisely unsettling. The show makes the case that corporate culture, at its extremes, "
    "functions identically to religious devotion."
)

add_para(
    "Identity and consciousness: The innies' growing awareness mirrors Plato's Allegory of the Cave. They "
    "begin accepting their reality, then start questioning it, then revolt. Season 2 moves from workplace "
    "satire into deeper philosophical territory about memory, selfhood, and whether a manufactured person can "
    "claim autonomy."
)

add_para(
    "Free will and resistance: Season 1 asks \"what is happening?\" Season 2 asks \"why is it happening and "
    "can it be stopped?\" The innies' rebellion follows the arc of every great dystopian narrative, from "
    "ignorance to awareness to action."
)

add_quote(
    "\"It's scary when you really look at how much of our self-worth and identity are wrapped up in our jobs. "
    "This feeling of slowly losing who you are.\" -- Dan Erickson"
)

# ============================================================
# 4. CRITICAL RECEPTION
# ============================================================
add_heading("Critical Reception", 1)

add_heading("The Numbers", 2)

# Scores table
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'
headers = table.rows[0].cells
headers[0].text = ""
headers[1].text = "Season 1"
headers[2].text = "Season 2"
s1 = table.rows[1].cells
s1[0].text = "Rotten Tomatoes"
s1[1].text = "97% (117 reviews)"
s1[2].text = "94% (191+ reviews)"
s2 = table.rows[2].cells
s2[0].text = "Metacritic"
s2[1].text = "83/100 (Universal Acclaim)"
s2[2].text = "87/100 (Universal Acclaim)"

doc.add_paragraph()  # spacer

add_para(
    "Season 2's Metacritic score of 87 actually improved on the first season by four points, a rare achievement "
    "for a sophomore season of an already acclaimed show. The initial reviews were even stronger: Season 2 "
    "debuted at a perfect 100% on Rotten Tomatoes before settling as more critics weighed in across the full run."
)

add_heading("What the Critics Said", 2)

add_quote(
    "\"It is mesmerising, gorgeous, heartbreaking and triumphant.\" -- The Guardian"
)
add_quote(
    "\"Not only has the series retained its beguiling, head-spinning power -- it's upped the mindf---ery "
    "to astronomical levels.\" -- The Daily Beast"
)
add_quote(
    "\"Every bit as trippy as Season 1 and even more emotionally resonant.\" -- TVLine (Grade: A)"
)
add_quote(
    "\"It's a tall order to follow up on one of the best first seasons ever, but Severance season 2 exceeds "
    "expectations in every way.\" -- Total Film (5/5)"
)

add_para(
    "Not everyone was swept up. The most common criticism of Season 2 centered on pacing, particularly in the "
    "middle episodes, and the show's tendency to add mysteries without resolving enough of them. TechRadar ran "
    "a piece titled \"I hate to say it Apple TV+, but Severance season 2 has a pacing problem,\" noting that "
    "\"the last few episodes felt like fillers rather than explosive plot advancers.\" The Hollywood Reporter "
    "called Season 2 \"frequently darker, less frequently amusing, and not necessarily more satisfying.\" These "
    "criticisms are fair, though they speak more to the challenge of serialized mystery storytelling than to any "
    "fundamental decline in quality."
)

add_heading("Awards", 2)

add_para(
    "Across the 74th through 77th Primetime Emmy Awards, Severance accumulated 41 nominations and 10 wins. "
    "Season 2 alone earned 27 nominations at the 2025 ceremony, just five short of the all-time single-year "
    "record set by Game of Thrones in 2019. Key wins include:"
)

awards_list = [
    "Britt Lower: Outstanding Lead Actress in a Drama Series (her first nomination and win)",
    "Tramell Tillman: Outstanding Supporting Actor in a Drama Series",
    "Creative Arts wins for Cinematography, Music Composition, Main Title Design, Sound Mixing, and Production Design",
    "Peabody Award for Entertainment (2023)",
    "Adam Scott: Critics Choice Super Award for Best Actor in a Science Fiction/Fantasy Series (2023)",
    "Multiple Golden Globe nominations across both seasons, including Best Drama Series, Best Actor, and Best Actress"
]
for item in awards_list:
    p = doc.add_paragraph(item, style='List Bullet')

add_heading("Viewership", 2)

add_para(
    "Season 2 transformed Severance from cult hit to mainstream phenomenon. It became Apple TV+'s most-watched "
    "series in history, surpassing Ted Lasso's three-season record. Nielsen reported 6.4 billion streaming "
    "minutes across the season, with a 218% increase over Season 1's initial run. The Season 2 finale was the "
    "first Severance episode to crack Nielsen's Overall Top 10, a list typically dominated by Netflix. Apple "
    "subsequently acquired full production rights for an estimated $70 million and greenlit a four-season arc. "
    "Season 3 begins filming in April 2026 with a Summer 2027 release expected."
)

# ============================================================
# 5. WHAT YOU MIGHT NOT KNOW
# ============================================================
add_heading("What You Might Not Know", 1)

add_heading("The Casting Wars", 2)

add_para(
    "Ben Stiller and Dan Erickson independently had Adam Scott in mind for Mark Scout. Stiller's conviction "
    "came from an unlikely place: Step Brothers. \"That's when I knew he was an incredible actor,\" Stiller "
    "said. But Apple didn't initially believe in the casting and demanded an audition, creating a year-long "
    "standoff. Stiller held firm, Scott eventually auditioned, and the rest is television history. Stiller "
    "later admitted he was \"terrified\" Scott would walk over the ordeal."
)

add_para(
    "John Turturro personally recruited Christopher Walken for the role of Burt. \"I've known Chris a long "
    "time and I don't have to really act like we're friends,\" Turturro said. Their real-life friendship, "
    "built over decades and prior collaborations, became the foundation for one of the show's most tender "
    "on-screen relationships. Patricia Arquette skipped the audition process entirely, cast directly from "
    "her work with Stiller on Escape at Dannemora."
)

add_heading("The Music/Dance Experience Was Improvised", 2)

add_para(
    "One of Season 1's most iconic and disturbing scenes nearly didn't happen the way we saw it. Tramell "
    "Tillman almost missed his audition entirely, making a split-second decision to run 15 minutes through "
    "the South Bronx rather than wait for a delayed train. When it came time to film the Music/Dance "
    "Experience, the script said little more than those three words. \"I didn't even know he was going to "
    "dance like that,\" Stiller said. The strobe lighting effect was discovered by accident when the lighting "
    "team found the dimmer board could create the effect, and cinematographer Jessica Lee Gagne loved it enough "
    "to incorporate it. The result is a scene that feels meticulously choreographed but was born from "
    "improvisation and happy accidents."
)

add_heading("The Elevator Uses a Hitchcock Trick", 2)

add_para(
    "The transition between innie and outie in the elevator uses a \"zolly\" (dolly zoom), the camera "
    "technique first made famous by Hitchcock in Vertigo. Cinematographer Gagne used a 19-90mm Panavision "
    "lens and a Kuper motion control computer to precisely synchronize the zoom with dolly movement. Two "
    "distinct visual languages govern the show: the severed floor uses static, dolly-tracked cameras with no "
    "Steadicam (creating an \"impartial observer\" feeling), while the outside world uses longer lenses and "
    "more naturalistic movement. The camera itself is a storytelling tool."
)

add_heading("The Hallway Shot Took Five Months", 2)

add_para(
    "The Season 2 premiere's two-minute continuous hallway sequence (Mark running through Lumon corridors) "
    "was actually filmed in ten separate parts over five months as a \"side project\" between other scenes. "
    "Techniques included a Bolt X Cinebot robotic arm, green screen with Scott on a treadmill wearing a stunt "
    "safety wire, and ILM-generated CG hallway extensions seamlessly blended with practical sets. One segment "
    "required pulling out an entire wall to accommodate the robotic arm."
)

add_heading("No Control, No Escape", 2)

add_para(
    "The keyboards used by Macrodata Refinement employees are deliberately missing the Control and Escape "
    "keys, a visual metaphor so on-the-nose it circles back around to brilliant. Other Easter eggs reward "
    "close viewing: goat imagery is scattered throughout Season 1 (shop windows, clothing stores, a statue in "
    "Ricken's home). The paintings on the severed floor track character arcs, with \"The Macrodata Refinement "
    "Calamity\" depicting an interdepartmental conspiracy. In Season 2, Dylan, Irving, and Helly's locker "
    "numbers are Lost numbers (23, 4, and 16). And every computer on the MDR floor is fully functional with "
    "working keypads, trackballs, and programs."
)

add_heading("The Green Carpet Is a Playground", 2)

add_para(
    "Production designer Jeremy Hindle chose the distinctive green floor of the Lumon offices to represent "
    "\"the grass of a children's playground,\" reinforcing the infantilization of severed workers. He "
    "custom-designed every piece of furniture in Gemma's office on the Testing Floor, always built visible "
    "ceilings on the sets (unusual and a fire hazard concern), and constructed 140-foot hallway runs around "
    "the perimeter of two soundstages at York Studios in the Bronx. The practical sets were large enough that "
    "actors reportedly got lost in the actual maze."
)

add_heading("Bell Labs Lives On as Lumon", 2)

add_para(
    "Lumon's exterior is the former Bell Laboratories complex in Holmdel, New Jersey, designed in 1958 by "
    "Finnish-American architect Eero Saarinen. The 460-acre complex features a 700-foot-long building with "
    "a mirrored \"curtain wall\" of thousands of individual glass panes. A cutting-edge research facility in "
    "its heyday, its corporate grandeur now serves as the perfect shell for Lumon's manufactured prestige."
)

add_heading("The $200 Million Season", 2)

add_para(
    "Season 2's budget ballooned to approximately $20 million per episode ($200 million for 10 episodes), "
    "making it one of the most expensive television seasons ever produced. Contributing factors included "
    "reported creative tensions between co-showrunners, extensive script rewrites, the 2023 WGA/SAG-AFTRA "
    "strikes pausing production, and the addition of Beau Willimon (House of Cards) to the writing team. "
    "Despite the turbulent production, the finished product shows where every dollar went."
)

add_heading("The Score Was Built from a Broken Piano", 2)

add_para(
    "Composer Theodore Shapiro and Stiller developed the main theme by isolating four chords from a draft "
    "piece, which became the foundation for a pre-shot music library. Shapiro drew inspiration from David "
    "Shire's score for The Conversation (1974). For Season 2, sound designer Chris Lane developed percussion "
    "using only the piano as a source, including beating a piano and creating loops. Contact microphones were "
    "used to build an \"intimate and internal sound\" from found objects. The approach mirrors the show's own "
    "method: taking something familiar and making it deeply unsettling."
)

# ============================================================
# 6. CULTURAL CONTEXT
# ============================================================
add_heading("Cultural Context", 1)

add_para(
    "Severance arrived at exactly the right cultural moment. Season 1 premiered in February 2022, as "
    "corporations pushed return-to-office mandates and the concepts of \"quiet quitting\" and work-life "
    "balance dominated cultural conversation. The pandemic had already blurred the line between home and "
    "office for millions of workers, making the show's premise feel less like science fiction and more like "
    "commentary. By the time Season 2 premiered in January 2025, the \"innie\" and \"outie\" terminology "
    "had entered the cultural vocabulary, and memes flooded social media."
)

add_para(
    "Slate called it \"the most vicious satire of corporate life since Office Space, and possibly ever.\" "
    "HuffPost argued the show \"nails how absolutely inhuman our work culture is,\" pointing out that workers "
    "already feel like two separate people, a professional self who suppresses emotions and a personal self who "
    "only emerges outside work. The severance procedure just makes the metaphor literal."
)

add_para(
    "The show has generated peer-reviewed academic scholarship. A 2025 article in MDPI's Administrative "
    "Sciences journal analyzed Severance as \"speculative organizational critique,\" examining themes of "
    "control, consent, and identity at work through formal organizational theory. That a TV show about "
    "sorting scary numbers on a computer screen has produced academic papers tells you something about how "
    "deeply it has burrowed into the cultural conversation about work."
)

add_para(
    "Within the genre, Severance occupies a distinctive position. It is grounded sci-fi with no spectacle: "
    "no distant planets, no futuristic cityscapes, no laser battles. One speculative premise, ruthlessly "
    "explored. Critics have called it a \"watershed show\" that marries the anthology-style psychological "
    "horror of Black Mirror with persistent characters and long-form storytelling. It proves that "
    "psychologically driven science fiction can achieve massive commercial success without a single explosion."
)

# ============================================================
# 7. IF YOU LIKED THIS
# ============================================================
add_heading("If You Liked This", 1)

add_heading("Shows", 2)

recs_shows = [
    ("Counterpart (Starz, 2017-2019)",
     "The closest match. J.K. Simmons plays dual roles as a bureaucrat who discovers a gateway to a parallel "
     "universe where his counterpart is a ruthless operative. Same dual-identity premise, institutional secrecy, "
     "and bureaucratic drudgery masking existential stakes. Two seasons, criminally underwatched."),
    ("Homecoming (Amazon, 2018-2020)",
     "Sam Esmail-directed thriller where government employees at a facility for returning soldiers suffer memory "
     "gaps. Shares Severance's core mechanic of institutional memory manipulation and Hitchcockian visual precision. "
     "Julia Roberts stars in Season 1, Janelle Monae in Season 2."),
    ("Dark Matter (Apple TV+, 2024)",
     "Based on Blake Crouch's novel. A man is kidnapped by an alternate version of himself and forced into an "
     "unfamiliar timeline. Where Severance splits one person into innie/outie, Dark Matter shows what happens when "
     "entirely different life paths collide. Same platform, strong crossover audience."),
    ("Devs (Hulu/FX, 2020)",
     "Alex Garland's miniseries set inside a secretive Silicon Valley quantum computing division. Shares Severance's "
     "claustrophobic corporate aesthetics, slow-burn pacing, and questions about determinism and corporate god "
     "complexes. If you liked the vibe, you'll love this."),
    ("Pluribus (Apple TV+, 2025)",
     "From Vince Gilligan, starring Rhea Seehorn battling a cheerful hive mind. The newest entry on this list "
     "and Apple TV+'s next prestige sci-fi play. The \"cheerful collective\" dynamic parallels Lumon's cult-like "
     "corporate warmth."),
]

for title, desc in recs_shows:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"  {desc}")
    p.paragraph_format.space_after = Pt(8)

add_heading("Movies", 2)

recs_movies = [
    ("Eternal Sunshine of the Spotless Mind (2004)",
     "The most direct thematic comparison. A couple undergoes a procedure to erase memories of each other. Both works "
     "ask: what happens when you sever memory from identity? The clinical, casually dehumanizing procedure echoes "
     "Lumon's severance chip."),
    ("Sorry to Bother You (2018)",
     "A telemarketer discovers that using a \"white voice\" launches him up the corporate ladder, leading to surreal "
     "and horrifying revelations. Takes corporate dystopia and cranks it to absurdist extremes. The tonal shift from "
     "comedy to body horror mirrors Severance's own genre-blending."),
    ("Brazil (1985)",
     "Terry Gilliam's surreal dystopia about a bureaucrat in a totalitarian state. The aesthetic and thematic "
     "grandfather of Severance: endless cubicles, meaningless tasks, institutional absurdity as control mechanism."),
]

for title, desc in recs_movies:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"  {desc}")
    p.paragraph_format.space_after = Pt(8)

add_heading("Books", 2)

recs_books = [
    ("The Factory by Hiroko Oyamada (2013)",
     "The most tonally precise book match. Three workers are hired by a massive, unnamed factory so large it has "
     "its own ecosystem. The disorientation of purposeless work, the creeping surrealism, the sense that the "
     "workplace is a living organism: this is literary Severance in 120 pages."),
    ("Never Let Me Go by Kazuo Ishiguro (2005)",
     "Students at an isolated boarding school gradually discover the tragic purpose of their existence. Shares "
     "Severance's emotional register: the horror comes not from spectacle but from characters' acceptance of "
     "dehumanizing systems. The innies' limited understanding of their own existence mirrors the students' slow "
     "awakening."),
    ("The Circle by Dave Eggers (2013)",
     "A young woman joins a powerful tech company that promotes radical transparency with dystopian consequences. "
     "Where Severance explores corporate control through memory erasure, The Circle explores it through total "
     "surveillance. The cultish corporate mantras will feel very familiar."),
]

for title, desc in recs_books:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"  {desc}")
    p.paragraph_format.space_after = Pt(8)

add_heading("Podcasts", 2)

recs_pods = [
    ("The Severance Podcast with Ben Stiller & Adam Scott (Official)",
     "The definitive companion. Episode-by-episode rewatch with behind-the-scenes analysis from the director "
     "and star. Features cast, crew, and celebrity superfans as guests."),
    ("Severed: The Ultimate Severance Podcast",
     "Scene-by-scene, moment-by-moment rewatch for viewers who want to catch every detail and Easter egg."),
    ("Waffle Poddy: A Severance Podcast",
     "From the Bald Move podcast network. Weekly recaps, analysis, and listener feedback with community-driven "
     "discussion."),
]

for title, desc in recs_pods:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"  {desc}")
    p.paragraph_format.space_after = Pt(8)

add_heading("Fan Community", 2)

add_para(
    "The r/SeveranceAppleTVPlus subreddit has grown to over 700,000 members and is one of the largest TV fan "
    "communities on Reddit. Creator Dan Erickson admitted he was warned to stay away but lasted \"about six "
    "minutes\" before visiting daily, calling it \"kind of addicting.\" The community's track record of correct "
    "predictions is remarkable: fans correctly identified Helly R as Helena Eagan from the Season 2 premiere and "
    "predicted that a severed human could be severed multiple times. The subreddit has become a genuine part of "
    "the show's cultural ecosystem."
)

# ============================================================
# 8. SOURCES
# ============================================================
add_heading("Sources", 1)

sources = [
    "ASC (2025). \"Ben Stiller on Severance.\" theasc.com",
    "Awards Focus (2025). \"Inside the Casting of 'Severance': Rachel Tenner Breaks Down Season 2.\" awardsfocus.com",
    "Bell Works (2025). \"Where Was Severance Filmed?\" bell.works",
    "Betches (2025). \"'Severance' Season 2 Easter Eggs, Unpacked.\" betches.com",
    "CBR (2025). \"Severance Star Says Season 1's Most Disturbing Scene Was Unscripted.\" cbr.com",
    "CinemaBlend (2025). \"Adam Scott and Ben Stiller on Severance's Season 2 Opening Scene.\" cinemablend.com",
    "Collider (2025). \"13 Movies To Watch if You Like 'Severance'.\" collider.com",
    "Collider (2025). \"The Creator of 'Severance' Literally Put His Own Blood Into the Work.\" collider.com",
    "ComicBook.com (2025). \"Severance Easter Eggs: Lumon Computers No Control Escape Keys.\" comicbook.com",
    "Creative Screenwriting. \"Dan Erickson Talks Work/Life Balance In 'Severance'.\" creativescreenwriting.com",
    "Dazed Digital (2025). \"Severance creator Dan Erickson unpacks the show's secrets.\" dazeddigital.com",
    "Deadline (2025). \"'Severance' Surpasses 'Ted Lasso' As Apple TV+'s Most Watched Series.\" deadline.com",
    "Deadline (2025). \"How 'Severance' Composer Theodore Shapiro Evolved the Score in Season 2.\" deadline.com",
    "Deadline (2026). \"Apple Acquires 'Severance', Eyes Season 3 Start & Season 4.\" deadline.com",
    "Dezeen (2025). \"'Offices now are so uncreative' says Severance designer Jeremy Hindle.\" dezeen.com",
    "Esquire Singapore (2025). \"Severance Season 2: 6 Best Fan Theories Before the Finale.\" esquiresg.com",
    "Fast Company. \"3 hard truths about work-life balance that 'Severance' eerily illustrates.\" fastcompany.com",
    "Film Independent (2025). \"Creating the Singular, Disquieting Aesthetic of 'Severance'.\" filmindependent.org",
    "Final Draft (2025). \"How 'Severance' Changed a Writer's Career.\" finaldraft.com",
    "GamesRadar+ (2025). \"Ben Stiller says Step Brothers led him to cast Adam Scott.\" gamesradar.com",
    "Gold Derby (2025). \"How 'Severance' assembled its 'out-of-the-box' Emmy-worthy cast.\" goldderby.com",
    "Hollywood Reporter (2025). \"Emmys Analysis: How 'The Pitt' Beat 'Severance'.\" hollywoodreporter.com",
    "HuffPost. \"Apple TV+'s 'Severance' Nails How Absolutely Inhuman Our Work Culture Is.\" huffpost.com",
    "IndieWire (2024). \"Ben Stiller Had to Battle Apple to Cast Adam Scott in 'Severance'.\" indiewire.com",
    "MDPI Administrative Sciences (2025). \"The TV Series Severance as Speculative Organizational Critique.\" mdpi.com",
    "MRMC (2025). \"The Making Of Severance S2 EP1.\" mrmoco.com",
    "NBC News (2025). \"'Severance' memes flood social media.\" nbcnews.com",
    "NYU Tisch (2022). \"Dan Erickson '10 and the 10 Year Journey to Creating 'Severance'.\" tisch.nyu.edu",
    "Peabody Awards. \"Severance.\" peabodyawards.com",
    "PetaPixel (2025). \"The Clever Camera Effect Used on 'Severance'.\" petapixel.com",
    "Rolling Stone (2022). \"Adam Scott Knows You Have Feelings About the 'Severance' Finale.\" rollingstone.com",
    "Rotten Tomatoes. \"Severance: Season 1.\" rottentomatoes.com",
    "Rotten Tomatoes. \"Severance: Season 2.\" rottentomatoes.com",
    "Rotten Tomatoes Editorial (2025). \"Severance Season 2 First Reviews.\" rottentomatoes.com",
    "ScreenRant (2025). \"Severance Season 2's Budget Ballooned During Filming Chaos.\" screenrant.com",
    "ScreenRant (2025). \"Severance Star Pushed For Christopher Walken To Play Burt.\" screenrant.com",
    "Slate (2022). \"Severance finale: a satire of office life.\" slate.com",
    "SlashFilm (2025). \"One Severance Actor Convinced Creators To Cast Christopher Walken.\" slashfilm.com",
    "SlashFilm. \"Severance Fans Need To Watch Counterpart.\" slashfilm.com",
    "Sound of Life (2025). \"Theodore Shapiro's Beautifully Erratic Score for 'Severance'.\" soundoflife.com",
    "TechRadar (2025). \"Severance season 2 has a pacing problem.\" techradar.com",
    "TellTale TV (2025). \"How 'Severance' is Redefining the Sci-Fi Genre.\" telltaletv.com",
    "The Direct. \"Severance: What Lumon Actually Does, Explained.\" thedirect.com",
    "The Ringer (2022). \"'Severance' Offers a Surreal, Allegorical Twist on Work-Life Balance.\" theringer.com",
    "The Wrap (2025). \"'Severance' Season 2 Sees Over 6.4 Billion Streaming Minutes.\" thewrap.com",
    "Variety (2022). \"'Severance' Production Designer Jeremy Hindle on Creating Lumon.\" variety.com",
    "Variety (2024). \"'Severance' Podcast Launched by Ben Stiller and Adam Scott.\" variety.com",
    "Variety (2025). \"'Severance' Renewed for Season 3.\" variety.com",
    "Variety (2025). \"Britt Lower Wins Emmy for Lead Actress in a Drama.\" variety.com",
    "Variety (2025). \"Emmy Nominations 2025.\" variety.com",
    "Wikipedia. \"Severance (TV series).\" en.wikipedia.org",
]

for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9.5)

# -- Save --
output_path = f"/Users/hornej/Documents/Research/severance/{today} Severance Companion Guide.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
