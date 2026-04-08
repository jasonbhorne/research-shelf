#!/usr/bin/env python3
"""Generate Deep Research Report: K-12 Education Administrators Using Claude Code"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.expanduser("~/Documents/Research/edu-admins-claude-code")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{date.today()} Education Administrators Using Claude Code Deep Research Report.docx")


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_colored_run(paragraph, text, color_rgb, bold=False, size=None):
    """Add a colored text run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.color.rgb = color_rgb
    run.bold = bold
    if size:
        run.font.size = size
    return run


def set_table_style(table):
    """Apply consistent formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(9)
                p.style.font.name = "Calibri"


def make_header_row(table, headers, color="1F4E79"):
    """Format the first row of a table as a header."""
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        set_cell_shading(cell, color)


# Confidence indicators
GREEN = RGBColor(0x2E, 0x7D, 0x32)
YELLOW = RGBColor(0xF9, 0xA8, 0x25)
RED = RGBColor(0xC6, 0x28, 0x28)

# ── Sources ──────────────────────────────────────────────────────────────
sources = [
    {
        "id": 1, "author": "Peninsula School District", "year": 2025,
        "title": "psd-claude-coding-system (GitHub Repository)",
        "url": "https://github.com/psd401/psd-claude-coding-system",
        "category": "Industry", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.77, "reliability": "Medium", "cited_by": ["Industry"]
    },
    {
        "id": 2, "author": "AIEmpoweredEDU2025", "year": 2025,
        "title": "Enterprise-Level AI in Education: PSD's psd401ai Framework",
        "url": "https://aiempowerededu2025.sched.com/event/1zZfJ/",
        "category": "Industry", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.70, "reliability": "Medium", "cited_by": ["Industry"]
    },
    {
        "id": 3, "author": "U.S. Department of Education", "year": 2025,
        "title": "Artificial Intelligence (AI) Guidance",
        "url": "https://www.ed.gov/about/ed-overview/artificial-intelligence-ai-guidance",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
    {
        "id": 4, "author": "Tennessee SCORE", "year": 2025,
        "title": "Survey Findings Capture AI Use in Tennessee School Districts",
        "url": "https://tnscore.org/perspectives-and-press/perspectives/survey-findings-capture-ai-use-in-tennessee-school-districts",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov", "Industry"]
    },
    {
        "id": 5, "author": "Tennessee General Assembly", "year": 2024,
        "title": "Public Chapter 550: AI Policies for LEAs",
        "url": "https://www.govtech.com/education/k-12/tennessee-bill-on-ai-in-schools-a-hint-of-whats-to-come",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
    {
        "id": 6, "author": "Tennessee AI Advisory Council", "year": 2025,
        "title": "Action Plan",
        "url": "https://www.tn.gov/content/dam/tn/finance/aicouncil/documents/TN%20AI%20Advisory%20Council%20Action%20Plan%20-%20November%202025.pdf",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
    {
        "id": 7, "author": "U.S. Department of Education", "year": 2023,
        "title": "Artificial Intelligence and the Future of Teaching and Learning",
        "url": "https://www.ed.gov/sites/ed/files/documents/ai-report/ai-report.pdf",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
    {
        "id": 8, "author": "Anthropic", "year": 2025,
        "title": "Anthropic Education Report: How Educators Use Claude",
        "url": "https://www.anthropic.com/news/anthropic-education-report-how-educators-use-claude",
        "category": "Industry", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.70, "reliability": "Medium", "cited_by": ["Gov", "Industry", "Academic"]
    },
    {
        "id": 9, "author": "RAND Corporation", "year": 2025,
        "title": "Uneven Adoption of AI Tools Among U.S. Teachers and Principals",
        "url": "https://www.rand.org/pubs/research_reports/RRA134-25.html",
        "category": "Academic", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Academic", "Industry"]
    },
    {
        "id": 10, "author": "Riegel, C., Ford, A., Brinkman, J., et al.", "year": 2025,
        "title": "Exploring AI in Education: A Multi-State Study on K-12 Teachers' and Administrators' Knowledge, Use, and Perceptions",
        "url": "https://journals.librarypublishing.arizona.edu/itlt/article/id/7512/",
        "category": "Academic", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Academic"]
    },
    {
        "id": 11, "author": "Berkovich, I.", "year": 2025,
        "title": "The Rise of AI-Assisted Instructional Leadership: GenAI Integration in School Leadership",
        "url": "https://doi.org/10.3389/feduc.2025.1643023",
        "category": "Academic", "agent_conf": "Medium", "url_status": "Live",
        "agg_score": 0.70, "reliability": "Medium", "cited_by": ["Academic"]
    },
    {
        "id": 12, "author": "Various Authors", "year": 2025,
        "title": "Artificial Intelligence in K-12 Education: An Umbrella Review",
        "url": "https://www.sciencedirect.com/science/article/pii/S2666920X25001596",
        "category": "Academic", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Academic"]
    },
    {
        "id": 13, "author": "Tyson, M. M. & Sauers, N. J.", "year": 2021,
        "title": "School Leaders' Adoption and Implementation of Artificial Intelligence",
        "url": "https://eric.ed.gov/?id=EJ1296923",
        "category": "Academic", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Academic"]
    },
    {
        "id": 14, "author": "CRPE", "year": 2025,
        "title": "Districts and AI: Early Adopters Focus More on Students in 2025-26",
        "url": "https://crpe.org/districts-and-ai-early-adopters-focus-more-on-students-in-2025-26/",
        "category": "Academic", "agent_conf": "Medium", "url_status": "Live",
        "agg_score": 0.80, "reliability": "High", "cited_by": ["Academic", "Industry"]
    },
    {
        "id": 15, "author": "U.S. DOE, Office of Ed. Technology", "year": 2024,
        "title": "Empowering Education Leaders: A Toolkit for Safe, Ethical, and Equitable AI Integration",
        "url": "https://eric.ed.gov/?id=ED661924",
        "category": "Government", "agent_conf": "Medium", "url_status": "Live",
        "agg_score": 0.70, "reliability": "Medium", "cited_by": ["Academic", "Gov"]
    },
    {
        "id": 16, "author": "Fortune", "year": 2026,
        "title": "Claude Code Gives Anthropic Its Viral Moment",
        "url": "https://fortune.com/2026/01/24/anthropic-boris-cherny-claude-code-non-coders-software-engineers/",
        "category": "News", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.80, "reliability": "High", "cited_by": ["News"]
    },
    {
        "id": 17, "author": "TechBuzz AI", "year": 2026,
        "title": "Claude Code Breaks Out: How Anthropic's Dev Tool Found Mass Appeal",
        "url": "https://www.techbuzz.ai/articles/claude-code-breaks-out-how-anthropic-s-dev-tool-found-mass-appeal",
        "category": "News", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.80, "reliability": "High", "cited_by": ["News"]
    },
    {
        "id": 18, "author": "Education Week Market Brief", "year": 2025,
        "title": "What Superintendents Want From AI: Build for Central Office Efficiency",
        "url": "https://marketbrief.edweek.org/meeting-district-needs/what-superintendents-want-from-ai-build-for-central-office-efficiency/2025/12",
        "category": "News", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.80, "reliability": "High", "cited_by": ["News", "Industry"]
    },
    {
        "id": 19, "author": "CoSN", "year": 2025,
        "title": "Operational AI in Education: A CoSN 2025 Member Survey",
        "url": "https://www.cosn.org/wp-content/uploads/2025/09/2025-HPE-Report_F2.pdf",
        "category": "Industry", "agent_conf": "High", "url_status": "PDF (unreadable via web)",
        "agg_score": 0.56, "reliability": "Medium", "cited_by": ["Industry"]
    },
    {
        "id": 20, "author": "Anthropic", "year": 2025,
        "title": "Introducing Claude for Education",
        "url": "https://www.anthropic.com/news/introducing-claude-for-education",
        "category": "Industry", "agent_conf": "High", "url_status": "Live",
        "agg_score": 0.70, "reliability": "Medium", "cited_by": ["News", "Industry"]
    },
    {
        "id": 21, "author": "Anthropic", "year": 2026,
        "title": "Customer Story: MagicSchool",
        "url": "https://claude.com/customers/magicschool",
        "category": "Industry", "agent_conf": "Medium", "url_status": "Live",
        "agg_score": 0.49, "reliability": "Low", "cited_by": ["Industry"]
    },
    {
        "id": 22, "author": "White House", "year": 2025,
        "title": "Major Organizations Commit to Supporting AI Education",
        "url": "https://www.whitehouse.gov/articles/2025/09/major-organizations-commit-to-supporting-ai-education/",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
    {
        "id": 23, "author": "AASA", "year": 2026,
        "title": "AASA and Day of AI Launch National Leadership & Innovation Fellowship",
        "url": "https://www.aasa.org/news-media/news/2026/03/19/aasa-and-day-of-ai-launch-national-fellowship",
        "category": "News", "agent_conf": "Medium", "url_status": "Live",
        "agg_score": 0.56, "reliability": "Medium", "cited_by": ["News"]
    },
    {
        "id": 24, "author": "RAND Corporation", "year": 2025,
        "title": "More Districts Are Training Teachers on Artificial Intelligence",
        "url": "https://www.rand.org/pubs/research_reports/RRA956-31.html",
        "category": "Academic", "agent_conf": "Medium", "url_status": "Gov-blocked (403)",
        "agg_score": 0.49, "reliability": "Low", "cited_by": ["Gov", "Industry"]
    },
    {
        "id": 25, "author": "U.S. Department of Education", "year": 2025,
        "title": "Dear Colleague Letter on AI and Federal Grant Funds",
        "url": "https://www.ed.gov/media/document/opepd-ai-dear-colleague-letter-7222025-110427.pdf",
        "category": "Government", "agent_conf": "High", "url_status": "Live",
        "agg_score": 1.0, "reliability": "High", "cited_by": ["Gov"]
    },
]


def build_document():
    doc = Document()

    # ── Default font ─────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ── Title Page ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n")
    run = p.add_run("Deep Research Report")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Directors of Schools and Assistant Directors of Schools\nUsing Claude Code")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")
    run = p.add_run("4 source categories | 25 sources | 22 claims analyzed")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "Calibri"

    doc.add_page_break()

    # ── Executive Summary ────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This report investigates whether any K-12 director of schools, assistant director of schools, "
        "superintendent, or assistant superintendent in the United States is using Claude Code, "
        "Anthropic's command-line AI coding tool, for district operations work. Four parallel research "
        "agents searched government, news, academic, and industry sources independently, producing "
        "25 unique sources across all four categories."
    )

    p = doc.add_paragraph()
    add_colored_run(p, "Bottom line: ", GREEN, bold=True)
    p.add_run(
        "No documented case exists of a K-12 director of schools or assistant director of schools "
        "using Claude Code for district operations. The closest documented case is Peninsula School "
        "District (WA), where CIO Kris Hagel built an open-source Claude Code plugin system for "
        "district operations, but Hagel is a technology leader, not a superintendent or instructional "
        "administrator."
    )

    doc.add_paragraph(
        "The broader context is significant: AI adoption among K-12 administrators is accelerating "
        "rapidly. RAND data shows 58% of principals used AI tools in 2023-24 (compared to 25% of "
        "teachers). In Tennessee, 85% of districts report educator AI use, with 84% citing reduced "
        "administrative task time as the primary benefit. However, this adoption is occurring through "
        "consumer-grade, no-code AI platforms (ChatGPT, MagicSchool, Google Gemini), not through "
        "developer-oriented tools like Claude Code."
    )

    p = doc.add_paragraph()
    add_colored_run(p, "Confidence in this conclusion: High. ", GREEN, bold=True)
    p.add_run(
        "All four source categories independently confirmed the null finding. The absence of evidence "
        "is itself significant given the breadth of search: federal and state government databases, "
        "major education trade press, peer-reviewed academic journals, professional organization "
        "publications, and Anthropic's own education reports all returned no results for this "
        "specific use case."
    )

    # ── Background & Context ─────────────────────────────────────────
    doc.add_heading("1. Background and Context", level=1)

    doc.add_paragraph(
        "Claude Code is Anthropic's agentic command-line interface (CLI) tool, launched in early 2025. "
        "Unlike Claude.ai (the chatbot) or the Claude API, Claude Code operates in a terminal environment, "
        "enabling users to write, edit, and execute code through natural language conversation. It has "
        "gained significant traction among software developers and, increasingly, non-technical professionals "
        "in industries like healthcare (Epic Systems), finance, marketing, and design."
    )

    doc.add_paragraph(
        "Directors of schools (the Tennessee term for superintendents) and assistant directors of schools "
        "are the chief executive and senior administrative officers of K-12 school districts. Their work "
        "spans budgeting, human resources, facilities management, data analysis, policy compliance, "
        "board communications, and strategic planning. Many of these tasks involve structured data, "
        "repetitive document generation, and analytical workflows that could benefit from AI-assisted coding."
    )

    doc.add_paragraph(
        "The question of whether any such administrator has adopted Claude Code is timely for several "
        "reasons: (1) Claude Code's user base has expanded rapidly to non-developers since early 2026, "
        "(2) K-12 districts are under increasing pressure to demonstrate AI adoption for operational "
        "efficiency, and (3) state legislatures (notably Tennessee's Public Chapter 550) now require "
        "districts to have formal AI policies."
    )

    # ── Key Findings ─────────────────────────────────────────────────
    doc.add_heading("2. Key Findings", level=1)

    # Finding group 1
    doc.add_heading("2.1 No K-12 Education Executive Uses Claude Code (Documented)", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[HIGH CONFIDENCE] ", GREEN, bold=True)
    p.add_run(
        "Across all four source categories, no government report, news article, academic study, or "
        "industry publication documents a superintendent, director of schools, assistant superintendent, "
        "or assistant director of schools using Claude Code for district operations work. This finding "
        "was confirmed independently by all four research agents."
    )

    p = doc.add_paragraph()
    p.add_run(
        "Anthropic's own education report explicitly excludes K-12 educators from analysis, focusing "
        "only on higher education accounts. Claude for Education, launched in 2025, targets university "
        "settings. Anthropic's K-12 partnerships (Teach For All, Iceland pilot) involve Claude the "
        "chatbot for classroom instruction, not Claude Code for administrative operations."
    )

    # Finding group 2
    doc.add_heading("2.2 Peninsula School District: The Closest Case", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[MEDIUM CONFIDENCE] ", YELLOW, bold=True)
    p.add_run(
        "Peninsula School District (PSD401) in Gig Harbor, WA maintains an open-source Claude Code "
        "plugin system on GitHub (psd401/psd-claude-coding-system) with 42 specialized agents and "
        "25 productivity workflows for district operations. This includes district-specific skills "
        "for athletics, brand guidelines, FreshService ticket integration, and RedRover systems. "
        "The system was built and is maintained by CIO Kris Hagel, a CoSN Board member and 2024 "
        "CoSN Volunteer of the Year. Hagel presented the framework at AIEmpoweredEDU2025."
    )

    p = doc.add_paragraph()
    p.add_run(
        "This is the only documented case of Claude Code being used for K-12 district operations "
        "anywhere in the United States. However, it was built by a technology leader (CIO), not by "
        "the superintendent or an instructional administrator. The distinction matters: the research "
        "question is whether education executives, not IT staff, are personally using Claude Code."
    )

    # Finding group 3
    doc.add_heading("2.3 K-12 Administrators Are Rapidly Adopting AI, but Not Developer Tools", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[HIGH CONFIDENCE] ", GREEN, bold=True)
    p.add_run(
        "Multiple independent data sources confirm that K-12 administrators use AI at significantly "
        "higher rates than teachers. RAND (2025) found 58% of U.S. principals used AI in 2023-24 "
        "vs. 25% of teachers. Riegel et al. (2025) found administrators scored 3.41 on an AI use "
        "scale vs. 2.48 for teachers (p<.001). Berkovich (2025) found 50% of school leaders are in "
        "the 'early majority' adoption stage."
    )

    p = doc.add_paragraph()
    p.add_run(
        "However, the tools being used are uniformly consumer-grade: ChatGPT, Google Gemini, "
        "MagicSchool, SchoolAI, and similar no-code platforms. Superintendent Corey Smith (South "
        "Putnam, IN) described using generative AI for RFP analysis and budget spreadsheet review, "
        "but through copy-paste workflows, not coding tools. No administrator in any source described "
        "using a CLI, writing scripts, or engaging with developer-oriented AI tools."
    )

    # Finding group 4
    doc.add_heading("2.4 Tennessee Context: Strong AI Momentum, No Claude Code Adoption", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[HIGH CONFIDENCE] ", GREEN, bold=True)
    p.add_run(
        "Tennessee is among the most proactive states in K-12 AI policy. Public Chapter 550 (2024) "
        "requires all LEAs to adopt AI policies. SCORE/TOSS survey data shows 85% of TN districts "
        "report educator AI use, with 84% identifying reduced administrative task time as the top "
        "benefit. The Tennessee AI Advisory Council published an action plan in November 2025 that "
        "includes education. None of these sources mention Claude Code or any AI coding tool for "
        "administrators."
    )

    # Finding group 5
    doc.add_heading("2.5 The Academic Literature Has a Complete Blind Spot", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[HIGH CONFIDENCE] ", GREEN, bold=True)
    p.add_run(
        "The academic research on AI in K-12 education focuses almost entirely on instructional "
        "applications (personalized learning, assessment, tutoring) and teacher adoption. An umbrella "
        "review of 102 systematic reviews (2025) found that 'preparing administrators and leaders to "
        "understand and ethically use AI has received much less attention.' No academic study examines "
        "whether school administrators possess programming skills, use coding tools, or leverage "
        "developer-oriented AI for operational tasks."
    )

    # Finding group 6
    doc.add_heading("2.6 Claude Code's Non-Developer Adoption Is Growing, but Education Is Absent", level=2)

    p = doc.add_paragraph()
    add_colored_run(p, "[HIGH CONFIDENCE] ", GREEN, bold=True)
    p.add_run(
        "Fortune (Jan 2026) and TechBuzz documented Claude Code's expansion to non-developers: "
        "marketing managers, designers, executives, healthcare workers at Epic Systems, and small "
        "business owners. The sectors named include healthcare, big tech (Netflix, Uber, Spotify), "
        "finance, and consulting. K-12 education is entirely absent from this adoption narrative."
    )

    # ── Points of Contradiction ──────────────────────────────────────
    doc.add_heading("3. Points of Contradiction", level=1)

    doc.add_paragraph(
        "No contradictions were found across the four source categories. All agents independently "
        "confirmed the same null finding regarding Claude Code use by K-12 education executives. "
        "The only area of potential tension is between the narrative that AI adoption among "
        "administrators is 'rapid' (RAND, Riegel) and the reality that this adoption is limited "
        "to consumer-grade tools. Whether administrators would benefit from developer-level tools "
        "like Claude Code is an open question that no source addresses."
    )

    # ── Single-Source Claims ─────────────────────────────────────────
    doc.add_heading("4. Single-Source Claims", level=1)

    doc.add_paragraph("The following claims are supported by only one source category and warrant additional scrutiny:")

    claims = [
        (
            "Peninsula School District's Claude Code plugin system is the only documented K-12 "
            "Claude Code deployment in the US.",
            "Industry sources only. No government, news, or academic source references this system. "
            "Additional verification would require examining other district GitHub organizations and "
            "interviewing Hagel directly."
        ),
        (
            "Budget planning/resource allocation has the lowest AI adoption rate among school "
            "leadership tasks (25.8%).",
            "Academic sources only (Berkovich, 2025). Based on an Israeli sample (n=302), which "
            "limits generalizability to U.S. districts. No U.S.-focused study has replicated this "
            "finding."
        ),
        (
            "Only 8% of K-12 IT workforce is 'well prepared' for AI (CoSN survey).",
            "Industry sources only. CoSN membership may skew toward more tech-engaged districts, "
            "meaning the actual number could be lower nationally."
        ),
    ]

    for claim, note in claims:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(claim)
        run.bold = True
        p.add_run(f" {note}")

    # ── Practical Implications ───────────────────────────────────────
    doc.add_heading("5. Practical Implications", level=1)

    doc.add_paragraph("For a K-12 administrator considering Claude Code:")

    implications = [
        (
            "You would be a genuine first mover.",
            "No documented superintendent or director of schools is using Claude Code for district "
            "operations. Peninsula School District's CIO has built the infrastructure, but no education "
            "executive has publicly adopted it. This represents both opportunity and risk."
        ),
        (
            "The use cases are strong but undocumented.",
            "Superintendent Smith's described workflows (RFP analysis, budget trend analysis, email triage, "
            "document generation) are precisely the tasks where Claude Code excels. The gap between "
            "'I upload spreadsheets to a chatbot' and 'Claude Code automates my enrollment pipeline' "
            "is significant but unexplored in practice by education leaders."
        ),
        (
            "Anthropic has a K-12 blind spot.",
            "Anthropic's education report explicitly excludes K-12. Claude for Education targets "
            "universities. Their K-12 partnerships focus on classroom instruction, not district operations. "
            "An administrator using Claude Code for district work would be outside Anthropic's current "
            "K-12 strategy."
        ),
        (
            "The policy environment is favorable but tool-agnostic.",
            "Tennessee's AI requirements (Public Chapter 550) and federal guidance encourage AI adoption "
            "but deliberately avoid naming specific tools. A district using Claude Code would need to "
            "ensure compliance with its own AI policy, data governance, and PII protections."
        ),
        (
            "Publication and speaking opportunities are wide open.",
            "The AASA/MIT RAISE fellowship (July 2026), CoSN conferences, and education trade press "
            "are actively seeking superintendent AI stories. A director of schools using Claude Code "
            "for data automation, enrollment projections, or board document generation would be a "
            "compelling case study with no competition in the space."
        ),
    ]

    for title, detail in implications:
        p = doc.add_paragraph()
        run = p.add_run(title + " ")
        run.bold = True
        p.add_run(detail)

    doc.add_heading("Unanswered Questions", level=2)
    questions = [
        "Are there administrators using Claude Code who simply haven't published about it?",
        "Would Anthropic be interested in a K-12 administrator case study for Claude Code?",
        "What data governance considerations are unique to running Claude Code on district data?",
        "Could Claude Code's multi-agent capabilities replace multiple point solutions (MagicSchool, SchoolAI, etc.) for a district?",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Bullet")

    # ── Federal & State Policy Context ───────────────────────────────
    doc.add_heading("6. Federal and State Policy Context", level=1)

    doc.add_paragraph(
        "The federal government has issued multiple guidance documents on AI in education. The USED's "
        "2023 report 'Artificial Intelligence and the Future of Teaching and Learning' identified "
        "administrative AI applications including resource allocation and student risk prediction. "
        "A July 2025 Dear Colleague Letter encourages grantees to use AI for improved education "
        "outcomes. The Office of Educational Technology's 2024 toolkit specifically targets education "
        "leaders for 'safe, ethical, and equitable AI integration.' None of these documents mention "
        "coding tools or administrators writing code."
    )

    doc.add_paragraph(
        "At the state level, Tennessee is a national leader. Public Chapter 550 (2024) requires all "
        "LEAs to adopt AI policies. The Tennessee AI Advisory Council's November 2025 action plan "
        "includes education as a focus area. TSBA's model policy (4.214) provides a framework for "
        "districts. Two additional 2025 bills direct TDOE to provide guidance and fund professional "
        "development. These policies are uniformly tool-agnostic, creating space for districts to "
        "adopt any AI tool, including Claude Code, within their governance framework."
    )

    # ── Tennessee Data Appendix ──────────────────────────────────────
    doc.add_heading("7. Tennessee Data Appendix", level=1)

    tn_data = [
        ["Metric", "Value", "Source"],
        ["Districts reporting AI use", "85% (of 86 responding districts)", "SCORE/TOSS 2025"],
        ["Districts offering AI training", "~64%", "SCORE/TOSS 2025"],
        ["Top benefit cited", "Reduced admin task time (84%)", "SCORE/TOSS 2025"],
        ["Top concern", "Cheating/plagiarism (84%)", "SCORE/TOSS 2025"],
        ["Need additional PD", "85%", "SCORE/TOSS 2025"],
        ["AI policy required by law", "Yes (Public Chapter 550, 2024-25)", "TN General Assembly"],
        ["State AI Advisory Council", "Active, action plan Nov 2025", "tn.gov"],
        ["Districts with Claude Code adoption", "0 documented", "All sources"],
    ]

    table = doc.add_table(rows=len(tn_data), cols=3)
    table.style = "Table Grid"
    make_header_row(table, tn_data[0])
    for i, row_data in enumerate(tn_data[1:], 1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
    set_table_style(table)

    # ── Source Reliability Matrix ─────────────────────────────────────
    doc.add_heading("8. Source Reliability Matrix", level=1)

    headers = ["#", "Author/Org", "Year", "Category", "Agent Conf.", "URL Status", "Agg. Score", "Reliability"]
    sorted_sources = sorted(sources, key=lambda s: s["agg_score"], reverse=True)

    table = doc.add_table(rows=len(sorted_sources) + 1, cols=len(headers))
    table.style = "Table Grid"
    make_header_row(table, headers)

    for i, s in enumerate(sorted_sources, 1):
        row = table.rows[i]
        vals = [
            str(s["id"]), s["author"], str(s["year"]), s["category"],
            s["agent_conf"], s["url_status"], f"{s['agg_score']:.2f}", s["reliability"]
        ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"
        # Color the reliability cell
        rel_cell = row.cells[7]
        if s["reliability"] == "High":
            set_cell_shading(rel_cell, "C6EFCE")
        elif s["reliability"] == "Medium":
            set_cell_shading(rel_cell, "FFEB9C")
        else:
            set_cell_shading(rel_cell, "FFC7CE")

    set_table_style(table)

    # Summary stats
    high_count = sum(1 for s in sources if s["reliability"] == "High")
    med_count = sum(1 for s in sources if s["reliability"] == "Medium")
    low_count = sum(1 for s in sources if s["reliability"] == "Low")

    p = doc.add_paragraph()
    p.add_run(f"\n{len(sources)} sources total: {high_count} high reliability, {med_count} medium, {low_count} low. ")
    p.add_run("All URLs verified live except 1 PDF (CoSN, unreadable via web fetch) and 1 gov-blocked (RAND, 403). 0 dead URLs replaced.")

    # ── Claim Cross-Reference Matrix ─────────────────────────────────
    doc.add_heading("9. Claim Cross-Reference Matrix", level=1)

    claim_headers = ["Claim", "Gov", "News", "Acad.", "Ind.", "Agreement"]
    claim_data = [
        ["No K-12 exec uses Claude Code", "\\u2713", "\\u2713", "\\u2713", "\\u2713", "Full"],
        ["PSD CIO built Claude Code plugins", "\\u2014", "\\u2014", "\\u2014", "\\u2713", "Single"],
        ["Admins use AI more than teachers", "\\u2014", "\\u2014", "\\u2713", "\\u2014", "Single"],
        ["58% of principals used AI (2023-24)", "\\u2014", "\\u2014", "\\u2713", "\\u2014", "Single"],
        ["85% TN districts report AI use", "\\u2713", "\\u2014", "\\u2014", "\\u2713", "Partial"],
        ["TN requires AI policies (PC 550)", "\\u2713", "\\u2014", "\\u2014", "\\u2014", "Single"],
        ["Anthropic excludes K-12 from research", "\\u2014", "\\u2713", "\\u2713", "\\u2713", "Full"],
        ["Claude Code adopted by non-devs", "\\u2014", "\\u2713", "\\u2014", "\\u2713", "Partial"],
        ["MagicSchool uses Claude backend", "\\u2014", "\\u2014", "\\u2014", "\\u2713", "Single"],
        ["Supts want AI for central office", "\\u2014", "\\u2713", "\\u2014", "\\u2713", "Partial"],
        ["Budget/finance is lowest AI adoption", "\\u2014", "\\u2014", "\\u2713", "\\u2014", "Single"],
        ["79 early-adopter districts (CRPE)", "\\u2014", "\\u2014", "\\u2713", "\\u2713", "Partial"],
    ]

    # Replace unicode escapes
    for row in claim_data:
        for i in range(len(row)):
            row[i] = row[i].replace("\\u2713", "\u2713").replace("\\u2014", "\u2014")

    table = doc.add_table(rows=len(claim_data) + 1, cols=len(claim_headers))
    table.style = "Table Grid"
    make_header_row(table, claim_headers)

    for i, row_data in enumerate(claim_data, 1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
        # Color agreement column
        agree_cell = table.rows[i].cells[5]
        agreement = row_data[5]
        if agreement == "Full":
            set_cell_shading(agree_cell, "C6EFCE")
        elif agreement == "Partial":
            set_cell_shading(agree_cell, "FFEB9C")
        else:
            set_cell_shading(agree_cell, "FFC7CE")

    set_table_style(table)

    # ── References ───────────────────────────────────────────────────
    doc.add_heading("10. References", level=1)

    sorted_refs = sorted(sources, key=lambda s: s["author"])
    for s in sorted_refs:
        p = doc.add_paragraph()
        p.add_run(f"{s['author']} ({s['year']}). ").bold = False
        run = p.add_run(f"{s['title']}. ")
        run.italic = True
        p.add_run(f"{s['url']} ")
        rel_color = GREEN if s["reliability"] == "High" else (YELLOW if s["reliability"] == "Medium" else RED)
        add_colored_run(p, f"[{s['reliability']} Reliability]", rel_color, bold=True, size=Pt(9))

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_document()
