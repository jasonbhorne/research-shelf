#!/usr/bin/env /opt/anaconda3/bin/python3
"""Generate companion guide for Following the Sun by Margaret Bensfield Sullivan."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
import os

OUTPUT_DIR = os.path.expanduser("~/Documents/Research/following-the-sun")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

# -- Style setup --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.name = "Calibri"
    heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    if level == 1:
        heading_style.font.size = Pt(22)
    elif level == 2:
        heading_style.font.size = Pt(16)
    else:
        heading_style.font.size = Pt(13)

# Quote style
quote_style = doc.styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
quote_style.font.name = "Calibri"
quote_style.font.size = Pt(11)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
quote_style.paragraph_format.left_indent = Inches(0.5)
quote_style.paragraph_format.right_indent = Inches(0.5)
quote_style.paragraph_format.space_before = Pt(6)
quote_style.paragraph_format.space_after = Pt(6)

# -- Title Page --
for _ in range(6):
    doc.add_paragraph()

title = doc.add_heading("Following the Sun", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Tales (and Fails) From a Year Around the World With Our Kids")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.runs[0]
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

doc.add_paragraph()

meta = doc.add_paragraph("A Media Companion Guide")
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.runs[0]
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

doc.add_paragraph()

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = details.add_run(f"Book by Margaret Bensfield Sullivan\nGenerated {date.today().strftime('%B %d, %Y')}\nFull Spoilers")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# -- Quick Take --
doc.add_heading("Quick Take", level=2)

doc.add_paragraph(
    "Following the Sun is a self-published family travel memoir that punches well above its weight. "
    "In January 2019, Margaret and Teddy Sullivan quit their careers, gave up their Manhattan apartment, "
    "pulled their two kids (Willa, 6, and James, 4) out of school, packed nothing but carry-on bags, "
    "and spent a year visiting 29 countries across six continents. They took 79 flights and slept in "
    "roughly 100 different places. The subtitle tells you everything about the book's tone: this is not "
    "a curated Instagram travelogue but a frank, funny, beautifully written account of what happens when "
    "two Type-A New Yorkers try to show their small children the world, complete with altitude sickness in "
    "Peru, a lice infestation in Berlin, a terrible stomach bug in Beijing, and an ethically disastrous "
    "orphanage visit in Vietnam. The top reviewer in the family-gap-year genre on Goodreads called it her "
    'favorite book in the category. With a 4.56/5 Goodreads average, major features in CNN and CNBC, and '
    "the unintentional poignancy of a 2019 trip completed just before COVID shut down global travel, this "
    "is one of the best entries in its niche."
)

# -- About the Author --
doc.add_heading("About the Author", level=2)

doc.add_paragraph(
    "Margaret Bensfield Sullivan is a Washington, D.C., native who holds a B.A. from Vanderbilt University "
    "and an M.A. in French from Middlebury College. Before writing the book, she spent nearly 15 years in "
    "corporate communications and brand marketing, rising to partner at WPP's Group SJR, where she designed "
    "storytelling campaigns for clients including TED, Target, Disney, and USAID. She is also an illustrator "
    "whose work has appeared in Tyler Anbinder's Plentiful Country: The Great Potato Famine and the Making "
    "of Irish New York (Little, Brown, 2024). She lives in Lower Manhattan."
)

doc.add_paragraph(
    "Her husband, Teddy Sullivan, brings his own interesting backstory. A former pitcher at Duke University "
    "who played in the minor leagues for the Cleveland Indians, Ted earned his MBA from Harvard Business "
    "School and co-founded GameChanger, a digital scorekeeping platform that was acquired by DICK'S Sporting "
    "Goods in late 2016. Both parents were at career inflection points when they decided to take the trip, "
    "which matters: this was not a sabbatical or a remote-work arrangement. They walked away from everything."
)

# -- The Trip --
doc.add_heading("The Trip That Became the Book", level=2)

doc.add_paragraph(
    "The origin story is the kind of thing that makes you either nod in recognition or roll your eyes, "
    "depending on your tolerance for epiphanies. In fall 2017, Margaret flew to Arusha, Tanzania, to "
    "oversee media relations for a TED conference through her work at Group SJR. Somewhere between the "
    "conference and the 20-hour flight home, she had a realization:"
)

doc.add_paragraph(
    '"The world is big, I am small, my life is small."',
    style="BlockQuote"
)

doc.add_paragraph(
    "She feared that her New York routine had become \"transactional to the point of being forgettable\" "
    "and that she might \"wake up in 15 years with kids in college, wondering where the time went.\" By "
    "the time the plane landed, she had formulated a pitch for Teddy: a full year of global travel, with "
    "the kids, without jobs, on the move, seeing and learning as much as they could about the rest of the world."
)

doc.add_paragraph(
    "In January 2019, they left. Their route followed warm weather seasonally around the globe (hence the "
    "title): South America first, then Africa, the Middle East, Europe, Asia, Australia, and New Zealand. "
    "This was not a winging-it adventure. They hired travel agents and had three months of plans mapped out "
    "at departure. The family packed only carry-on bags for the entire year."
)

doc.add_paragraph(
    "Specific highlights and lowlights from the 29-country itinerary include:"
)

# Highlights table
table = doc.add_table(rows=8, cols=2)
table.style = "Light Shading Accent 1"
table.cell(0, 0).text = "Destination"
table.cell(0, 1).text = "Notable Episode"
highlights = [
    ("Peru", "Cloud forests and ascending Machu Picchu, plus altitude sickness"),
    ("Mongolia", "Horse races and vast open landscape"),
    ("Zimbabwe", "Sunsets that define the trip's visual memory"),
    ("Beijing, China", "A terrible stomach bug that flattened the family"),
    ("Berlin, Germany", "A lice infestation, because of course"),
    ("Vietnam", "An ethically problematic orphanage visit Sullivan has spoken about forcefully"),
    ("29 countries total", "79 flights, ~100 accommodations, 6 continents, 1 year"),
]
for i, (dest, episode) in enumerate(highlights):
    table.cell(i + 1, 0).text = dest
    table.cell(i + 1, 1).text = episode

doc.add_paragraph()

# -- What the Critics Say --
doc.add_heading("What the Critics Say", level=2)

doc.add_paragraph(
    "Following the Sun is self-published (December 2023, 268 pages, ISBN 9798218280505), which means "
    "it did not receive coverage from the traditional literary gatekeepers: no New York Times review, no "
    "Kirkus, no Publishers Weekly. What it got instead was something arguably more valuable for a travel "
    "memoir: major mainstream media features."
)

doc.add_heading("Media Coverage", level=3)

doc.add_paragraph(
    "CNN Travel ran a feature with the headline: \"They'd never seen the point of traveling with their "
    "young children. Then they hit the road for a year and discovered it came with unexpected perks.\" "
    "CNBC published \"This family traveled for a year. Here are the biggest mistakes they made,\" which "
    "was syndicated across NBC affiliates in Washington, Los Angeles, New York, Dallas, and San Diego. "
    "Additional features appeared in Silicon Valley Times, Authority Magazine, Confetti Travel Cafe, "
    "and InnerSelf.com. For a self-published debut, the PR campaign has been remarkably effective, almost "
    "certainly reflecting Sullivan's marketing expertise."
)

doc.add_heading("Reader Reception", level=3)

doc.add_paragraph(
    "On Goodreads, the book holds a 4.56/5 average from 39 ratings. The sample is small, but the "
    "reviews are telling:"
)

doc.add_paragraph(
    '"Absolutely excellent. Fantastically written and great mix between logistics and local color. '
    "I try to read every single book in this genre (family gap year) and this is my favorite one.\" "
    "-- Sarah (5 stars, Goodreads)",
    style="BlockQuote"
)

doc.add_paragraph(
    '"Margaret is a lovely writer who exposes all the good, bad and wonderful things that can happen '
    'when you break free of your comfort zone." -- Michael Gilbert (4 stars, Goodreads)',
    style="BlockQuote"
)

doc.add_paragraph(
    "The most critical Goodreads review (2 stars, Annie Sparks) called it \"kind of a play by play diary\" "
    "and \"a little mundane,\" but still pulled out a memorable line:"
)

doc.add_paragraph(
    '"An autopilot routine can make life feel short, whereas memories make life fuller."',
    style="BlockQuote"
)

doc.add_paragraph(
    "Colin Jordan's Medium review praised Sullivan's \"ability to blend personal experience with anecdotes "
    'on carpe diem, and subjective, quality of life Taoist principles" and noted "considerable craft" in '
    "the writing."
)

# -- Themes and What Makes It Work --
doc.add_heading("Themes and What Makes It Work", level=2)

doc.add_paragraph(
    "The book operates on several levels simultaneously, which is what elevates it beyond a standard "
    "travel diary."
)

doc.add_heading("Autopilot vs. Intentional Living", level=3)
doc.add_paragraph(
    "The central thesis, stated plainly, is that routine can make life feel short while memories make it "
    "feel full. Sullivan frames the trip as a deliberate break from autopilot: the New York grind of "
    "commute, daycare, work, repeat. This is not a new idea, but the scale of her response to it (quit "
    "everything, travel the world) gives it dramatic weight."
)

doc.add_heading("Honest Parenting on the Road", level=3)
doc.add_paragraph(
    "The \"Fails\" in the subtitle do real work. Sullivan does not pretend that traveling with a 4-year-old "
    "and a 6-year-old across 29 countries is glamorous. The kids got sick. They got lice. They melted down. "
    "But they also adapted better than the adults expected, and those moments of resilience are some of the "
    "book's most satisfying passages."
)

doc.add_heading("Ethics and Self-Awareness", level=3)
doc.add_paragraph(
    "The Vietnam orphanage episode is the book's most uncomfortable moment, and Sullivan handles it by "
    "being direct:"
)
doc.add_paragraph(
    '"I cannot state strongly enough that this kind of tourist destination is completely irresponsible, '
    'and no one should visit an orphanage."',
    style="BlockQuote"
)
doc.add_paragraph(
    "This willingness to admit mistakes, not just inconveniences but genuine ethical misjudgments, gives "
    "the book credibility that a more polished account would lack."
)

doc.add_heading("Privilege, Acknowledged", level=3)
doc.add_paragraph(
    "A year of global travel with two kids requires substantial financial resources and professional "
    "flexibility. The Sullivans had both: her career at WPP, his GameChanger exit. The book does not "
    "pretend otherwise, and this honesty is important. The experience is aspirational but not replicable "
    "for most families, and readers appreciate when an author knows that."
)

doc.add_heading("Pre-COVID Timing", level=3)
doc.add_paragraph(
    "The family traveled in 2019 and returned to the US in early 2020, just before the pandemic shut "
    "down international travel. This timing was pure accident, but it gives the book an unintentional "
    "elegiac quality: a portrait of a world that was about to become, for a while, unreachable. Multiple "
    "reviewers have noted this poignancy."
)

doc.add_heading("The Family Unit as Project", level=3)
doc.add_paragraph(
    "Colin Jordan's Medium review identified the book's \"ideological crux\" as \"the pursuit of "
    "self-determining the family unit.\" By removing every external structure (jobs, school, apartment, "
    "routine), the Sullivans turned their family into a self-contained system that had to figure out its "
    "own rhythms. The book is, at its core, a relationship story: two adults and two small children "
    "negotiating a year of intense togetherness."
)

# -- What You Might Not Know --
doc.add_heading("What You Might Not Know", level=2)

doc.add_paragraph(
    "Margaret's French M.A. from Middlebury was not just an academic credential. She credits the program "
    "with giving her \"the confidence to see the world this way.\" Language skills opened doors and shaped "
    "the family's experience in Francophone countries."
)

doc.add_paragraph(
    "The book took nearly four years to write. The family returned in early 2020, and the book was "
    "published in December 2023. Sullivan was \"constantly being asked for advice from other families "
    "hoping to take a similar trip,\" which became the catalyst to write it down. The long gestation "
    "period shows: this reads as a shaped narrative, not a rushed travel diary."
)

doc.add_paragraph(
    "Teddy Sullivan pitched for the Cleveland Indians' minor league system before pivoting to business. "
    "His co-founding of GameChanger, a scorekeeping app acquired by DICK'S Sporting Goods in 2016, is "
    "barely mentioned in the book's promotional materials but is essential context for understanding how "
    "the trip was financially possible."
)

doc.add_paragraph(
    "Sullivan's branding career shows in the book's packaging. The \"follow the sun\" concept, chasing "
    "warm weather seasonally to pack light, is both a practical strategy and a marketing-sharp metaphor. "
    "It gave the book its title, its organizing logic, and its emotional register."
)

doc.add_paragraph(
    "The self-publishing route was likely a deliberate choice, not a last resort. Given Sullivan's "
    "marketing background and the niche audience for family gap year memoirs, indie publishing with a "
    "strong PR campaign may have been the smarter business decision. The CNN and CNBC placements suggest "
    "she knew exactly what she was doing."
)

# -- On Screen --
doc.add_heading("On Screen", level=2)

doc.add_paragraph(
    "No film, TV, documentary, or stage adaptations exist, which is unsurprising for a 2023 self-published "
    "debut. The book's episodic structure (29 countries in a year) would lend itself naturally to a "
    "docuseries format. Whether that happens will depend on the book's continued visibility."
)

doc.add_paragraph(
    "The closest existing documentary parallel is Given (2016, dir. Jess Bianchi, IMDb 7.7/10), which "
    "follows legendary surfers Aamion and Daize Goodwin and their two young children on a 14-month journey "
    "through 15 countries across 6 continents, narrated through the perspective of their 6-year-old son. "
    "The surfing angle differentiates it, but the core themes, family bonding, seeing the world through "
    "children's eyes, choosing adventure over stability, are identical."
)

# -- Author Interviews and Appearances --
doc.add_heading("Hear More From the Author", level=2)

doc.add_paragraph(
    "Sullivan has been active on the podcast and interview circuit. Key appearances:"
)

appearances = [
    ("A Dad's Path, Episode 102", "Practical travel hacks, balancing career and family, stepping out of comfort zones"),
    ("I Want What She Has, Episode 328", "Paired with Tami Lynn Kent (Wild Mothering); the year-long journey and family dynamics"),
    ("James Miller LIFEOLOGY", "Author/illustrator background, brand marketing career, life design"),
    ("eHealth Radio Network", "What a family gap year is, inspiration, lessons learned"),
    ("Silicon Valley Times (written)", "The Tanzania epiphany, Einstein on stepping away from routine"),
    ("Authority Magazine", "Travel and personal growth"),
]

table2 = doc.add_table(rows=len(appearances) + 1, cols=2)
table2.style = "Light Shading Accent 1"
table2.cell(0, 0).text = "Outlet / Episode"
table2.cell(0, 1).text = "Focus"
for i, (outlet, focus) in enumerate(appearances):
    table2.cell(i + 1, 0).text = outlet
    table2.cell(i + 1, 1).text = focus

doc.add_paragraph()

# -- If You Liked This --
doc.add_heading("If You Liked This", level=2)

doc.add_paragraph(
    "Five books in the same orbit, ranked by how closely they match:"
)

recs = [
    (
        "How to Be a Family by Dan Kois (2019)",
        "The closest thematic match. A Slate editor drags his wife and two pre-teen daughters from "
        "suburban Virginia to live in New Zealand, the Netherlands, Costa Rica, and Kansas to understand "
        "how different cultures approach family life. Same \"quit the routine\" DNA, but with a more "
        "sociological lens. Where Sullivan covers 29 countries in a year, Kois goes deep in four."
    ),
    (
        "Excess Baggage by Tracey Carisch (2018)",
        "The closest structural match. A high-achieving executive has a panic attack, questions everything, "
        "and convinces her husband to sell their possessions and travel 24 countries on 6 continents for "
        "18 months with three young daughters. The catalyst is burnout rather than wanderlust."
    ),
    (
        "At Home in the World by Tsh Oxenreider (2017)",
        "Same family size, same duration, more reflective and philosophical in tone. Tsh, her husband, "
        "and three kids under ten circumnavigate the globe for nine months out of backpacks. A meditation "
        "on belonging and home."
    ),
    (
        "One Year Off by David Elliot Cohen (1999)",
        "The original in the genre. The co-creator of the \"Day in the Life\" photo book series sells "
        "everything and takes his family around the world. Written as email dispatches to friends, which "
        "is a charming late-'90s artifact."
    ),
    (
        "The Passport Project by Kellie McIntyre (2022)",
        "A mother and her two teen daughters ditch middle school for five months around the world. The "
        "daughters' own blog entries and journal excerpts are woven into the narrative. Good for readers "
        "who want the kids' perspective."
    ),
]

for title_text, description in recs:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(f"\n{description}")
    p.paragraph_format.space_after = Pt(8)

doc.add_heading("Companion Reading", level=3)

doc.add_paragraph(
    "For readers who finish Following the Sun and think, \"Could we actually do this?\":"
)

companions = [
    "The Family Sabbatical Handbook by Elisa Bernick -- the practical how-to for families considering extended travel abroad",
    "World Schooling by Ashley Dymock de Tello -- the educational philosophy behind pulling kids from school to travel",
    "Wonder Year (wonderyear.com) -- an online guide to long-term family travel and worldschooling",
]
for c in companions:
    doc.add_paragraph(c, style="List Bullet")

doc.add_heading("Related Media", level=3)

doc.add_paragraph(
    "Given (2016, documentary) -- a family with young children traveling the world for 14 months, "
    "narrated through a child's eyes. IMDb 7.7/10."
)
doc.add_paragraph(
    "360 Degrees Longitude by John Higham -- a year-long, 28-country family trip with an interactive "
    "Google Earth companion."
)
doc.add_paragraph(
    "Podcasts: The Jetsetting Family Travel Podcast (Rod and Jess, 50+ countries with two kids), "
    "Travel with Kids (Emily Krause), Family Travel Podcast (A Big Peachey Adventure)."
)

# -- Book Club Questions --
doc.add_heading("Book Club Discussion Questions", level=2)

questions = [
    "The Sullivans quit their NYC jobs, gave up their apartment, and left with kids ages 4 and 6. "
    "At what point does \"adventurous parenting\" cross into recklessness? Did the Sullivans cross that line?",

    "Twenty-nine countries in one year is roughly a new country every 12-13 days. How does this pace "
    "shape the kind of experience a family can have? Would fewer countries and longer stays have produced "
    "a different book and a different family?",

    "The book was written by Margaret, but the trip was a joint decision with Teddy. How does a single "
    "narrator shape a family story? What might Teddy's version look like? The kids' version in 20 years?",

    "The Sullivans traveled in 2019, just before COVID-19 shut down global travel. How does knowing what "
    "came next change how you read the book?",

    "Extended family travel of this kind requires significant financial resources. How does the book handle "
    "the privilege required to do what the Sullivans did? Does acknowledging privilege make a difference, "
    "or is the acknowledgment itself a kind of performance?",

    "What is the youngest age at which children actually benefit from international travel, versus simply "
    "being along for the ride? Will the Sullivan kids remember any of this, and does it matter if they don't?",

    "If you had one year and no obligations, would you do what the Sullivans did? If not, what would you "
    "do instead, and what does your answer reveal about your values?"
]

for i, q in enumerate(questions, 1):
    doc.add_paragraph(f"{i}. {q}")

# -- Sources --
doc.add_heading("Sources", level=2)

sources = [
    "Authority Magazine. \"Travel and Personal Growth: Author Margaret Sullivan.\" Medium. https://medium.com/authority-magazine/travel-and-personal-growth-author-margaret-sullivan-on-why-how-traveling-can-help-us-become-f4032d45e8ba",
    "BetterPitching.com. \"Talking Pitching with GameChanger CEO Ted Sullivan.\" https://betterpitching.com/talking-pitching-with-gamechanger-ceo-ted-sullivan/",
    "CNN Travel. \"They'd never seen the point of traveling with their young children. Then they hit the road for a year.\" https://www.cnn.com/travel/follow-the-sun-traveling-with-kids",
    "CNBC (2024). \"This family traveled for a year. Here are the biggest mistakes they made.\" https://www.cnbc.com/2024/05/02/this-family-traveled-for-a-year-these-are-their-biggest-mistakes.html",
    "Confetti Travel Cafe. \"Worried They Might Not Remember? Travel With Little Kids Anyway.\" https://www.confettitravelcafe.com/worried-they-might-not-remember-travel-with-little-kids-anyway/",
    "Feedspot. \"30 Best Family Travel Podcasts to Listen to in 2026.\" https://podcast.feedspot.com/family_travel_podcasts/",
    "Goodreads. Following the Sun listing and reviews. https://www.goodreads.com/book/show/203437832-following-the-sun",
    "InnerSelf.com. \"Wandering Without Whining: How to Keep Kids Engaged When Traveling.\" https://innerself.com/personal/relationships/parenting/30019-wandering-without-whining.html",
    "James Miller LIFEOLOGY. \"Following The Sun | Margaret Bensfield Sullivan.\" https://jamesmillerlifeology.com/following-the-sun-margaret-bensfield-sullivan/",
    "Jordan, Colin. \"REVIEW: Margaret Bensfield Sullivan -- Following the Sun (BOOK).\" Medium. https://colin-jordan524.medium.com/review-margaret-bensfield-sullivan-following-the-sun-book-5d761c2a826a",
    "Middlebury Language Schools. \"Margaret Sullivan: Middlebury French MA Alumna, Author and Illustrator.\" https://www.middlebury.edu/language-schools/blog/margaret-sullivan-middlebury-french-ma-alumna-author-and-illustrator",
    "Silicon Valley Times. \"Exploring the Globe with Margaret Sullivan.\" https://siliconvalleytime.com/interview/exploring-the-globe-with-margaret-sullivan-inside-following-the-sun-tales-and-fails-from-a-year-around-the-world-with-our-kids/",
    "Sullivan, Margaret Bensfield. Following the Sun: Tales (and Fails) From a Year Around the World With Our Kids. Self-published, 2023. ISBN 9798218280505.",
    "TechCrunch (2016). \"Dick's Sporting Goods acquires mobile scorekeeping company GameChanger Media.\" https://techcrunch.com/2016/11/28/dicks-sporting-goods-acquires-mobile-scorekeeping-company-gamechanger-media/",
]

for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# -- Save --
output_path = os.path.join(OUTPUT_DIR, f"{date.today().isoformat()} Following the Sun Companion Guide.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
