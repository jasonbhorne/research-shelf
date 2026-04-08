"""
Companion Guide: The Hard Thing About Hard Things by Ben Horowitz
Generated: March 18, 2026
Output: ~/Documents/Research/hard-thing-about-hard-things/
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.expanduser(
    "~/Documents/Research/hard-thing-about-hard-things"
)
OUTPUT_FILE = os.path.join(
    OUTPUT_DIR,
    "2026-03-18 The Hard Thing About Hard Things Companion Guide.docx"
)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, italic=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=11, italic=True)
    return p


def add_recommendation(doc, title, author_and_desc):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    bold_run = p.add_run(title)
    set_font(bold_run, size=11, bold=True)
    rest_run = p.add_run(" " + author_and_desc)
    set_font(rest_run, size=11)
    return p


def add_source(doc, text):
    """Hanging indent: first line flush, continuation indented."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10)
    return p


def add_page_break(doc):
    doc.add_page_break()


def set_default_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def build_document():
    doc = Document()
    set_default_font(doc)

    # ------------------------------------------------------------------ #
    # Title Page
    # ------------------------------------------------------------------ #
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(6)
    r = p_title.add_run("The Hard Thing About Hard Things")
    set_font(r, size=18, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(4)
    r2 = p_sub.add_run("by Ben Horowitz")
    set_font(r2, size=13, italic=True)

    for meta_line in ["Book", "March 18, 2026", "Full detail \u2014 nonfiction, no spoiler concerns."]:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_after = Pt(3)
        rm = pm.add_run(meta_line)
        set_font(rm, size=11)

    add_page_break(doc)

    # ------------------------------------------------------------------ #
    # Section 1: Quick Take
    # ------------------------------------------------------------------ #
    add_heading2(doc, "1. Quick Take")
    add_body(
        doc,
        "One of the most honest books about leadership ever written. Ben Horowitz doesn\u2019t tell you "
        "how to succeed \u2014 he tells you what it feels like to almost fail, repeatedly, and what he did "
        "about it. Built from his experience co-founding Loudcloud in 1999, nearly going bankrupt in the "
        "dot-com crash, pivoting to Opsware, and eventually selling to HP for $1.6 billion. The book "
        "originated as blog posts that went viral in Silicon Valley, and the raw, unpolished voice survived "
        "the transition to print. If most business books are written by professors studying war from a "
        "distance, this one is written by a soldier still covered in mud. Goodreads 4.2/5 across 110,000+ "
        "ratings. NYT bestseller. Still actively recommended 12 years after publication \u2014 it\u2019s "
        "crossed from \u201cstartup book\u201d to management canon. Lenny Rachitsky called it \u201cperhaps "
        "the single most-cited source of inspiration for practicing CEOs.\u201d"
    )

    # ------------------------------------------------------------------ #
    # Section 2: About the Author
    # ------------------------------------------------------------------ #
    add_heading2(doc, "2. About the Author")
    add_body(
        doc,
        "Ben Horowitz was born in London and raised in Berkeley. He holds a B.S. in Computer Science from "
        "Columbia and an M.S. from UCLA. He ran product at Netscape before co-founding Loudcloud with Marc "
        "Andreessen in 1999. The company went public with $2 million in revenue \u2014 an event Horowitz "
        "himself called \u201cthe IPO from hell.\u201d After surviving the dot-com crash, he pivoted to "
        "Opsware and sold it to HP for $1.6 billion in 2007."
    )
    add_body(
        doc,
        "In 2009, he co-founded Andreessen Horowitz (a16z), now managing $46 billion+ in committed capital "
        "and one of the most influential venture firms in the world. His father is conservative writer David "
        "Horowitz. Ben\u2019s Berkeley upbringing, hip-hop fandom, and political independence make him an "
        "unusual figure in venture capital. The book fits his arc as a practitioner-turned-investor who "
        "wanted to share what no one told him when he was in the trenches."
    )
    add_body(
        doc,
        "His second book, \u201cWhat You Do Is Who You Are\u201d (2019), addresses building company culture "
        "through the lens of Toussaint Louverture, samurai bushido, and prison gang leadership. Two mentors "
        "shaped his management philosophy above all others: Andy Grove, Intel\u2019s legendary CEO and author "
        "of \u201cHigh Output Management,\u201d and Bill Campbell, known as \u201cThe Coach of Silicon "
        "Valley,\u201d who also mentored Steve Jobs, Larry Page, and Jeff Bezos."
    )

    # ------------------------------------------------------------------ #
    # Section 3: What the Critics Say
    # ------------------------------------------------------------------ #
    add_heading2(doc, "3. What the Critics Say")
    add_body(
        doc,
        "The critical reception was strong. The New York Times called the honesty \u201cboth refreshing and "
        "compelling.\u201d Fortune declared it essential reading for CEOs. The Economist praised it for "
        "\u201can insider\u2019s perspective on what it\u2019s like to lead and scale a startup.\u201d The "
        "book was longlisted for the Financial Times Business Book of the Year 2014 and named to Business "
        "Insider\u2019s and Entrepreneur\u2019s best business books of the year. On Goodreads: 4.2/5, with "
        "48% five-star ratings and 31% four-star. Goodreads Choice Award nominee."
    )
    add_body(
        doc,
        "Criticisms are real and worth knowing. First, the structure is uneven \u2014 half memoir, half "
        "reformatted blog posts. The blog-post origins mean it sometimes reads as a greatest-hits "
        "compilation rather than a unified narrative. Second, the advice is tuned tightly to VC-backed tech "
        "startups; running a school district or a manufacturing plant is a different kind of hard. Third, "
        "some found Horowitz\u2019s tone self-congratulatory in places. Fourth, readers already familiar "
        "with the blog found the book added little new material."
    )
    add_body(
        doc,
        "The structural criticism is fair but beside the point for most readers. The war stories carry the "
        "book regardless of how it\u2019s organized."
    )

    # ------------------------------------------------------------------ #
    # Section 4: Themes and What Makes It Work
    # ------------------------------------------------------------------ #
    add_heading2(doc, "4. Themes and What Makes It Work")

    add_heading3(doc, "The Struggle")
    add_body(
        doc,
        "The book\u2019s emotional core. Horowitz describes it directly, and the passage is worth quoting "
        "in full:"
    )
    add_blockquote(
        doc,
        "\u201cThe Struggle is when you wonder why you started the company in the first place. The Struggle "
        "is when you go on vacation to feel better and you feel worse. The Struggle is when you are "
        "surrounded by people and you are all alone. The Struggle has no mercy.\u201d"
    )
    add_body(
        doc,
        "Horowitz says managing his own psychology was \u201cby far the most difficult skill I learned as "
        "CEO\u201d \u2014 harder than organizational design, hiring, or metrics. His core advice: \u201cLife "
        "is struggle. Embrace the struggle.\u201d This passage resonated because business books almost never "
        "discuss the emotional toll of leadership with this level of raw honesty. Most management literature "
        "presents leadership as a series of strategic frameworks. Horowitz presents it as a psychological "
        "survival exercise."
    )

    add_heading3(doc, "Wartime CEO vs. Peacetime CEO")
    add_body(
        doc,
        "His most famous framework. The peacetime CEO \u201cknows that proper protocol leads to winning,\u201d "
        "focuses on the big picture, empowers decentralized decisions, and deliberately builds culture. The "
        "wartime CEO \u201cviolates protocol in order to win,\u201d cares about details down to \u201ca speck "
        "of dust on a gnat\u2019s ass,\u201d lets the war define the culture, and \u201cthinks the competition "
        "is sneaking into her house and trying to kidnap her children.\u201d"
    )
    add_body(
        doc,
        "The key insight: most management advice assumes peacetime, but startups spend much of their time at "
        "war. The framework gave language to something founders already felt but couldn\u2019t articulate. It "
        "also provides cover for aggressive, top-down management during crises \u2014 which is both its "
        "practical strength and the source of its criticism. Applied indiscriminately, it can justify "
        "authoritarian leadership well past the crisis that warranted it."
    )

    add_heading3(doc, "Lead Bullets, Not Silver Bullets")
    add_body(
        doc,
        "When Opsware was losing market share to competitor BladeLogic, Horowitz\u2019s team kept proposing "
        "workarounds \u2014 partnerships, acquisitions, repositioning. His conclusion: \u201cThere are no "
        "silver bullets for this, only lead bullets.\u201d They had to build a better product. No shortcut, "
        "no hack, no clever pivot. This counters the \u201cgrowth hack\u201d mentality prevalent in startup "
        "culture and delivers a fundamentally conservative, craft-oriented message: do the hard work of "
        "building something better."
    )

    add_heading3(doc, "The Right Kind of Ambition")
    add_body(
        doc,
        "Horowitz defines the right kind of ambition as \u201cambition for the company\u2019s success, with "
        "the executive\u2019s own success only coming as a by-product of the company\u2019s victory.\u201d "
        "The wrong kind: ambition for personal success regardless of company outcomes. Hiring people with "
        "wrong-kind ambition is, in his words, \u201cthe surest way to turn your company into the political "
        "equivalent of the U.S. Senate.\u201d"
    )

    add_heading3(doc, "Firing and Demoting")
    add_body(
        doc,
        "The sections on firing executives and demoting friends are where the book earns its title. His "
        "mentor Bill Campbell\u2019s line: \u201cYou cannot let him keep his job, but you absolutely can let "
        "him keep his respect.\u201d Horowitz gives specific, step-by-step guidance on how to do both "
        "\u2014 most business books skip this entirely or reduce it to platitudes. Demoting loyal friends "
        "who can\u2019t scale with the company is acknowledged as one of the hardest things a CEO does, "
        "and he treats it as such."
    )

    add_heading3(doc, "Hip-Hop as Management Philosophy")
    add_body(
        doc,
        "Every chapter opens with hip-hop lyrics \u2014 Nas, Kanye, Jay-Z, DMX. Chapter 1 opens with "
        "DMX\u2019s \u201cWho We Be.\u201d Horowitz told Fortune: \u201cThe confrontations and the "
        "conversations that I write about in the book are hard. The lyrics I use express the emotional "
        "intensity that goes with the logic in the book.\u201d He sees a parallel between rappers who "
        "bootstrapped careers from nothing and entrepreneurs building companies. This isn\u2019t "
        "decorative \u2014 it signals a different kind of business book, grounded in struggle rather than "
        "Harvard case studies."
    )

    # ------------------------------------------------------------------ #
    # Section 5: What You Might Not Know
    # ------------------------------------------------------------------ #
    add_heading2(doc, "5. What You Might Not Know")
    add_body(
        doc,
        "The Loudcloud survival story is genuinely dramatic. The company was valued at $720 million in 2000. "
        "After the dot-com crash, the stock hit $0.35 per share \u2014 a $30 million market cap, roughly "
        "half the company\u2019s remaining cash on hand. The market was pricing in bankruptcy. With 440 of "
        "450 employees in the cloud business, Horowitz sold that unit to EDS for $63.5 million, kept the "
        "automation software (Opsware), and rebuilt from near-nothing to a $1.6 billion exit. That sequence "
        "\u2014 near-death, radical pivot, billion-dollar outcome \u2014 is not a metaphor. It happened."
    )
    add_body(
        doc,
        "The book originated from blog posts at bhorowitz.com that attracted nearly 10 million readers before "
        "publication. The blog-to-book pipeline meant the material was battle-tested with a real audience "
        "before it was ever edited for print. Some posts weren\u2019t included in the book and remain "
        "blog-only \u2014 worth reading if you want more than the published version contains."
    )
    add_body(
        doc,
        "Andy Grove made \u201cHigh Output Management\u201d mandatory reading at Intel. Horowitz made it "
        "mandatory for every manager at his companies and later wrote the foreword for the reissue. He called "
        "it \u201cthe only management book I ever read that I liked.\u201d The management lineage is "
        "direct: Grove mentored Campbell, Campbell mentored Horowitz, and all three shaped a generation of "
        "Silicon Valley leaders."
    )
    add_body(
        doc,
        "Bill Campbell mentored not just Horowitz but Steve Jobs, Larry Page, Eric Schmidt, and Jeff Bezos. "
        "His story is told in \u201cTrilliondollar Coach\u201d by Eric Schmidt, Jonathan Rosenberg, and "
        "Alan Eagle \u2014 a book that would not exist without the relationships Horowitz describes here."
    )
    add_body(
        doc,
        "Complex ran a feature cataloging Horowitz\u2019s best rap references. The Mercury News ran a "
        "profile headlined \u201cHip-Hop Venture Capitalist.\u201d His second book uses Toussaint "
        "Louverture\u2019s slave revolution, Genghis Khan, and a prison gang leader as management case "
        "studies. He is not playing it safe."
    )

    # ------------------------------------------------------------------ #
    # Section 6: On Screen
    # ------------------------------------------------------------------ #
    add_heading2(doc, "6. On Screen")
    add_body(
        doc,
        "No adaptations exist. No documentary, film, or television treatment of Horowitz or the "
        "Loudcloud/Opsware story has been produced. This is notable given how cinematic the arc is \u2014 "
        "an IPO during the dot-com bust, near-bankruptcy, a radical pivot, and a billion-dollar exit. The "
        "Loudcloud survival sequence has more narrative tension than most business stories that did get "
        "adapted (Theranos, Uber, WeWork). If a producer is looking for the next \u201cAir\u201d or "
        "\u201cTetris,\u201d this story is sitting unclaimed."
    )

    # ------------------------------------------------------------------ #
    # Section 7: If You Liked This
    # ------------------------------------------------------------------ #
    add_heading2(doc, "7. If You Liked This")
    add_body(
        doc,
        "These are the books most likely to reward you if this one did."
    )

    add_recommendation(
        doc,
        '"High Output Management" by Andy Grove',
        "\u2014 The intellectual foundation for Horowitz\u2019s management thinking. Grove was his mentor; "
        "this was mandatory reading at all his companies. If you want the framework behind the war stories, "
        "start here. Horowitz wrote the foreword for the reissue."
    )
    add_recommendation(
        doc,
        '"Shoe Dog" by Phil Knight',
        "\u2014 Nike\u2019s founder memoir. Same DNA: a founder who nearly went bankrupt repeatedly, fought "
        "dirty battles, and tells it without sugarcoating. Bill Gates and Warren Buffett both named it a "
        "favorite. The honesty is comparable; the industry is completely different."
    )
    add_recommendation(
        doc,
        '"Only the Paranoid Survive" by Andy Grove',
        "\u2014 Grove on \u201cStrategic Inflection Points,\u201d the moments when a company must pivot or "
        "die. Horowitz lived through exactly this with Loudcloud\u2019s transformation into Opsware. "
        "Reading both books together makes each one richer."
    )
    add_recommendation(
        doc,
        '"Creativity, Inc." by Ed Catmull',
        "\u2014 Pixar\u2019s co-founder on building a culture where honest feedback survives hierarchy. Less "
        "about crisis, more about sustained creative leadership, but equally candid about organizational "
        "dysfunction and the ego traps that kill good work."
    )
    add_recommendation(
        doc,
        '"Bad Blood" by John Carreyrou',
        "\u2014 The Theranos investigation. If Horowitz\u2019s book is the honest founder\u2019s memoir, "
        "this is the cautionary tale of what happens when a founder chooses deception over transparency. "
        "A dark mirror that makes the values in \u201cHard Thing\u201d clearer by contrast."
    )
    add_recommendation(
        doc,
        '"The Founder\'s Dilemmas" by Noam Wasserman',
        "\u2014 Research-based examination of the early decisions (co-founder splits, hiring, equity) that "
        "make or break startups. The academic complement to Horowitz\u2019s street-level perspective. "
        "Better together than either is alone."
    )
    add_recommendation(
        doc,
        '"What You Do Is Who You Are" by Ben Horowitz',
        "\u2014 His follow-up on building company culture using Toussaint Louverture, Genghis Khan, and "
        "Shaka Senghor as case studies. If \u201cHard Thing\u201d is about surviving as a CEO, this is "
        "about what you build once you survive."
    )

    add_body(doc, "Also worth your time:")

    for item in [
        "The original blog posts at a16z.com \u2014 especially \u201cThe Struggle,\u201d \u201cPeacetime "
        "CEO/Wartime CEO,\u201d and \u201cLead Bullets.\u201d The a16z \u201cReading List for Leaders in "
        "Uncertain Times\u201d collects the key essays.",
        "Stanford eCorner talk: \u201cNailing the Hard Things\u201d with Prof. Tom Byers.",
        "Tim Ferriss Show #392: Deep interview covering Andy Grove\u2019s influence and management philosophy.",
        'a16z Podcast: "Wartime vs. Peacetime" \u2014 dedicated episode expanding the framework.',
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_font(run, size=11)

    # ------------------------------------------------------------------ #
    # Section 8: Sources
    # ------------------------------------------------------------------ #
    add_heading2(doc, "8. Sources")

    sources = [
        "Andreessen Horowitz. \u201cBen Horowitz\u2019s Blog, from Book to Ebook.\u201d a16z.com",
        "Andreessen Horowitz. \u201cPeacetime CEO/Wartime CEO.\u201d a16z.com",
        "Andreessen Horowitz. \u201cPreparing to Fire an Executive.\u201d a16z.com",
        "Andreessen Horowitz. \u201cReading List for Leaders in Uncertain Times.\u201d a16z.com",
        "Andreessen Horowitz. \u201cThe Hard Thing About Hard Things.\u201d a16z.com",
        "Aure\u2019s Notes. \u201cSummary: The Hard Thing About Hard Things.\u201d",
        "Complex (2013). \u201cA History of Ben Horowitz\u2019s Best Rap References.\u201d",
        "Dream10x. \u201cHard Thing About Hard Things.\u201d",
        "Farnam Street. \u201cBen Horowitz: The Struggle.\u201d",
        "Farnam Street. \u201cThe Hard Thing About Hard Things.\u201d",
        "Fortune (2014). \u201cBen Horowitz schools you on hip-hop.\u201d",
        "Fortune (2019). \u201cQ\u0026A: Ben Horowitz on His New Book.\u201d",
        "Goodreads. \u201cThe Hard Thing About Hard Things.\u201d",
        "Lenny\u2019s Newsletter (2025). \u201c$46B of hard truths from Ben Horowitz.\u201d",
        "Mercury Blog. \u201cBen Horowitz.\u201d",
        "Mercury News (2014). \u201cHip-Hop Venture Capitalist.\u201d",
        "Rise with Drew (2025). \u201cFrom Death Spiral to $63.5 Million Deal.\u201d",
        "Ryan Craggs. \u201cThe Hard Thing About Hard Things Review.\u201d",
        "SuperSummary. \u201cThe Hard Thing About Hard Things Study Guide.\u201d",
        "The Psmiths. \u201cReview: The Hard Thing About Hard Things.\u201d",
        "Teamly. \u201cThe Hard Thing About Hard Things Chapter 6.\u201d",
        "Thrive Street Advisors. \u201cThe Hard Thing About Hard Things.\u201d",
        "Time (2015). \u201cThe Struggle, According to Ben Horowitz.\u201d",
        "Tyler DeVries. \u201cBook Summary: The Hard Thing About Hard Things.\u201d",
        "Wikipedia. \u201cBen Horowitz.\u201d",
    ]

    for i, source in enumerate(sources, 1):
        add_source(doc, f"{i}. {source}")

    return doc


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
