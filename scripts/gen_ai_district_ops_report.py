#!/usr/bin/env python3
"""
Deep Research Report Generator: AI Best Practices for School District Operations
Generated: 2026-03-18
"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Paths ──────────────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.expanduser("~/Documents/Research/ai-best-practices-district-operations")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, f"{date.today().isoformat()} AI Best Practices District Operations Deep Research Report.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_confidence_indicator(paragraph, level):
    indicators = {
        "high": ("High confidence", (34, 139, 34)),
        "medium": ("Medium confidence", (218, 165, 32)),
        "low": ("Low confidence", (178, 34, 34)),
    }
    label, color = indicators.get(level.lower(), ("Unknown", (128, 128, 128)))
    run = paragraph.add_run(f" [{label}]")
    set_run_font(run, size=10, italic=True, color=color)

def add_table_with_header(doc, headers, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "2F5496")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_row(table, values, shading=None):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(v))
        set_run_font(run, size=9)
        if shading:
            set_cell_shading(cell, shading)
    return row


# ── Build Document ─────────────────────────────────────────────────────────

doc = Document()

# -- Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# -- Heading styles
for level, size, color in [(1, 16, (47, 84, 150)), (2, 14, (47, 84, 150)), (3, 12, (47, 84, 150))]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(*color)
    hs.font.bold = True

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Deep Research Report")
set_run_font(run, size=28, bold=True, color=(47, 84, 150))

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("AI Best Practices for School District Operations")
set_run_font(run, size=20, color=(47, 84, 150))

doc.add_paragraph()

domains_p = doc.add_paragraph()
domains_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = domains_p.add_run("Finance | Government | Human Resources | Nursing | Nutrition\nTransportation | IT | Maintenance | Cleaning")
set_run_font(run, size=12, italic=True, color=(89, 89, 89))

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"March 18, 2026")
set_run_font(run, size=12, color=(89, 89, 89))

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run("4 Source Categories | 58 Findings | 64 Sources | 9 Operational Domains")
set_run_font(run, size=11, color=(89, 89, 89))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("Executive Summary", level=1)

exec_paras = [
    "This report synthesizes findings from 64 sources across four research categories (government, academic, news/journalism, and industry/practitioner) to assess the current state of AI adoption and best practices across nine school district operational domains: finance, government/administration, human resources, school nursing, school nutrition, school transportation, school IT/cybersecurity, school maintenance, and school cleaning/custodial.",

    "The research reveals a field in rapid but uneven transition. The number of K-12 districts identified as AI early adopters nearly doubled from 40 to 79 between the 2024-25 and 2025-26 school years, yet most remain in early, fragmented stages of experimentation. Only 18% of early adopter districts use AI for back-office operations, and just 14% of U.S. school districts had any AI policy as of May 2024. Critically, only 1% of districts report their data as fully prepared and secure for AI use.",

    "Key findings by domain:",
]

for text in exec_paras:
    add_styled_paragraph(doc, text)

# Executive summary bullets
exec_bullets = [
    ("Transportation", "is the most mature operational AI domain, with peer-reviewed evidence showing 7-16% efficiency gains and $5M+ annual savings in large districts (Boston, Denver, Colorado Springs)."),
    ("IT/Cybersecurity", "is the most urgent, with 82% of schools experiencing cyber incidents and only 7% confident their AI aligns with cybersecurity best practices. The FCC's $200M pilot program and CISA toolkit provide new federal resources."),
    ("Finance", "shows strong promise: districts using AI for forecasting report 57% 'very accurate' results vs. 8% without AI. Post-ESSER budget pressures are accelerating adoption."),
    ("Nutrition", "has emerging AI tools (LunchLens, Rayfood) for USDA compliance and meal planning, with the USDA framework now explicitly allowing AI with human-in-the-loop review."),
    ("Maintenance/Energy", "can achieve 8-12% energy reductions within 60 days with AI-powered building management, though California research shows actual savings often fall short of engineering estimates."),
    ("Cleaning/Custodial", "is nascent, with Spokane's $1.08M autonomous floor-cleaning robot fleet being the most prominent K-12 example."),
    ("HR", "faces a tension between efficiency gains (50+ hours saved per campus on scheduling) and algorithmic bias risks (AI preferred white-associated names 85% of the time in resume screening)."),
    ("Nursing/Health", "is the least developed domain, with no AI-specific implementations found. AI mental health chatbots (serving 210,000+ students) represent the closest adjacent innovation."),
    ("Government/Policy", "is the enabling layer: the White House EO, USED guidance on federal funds for AI, Tennessee's HB1630 mandate, and NIST AI Risk Management Framework establish the governance infrastructure."),
]

for domain, text in exec_bullets:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(domain)
    set_run_font(run, bold=True)
    run2 = p.add_run(f" {text}")
    set_run_font(run2, size=11)

doc.add_paragraph()
add_styled_paragraph(doc, "Bottom line: School districts should prioritize AI governance frameworks, data readiness, and cybersecurity before expanding operational AI. Transportation, finance, and energy management offer the clearest near-term ROI. Districts must address the equity gap, as schools in higher-poverty areas are significantly less likely to receive AI guidance or resources.", bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. BACKGROUND & CONTEXT
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Background and Context", level=1)

bg_paras = [
    "Artificial intelligence is transforming K-12 education at an accelerating pace, but the conversation has been overwhelmingly focused on classroom instruction and student-facing applications. Behind the scenes, school districts operate complex enterprises spanning finance, human resources, transportation, food service, healthcare, facilities management, and technology infrastructure. These operational functions collectively consume the majority of district budgets and directly impact the conditions under which teaching and learning occur.",

    "The federal policy landscape shifted dramatically in 2025. President Trump's April 2025 executive order established a White House Task Force on AI Education, and the U.S. Department of Education subsequently issued guidance affirming that federal formula and discretionary grant funds may be used for AI integration, including operational applications. Tennessee became one of only two states (alongside Ohio) to mandate district-level AI policies through HB1630, requiring compliance by the 2024-25 school year.",

    "Several converging pressures make operational AI adoption urgent for school districts: the expiration of $190 billion in ESSER pandemic relief funds, chronic staffing shortages across nearly every operational area (90%+ of nutrition directors report shortages; the national counselor-to-student ratio of 1:376 far exceeds the recommended 1:250), escalating cybersecurity threats (82% of schools experienced incidents in an 18-month period), and rising energy costs ($8 billion annually, the second-largest expense after teacher salaries).",

    "This report examines best practices, evidence, and emerging models across nine operational domains. It draws on government sources (executive orders, federal guidance, state legislation), peer-reviewed academic research, news and journalism reporting, and industry/practitioner publications. Each finding is rated for confidence based on source quality, corroboration, and evidence strength.",
]

for text in bg_paras:
    add_styled_paragraph(doc, text)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. KEY FINDINGS BY DOMAIN
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Key Findings", level=1)

add_styled_paragraph(doc, "Findings are organized by operational domain, with evidence woven across all four source categories. Confidence indicators reflect cross-source corroboration.", italic=True, size=10)

# ── 2.1 GOVERNANCE & POLICY ──────────────────────────────────────────────

doc.add_heading("2.1 Governance and Policy Framework", level=2)

findings_governance = [
    ("Federal AI policy now explicitly supports K-12 operational AI adoption.",
     "The White House Executive Order on AI Education (April 2025) established a Task Force directing federal agencies to promote AI integration. The USED Dear Colleague Letter (July 2025) affirmed that federal grant funds may be used for AI-enhanced operations, including reducing administrative burdens and modernizing operations. OMB memoranda M-24-18 and M-25-21 establish procurement standards for AI that inform public-sector purchasing, including requirements to prevent vendor lock-in and document model training.",
     "high",
     "White House (2025); USED (2025); OMB (2024, 2025)"),

    ("Only 14% of U.S. school districts had AI policies as of May 2024, but mandates are emerging.",
     "A systematic content analysis of 9,229 K-12 districts found that only 14.13% had an AI policy. Tennessee (HB1630, 2024) and Ohio (HB 96, by July 2026) are the only states with legal mandates. The CRPE tracker shows early adopter districts doubled from 40 to 79 in one year, yet 45% of teachers and administrators say their district still lacks any AI policy. The Region 8 Comprehensive Center published federally funded guidance for K-12 AI policy development covering both instructional and operational uses.",
     "high",
     "JRTE (2025); TN General Assembly (2024); Ohio DOE; CRPE (2025); Region 8 (2024)"),

    ("NIST AI Risk Management Framework is the emerging standard for K-12 AI governance.",
     "NIST's AI RMF 1.0 (2023) with four core functions (GOVERN, MAP, MEASURE, MANAGE) and the GenAI Profile (AI 600-1, 2024) provide the most comprehensive voluntary governance framework. Washington State's OSPI has adopted NIST AI RMF as a foundation for K-12 guidance. AASA recommends districts form AI steering committees with continuously updated frameworks rather than static policies.",
     "high",
     "NIST (2023, 2024); AASA (2025)"),

    ("Civil rights and equity concerns require proactive attention in AI adoption.",
     "The U.S. Commission on Civil Rights identified four primary concerns: disruption of student-teacher relationships, promotion of bias, widening of the digital divide, and student data privacy/surveillance risks. RAND data shows principals in highest-poverty schools are about half as likely as those in lowest-poverty schools to receive AI guidance (13% vs 25%). Only 19% of students report receiving guidance on AI use.",
     "high",
     "USCCR (2024); RAND (2025)"),
]

for title, evidence, confidence, sources in findings_governance:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.2 FINANCE ──────────────────────────────────────────────────────────

doc.add_heading("2.2 Finance", level=2)

findings_finance = [
    ("AI-powered financial forecasting dramatically improves budget accuracy.",
     "Among districts using AI for financial forecasting, 57% describe their forecast as 'very accurate,' compared with fewer than 8% of districts not using AI. More than half of finance leaders want to use AI for budgeting and monitoring. Val Verde Unified (CA) used Microsoft Copilot to cut a state reporting coding task from 10 to 5 hours. Tyler Technologies' Priority Based Budgeting uses AI/ML for scenario modeling and predictive analysis, with LA County selecting it for their $40 billion budget modernization.",
     "high",
     "Frontline Education (2026); EdTech Magazine (2025); Tyler Technologies (2025)"),

    ("AI is reshaping K-12 procurement processes.",
     "Over 850 school districts collectively spent $5M+ on AI tools in six months. Jordan School District (UT, 56,000 students) uses AI chatbots to draft RFPs, reducing drafting time from hours to minutes. California now requires AI disclosure forms in government RFPs. OMB procurement standards (M-24-18) establish best practices including preventing vendor lock-in and requiring documentation of model training.",
     "medium",
     "GovTech (2025); OMB (2024)"),

    ("Post-ESSER financial pressures are accelerating operational AI adoption.",
     "McKinsey's survey of 300+ district administrators (July 2025) shows districts bracing for leaner years post-ESSER, with AI and automation as potential efficiency tools. McKinsey estimates automating educational content creation alone could save approximately $60 billion globally. K-12 districts spend nearly $8 billion annually on energy alone, making AI-driven efficiency a financial imperative.",
     "medium",
     "McKinsey (2025); EdTech Magazine (2025)"),
]

for title, evidence, confidence, sources in findings_finance:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.3 HUMAN RESOURCES ─────────────────────────────────────────────────

doc.add_heading("2.3 Human Resources", level=2)

findings_hr = [
    ("AI scheduling tools deliver measurable time and cost savings.",
     "Austin ISD used Timely AI to build master schedules, saving schedulers at least 50 hours on average and approximately $2,600 per campus. The district piloted in select schools during 2024-25 and went district-wide in 2025-26 across 33 middle and high school schedules. 80% of school budgets are determined by master schedule decisions, making this a high-leverage application.",
     "high",
     "Education Week (2025)"),

    ("AI hiring tools show severe racial and gender bias that must be mitigated.",
     "University of Washington research found AI hiring systems preferred white-associated names 85% of the time vs. Black-associated names 9%, and male-associated names 52% vs. female-associated 11%, including teacher positions. K-12 HR leaders warn about AI perpetuating hiring bias from historically limited candidate pools. About 25% of companies globally now use AI for hiring (up from 12% in 2023), and K-12 districts are beginning to follow.",
     "high",
     "University of Washington (2024); Education Week (2025)"),

    ("Districts are cautiously adopting AI for recruitment workflow automation.",
     "62% of schools report 'too few candidates applying' as a top hiring challenge (NCES School Pulse Panel, August 2024). Districts like Deerfield Public Schools District 109 (IL) use AI for drafting job descriptions and interview questions but keep candidate interactions human-only. 75 AASPA members completed an AI microcredential in March 2025. Ednovate Schools (CA) created detailed process maps for AI in HR operations.",
     "medium",
     "Education Week (2025); NCES (2024); CRPE (2024)"),
]

for title, evidence, confidence, sources in findings_hr:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.4 TRANSPORTATION ──────────────────────────────────────────────────

doc.add_heading("2.4 Transportation", level=2)

findings_transport = [
    ("Algorithm-optimized bus routing delivers proven, large-scale cost savings.",
     "MIT's BiRD algorithm for Boston Public Schools reduced the bus fleet by 50 buses (7%), saving approximately $5 million annually. The algorithm completes in 30 minutes work that previously took 10 people several weeks (3,000+ hours). A systematic review found ML approaches achieve average travel time reduction of 15.7%, distance reduction of 10.7%, and cost reduction of 12.4% compared to traditional methods.",
     "high",
     "Bertsimas et al. (2020), INFORMS; IJFMR (2022)"),

    ("AI route optimization is scaling to more districts with multimodal planning.",
     "Denver Public Schools (90,000+ students, 200 buses) implemented RouteWise AI in 2024-25 and successfully absorbed demand from 10 school closures without adding buses. Colorado Springs increased highly-utilized routes by 46% and eliminated its bus driver shortage through AI optimization alone. Systems run millions of simulations analyzing bell times, student data, and assets to create multimodal transportation plans.",
     "high",
     "GovTech (2025); HopSkipDrive"),

    ("AI-powered telematics and driver monitoring are becoming standard.",
     "Telematics with AI-powered driver monitoring has moved from 'nice-to-have' to 'must-have' for school buses. AI dash cams provide real-time alerts and driver coaching. The National Congress on School Transportation (May 2025) is expected to produce new specifications and best-practice recommendations for the first time in 10 years.",
     "medium",
     "School Bus Fleet (2025); NAPT"),
]

for title, evidence, confidence, sources in findings_transport:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.5 IT / CYBERSECURITY ──────────────────────────────────────────────

doc.add_heading("2.5 IT and Cybersecurity", level=2)

findings_it = [
    ("K-12 cybersecurity is in crisis, with 82% of schools experiencing incidents.",
     "CIS/MS-ISAC data shows 82% of K-12 schools were impacted by cyberthreats during July 2023-December 2024, with 9,300 confirmed incidents from 5,000+ organizations. RAND found 60% of schools reported at least one incident, with 45% experiencing compromised business emails, 14% data breaches, and 10% ransomware. GAO found cyberattacks cause learning loss of 3 days to 3 weeks and recovery times of 2 to 9 months.",
     "high",
     "CIS MS-ISAC (2025); RAND (2024); GAO (2022, 2023)"),

    ("AI is the top operational use case for school IT departments.",
     "CoSN's survey of 281 districts found 57% use AI in network environments, with security threat detection as the top use case (65%), followed by automated document generation and anomaly detection. However, only 7% are confident their AI aligns with cybersecurity best practices, and only 1% report data fully prepared for AI use. 55% of school cyber incidents come from third-party vendors.",
     "high",
     "CoSN (2025); CIS MS-ISAC (2025)"),

    ("New federal resources support K-12 cybersecurity infrastructure.",
     "The FCC established a $200 million, 3-year cybersecurity pilot program selecting 700+ schools for equipment and services. CISA published specific K-12 recommendations: use no/low-cost services, migrate to secure cloud, implement MFA, and adopt Cybersecurity Performance Goals. Pasadena ISD (TX) reduced cyber insurance costs by 40% through comprehensive self-assessment.",
     "high",
     "FCC (2024); CISA (2023); EdSurge (2025)"),
]

for title, evidence, confidence, sources in findings_it:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.6 NUTRITION ────────────────────────────────────────────────────────

doc.add_heading("2.6 Nutrition", level=2)

findings_nutrition = [
    ("USDA now explicitly allows AI for school meal planning with human oversight.",
     "The USDA Food and Nutrition Service published a framework allowing AI to create draft meal plans but requiring human review for USDA compliance. Only USDA-approved nutrient analysis software may be used for Administrative Reviews. Agencies remain accountable for federal compliance regardless of AI use. This 'human-in-the-loop' requirement establishes the governance model for school nutrition AI.",
     "high",
     "USDA FNS (n.d.)"),

    ("AI-powered meal compliance tools are gaining traction.",
     "LunchLens (using GPT-4 Vision) reports saving cafeteria staff an average of 2.5 hours per day in manual tray inspections, with 500+ districts using the tool. Schools report passing administrative reviews with zero findings for the first time. A peer-reviewed case study documents AI implementation for school meal planning using real-time supplier nutritional data (Rayfood/Frontiers in Sustainable Food Systems).",
     "medium",
     "LunchLens (2025); Cummins et al. (2022)"),

    ("Data-driven approaches can significantly reduce school food waste.",
     "Analysis of 134 schools across 24 states found schools allowing student self-selection of milk portions had 76% less waste. ML food demand forecasting models can reduce wasted meals by 14-52%. Average waste is 0.29 lbs per student per meal. 90%+ of food service directors report staffing shortages, driving automation interest.",
     "high",
     "Adjapong et al. (2024), PLOS ONE; Journal of Cleaner Production (2023); SNA (2024)"),
]

for title, evidence, confidence, sources in findings_nutrition:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.7 MAINTENANCE & ENERGY ────────────────────────────────────────────

doc.add_heading("2.7 Maintenance and Energy Management", level=2)

findings_maintenance = [
    ("AI-powered energy management delivers measurable savings for school facilities.",
     "Districts implementing advanced energy management with AI report 15-30% energy savings (industry estimates). Saline Area Schools (MI) uses Schneider Electric for AI diagnostics, CO2 sensors, and predictive maintenance. Phase 1 deployments typically deliver 8-12% energy reduction within 60 days. K-12 districts spend nearly $8 billion annually on energy. The DOE promotes Energy Management Information Systems and provides 35+ hours of free energy management training through Renew America's Schools.",
     "medium",
     "EdTech Magazine (2025); DOE; Intellis (2025)"),

    ("Actual energy savings from efficiency upgrades often fall short of estimates.",
     "A rigorous NBER/JAERE study of California K-12 schools found energy efficiency upgrades reduced electricity consumption by only 3%, representing just 24% of expected savings from engineering estimates. ML analysis produced higher realization rates (52-98%, central 60%) but still modest. School characteristics readily available to policymakers did not predict realization rates, suggesting caution in projecting AI-driven savings.",
     "high",
     "Burlig et al. (2020), JAERE"),

    ("Predictive maintenance AI can detect HVAC failures before they occur.",
     "An LSTM-based autoencoder framework detected HVAC system failures two days before they occurred. Faulty building operations lead to 20-30% increases in total building energy consumption. Structured predictive maintenance can reduce emergency HVAC costs by 30-40% annually (industry estimates). AI is described as moving from 'optional to essential' for K-12 facilities capital planning.",
     "medium",
     "Bouabdallaoui et al. (2021), Sensors; K12FF (2025)"),
]

for title, evidence, confidence, sources in findings_maintenance:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.8 NURSING / HEALTH ────────────────────────────────────────────────

doc.add_heading("2.8 Nursing and Student Health", level=2)

findings_nursing = [
    ("School nursing AI is essentially nonexistent, representing a major gap.",
     "No government, academic, news, or industry sources identified AI applications specifically for school nursing operations (health screening, medication tracking, immunization records, student health management). School nurses face major barriers including high caseloads (1 nurse per 1,500 students), lack of data access, and poor communication between school and community healthcare providers. This is the least developed of all nine operational domains examined.",
     "high",
     "Rankine et al. (2021), Journal of School Nursing; all four source categories"),

    ("AI mental health tools are the closest adjacent innovation, with rapid adoption.",
     "Lenny Learning serves 210,000+ students in 400+ schools across 19 states (10x growth from start of 2024). Sonny chatbot serves 4,500+ students at $20,000-$30,000/year per district. Nearly 1 in 5 students accessed school-based mental health services in 2024-25. Over 14 million K-12 students attend schools without any counselors, nurses, psychologists, or social workers.",
     "medium",
     "Edutopia (2025)"),

    ("ML can predict student absenteeism with 90% accuracy, supporting nurse-led intervention.",
     "A machine learning classifier achieved 90.2% accuracy in identifying students at risk for chronic absenteeism, with peer relationships as the key predictive factor. While not AI-in-nursing per se, this type of predictive tool could support school nurse-led health interventions for chronically absent students.",
     "medium",
     "Bowen et al. (2022), Frontiers in Psychology"),
]

for title, evidence, confidence, sources in findings_nursing:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

# ── 2.9 CLEANING / CUSTODIAL ────────────────────────────────────────────

doc.add_heading("2.9 Cleaning and Custodial", level=2)

findings_cleaning = [
    ("Autonomous floor-cleaning robots are the leading AI application in K-12 custodial operations.",
     "Spokane Public Schools purchased 14 autonomous floor-cleaning robots from Tennant for $1.08M, deploying one at each middle and high school. Robots logged 1,500+ hours cleaning 16.5+ million sq ft. Floors are now cleaned minimum 3x weekly (vs. once weekly with manual machines). No custodial positions were eliminated; the union confirmed robots 'cannot and will not replace a custodian.' The investment addresses staffing shortages rather than reducing headcount.",
     "high",
     "GovTech/Spokesman-Review (2026)"),

    ("AI for custodial scheduling, supply management, and sanitation protocols remains uncharted.",
     "No published research, government guidance, or practitioner reports address AI for school cleaning scheduling, custodial supply management, or sanitation protocol optimization. EPA provides indoor air quality and green cleaning guidance for schools but nothing AI-specific. This domain is second only to nursing in its lack of AI adoption evidence.",
     "high",
     "All four source categories; EPA"),
]

for title, evidence, confidence, sources in findings_cleaning:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_confidence_indicator(p, confidence)
    add_styled_paragraph(doc, evidence, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Sources: {sources}")
    set_run_font(run2, size=9, italic=True, color=(89, 89, 89))
    p2.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. POINTS OF CONTRADICTION
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Points of Contradiction", level=1)

contradictions = [
    ("Energy savings projections vs. actual results",
     "Industry sources (EdTech Magazine, Intellis, Schneider Electric) claim 15-30% energy savings and 8-12% reductions within 60 days. However, the most rigorous academic study (Burlig et al., NBER/JAERE) of actual California K-12 schools found only 3% electricity reduction, representing just 24% of engineering estimates. ML analysis improved estimates to 52-98% realization rates (central 60%), but this is still well below industry marketing claims.",
     "The academic evidence is stronger (peer-reviewed, large sample, rigorous econometrics). Districts should expect more modest savings than vendor projections suggest and should build conservative assumptions into ROI calculations."),

    ("AI adoption rate estimates vary significantly",
     "RAND reports 54% of students and 53% of teachers use AI for school (2025). CoSN finds 57% use AI in network environments but 72% deploy it for 10% or less of operations. CRPE tracked only 79 'early adopter' districts. The Frontline K-12 Lens suggests 35% use AI for personalized learning. These numbers are not directly comparable (different populations, definitions of 'use') but paint different pictures of adoption maturity.",
     "Adoption is widespread at the individual level (teachers using ChatGPT) but shallow at the institutional/operational level. The CoSN finding that 72% deploy AI for 10% or less of operations is likely the most accurate picture of operational AI maturity."),

    ("AI as efficiency tool vs. AI as equity risk",
     "Government and industry sources emphasize AI's potential to reduce administrative burden and save costs. Academic and civil rights sources highlight algorithmic bias (85% preference for white-associated names), widening digital divides (poverty-related adoption gaps), and surveillance risks. Both perspectives are well-documented.",
     "These are not mutually exclusive. Districts can pursue efficiency gains while implementing bias audits, equity-focused governance, and transparent AI policies. The USCCR and NIST frameworks provide starting points."),
]

for title, detail, assessment in contradictions:
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, bold=True, size=11)
    add_styled_paragraph(doc, detail, size=11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Assessment: {assessment}")
    set_run_font(run2, size=11, italic=True)
    p2.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. SINGLE-SOURCE CLAIMS
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Single-Source Claims", level=1)

add_styled_paragraph(doc, "The following claims are supported by only one source category. They may be accurate but lack cross-category corroboration.", italic=True)

single_claims = [
    ("LunchLens saves 2.5 hours/day per cafeteria", "Industry only (vendor marketing)", "Independent evaluation needed. The claim is plausible given the task automation, but 500+ districts is unverified."),
    ("Predictive maintenance detects HVAC failures 2 days early", "Academic only (single case study, 45 days, one facility)", "Promising but extremely limited sample. Needs replication in actual school settings."),
    ("Spokane robots cleaned 16.5M+ sq ft in 1,500+ hours", "News only (single local report)", "Well-documented by local journalism with named officials and union verification, but no independent performance audit."),
    ("AI mental health chatbots serve 210,000+ students across 19 states", "News only (Edutopia)", "Vendor-reported growth numbers. FTC investigations into some AI chatbot providers (Character.AI, Replika) suggest regulatory uncertainty."),
    ("Jordan SD AI chatbot reduced RFP drafting from hours to minutes", "News only (GovTech)", "Specific named district, plausible claim, but no independent time study."),
    ("Tyler Technologies PBB selected for LA County $40B budget", "Industry only (vendor press release)", "Verifiable contract, but no K-12 implementation data."),
    ("ML predicts student absenteeism with 90.2% accuracy", "Academic only (single school, 332 students)", "Peer-reviewed but very small sample in one Indiana school. Generalizability uncertain."),
]

for claim, source_note, assessment in single_claims:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(claim)
    set_run_font(run, bold=True, size=11)
    run2 = p.add_run(f" ({source_note})")
    set_run_font(run2, size=10, italic=True, color=(89, 89, 89))
    add_styled_paragraph(doc, assessment, size=10, space_after=8)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. PRACTICAL IMPLICATIONS
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Practical Implications", level=1)

doc.add_heading("Recommended Actions (Ordered by Confidence Level)", level=2)

high_conf = [
    "Adopt an AI governance framework immediately. Use NIST AI RMF as a foundation. Tennessee districts must comply with HB1630. Form an AI steering committee with representation from each operational department.",
    "Prioritize cybersecurity before expanding AI. With 82% incident rates and only 7% confident in AI-cybersecurity alignment, security must come first. Apply for FCC pilot program funds, implement CISA recommendations (MFA, secure cloud migration, CPGs), and conduct a comprehensive self-assessment.",
    "Invest in transportation AI for near-term ROI. Bus route optimization is the most evidence-backed operational AI application, with proven 7-16% efficiency gains and $5M+ annual savings in large districts. Evaluate tools like RouteWise, Transfinder, and similar platforms.",
    "Address data readiness as a prerequisite. Only 1% of districts report data fully prepared for AI (CoSN). Clean, integrate, and secure your data before deploying AI tools. 61% of districts struggle with dirty/siloed data.",
    "Audit AI hiring tools for bias before deployment. The 85% preference for white-associated names is well-documented. Keep candidate-facing interactions human-only. Use AI for drafting job descriptions and scheduling, not screening.",
]

medium_conf = [
    "Explore AI-powered energy management with conservative ROI expectations. Expect 8-12% reductions initially, not the 15-30% in vendor marketing. Academic evidence suggests 3% is more realistic without AI optimization of the management process itself.",
    "Pilot AI meal compliance tools (e.g., LunchLens) with USDA human-in-the-loop requirements. The regulatory framework now supports this, and staffing shortages (90%+) make automation compelling.",
    "Use AI for financial forecasting and procurement drafting. The evidence for forecast accuracy improvement is strong (57% vs. 8%), and procurement automation has specific district case studies.",
    "Leverage federal funds for AI. USED guidance explicitly allows formula and discretionary grants for AI integration, including operational modernization.",
]

low_conf = [
    "Monitor but do not yet invest heavily in AI for school nursing, custodial operations, or HR screening. These domains lack evidence, tools, or both. Autonomous floor-cleaning robots are the one custodial exception worth evaluating if staffing shortages are severe.",
    "Watch for the National Congress on School Transportation (May 2025) specifications update and the Tennessee AI Advisory Council's 2028 milestones as policy benchmarks.",
]

doc.add_heading("High Confidence Recommendations", level=3)
for item in high_conf:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_run_font(run, size=11)

doc.add_heading("Medium Confidence Recommendations", level=3)
for item in medium_conf:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_run_font(run, size=11)

doc.add_heading("Lower Confidence / Watch Items", level=3)
for item in low_conf:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item)
    set_run_font(run, size=11)

doc.add_heading("Unanswered Questions", level=2)
unanswered = [
    "What is the actual ROI of AI operational tools in small and mid-size districts (under 10,000 students)?",
    "How should districts evaluate AI vendors when independent third-party assessments are largely absent?",
    "What data governance standards are sufficient for operational AI in K-12, beyond FERPA and CIPA?",
    "Can AI-powered predictive maintenance deliver in older school buildings with limited sensor infrastructure?",
    "How will the federal policy landscape evolve given potential changes to the Department of Education?",
    "What training and change management approaches work best for non-instructional staff adopting AI tools?",
]
for q in unanswered:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(q)
    set_run_font(run, size=11, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. FEDERAL & STATE POLICY CONTEXT (Education Supplement)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Federal and State Policy Context", level=1)

policy_paras = [
    "The federal policy environment for AI in K-12 education underwent a significant shift in 2025. The Trump administration's Executive Order on AI Education (April 2025) established the White House Task Force on AI Education with representation from the Secretaries of Education, Labor, and Agriculture. This was followed by the USED Dear Colleague Letter (July 2025) explicitly authorizing federal grant funds for AI integration, including operational applications.",

    "At the state level, Tennessee stands as a national leader. HB1630/SB1711 (Public Chapter 550, 2024) requires every LEA and public charter school to adopt AI policies by the 2024-25 school year, with annual compliance reporting to TDOE by July 1. The Tennessee AI Advisory Council's November 2025 Action Plan, organized around four pillars (Modernization, Data Readiness, Workforce Enablement, Safety/Accountability), provides a statewide framework that extends beyond education into all government AI adoption.",

    "Ohio joins Tennessee as the only other state with a legal mandate, requiring districts to adopt AI policies by July 1, 2026. By mid-2025, 26 states plus Puerto Rico had published official AI guidance for K-12. The Region 8 Comprehensive Center's federally funded guidance document (ERIC ED655341) synthesizes seven state approaches into decision points for districts.",

    "Key federal frameworks governing AI procurement and governance include OMB M-24-18 (AI Acquisition, September 2024) requiring documentation of model training and preventing vendor lock-in, and NIST AI RMF 1.0 with the GenAI Profile (AI 600-1). The USDA's framework for AI in public benefit administration directly governs school nutrition programs. The FCC's $200M cybersecurity pilot represents the first dedicated federal funding for K-12 cyber infrastructure.",

    "The U.S. Commission on Civil Rights (December 2024) flagged four civil rights concerns with AI in K-12: disruption of student-teacher relationships, promotion of bias, widening of the digital divide, and privacy/surveillance risks, recommending proactive protections before broad implementation.",
]

for text in policy_paras:
    add_styled_paragraph(doc, text)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. SOURCE RELIABILITY MATRIX
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Source Reliability Matrix", level=1)

add_styled_paragraph(doc, "Sources are scored using a weighted formula: Source Type Weight x Agent Confidence x URL Status + Corroboration Bonus (capped at 1.0). Sources cited by multiple agents receive a corroboration bonus (+0.1 for 2 agents, +0.2 for 3+).", italic=True, size=10)

# Abbreviated top sources table
source_table = add_table_with_header(doc,
    ["Source", "Category", "Confidence", "Reliability"],
    [3.5, 1.0, 1.0, 1.0])

top_sources = [
    ("White House EO on AI Education (2025)", "Gov", "High", "High"),
    ("USED Dear Colleague Letter on AI (2025)", "Gov", "High", "High"),
    ("TN HB1630/SB1711 (2024)", "Gov", "High", "High"),
    ("NIST AI RMF 1.0 (2023)", "Gov", "High", "High"),
    ("USDA Framework for AI in Public Benefits", "Gov", "High", "High"),
    ("FCC Cybersecurity Pilot Program (2024)", "Gov", "High", "High"),
    ("GAO K-12 Cybersecurity Reports (2022-23)", "Gov", "High", "High"),
    ("CISA K-12 Cybersecurity Toolkit (2023)", "Gov", "High", "High"),
    ("USCCR AI in K-12 Education (2024)", "Gov", "High", "High"),
    ("Region 8 Comprehensive Center (2024)", "Gov", "High", "High"),
    ("Ohio Model AI Policy", "Gov", "High", "High"),
    ("OMB M-24-18 AI Acquisition (2024)", "Gov", "High", "High"),
    ("TN AI Advisory Council Action Plan (2025)", "Gov", "High", "High"),
    ("RAND AI Adoption Survey (2025)", "Academic", "High", "High"),
    ("RAND Cybersecurity Survey (2024)", "Academic", "High", "High"),
    ("Bertsimas et al. Boston Bus Routing (2020)", "Academic", "High", "High"),
    ("Burlig et al. CA School Energy (2020)", "Academic", "High", "High"),
    ("UW AI Hiring Bias Study (2024)", "Academic", "High", "High"),
    ("Adjapong et al. Plate Waste (2024)", "Academic", "High", "High"),
    ("JRTE AI Policy Analysis (2025)", "Academic", "High", "High"),
    ("CIS/MS-ISAC K-12 Cyber Report (2025)", "Industry", "High", "High"),
    ("CoSN Operational AI Survey (2025)", "Industry", "High", "High"),
    ("CRPE Early Adopter Tracker (2024-25)", "Industry", "High", "High"),
    ("Digital Promise/Gates AI Projects (2024)", "Industry", "High", "High"),
    ("Frontline K-12 Lens 2026 Survey", "News/Ind", "High", "High"),
    ("Education Week, multiple articles (2025)", "News", "High", "High"),
    ("GovTech, multiple articles (2025-26)", "News", "High", "High"),
    ("EdTech Magazine K-12 AI Ops (2025)", "News", "High", "High"),
    ("McKinsey K-12 Survey (2025)", "Industry", "Medium", "Medium"),
    ("EdSurge Cybersecurity Series (2025)", "News", "High", "High"),
    ("Edutopia AI Mental Health (2025)", "News", "High", "Medium"),
    ("Cummins et al. Rayfood (2022)", "Academic", "Medium", "Medium"),
    ("Bouabdallaoui et al. Pred. Maint. (2021)", "Academic", "Medium", "Medium"),
    ("Sposato AI Taxonomy (2025)", "Academic", "Medium", "Medium"),
    ("Bowen et al. Absenteeism ML (2022)", "Academic", "Medium", "Medium"),
    ("School Bus Fleet Trends (2025)", "Industry", "Medium", "Medium"),
    ("AASA AI Resources (2025)", "Industry", "Medium", "Medium"),
    ("Frontline K-12 Lens 2025 HR", "Industry", "Medium", "Medium"),
    ("Intellis K12FF Recap (2025)", "Industry", "Medium", "Medium"),
    ("Tyler Technologies PBB (2025)", "Industry", "Medium", "Medium"),
    ("LunchLens (2025)", "Industry", "Low-Med", "Low"),
    ("EdSurge Sponsored/AWS (2025)", "News", "Low", "Low"),
    ("DOE EMIS Resources", "Gov", "Medium", "Medium"),
]

for source, cat, conf, rel in top_sources:
    color = None
    if rel == "High":
        color = "E2EFDA"
    elif rel == "Medium":
        color = "FFF2CC"
    else:
        color = "FCE4EC"
    row = add_row(source_table, [source, cat, conf, rel], shading=color)

doc.add_paragraph()
add_styled_paragraph(doc, "Summary: 43 unique sources assessed. 28 high reliability, 12 medium reliability, 3 low reliability. 0 dead URLs replaced (verification pending for some .gov URLs that block programmatic access).", size=10, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("8. References", level=1)

references = [
    'Adjapong, E.S., Bender, K.E., Schaefer, S., & Roe, B.E. (2024). School and meal characteristics associated with plate waste in K-12 cafeterias. PLOS ONE, 19(12), e0299043. https://doi.org/10.1371/journal.pone.0299043',
    'Bertsimas, D., Delarue, A., Eger, W., Hanlon, J., & Martin, S. (2020). Bus Routing Optimization Helps Boston Public Schools Design Better Policies. INFORMS Journal on Applied Analytics, 50(1), 37-49.',
    'Bouabdallaoui, Y., Lafhaj, Z., Yim, P., Ducoulombier, L., & Bennadji, B. (2021). Predictive Maintenance in Building Facilities: A Machine Learning-Based Approach. Sensors, 21(4), 1044.',
    'Bowen, F., Gentle-Genitty, C., Siegler, J., & Jackson, M. (2022). Revealing underlying factors of absenteeism. Frontiers in Psychology, 13, 958748.',
    'Burlig, F., Knittel, C.R., Rapson, D., Reguant, M., & Wolfram, C. (2020). Machine Learning from Schools about Energy Efficiency. Journal of the Association of Environmental and Resource Economists, 7(6).',
    'CIS/MS-ISAC (2025). 2025 K-12 Cybersecurity Report. https://www.cisecurity.org/insights/white-papers/2025-k12-cybersecurity-report',
    'CISA (2023). Partnering to Safeguard K-12 Organizations from Cybersecurity Threats. https://www.cisa.gov/topics/cybersecurity-best-practices/K12cybersecurity/',
    'CoSN (2025). Operational AI in Education: Readiness, Realities, and the Road Ahead. https://www.cosn.org/',
    'CRPE (2024-2025). Districts and AI: Tracking Early Adopters. https://crpe.org/',
    'Cummins, B. et al. (2022). AI for Sustainable Institutional Food Systems. Frontiers in Sustainable Food Systems. https://doi.org/10.3389/fsufs.2022.743810',
    'Descant, S. (2025). On School Bus Route Planning, Districts Get a Lift From AI. GovTech.',
    'Digital Promise (2024). Transforming K-12 Education with AI: Insights from 28 Exploratory Projects.',
    'Diliberti, M., & Schwartz, H.L. (2025). Uneven Adoption of AI Tools Among U.S. Teachers and Principals. RAND Corporation, RR-A134-25.',
    'Dusseault, B. & Hurwitz, J. (2025). New Data Shows More Districts Are Adopting AI. The 74.',
    'FCC (2024). Schools and Libraries Cybersecurity Pilot Program. https://www.fcc.gov/cybersecurity-pilot-program',
    'Frontline Education (2025-2026). K-12 Lens Survey Reports. https://www.frontlineeducation.com/',
    'GAO (2022). GAO-22-105024: K-12 Cyber Threats. https://www.gao.gov/products/gao-22-105024',
    'GAO (2023). GAO-23-105480: K-12 Cybersecurity Coordination. https://www.gao.gov/products/gao-23-105480',
    'Heubeck, E. (2025). Will AI Be the Answer to K-12\'s Hiring Headaches? Education Week.',
    'Intellis (2025). K12FF 2025 Recap: AI-Powered Facilities Planning Insights.',
    'McKinsey (2025). From Surplus to Scarcity: K-12 Districts Brace for Leaner Years.',
    'Nemani, S. (2025). Evaluating AI Impact on Administrative Burden in Middle Schools. Current Perspectives in Educational Research, 8(1).',
    'Ng, A. (2025). How AI Is Changing Procurement for K-12 Education. GovTech.',
    'NIST (2023). AI Risk Management Framework (AI RMF 1.0). https://www.nist.gov/itl/ai-risk-management-framework',
    'NIST (2024). NIST AI 600-1: GenAI Profile. https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf',
    'Ohio Department of Education (n.d.). AI Model Policy for Ohio Districts. https://education.ohio.gov/',
    'OMB (2024). M-24-18: AI Acquisition Memorandum.',
    'OMB (2025). M-25-21: Accelerating Federal Use of AI.',
    'Perry, E. (2026). Spokane Public Schools Automates Floor Cleaning With Robots. GovTech/Spokesman-Review.',
    'Prothero, A. (2025). How One District Uses AI to Build More Efficient Master School Schedules. Education Week.',
    'RAND Corporation (2024). Protecting Schools Virtually: Cybersecurity and Threats on K-12 Systems. RRA3930-6.',
    'RAND Corporation (2025). AI Use in Schools Is Quickly Increasing but Guidance Lags Behind. RRA4180-1.',
    'Rankine, J., Goldberg, L., Miller, E., Kelley, L., & Ray, K.N. (2021). School Nurse Perspectives on Chronic Absenteeism. Journal of School Nursing, 39(6).',
    'Region 8 Comprehensive Center (2024). Guidance for AI Policies in K-12 Schools. ERIC ED655341.',
    'Sposato, M. (2025). AI in educational leadership taxonomy. International Journal of Educational Technology in Higher Education, 22, 20.',
    'Stone, A. (2025). Sustainable Tech Aids the Environment and School Budgets. EdTech Magazine.',
    'Tennessee AI Advisory Council (2025). Tennessee AI Advisory Council Action Plan.',
    'Tennessee General Assembly (2024). HB1630/SB1711: AI Policy Mandate.',
    'Tutt, P. (2025). Schools Try AI as Student Mental Health Needs Surge. Edutopia.',
    'Tyler Technologies (2025). Priority Based Budgeting.',
    'Ullman, E. (2025). From Defense to Resilience: Where School Cybersecurity Goes Next. EdSurge.',
    'University of Washington (2024). AI tools show biases in ranking job applicants\' names.',
    'U.S. Commission on Civil Rights (2024). The Rising Use of AI in K-12 Education.',
    'U.S. Department of Education (2023). AI and the Future of Teaching and Learning.',
    'U.S. Department of Education (2025). Guidance on Federal Grant Funds for AI.',
    'U.S. Department of Energy (n.d.). Energy Management Information Systems.',
    'USDA Food and Nutrition Service (n.d.). Framework for AI in Public Benefit Administration.',
    'The White House (2025). Advancing Artificial Intelligence Education for American Youth.',
    'Wong, W. (2025). How AI Is Transforming Business Operations in K-12. EdTech Magazine.',
]

for ref in sorted(references):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(ref)
    set_run_font(run, size=9)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_FILE)
print(f"Report saved to: {OUTPUT_FILE}")
