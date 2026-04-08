#!/usr/bin/env python3
"""Generate research report: Thumbs Up as Passive-Aggressive Communication"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import date

def add_paragraph(doc, text, style='Normal', bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_citation(doc, text):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    return p

def build_report(output_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Thumbs Up as Passive-Aggressive Communication', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Research Report | {date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 1: Definition & Background
    # =========================================================================
    doc.add_heading('1. Definition & Background', level=1)

    doc.add_paragraph(
        'The thumbs up emoji has become one of the most contested symbols in digital communication. '
        'Once a universally understood gesture of approval, the simple act of sending a thumbs up in a text message or '
        'workplace chat now triggers genuine confusion, frustration, and even legal consequences. The core question is '
        'straightforward: when someone sends you a thumbs up, do they mean "great, got it" or "I do not care enough '
        'about you to type actual words"?'
    )

    doc.add_paragraph(
        'The debate went mainstream in October 2022, when a 24-year-old Reddit user posted about feeling "unsettled" '
        'by thumbs up emoji in workplace messages. Other users agreed it felt "hostile" and "hurtful." The Daily Mail '
        'amplified the story on October 12, 2022, and coverage exploded globally, from the New York Post ("Gen Z '
        'canceled the hostile thumbs-up emoji") to NBC\'s Today Show and international outlets. But the phenomenon '
        'predates the viral moment. Researchers had been documenting generational divergence in emoji interpretation '
        'for years.'
    )

    doc.add_heading('Historical Origins of the Gesture', level=2)

    doc.add_paragraph(
        'The thumbs up gesture has a complicated lineage. The popular story traces it to Roman gladiatorial combat, '
        'where "pollice verso" (turned thumb) supposedly determined whether a gladiator lived or died. The reality is '
        'murkier. Classical studies professor Anthony Corbeill argues that in Rome, thumbs up may have actually '
        'signaled killing a gladiator, while a closed fist with wrapped thumb meant sparing him. The popular '
        '"thumbs up = life" narrative was largely created by Jean-Leon Gerome\'s 1872 painting Pollice Verso, which '
        'likely reversed the actual meaning.'
    )

    doc.add_paragraph(
        'The modern positive connotation spread through military use. WWII pilots used thumbs up to communicate with '
        'ground crews before takeoff, possibly originating with the China-based Flying Tigers. This military usage '
        'carried the gesture\'s positive connotation into mainstream American culture.'
    )

    doc.add_heading('The Digital Thumbs Up', level=2)

    doc.add_paragraph(
        'Facebook launched its Like button on February 9, 2009, as a thumbs-up icon. Designer Soleio Cuervo tried '
        'green thumbs up, hearts, and stars before settling on the blue-and-white thumb. Internally, it was called '
        'the "awesome button" during development. The thumbs up was approved as part of Unicode 6.0 in 2010 and added '
        'to Emoji 1.0 in 2015, with skin tone variations introduced the same year. The Facebook Like button alone '
        'cemented the thumbs up as the default digital gesture of approval for over a billion users.'
    )

    # =========================================================================
    # SECTION 2: Current State of Knowledge
    # =========================================================================
    doc.add_heading('2. Current State of Knowledge', level=1)

    doc.add_heading('The Generational Divide: Strong Evidence', level=2)

    doc.add_paragraph(
        'The most rigorous study on this topic is Zhukova and Herring (2024), published in Language@Internet. Through '
        'an online survey, participants rated text-and-emoji combinations on scales of friendliness, playfulness, '
        'sarcasm, passive-aggression, offensiveness, and threat. The key finding: messages with the thumbs up emoji '
        'were perceived as more passive-aggressive compared with other emojis. Gen Z respondents and non-binary '
        'individuals perceived emoji-containing messages as significantly more sarcastic, passive-aggressive, and '
        'threatening than older generations. Males were more likely than females to perceive a thumbs up message as '
        'sarcastic.'
    )
    add_citation(doc,
        'Zhukova, M. & Herring, S.C. (2024). "Benign or Toxic? Differences in Emoji Interpretation by Gender, '
        'Generation, and Emoji Type." Language@Internet, 22. Indiana University.'
    )

    doc.add_paragraph(
        'A 2024 study in the Advance Social Science Archive Journal confirmed the pattern across three generations. '
        'Gen Z assigns different meanings to emojis than their face-value interpretations, Millennials tend to take '
        'them at face value, and Boomers use them plainly. The thumbs up, crying laughing face, and skull emojis '
        'cause the most inter-generational confusion.'
    )
    add_citation(doc,
        'ASSA Journal (2024). "Generational Differences in Emoji Interpretation: A Study of Millennial, Gen Z, '
        'and Baby Boomers."'
    )

    doc.add_paragraph(
        'Survey data reinforces the academic findings. A Perspectus Global survey of 2,000 respondents aged 16-29 '
        'found that 24% of Gen Z specifically identified the thumbs up as making users "look old." A 2024/2025 '
        'British survey of 2,000 respondents found thumbs up was the #1 most annoying emoji, with 22% finding it '
        'irritating. A Bitrix24 study found 44% of Gen Z workers prefer ironic emoji meanings, compared to 17% of '
        'Millennials, 14% of Boomers, and 12% of Gen X.'
    )

    doc.add_heading('Cross-Platform Interpretation: Strong Evidence', level=2)

    doc.add_paragraph(
        'The landmark Miller et al. (2016) GroupLens study surveyed users on emoji interpretations across five '
        'platforms (Apple, Google, Microsoft, Samsung, LG). For the same Unicode character, participants described '
        'the Google rendering as "blissfully happy" while the Apple rendering was described as "ready to fight." '
        'When participants rated the same emoji rendering, they disagreed on whether the sentiment was positive, '
        'neutral, or negative 25% of the time. This means that even before generational differences enter the '
        'picture, the platform you are using already introduces ambiguity.'
    )
    add_citation(doc,
        'Miller, H., Thebault-Spieker, J., Chang, S., Johnson, I., Terveen, L., & Hecht, B. (2016). "Blissfully '
        'happy or ready to fight: Varying Interpretations of Emoji." Proceedings of ICWSM 2016.'
    )

    doc.add_heading('The Counterpoint: Mixed Evidence', level=2)

    doc.add_paragraph(
        'Not all research supports a clean generational split. Kempe and Raviv (2025) found no evidence for '
        'generational differences in the conventionalization of face emojis, arguing that after a decade of use, '
        'face emojis have become a widely conventionalized semiotic system accessible regardless of age. However, '
        'this study focused on face emojis specifically and may not generalize to gesture emojis like the thumbs up, '
        'which carry additional cultural and contextual baggage.'
    )
    add_citation(doc,
        'Kempe, V. & Raviv, L. (2025). "No Evidence for Generational Differences in the Conventionalisation of '
        'Face Emojis." ScienceDirect/SSRN.'
    )

    doc.add_heading('Broader Emoji Interpretation Research', level=2)

    doc.add_paragraph(
        'A systematic review by Bai et al. (2019) in Frontiers in Psychology surveyed the full landscape of emoji '
        'research through 2019, covering interpretation, sentiment, cross-platform differences, and communication '
        'effects. The review established that emoji interpretation is influenced by individual differences (personality, '
        'gender, age), contextual factors (relationship, platform, message content), and cultural background. Chen et '
        'al. (2024) studied 253 Chinese and 270 UK adults and found that age, gender, and culture all significantly '
        'affect emoji comprehension, with older participants less likely to match intended meanings for emojis '
        'representing surprise, fear, sadness, and anger.'
    )
    add_citation(doc,
        'Bai, Q., Dan, Q., Mu, Z., & Yang, M. (2019). "A Systematic Review of Emoji: Current Research and Future '
        'Perspectives." Frontiers in Psychology, 10, 2221.'
    )
    add_citation(doc,
        'Chen, Y., Yang, X. et al. (2024). "Individual differences in emoji comprehension: Gender, age, and '
        'culture." PLOS ONE.'
    )

    # =========================================================================
    # SECTION 3: What Works vs. What Doesn't
    # =========================================================================
    doc.add_heading('3. What Works vs. What Doesn\'t', level=1)

    doc.add_heading('When Thumbs Up Works', level=2)

    doc.add_paragraph(
        'Thumbs up functions well as a quick acknowledgment in contexts where brevity is expected and the '
        'relationship is established. Slack data shows 69% of American workers say emoji allows more nuance with '
        'fewer words, and 67% say it speeds up communication. In fast-moving team chats where messages flow '
        'continuously, a thumbs up reaction to a status update or logistics message is generally read as intended: '
        '"got it, thanks."'
    )

    doc.add_paragraph(
        'The key conditions for thumbs up to land well:'
    )

    items = [
        'The sender and receiver have an established relationship with shared communication norms.',
        'The message being responded to is logistical or informational, not emotional or vulnerable.',
        'The platform culture supports reactions (e.g., Slack reacjis, Teams reactions).',
        'Both parties are from similar generational or cultural backgrounds.',
        'The thumbs up is one of several communication tools being used, not the only response.'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('When Thumbs Up Fails', level=2)

    doc.add_paragraph(
        'The thumbs up becomes problematic in several predictable contexts:'
    )

    items = [
        'Responding to emotional or vulnerable messages (reads as dismissive).',
        'Cross-generational communication without established norms.',
        'As the sole response to a detailed message that clearly requires substantive engagement.',
        'From a supervisor to a subordinate after a request or concern (reads as "noted, moving on").',
        'In cultures where the gesture itself carries negative connotations.',
        'When the sender has a pattern of using thumbs up to avoid real conversation.'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'A Preply survey of 1,001 Americans found that 78% had been confused by someone else\'s emoji use, and '
        '1 in 3 had seen a misinterpreted emoji create an uncomfortable situation. The Bitrix24 study found 65% '
        'of employees had avoided using an emoji at work for fear of misinterpretation, and 27% reported feeling '
        'offended by a workplace emoji.'
    )

    # =========================================================================
    # SECTION 4: Criticisms & Limitations
    # =========================================================================
    doc.add_heading('4. Criticisms & Limitations', level=1)

    doc.add_heading('Methodological Concerns', level=2)

    doc.add_paragraph(
        'Much of the popular discourse around the "thumbs up is passive-aggressive" narrative rests on surveys with '
        'methodological limitations. The Perspectus Global survey that launched the 2022 viral moment surveyed only '
        '2,000 respondents aged 16-29, meaning it captured Gen Z\'s perspective without a comparison group. Many '
        'media articles then extrapolated from this single-cohort survey to make claims about generational conflict '
        'that the data did not support.'
    )

    doc.add_paragraph(
        'The Kempe and Raviv (2025) counterpoint study directly challenges the generational narrative, finding no '
        'significant differences in how generations interpret face emojis. Their argument: the "generational emoji '
        'gap" may be partly a media construction amplified by viral content rather than a robust empirical finding.'
    )

    doc.add_heading('Context Dependency', level=2)

    doc.add_paragraph(
        'Researchers consistently note that context matters more than any single variable. The same thumbs up from '
        'your spouse after you text "picking up groceries" means something entirely different from a thumbs up from '
        'your boss after you send a three-paragraph email about a project concern. Studies that strip emoji of their '
        'conversational context (as survey-based research necessarily does) may overstate the degree of '
        'misinterpretation that occurs in real-world exchanges where relationship history provides interpretive '
        'scaffolding.'
    )

    doc.add_heading('Cultural Variability', level=2)

    doc.add_paragraph(
        'The thumbs up gesture is considered offensive in Iran, Iraq, Afghanistan, parts of West Africa, Greece, '
        'Russia, and Sardinia. In these regions, the gesture is roughly equivalent to the middle finger in American '
        'culture. This means that the digital thumbs up emoji carries culturally loaded associations that vary '
        'dramatically by geography, a dimension that most English-language research underexplores.'
    )

    doc.add_paragraph(
        'Bitrix24 data shows that high-context cultures (e.g., Brazil, Latin America) use emoji far more heavily '
        'than low-context cultures (e.g., Germany). A Frontiers in Psychology (2024) study of Chinese WeChat users '
        'found intergenerational discrepancies in emoji comprehension and aesthetic preference specific to that '
        'cultural context, suggesting the dynamics are not simply a Western phenomenon but manifest differently '
        'across cultures.'
    )

    # =========================================================================
    # SECTION 5: Measurable Outcomes
    # =========================================================================
    doc.add_heading('5. Measurable Outcomes', level=1)

    doc.add_heading('Workplace Communication', level=2)

    doc.add_paragraph(
        'Glikson, Cheshin, and van Kleef (2018) conducted three experiments with 549 participants from 29 countries '
        'and found that smiley emoticons in professional contexts do not increase perceptions of warmth and actually '
        'decrease perceptions of competence. Low competence perceptions undermined information sharing in formal work '
        'emails. This suggests that casual digital gestures (including thumbs up) in formal professional contexts can '
        'carry real costs to perceived credibility.'
    )
    add_citation(doc,
        'Glikson, E., Cheshin, A., & van Kleef, G.A. (2018). "The Dark Side of a Smiley: Effects of Smiling '
        'Emoticons on Virtual First Impressions." Social Psychological and Personality Science.'
    )

    doc.add_paragraph(
        'Courtice et al. (2026) studied 243 participants and found that while positive emojis can enhance '
        'competence perceptions when paired with positive or neutral sentences, negative emojis consistently '
        'decreased perceived competence. Messages without emojis or with positive emojis were viewed as most '
        'appropriate in workplace contexts.'
    )
    add_citation(doc,
        'Courtice, E.L. et al. (2026). "Emojis at Work: The Effects of Emoji Use on Perceptions of Competence '
        'and Appropriateness." Collabra: Psychology, University of California Press.'
    )

    doc.add_heading('Legal Consequences', level=2)

    doc.add_paragraph(
        'In Achter Land & Cattle Ltd. v. South West Terminal Ltd. (2023, upheld on appeal 2024), a Saskatchewan, '
        'Canada court ruled that a thumbs up emoji sent via text constituted a valid signature on a contract. A '
        'grain buyer texted a photo of a contract; the farmer replied with a thumbs up emoji. The court ordered '
        'the farmer to pay $82,000 CAD for failing to deliver the flax. The appeal court upheld the ruling 2-1, '
        'noting the parties\' history of informal text-based contracting. This contrasts with Lightstone v. Zinntex '
        '(2022, New York), where a thumbs up was ruled not to constitute contract acceptance because the party had '
        'earlier stated he would not sign.'
    )

    doc.add_heading('Relationship and Trust Effects', level=2)

    doc.add_paragraph(
        'Herring and Dainas (2020) found that texts with affectionate emojis were judged as more appropriate and '
        'likable when they came from women, while sender\'s emotional intent can differ from the receiver\'s '
        'interpretation by as much as 26% for the same negative emoji. Riordan (2017) framed emoji use in terms of '
        'sociological emotion work theory: the effort involved in selecting and using emojis helps maintain social '
        'relationships. A thumbs up, by requiring minimal effort, may signal minimal investment in the relationship.'
    )
    add_citation(doc,
        'Riordan, M.A. (2017). "Emojis as Tools for Emotion Work: Communicating Affect in Text Messages." '
        'Journal of Language and Social Psychology, 36(5).'
    )

    # =========================================================================
    # SECTION 6: Practical Applications
    # =========================================================================
    doc.add_heading('6. Practical Applications', level=1)

    doc.add_heading('Theoretical Framework: Why This Happens', level=2)

    doc.add_paragraph(
        'Linguist Vyvyan Evans offers the most coherent theoretical explanation through four mechanisms:'
    )

    items = [
        ('Semantic satiation: ', 'Repeated overuse of the thumbs up as a default positive response drained its sincerity. When every message gets a thumbs up, no message feels genuinely acknowledged.'),
        ('Pragmatic reanalysis: ', 'When the positive meaning faded through overuse, what remained was perceived insincerity, which reads as dismissive or passive-aggressive.'),
        ('Digital paralanguage: ', 'Emojis function as substitutes for face-to-face nonverbal cues (gestures, facial expressions, tone of voice), so their misinterpretation carries emotional weight comparable to misreading body language.'),
        ('Platform vernaculars: ', 'Emojis take on distinct meanings depending on the digital platform, making universal interpretation impossible. A thumbs up reaction on Slack means something different than a thumbs up text on iMessage.')
    ]
    for label, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(label)
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'Gretchen McCulloch, author of Because Internet (2019), provides additional framing. She argues that emojis '
        'function in digital writing the way gestures function in spoken conversation. They are not a language but a '
        'paralinguistic system conveying attitude and tone. She distinguishes between "emblematic" emoji (standalone '
        'replacements for words) and "co-speech" emoji (modifying the tone of accompanying text, like hand gestures '
        'in conversation). A standalone thumbs up is emblematic and therefore carries more interpretive weight than '
        'a thumbs up appended to a full sentence.'
    )

    doc.add_heading('The Period Parallel', level=2)

    doc.add_paragraph(
        'The thumbs up controversy mirrors earlier research on text-message punctuation. Klin and Huang (2015) at '
        'Binghamton University found that 126 undergraduates rated one-word text responses ending with a period '
        '(e.g., "Okay." vs "Okay") as significantly less sincere. This effect did not exist for handwritten notes. '
        'Follow-up work found that exclamation marks made messages seem more sincere. The pattern is the same: '
        'digital communication conventions evolve independently from their offline equivalents, and what seems '
        'neutral in one medium can read as hostile in another.'
    )
    add_citation(doc,
        'Klin, C.M. & Huang, D.H. (2015). "Texting insincerely: The role of the period in text messaging." '
        'Computers in Human Behavior.'
    )

    doc.add_heading('Practical Guidelines for Digital Communication', level=2)

    doc.add_paragraph(
        'Based on the research, several practical strategies emerge for navigating emoji ambiguity:'
    )

    items = [
        'Know your audience. If communicating across generations, default to words over reactions for substantive exchanges.',
        'Match the effort. If someone sends you a detailed message, a thumbs up is likely insufficient regardless of your intent.',
        'Use thumbs up for logistics, not emotions. "I will be there at 3" merits a thumbs up. "I have been struggling with this project" does not.',
        'Establish team norms. Erica Dhawan recommends teams explicitly discuss digital communication preferences, including reaction etiquette.',
        'When in doubt, add words. A thumbs up plus "sounds good" eliminates almost all ambiguity.',
        'Be aware of cultural context. In international teams, the gesture may carry unintended offense.'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # =========================================================================
    # SECTION 7: Bottom Line
    # =========================================================================
    doc.add_heading('7. Bottom Line', level=1)

    doc.add_paragraph(
        'The evidence is clear that the thumbs up emoji is genuinely interpreted differently across age groups, '
        'genders, cultures, and platforms. This is not a media-fabricated controversy. Peer-reviewed research from '
        'Indiana University, Binghamton University, and multiple international teams consistently documents that '
        'younger communicators (particularly Gen Z) are more likely to read thumbs up as dismissive, sarcastic, or '
        'passive-aggressive. The effect is amplified when the thumbs up serves as the sole response to a message '
        'that warranted substantive engagement.'
    )

    doc.add_paragraph(
        'However, the narrative that "Gen Z has canceled the thumbs up" is an oversimplification. The Kempe and '
        'Raviv (2025) counterpoint study found no generational differences in face emoji interpretation, and Slack\'s '
        'own survey data shows that most workers across generations find emoji helpful for efficient communication. '
        'The issue is not the thumbs up itself but the mismatch between sender intent and receiver interpretation '
        'that occurs when relationship context, generational norms, and platform conventions diverge. In workplaces '
        'with established communication cultures, thumbs up reactions remain a practical, efficient tool.'
    )

    doc.add_paragraph(
        'The deeper insight from this research is about digital paralanguage more broadly. As Erica Dhawan\'s survey '
        'of nearly 2,000 office workers found, over 70% experienced unclear communication from colleagues. Research '
        'by Kruger et al. (2005) in the Journal of Personality and Social Psychology found that email tone is '
        'correctly identified only about 56% of the time, barely better than chance. The thumbs up is a lightning rod for a much '
        'larger problem: text-based communication strips away the nonverbal cues that carry the majority of meaning '
        'in face-to-face interaction, and our digital substitutes (emoji, punctuation, formatting) are not yet stable '
        'enough to fill that gap reliably.'
    )

    doc.add_paragraph(
        'What remains unanswered: longitudinal studies tracking how emoji interpretation evolves within individuals '
        'over time (rather than cross-sectional snapshots); robust experimental research isolating the thumbs up '
        'specifically (most studies examine emoji as a class); and whether explicit norm-setting in teams actually '
        'reduces misinterpretation in practice. The legal dimension is also evolving rapidly, as the Canadian court '
        'ruling demonstrates, and deserves closer attention as emoji become more embedded in commercial and '
        'contractual communication.'
    )

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_heading('References', level=1)

    refs = [
        'Bai, Q., Dan, Q., Mu, Z., & Yang, M. (2019). A systematic review of emoji: Current research and future perspectives. Frontiers in Psychology, 10, 2221.',
        'Chen, Y., Yang, X. et al. (2024). Individual differences in emoji comprehension: Gender, age, and culture. PLOS ONE.',
        'Courtice, E.L., Lawrence, M., Collin, C.A., & Boutet, I. (2026). Emojis at work: The effects of emoji use on perceptions of competence and appropriateness. Collabra: Psychology, University of California Press.',
        'Dhawan, E. (2021). Digital body language: How to build trust and connection, no matter the distance. St. Martin\'s Press.',
        'Evans, V. (2017). The emoji code: The linguistics behind smiley faces and scaredy cats. Picador.',
        'Evans, V. (2022). What\'s so wrong about using the thumbs-up emoji? Psychology Today.',
        'Glikson, E., Cheshin, A., & van Kleef, G.A. (2018). The dark side of a smiley: Effects of smiling emoticons on virtual first impressions. Social Psychological and Personality Science.',
        'Herring, S.C. & Dainas, A.R. (2020). Receiver interpretations of emoji functions: A gender perspective.',
        'Kempe, V. & Raviv, L. (2025). No evidence for generational differences in the conventionalisation of face emojis. ScienceDirect/SSRN.',
        'Klin, C.M. & Huang, D.H. (2015). Texting insincerely: The role of the period in text messaging. Computers in Human Behavior.',
        'McCulloch, G. (2019). Because internet: Understanding the new rules of language. Riverhead Books.',
        'Miller, H., Thebault-Spieker, J., Chang, S., Johnson, I., Terveen, L., & Hecht, B. (2016). Blissfully happy or ready to fight: Varying interpretations of emoji. Proceedings of ICWSM 2016.',
        'Riordan, M.A. (2017). Emojis as tools for emotion work: Communicating affect in text messages. Journal of Language and Social Psychology, 36(5).',
        'Kruger, J., Epley, N., Parker, J., & Ng, Z. (2005). Egocentrism over e-mail: Can we communicate as well as we think? Journal of Personality and Social Psychology, 89(6), 925-936.',
        'Zhukova, M. & Herring, S.C. (2024). Benign or toxic? Differences in emoji interpretation by gender, generation, and emoji type. Language@Internet, 22.',
    ]

    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ref)
        run.font.size = Pt(10)

    doc.add_paragraph()

    surveys = [
        'Bitrix24 (2024). The emoji economy. https://www.bitrix24.com/resources/files/the_emoji_economy.pdf',
        'Perspectus Global (2021/2022). Survey of 2,000 respondents aged 16-29 on emoji perceptions.',
        'Preply (2023). Most confusing emojis survey (n=1,001). https://preply.com/en/blog/most-confusing-emojis/',
        'Slack x Duolingo (2022). World Emoji Day workplace survey. https://slack.com/blog/collaboration/emoji-use-at-work',
        'YouGov/Atlassian (2025). Workplace emoji survey (n=10,000 across 5 countries).',
    ]

    p = doc.add_paragraph()
    run = p.add_run('Survey Sources:')
    run.bold = True

    for s in surveys:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(s)
        run.font.size = Pt(10)

    # Save
    doc.save(output_path)
    print(f'Report saved to: {output_path}')

if __name__ == '__main__':
    output_dir = '/Users/hornej/Documents/Research/thumbs-up-passive-aggressive'
    drafts_dir = os.path.join(output_dir, 'drafts')
    os.makedirs(drafts_dir, exist_ok=True)

    # Save preliminary draft for fact-checking
    preliminary_path = os.path.join(drafts_dir, 'preliminary_report.docx')
    build_report(preliminary_path)

    # Save final report
    final_path = os.path.join(output_dir, '2026-03-29 Thumbs Up Passive-Aggressive Research Report.docx')
    build_report(final_path)
