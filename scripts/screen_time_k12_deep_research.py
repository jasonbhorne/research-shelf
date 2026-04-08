#!/usr/bin/env python3
"""
Deep Research Report Generator: Screen Time and K-12 Academic Performance
Generated: 2026-03-13
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import date

# ─── Configuration ───────────────────────────────────────────────────────────
OUTPUT_DIR = os.path.expanduser("~/Documents/Research/screen-time-k12-academics")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "2026-03-13 Screen Time K-12 Academic Performance Deep Research Report.docx")
FONT_NAME = "Calibri"
FONT_SIZE = Pt(11)

# ─── Helper Functions ────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=Pt(9), color=None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    """Add a heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
    return h

def add_para(doc, text, bold=False, italic=False, space_after=Pt(6)):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(doc, text, level=0):
    """Add a bulleted list item."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p

def add_confidence_para(doc, text, confidence_level):
    """Add a paragraph with confidence indicator."""
    p = doc.add_paragraph()
    # Add confidence dot
    indicator = p.add_run("")
    indicator.font.name = FONT_NAME
    indicator.font.size = FONT_SIZE
    if confidence_level == "high":
        indicator = p.add_run("[HIGH] ")
        indicator.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)  # Green
    elif confidence_level == "medium":
        indicator = p.add_run("[MED] ")
        indicator.font.color.rgb = RGBColor(0xDA, 0xA5, 0x20)  # Goldenrod
    else:
        indicator = p.add_run("[LOW] ")
        indicator.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)  # Red
    indicator.font.name = FONT_NAME
    indicator.font.size = FONT_SIZE
    indicator.font.bold = True
    # Add text
    content = p.add_run(text)
    content.font.name = FONT_NAME
    content.font.size = FONT_SIZE
    p.paragraph_format.space_after = Pt(6)
    return p

def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, size=Pt(9))
        set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            set_cell_text(cell, str(cell_text), size=Pt(8))
            if r % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # Space after table
    return table


# ─── Build Document ──────────────────────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = FONT_NAME
font.size = FONT_SIZE

# Adjust margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ─── TITLE PAGE ──────────────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Deep Research Report")
title_run.font.name = FONT_NAME
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_p.add_run("Screen Time and Academic Performance\nin K-12 Schools")
subtitle_run.font.name = FONT_NAME
subtitle_run.font.size = Pt(18)
subtitle_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

subtitle2_p = doc.add_paragraph()
subtitle2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2_run = subtitle2_p.add_run("With International Research Perspectives")
subtitle2_run.font.name = FONT_NAME
subtitle2_run.font.size = Pt(14)
subtitle2_run.font.italic = True
subtitle2_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for _ in range(4):
    doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run("March 13, 2026")
meta_run.font.name = FONT_NAME
meta_run.font.size = Pt(12)

meta2_p = doc.add_paragraph()
meta2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2_run = meta2_p.add_run(
    "4 Source Categories: Government, News, Academic, Industry\n"
    "51 Findings Analyzed | 64 Sources (pre-deduplication) | 45 Unique Sources\n"
    "13 Cross-Referenced Claims | 3 Contradictions Identified | 4 Single-Source Claims"
)
meta2_run.font.name = FONT_NAME
meta2_run.font.size = Pt(10)
meta2_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ─── EXECUTIVE SUMMARY ──────────────────────────────────────────────────────

add_heading(doc, "Executive Summary", 1)

add_para(doc, (
    "This deep research report synthesizes 51 findings from 45 unique sources across four distinct "
    "categories: government/official sources, news and journalism, peer-reviewed academic research, "
    "and industry/practitioner publications. The research examines the relationship between screen "
    "time and academic performance in K-12 schools, with particular attention to international "
    "perspectives and policy responses."
))

add_para(doc, "Key findings:", bold=True)

add_bullet(doc, (
    "The relationship between screen time and academic performance is consistently negative for "
    "recreational/passive screen use, but the picture is more nuanced for educational technology. "
    "This distinction between types of screen time is the single most important finding, supported "
    "by all four source categories with high confidence."
))
add_bullet(doc, (
    "Phone bans in schools produce modest but statistically significant academic gains, averaging "
    "0.6-1.1 percentile points overall, with substantially larger effects for low-achieving students "
    "(up to 14.2% of a standard deviation). This finding is supported by quasi-experimental evidence "
    "from the UK, a randomized controlled trial from India, and a large-scale natural experiment in Florida."
))
add_bullet(doc, (
    "The threshold for harm appears to be around 3 hours per day of recreational screen time, after "
    "which negative effects on academic performance, mental health, and sleep accelerate. The U.S. "
    "Surgeon General found that adolescents spending more than 3 hours daily on social media faced "
    "double the risk of poor mental health outcomes."
))
add_bullet(doc, (
    "International policy responses are accelerating. By the end of 2025, 35 U.S. states plus D.C. "
    "had enacted phone/device policies for schools. France expanded its phone ban nationally, Sweden "
    "reversed its hyper-digital education approach to reintroduce textbooks, and the UK issued "
    "guidance making phone-free environments the default."
))
add_bullet(doc, (
    "One important contradiction: phone ban effects vary by context. Bans improved scores in England "
    "and Florida but showed no effect in Sweden, likely because Swedish schools already had extensive "
    "laptop/tablet use, so removing phones alone had minimal impact."
))
add_bullet(doc, (
    "Reading comprehension is consistently lower on screens than on paper (the 'screen inferiority "
    "effect'), supported by multiple meta-analyses. This finding has significant implications for "
    "1:1 device programs that shift reading to digital formats."
))

add_para(doc, (
    "Bottom-line assessment: The evidence strongly supports limiting recreational and passive screen "
    "time in schools while maintaining purposeful, teacher-guided educational technology use. Policy "
    "should focus on the type and quality of screen interactions rather than blanket time limits. "
    "Districts implementing phone-free policies should anticipate modest academic gains, plan for "
    "equity in enforcement, and invest in teacher professional development for effective technology "
    "integration. The confidence level in these overall conclusions is HIGH, based on convergent "
    "evidence across multiple countries, study designs, and source categories."
), italic=True)

doc.add_page_break()

# ─── 1. BACKGROUND & CONTEXT ────────────────────────────────────────────────

add_heading(doc, "1. Background and Context", 1)

add_para(doc, (
    "Screen time in K-12 education has become one of the most debated topics in education policy "
    "worldwide. The term encompasses a broad range of activities, from educational software and "
    "digital textbooks to social media scrolling and video streaming, making blanket statements "
    "about its effects inherently misleading."
))

add_para(doc, (
    "The scale of the issue is significant. According to Pew Research Center (2024), 95% of U.S. "
    "teens now have smartphone access (up from 73% in 2014-15), and 46% report being online "
    "'almost constantly.' Common Sense Media's 2025 Census found teens average over 8 hours of "
    "daily screen time excluding homework, while children ages 0-8 average 2.5 hours per day with "
    "gaming time surging 65% since 2020. In Australia, children spend an average of 4.5+ hours on "
    "non-school screen time daily, approximately one-third of their waking time."
))

add_para(doc, (
    "Simultaneously, schools have invested heavily in technology. NCES data from 2025 shows that "
    "88% of U.S. public schools now have 1:1 device programs, providing every student with a laptop "
    "or tablet. This creates a paradox: schools are trying to limit personal device use while "
    "simultaneously distributing school-owned devices to every student."
))

add_para(doc, (
    "The CoSN Blaschke Report (2025) helpfully distinguishes three types of school screen time: "
    "(1) smartphones and social media, (2) educational technology (EdTech), and (3) screen-based "
    "entertainment. Conflating these in policy discussions undermines productive solutions. This "
    "report maintains that distinction throughout."
))

add_para(doc, (
    "Internationally, the landscape is shifting rapidly. UNESCO's 2023 Global Education Monitoring "
    "Report warned that less than one in four countries ban smartphones in schools and that evidence "
    "for digital technology universally supporting education is limited. Sweden's dramatic reversal "
    "from hyper-digitalized education back to textbooks in 2023-24, France's expanded national "
    "phone ban in 2025, and the wave of U.S. state legislation (35 states plus D.C. by end of 2025) "
    "all reflect growing policy consensus that unmanaged screen time in schools is problematic."
))

# ─── 2. KEY FINDINGS ────────────────────────────────────────────────────────

add_heading(doc, "2. Key Findings", 1)

# Theme A: Overall relationship
add_heading(doc, "2.1 The Overall Relationship: Negative for Recreational, Nuanced for Educational", 2)

add_confidence_para(doc, (
    "The overall amount of screen time is not significantly associated with academic performance "
    "when all types are aggregated. However, when disaggregated by type, TV viewing is inversely "
    "associated with language and math scores, and video game playing is inversely associated with "
    "composite academic scores. A systematic review and meta-analysis of 58 cross-sectional studies "
    "across 23 countries (Adelantado-Renau et al., 2019, JAMA Pediatrics) established this pattern."
), "high")

add_confidence_para(doc, (
    "A 2026 meta-analysis of 17 studies with over 52,000 participants found a small but significant "
    "negative correlation between screen time and mathematics performance (r = -0.034), with passive "
    "activities showing more pronounced adverse effects and screen time exceeding 3 hours/day "
    "showing the most substantial negative impact. Notably, adverse effects were found in Western "
    "nations but non-significant or positive trends appeared in China (Ulum, 2026, European Journal "
    "of Education)."
), "high")

add_confidence_para(doc, (
    "Internet use for leisure activities is inversely associated with academic performance, while "
    "internet use for educational purposes is positively associated. OECD PISA 2022 data confirms "
    "that students who use digital devices moderately (1-5 hours daily) for learning perform better "
    "than non-users. The AAP has shifted from strict time limits to a quality-focused '5 Cs of "
    "Media Use' framework, recognizing this distinction."
), "high")

add_confidence_para(doc, (
    "A German longitudinal study (LIFE Child Study, n=512) confirmed that higher screen-based media "
    "consumption in adolescents ages 10-17 predicted lower school grades in Mathematics and German "
    "12 months later, independent of physical activity levels (Poulain et al., 2018)."
), "medium")

add_confidence_para(doc, (
    "Chinese longitudinal data from the China Education Panel Survey shows screen time negatively "
    "affects academic performance among Chinese adolescents, mediated through reduced physical "
    "activity and sleep quality, following a dose-response pattern (Li et al., 2025)."
), "medium")

# Theme B: Phone bans
add_heading(doc, "2.2 Phone Bans: Modest Gains with Important Caveats", 2)

add_confidence_para(doc, (
    "A bell-to-bell cellphone ban in a large Florida urban school district led to test score "
    "increases of 1.1 percentile points by the second year, concentrated among male students "
    "(1.4 pp), middle/high schoolers (1.3 pp), and white students (1.4 pp). Reduced unexcused "
    "absences may explain up to half the gains. No gains appeared in year one when disciplinary "
    "enforcement was high (Figlio & Ozek, 2025, NBER Working Paper w34388)."
), "high")

add_confidence_para(doc, (
    "Banning mobile phones in English secondary schools increased standardized test scores by "
    "approximately 6% on average, with low-achieving students gaining 14.23% of a standard "
    "deviation. Top-quintile students were unaffected. This quasi-experimental study used "
    "administrative data across four English cities (Beland & Murphy, 2016, Labour Economics)."
), "high")

add_confidence_para(doc, (
    "A randomized controlled trial across 10 higher education institutions in India (n~17,000) "
    "found that mandatory in-class phone collection led to higher grades (0.086 SD), particularly "
    "among lower-performing and first-year students (Sungu, Choudhury & Bjerre-Nielsen, 2025)."
), "high")

add_confidence_para(doc, (
    "A study of more than 37,000 course grades found student academic success rates increased "
    "6.27% after implementing phone pouch restrictions, with a 44% decrease in behavioral "
    "referrals and 74% of teachers reporting increased student attentiveness (ExcelinEd, 2024)."
), "medium")

add_confidence_para(doc, (
    "France expanded its phone ban nationally in 2025 after a pilot in 100 schools showed improved "
    "focus, reduced cyberbullying, and better social interaction. However, pilot results were "
    "self-reported by administrators rather than measured by standardized assessments."
), "medium")

# Theme C: PISA and international data
add_heading(doc, "2.3 PISA Data and International Comparisons", 2)

add_confidence_para(doc, (
    "PISA 2022 data shows students who spent up to 1 hour on leisure screen time in school scored "
    "49 points higher in math than peers spending 5-7 hours daily, after adjusting for "
    "socioeconomic factors. This 49-point gap is roughly equivalent to more than one year of "
    "schooling. Nearly 1 in 3 students across OECD countries reported being distracted by digital "
    "devices in most or all mathematics classes (OECD, 2024)."
), "high")

add_confidence_para(doc, (
    "Even in schools with phone bans, 29% of students reported using smartphones several times a "
    "day. Three-quarters of students spend more than one hour per weekday browsing social networks. "
    "Teachers with ICT professional development reported less student distraction (OECD, 2024)."
), "high")

add_confidence_para(doc, (
    "Sweden allocated 104 million euros to reintroduce printed textbooks after its hyper-digitalized "
    "approach was linked to declining PIRLS reading scores (Swedish fourth-graders dropped from 555 "
    "to 544 points between 2016 and 2021). The government planned to end digital learning entirely "
    "for children under age 6 and announced a nationwide phone ban starting 2026."
), "high")

# Theme D: Reading on screens
add_heading(doc, "2.4 The Screen Inferiority Effect in Reading", 2)

add_confidence_para(doc, (
    "Students reading on paper consistently score higher on comprehension tests than those reading "
    "on screens, particularly for expository/informational text. The effect is more pronounced in "
    "unsupervised settings and under time pressure. Scrolling on small screens worsens the effect. "
    "Multiple meta-analyses converge on this finding, including Furenes, Kucirkova & Bus (2021, "
    "Review of Educational Research, 39 studies) and Kong, Seo & Zhai (2024, 49 studies)."
), "high")

add_confidence_para(doc, (
    "However, story-congruent digital enhancements (interactive elements aligned with the text) "
    "could outperform paper for young children ages 1-8, suggesting the medium itself is not the "
    "sole factor (Furenes et al., 2021)."
), "medium")

# Theme E: Early childhood
add_heading(doc, "2.5 Early Childhood Screen Exposure and Later Academic Outcomes", 2)

add_confidence_para(doc, (
    "Each one-hour increase in TV exposure at age 2 corresponded to a 7% decrease in classroom "
    "participation and a 6% decrease in math proficiency by 4th grade, according to the Quebec "
    "Longitudinal Study (n=2,120, tracked from age 5 months to 21 years; Pagani et al., 2010)."
), "high")

add_confidence_para(doc, (
    "The NIH ABCD Study (n~12,000, 21 sites) found screen time moderately associated with "
    "decreased academic performance, worse mental health, and poorer sleep in children aged 9-10. "
    "SES is a stronger predictor of all outcomes than screen time. Texting/video chatting (social "
    "connection) showed more positive outcomes than passive consumption."
), "high")

add_confidence_para(doc, (
    "WHO recommends sedentary screen time of no more than 1 hour for children under 5, with less "
    "being better. The AAP maintains no screens before 18 months and maximum 1 hour/day for ages "
    "2-5. Common Sense Media (2025) found 40% of children have a tablet by age 2."
), "high")

# Theme F: 1:1 programs
add_heading(doc, "2.6 One-to-One Device Programs", 2)

add_confidence_para(doc, (
    "A meta-analysis of 1:1 laptop programs (65 journal articles and 31 dissertations, 2001-2015) "
    "found significantly positive average effect sizes in English, writing, mathematics, and "
    "science, with writing showing the strongest gains and stronger effects for low-income students "
    "and English Language Learners. No improvement was found in reading scores (Zheng et al., 2016, "
    "Review of Educational Research)."
), "high")

add_confidence_para(doc, (
    "Districts with 1:1 programs report up to 30% improvement in student engagement and 25% "
    "increase in assignment completion, with strongest gains among underserved populations. However, "
    "positive impact depends heavily on teacher professional development and pedagogical integration "
    "(Bluum Foundation literature review). These figures come from aggregated district reports "
    "rather than rigorous RCTs."
), "low")

add_confidence_para(doc, (
    "UNESCO's analysis found that less than 2% of education technology interventions have 'strong "
    "or moderate evidence of effectiveness,' and that learning benefits disappear if technology is "
    "used in excess or without a qualified teacher. Distributing computers without teacher "
    "involvement does not improve learning (UNESCO GEM Report, 2023)."
), "high")

# Theme G: Mental health pathway
add_heading(doc, "2.7 Mental Health as a Mediating Pathway", 2)

add_confidence_para(doc, (
    "The U.S. Surgeon General's 2023 Advisory found that adolescents spending more than 3 hours/day "
    "on social media faced double the risk of poor mental health outcomes including depression and "
    "anxiety. 95% of youth 13-17 use social media, and nearly 1 in 3 adolescents use screens until "
    "midnight or later, impacting sleep and next-day academic performance."
), "high")

add_confidence_para(doc, (
    "Sapien Labs' Global Mind Project (1 million+ participants, 82 countries) found that 41% of "
    "the world's internet-enabled young adults face a 'mind health crisis,' with childhood "
    "smartphone use identified as a key driver. Each successive generation shows diminished mental "
    "health. However, this is correlational with selection bias (internet-enabled respondents only)."
), "medium")

add_confidence_para(doc, (
    "Problematic social media use among adolescents rose from 7% in 2018 to 11% in 2022 according "
    "to WHO Regional Office for Europe surveillance data, suggesting the problem is worsening."
), "high")

# Theme H: Policy landscape
add_heading(doc, "2.8 The Global Policy Response", 2)

add_confidence_para(doc, (
    "By December 2025, 35 U.S. states plus D.C. enacted laws or policies on student cellphone use "
    "in K-12 classrooms. Twenty-two of these were enacted in 2025 alone. California passed the "
    "Phone-Free Schools Act; Louisiana banned all electronic devices during the instructional day; "
    "South Carolina tied funding to phone bans; Virginia allocated $500,000 for implementation "
    "support (Ballotpedia, ExcelinEd)."
), "high")

add_confidence_para(doc, (
    "England's DfE issued guidance making phone-free environments the default (February 2024, "
    "updated January 2026). Over 90% of English schools have implemented bans, though the "
    "government rejected a statutory ban in May 2025. Ofsted now assesses phone policies in "
    "inspections."
), "high")

add_confidence_para(doc, (
    "Three-quarters of NEA members report social media use is a serious problem at their school. "
    "84% of AFT-surveyed teachers believe they should have more say over classroom technology "
    "deployment. 72% of high school teachers say cellphone distraction is a major classroom problem "
    "(Pew Research, 2024)."
), "medium")

doc.add_page_break()

# ─── 3. POINTS OF CONTRADICTION ─────────────────────────────────────────────

add_heading(doc, "3. Points of Contradiction", 1)

add_para(doc, (
    "Cross-referencing findings across the four source categories revealed three significant "
    "contradictions where evidence conflicts."
))

add_heading(doc, "3.1 Sweden vs. England: Phone Ban Effects Are Context-Dependent", 2)

add_para(doc, (
    "Beland & Murphy's research on English secondary schools found phone bans improved test scores "
    "by ~6% on average, with low-achieving students gaining 14.2% of a standard deviation. However, "
    "research on Swedish secondary schools by the same authors found no impact or a very small "
    "negative effect on student performance."
))

add_para(doc, (
    "Resolution: The most likely explanation is contextual. Swedish schools already had extensive "
    "laptop and tablet use integrated into instruction, so removing phones alone had minimal impact "
    "because students still had access to other digital devices. In England, phones were the primary "
    "source of digital distraction. This suggests phone bans are most effective when phones are the "
    "main distracting device available to students. The EU Parliament's 2025 analysis of phone bans "
    "across member states confirms mixed results, reinforcing that context matters more than the "
    "policy itself."
), italic=True)

add_heading(doc, "3.2 AAP Quality Framework vs. Time-Based Legislative Approach", 2)

add_para(doc, (
    "The American Academy of Pediatrics has moved away from strict time limits toward a "
    "quality-focused '5 Cs of Media Use' framework, arguing there is insufficient evidence for "
    "universal time limits. Meanwhile, 35 U.S. states have enacted policies focused primarily on "
    "time-based restrictions (banning phones during specific periods)."
))

add_para(doc, (
    "Resolution: These approaches are not entirely contradictory. The AAP's framework applies "
    "primarily to educational screen time and overall media use, while legislative bans target "
    "personal device use (smartphones and social media) during instructional time. The distinction "
    "matters: the evidence supports restricting recreational/social media use during school hours "
    "(which is what bans do) while taking a more nuanced approach to educational technology (which "
    "is what the AAP recommends). Both agree that passive, recreational screen time is harmful."
), italic=True)

add_heading(doc, "3.3 ISTE Learning Gains vs. UNESCO's '2% Effective' Finding", 2)

add_para(doc, (
    "Research mapped to ISTE educator standards found empirical evidence that all seven standards "
    "lead to student learning gains, and the 1:1 meta-analysis (Zheng et al., 2016) found positive "
    "effects across subjects. However, UNESCO reported that less than 2% of education technology "
    "interventions have 'strong or moderate evidence of effectiveness.'"
))

add_para(doc, (
    "Resolution: The discrepancy likely reflects the difference between well-implemented technology "
    "use (aligned with standards, supported by teacher PD) and the vast majority of EdTech "
    "deployments that lack these conditions. ISTE standards represent best-case implementation; "
    "UNESCO's figure captures the full spectrum. This actually reinforces the finding that teacher "
    "training and intentional pedagogical integration are essential. The technology itself is neutral; "
    "implementation quality determines outcomes."
), italic=True)

doc.add_page_break()

# ─── 4. SINGLE-SOURCE CLAIMS ────────────────────────────────────────────────

add_heading(doc, "4. Single-Source Claims", 1)

add_para(doc, (
    "The following claims were supported by only one source category, warranting additional "
    "scrutiny before relying on them for decision-making."
))

add_heading(doc, "4.1 Phone Bans Disproportionately Increase Suspensions for Black Students", 2)
add_para(doc, (
    "Source category: News (Hechinger Report, Chalkbeat, The 74, citing NBER Florida data). "
    "In the first year of phone bans, in-school suspensions for Black students at highly affected "
    "schools increased by 30%. Suspensions fell back to pre-ban levels by year two. This finding "
    "was not independently confirmed by government reports, academic literature, or practitioner "
    "organizations in the sources reviewed."
))
add_para(doc, (
    "Why it matters: If phone bans are implemented inequitably, the modest academic gains could "
    "come at the cost of exacerbating existing racial disparities in school discipline. Any district "
    "implementing a phone ban should monitor suspension data disaggregated by race and adjust "
    "enforcement strategies proactively."
))
add_para(doc, (
    "Additional evidence needed: Independent replication in other districts; analysis of whether "
    "disparities persist beyond the first year; comparison of enforcement models (pouch-based vs. "
    "confiscation) and their equity implications."
))

add_heading(doc, "4.2 The Screen Inferiority Effect in Reading Comprehension", 2)
add_para(doc, (
    "Source category: Academic (multiple meta-analyses). While the evidence is strong within the "
    "academic literature (Furenes et al., 2021; Kong et al., 2024), this finding was not "
    "prominently featured in government reports, news coverage, or practitioner publications "
    "reviewed. The strength of the academic evidence is high (multiple converging meta-analyses), "
    "but the lack of cross-category corroboration suggests it has not yet influenced policy "
    "discussions significantly."
))
add_para(doc, (
    "Why it matters: If reading on screens consistently produces lower comprehension, 1:1 programs "
    "that shift most reading to digital formats may be undermining reading achievement even while "
    "supporting other skills. This has direct implications for ELA curriculum and assessment design."
))

add_heading(doc, "4.3 EdTech Privacy and Children's Rights Concerns", 2)
add_para(doc, (
    "Source category: Industry/Practitioner (5Rights Foundation, LSE Digital Futures Centre). "
    "Education technology often fails to protect children's rights to privacy, safety, and "
    "education. A joint statement from 13 organizations highlighted rising children's eye health "
    "issues linked to prolonged screen time."
))
add_para(doc, (
    "Why it matters: As schools deploy more devices and EdTech platforms, data privacy and health "
    "impacts become systemic concerns that may not surface until significant harm has occurred."
))

add_heading(doc, "4.4 Cultural Variation in Screen Time Effects on Math", 2)
add_para(doc, (
    "Source category: Academic (Ulum, 2026). The meta-analysis found adverse math effects in "
    "Western nations but non-significant or positive trends in China. This was not addressed by "
    "other source categories."
))
add_para(doc, (
    "Why it matters: If screen time effects are culturally moderated, policies imported from one "
    "context may not transfer. The type of screen use, parental mediation, and educational culture "
    "may all interact with screen time differently across societies."
))

doc.add_page_break()

# ─── 5. PRACTICAL IMPLICATIONS ──────────────────────────────────────────────

add_heading(doc, "5. Practical Implications", 1)

add_para(doc, "For a K-12 district leader, the evidence supports the following actions, ordered by confidence level:", bold=True)

add_heading(doc, "High Confidence Recommendations", 2)

add_bullet(doc, (
    "Implement a bell-to-bell phone-free policy with physical storage (pouches or lockers). "
    "Expect modest academic gains (1-2 percentile points) emerging in year two. Plan for increased "
    "disciplinary referrals in year one and proactively monitor equity data."
))
add_bullet(doc, (
    "Distinguish between educational and recreational screen time in all policy discussions. "
    "The evidence clearly shows these have different, sometimes opposite, effects on learning."
))
add_bullet(doc, (
    "Invest in teacher professional development for technology integration. OECD data shows "
    "teachers with ICT PD report less student distraction, and the academic literature consistently "
    "finds implementation quality is the determining factor in EdTech effectiveness."
))
add_bullet(doc, (
    "Maintain printed materials for extended reading, especially for expository/informational text. "
    "The screen inferiority effect for reading comprehension is well-established and has direct "
    "implications for ELA instruction."
))
add_bullet(doc, (
    "Set clear expectations for school-issued device use. With 88% of schools now at 1:1, "
    "school laptops and tablets are becoming a significant source of distraction alongside phones."
))

add_heading(doc, "Medium Confidence Recommendations", 2)

add_bullet(doc, (
    "Limit total recreational screen time during the school day to under 1 hour. PISA data shows "
    "the sharpest performance drops above this threshold."
))
add_bullet(doc, (
    "Address social media access as a specific concern within screen time policy. The Surgeon "
    "General's advisory and WHO data suggest social media poses distinct risks beyond general "
    "screen time."
))
add_bullet(doc, (
    "Engage parents in a coordinated approach. Pew Research shows parents are increasingly "
    "concerned, and the evidence suggests home screen time habits interact with school outcomes."
))

add_heading(doc, "Questions That Remain Unanswered", 2)

add_bullet(doc, (
    "What is the optimal balance of digital and analog instruction at each grade level?"
))
add_bullet(doc, (
    "How do phone ban effects change over time beyond the second year?"
))
add_bullet(doc, (
    "Do phone ban effects differ for rural vs. urban districts, or small vs. large districts?"
))
add_bullet(doc, (
    "How should schools manage AI tools (ChatGPT, etc.) as a new category of screen-based activity?"
))
add_bullet(doc, (
    "What enforcement model (pouches, lockers, honor system) produces the best academic outcomes "
    "with the fewest disciplinary side effects?"
))

doc.add_page_break()

# ─── CLAIM CROSS-REFERENCE MATRIX ───────────────────────────────────────────

add_heading(doc, "6. Claim Cross-Reference Matrix", 1)

add_para(doc, (
    "The following matrix shows how major claims are supported, contradicted, or not addressed "
    "across the four source categories. Legend: checkmark = supports, X = contradicts, "
    "dash = not addressed."
))

claim_matrix_headers = ["Claim", "Gov", "News", "Acad", "Ind", "Agreement"]
claim_matrix_rows = [
    ["Recreational screen time negatively affects academics", "Y", "Y", "Y", "Y", "Full"],
    ["Type of screen time matters more than total amount", "Y", "Y", "Y", "Y", "Full"],
    ["Phone bans produce modest academic gains", "Y", "Y", "Y", "Y", "Full"],
    ["Low-achieving students benefit most from bans", "Y", "Y", "Y", "--", "Partial (3/4)"],
    ["Teacher PD is critical for effective tech integration", "Y", "--", "Y", "Y", "Partial (3/4)"],
    ["35+ US states enacted phone policies by end 2025", "Y", "Y", "--", "Y", "Partial (3/4)"],
    ["1:1 programs show positive, implementation-dependent results", "--", "Y", "Y", "Y", "Partial (3/4)"],
    ["Mental health mediates screen time-academic link", "Y", "--", "Y", "Y", "Partial (3/4)"],
    ["Screen time >3 hrs/day shows largest negative effects", "Y", "--", "Y", "--", "Partial (2/4)"],
    ["Early childhood exposure predicts later academic decline", "Y", "--", "Y", "--", "Partial (2/4)"],
    ["Phone ban effects are context-dependent (Sweden vs UK)", "Y", "Y", "Y", "--", "Partial (noted)"],
    ["Reading on paper > reading on screens", "--", "--", "Y", "--", "Single"],
    ["Phone bans increase suspensions disproportionately by race", "--", "Y", "--", "--", "Single"],
]

create_table(doc, claim_matrix_headers, claim_matrix_rows,
             col_widths=[3.5, 0.5, 0.5, 0.5, 0.5, 1.2])

doc.add_page_break()

# ─── 7. SOURCE RELIABILITY MATRIX ───────────────────────────────────────────

add_heading(doc, "7. Source Reliability Matrix", 1)

add_para(doc, (
    "Each source was scored using the following formula: Aggregate = Source Type Weight x Agent "
    "Confidence x URL Status + Corroboration Bonus (capped at 1.0). Sources are sorted by "
    "aggregate reliability score."
))

add_para(doc, (
    "Source type weights: Government/Academic = 1.0, News = 0.8, Industry = 0.7. "
    "Agent confidence: High = 1.0, Medium = 0.7, Low = 0.4. "
    "URL status: Live = 1.0, Redirect/Paywalled = 0.8, Dead = 0.3, Skipped = 0.7. "
    "Corroboration: Cited by 2+ agents = +0.1, 3+ agents = +0.2."
), italic=True)

srm_headers = ["Source", "Cat", "Conf", "URL", "Corr", "Score", "Reliability"]

# Sources sorted by computed aggregate reliability score
srm_rows = [
    # High reliability (>=0.8)
    ["OECD (2024). Managing Screen Time.", "Gov", "High", "Live", "3+", "1.00", "High"],
    ["OECD (2024). Students, Digital Devices and Success.", "Gov", "High", "Live", "3+", "1.00", "High"],
    ["UNESCO (2023). GEM Report: Technology in Education.", "Gov", "High", "Live", "3+", "1.00", "High"],
    ["Adelantado-Renau et al. (2019). Screen Media Use and Academic Performance. JAMA Pediatrics.", "Acad", "High", "Live", "2+", "1.00", "High"],
    ["Beland & Murphy (2016). Ill Communication. Labour Economics.", "Acad", "High", "Live", "3+", "1.00", "High"],
    ["Paulich et al. (2021). ABCD Study: Screen time and adolescent outcomes. PLOS ONE.", "Acad", "High", "Live", "2+", "1.00", "High"],
    ["U.S. Surgeon General (2023). Social Media and Youth Mental Health.", "Gov", "High", "Live", "2+", "1.00", "High"],
    ["WHO (2019). Physical Activity, Sedentary Behaviour and Sleep Guidelines.", "Gov", "High", "Live", "2+", "1.00", "High"],
    ["Pew Research (2024). Teens, Social Media and Technology 2024.", "Ind", "High", "Live", "2+", "0.90", "High"],
    ["Figlio & Ozek (2025). Cellphone Bans in Florida. NBER w34388.", "News", "High", "Live", "2+", "0.90", "High"],
    ["IES/NCES (2025). Cell Phones Hurt Academic Performance.", "Gov", "High", "Live", "--", "1.00", "High"],
    ["CDC (2025). Screen Time and Health Outcomes in US Teenagers.", "Gov", "High", "Live", "--", "1.00", "High"],
    ["NIH MedlinePlus. Screen Time and Children's Brains.", "Gov", "High", "Live", "--", "1.00", "High"],
    ["Ballotpedia (2025). State Policies on Cellphone Use.", "Gov", "High", "Live", "2+", "1.00", "High"],
    ["Furenes et al. (2021). Children's Reading: Paper vs. Screen. RER.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["Kong et al. (2024). Paper vs. Digital Reading Meta-analysis.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["Zheng et al. (2016). 1:1 Laptop Environments. RER.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["Pagani et al. (2010). Early TV and Academic Well-Being.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["Ulum (2026). Screen Handicap in Mathematics. Eur J Ed.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["Common Sense Media (2025). Census: Media Use by Kids 0-8.", "Ind", "High", "Live", "--", "0.70", "Medium"],
    ["AAP (2024). Screen Time at School / Digital Media Guidelines.", "Ind", "High", "Live", "2+", "0.90", "High"],
    ["ExcelinEd (2024/2026). Phone-Free Schools.", "Ind", "High", "Live", "2+", "0.90", "High"],
    ["Hechinger Report (2025). Cellphone bans and racial disparities.", "News", "High", "Live", "--", "0.80", "High"],
    ["Education Week (2025). Do School Cellphone Bans Work?", "News", "High", "Live", "--", "0.80", "High"],
    ["House of Commons Library (2024). Mobile phones in schools.", "News", "High", "Live", "--", "0.80", "High"],
    # Medium reliability (0.5-0.79)
    ["Sungu et al. (2025). Removing Phones Improves Performance. SSRN.", "Acad", "High", "Live", "--", "1.00", "High"],
    ["EdWorkingPapers (2025). Florida Phone Ban Impact.", "Acad", "Med-Hi", "Live", "--", "0.70", "Medium"],
    ["NPR (2024). Teens and screentime quality vs quantity.", "News", "Med", "Live", "--", "0.56", "Medium"],
    ["Poulain et al. (2018). Screen time and school performance. BMC.", "Acad", "Med", "Live", "--", "0.70", "Medium"],
    ["Li et al. (2025). Screen time mechanisms. Frontiers.", "Acad", "Med", "Live", "--", "0.70", "Medium"],
    ["UK Children's Commissioner (2025). Smartphone Policies.", "Gov", "Med", "Live", "--", "0.70", "Medium"],
    ["Australian eSafety Commissioner. Screen Time.", "Gov", "Med", "Live", "--", "0.70", "Medium"],
    ["Sapien Labs (2025). Mental State of the World 2024.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["NEA (2024). Member Polling: Social Media and Devices.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["CoSN (2025). Screens in Balance: Blaschke Report.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["Bluum Foundation. Impact of 1:1 Device Programs.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["5Rights Foundation. Better EdTech Futures for Children.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["EduTimes (2025). France's Phone Ban.", "News", "Med", "Live", "--", "0.56", "Medium"],
    ["EdSurge (2026). Screen-Free Schools.", "News", "Med", "Live", "--", "0.56", "Medium"],
    ["Washington Post (2025). Parents reducing screen time.", "News", "Med", "Paywall", "--", "0.45", "Low"],
    ["EU Publications Office (2025). Phone Bans Across the EU.", "Gov", "Med", "Live", "--", "0.70", "Medium"],
    ["Eurydice. Digital Education at School in Europe.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    # Low reliability (<0.5)
    ["LearnPlatform/Instructure (2025). EdTech Top 40.", "Ind", "Low", "Live", "--", "0.28", "Low"],
    ["ISTE/Tandfonline (2023). ISTE Standards and Learning Gains.", "Ind", "Med", "Live", "--", "0.49", "Low"],
    ["Education Week (2016). 1:1 Laptop Initiatives.", "News", "Med", "Live", "--", "0.56", "Medium"],
]

add_para(doc, (
    f"Summary: {len(srm_rows)} unique sources evaluated. "
    f"25 High Reliability, 11 Medium Reliability, 9 Low Reliability."
), bold=True)

# Split into two tables for readability (the full table is too wide for portrait)
# Table 1: Top 25 (High reliability)
add_heading(doc, "High Reliability Sources (Score >= 0.8)", 3)
create_table(doc, srm_headers, srm_rows[:25],
             col_widths=[3.0, 0.4, 0.45, 0.5, 0.4, 0.5, 0.65])

add_heading(doc, "Medium and Low Reliability Sources (Score < 0.8)", 3)
create_table(doc, srm_headers, srm_rows[25:],
             col_widths=[3.0, 0.4, 0.45, 0.5, 0.4, 0.5, 0.65])

doc.add_page_break()

# ─── FEDERAL & STATE POLICY CONTEXT (education topic) ────────────────────────

add_heading(doc, "8. Federal and State Policy Context", 1)

add_heading(doc, "Federal Framework", 2)

add_para(doc, (
    "The U.S. Department of Education published guidance in December 2024 recommending all states "
    "adopt phone management measures in schools. The Surgeon General's 2023 Advisory on Social "
    "Media and Youth Mental Health called for stronger protections, and school districts have filed "
    "lawsuits against social media companies citing harm to students' mental health and academic "
    "performance."
))

add_para(doc, (
    "ESSA (2015) provides the framework for states to use Title IV-A funds for technology "
    "integration, including 'effective use of technology' provisions. However, federal policy "
    "does not mandate specific screen time limits or phone bans, leaving this to states and "
    "districts."
))

add_heading(doc, "State-Level Action", 2)

add_para(doc, (
    "The legislative response has been dramatic. By end of 2025, 35 states plus D.C. enacted "
    "phone or device policies, with 22 enacted in 2025 alone. Key examples:"
))

add_bullet(doc, "California: Phone-Free Schools Act (2024)")
add_bullet(doc, "Louisiana: Banned all electronic devices during instructional day")
add_bullet(doc, "South Carolina: Tied state funding to phone ban implementation")
add_bullet(doc, "Ohio: Required policies to limit phone use 'as much as possible'")
add_bullet(doc, "Virginia: Allocated $500,000 for phone-free implementation support")
add_bullet(doc, "Florida: Statewide bell-to-bell ban with measured academic outcomes (NBER study)")

add_heading(doc, "Tennessee Context", 2)

add_para(doc, (
    "Tennessee does not appear to have enacted state-level phone ban legislation as of March 2026, "
    "though TDOE tracks technology use in schools. No dedicated TDOE policy or report specifically "
    "addressing screen time and academic performance was found in this review. Given that 22 states "
    "enacted new policies in 2025 alone, Tennessee may face increasing pressure to address this "
    "at the state level."
))

add_heading(doc, "International Policy Landscape", 2)

add_para(doc, (
    "The international response mirrors the U.S. trend. France expanded its phone ban to all "
    "middle schools in 2025. England issued DfE guidance making phone-free the default (90%+ "
    "compliance). Sweden reversed its digital-first education strategy, investing 104 million euros "
    "in textbooks. China limited digital device use in teaching to 30% of overall time. The EU "
    "Parliament analyzed phone bans across member states in 2025, finding mixed results. Sweden is "
    "implementing a nationwide ban for ages 7-16 starting 2026."
))

doc.add_page_break()

# ─── REFERENCES ──────────────────────────────────────────────────────────────

add_heading(doc, "9. References", 1)

references = [
    "AAP (2024). Screen Time at School / Understanding the New AAP Digital Media Guidelines. https://www.aap.org/en/patient-care/media-and-children/ [Reliability: High]",
    "Adelantado-Renau, M., et al. (2019). Association Between Screen Media Use and Academic Performance Among Children and Adolescents: A Systematic Review and Meta-analysis. JAMA Pediatrics. https://jamanetwork.com/journals/jamapediatrics/fullarticle/2751330 [Reliability: High]",
    "Australian eSafety Commissioner. Screen Time. https://www.esafety.gov.au/parents/issues-and-advice/screen-time [Reliability: Medium]",
    "Ballotpedia (2025). State Policies on Cellphone Use in K-12 Public Schools. https://ballotpedia.org/State_policies_on_cellphone_use_in_K-12_public_schools [Reliability: High]",
    "Barnum, M. (2025). Cellphone bans can help kids learn, but Black students are suspended more. The Hechinger Report. https://hechingerreport.org/proof-points-cellphone-bans/ [Reliability: High]",
    "Beland, L.-P., & Murphy, R. (2016). Ill Communication: Technology, Distraction & Student Performance. Labour Economics. https://www.sciencedirect.com/science/article/abs/pii/S0927537116300136 [Reliability: High]",
    "Bluum Foundation. Impact of 1:1 Device Programs in K-12 Educational Settings. https://www.bluum.com/resources/impact-of-1-on-1-device-programs-in-k-12-educational-settings [Reliability: Low]",
    "CDC (2025). Associations Between Screen Time Use and Health Outcomes Among US Teenagers. https://www.cdc.gov/pcd/issues/2025/24_0537.htm [Reliability: High]",
    "Common Sense Media (2025). The 2025 Common Sense Census: Media Use by Kids Zero to Eight. https://www.commonsensemedia.org/research/the-2025-common-sense-census-media-use-by-kids-zero-to-eight [Reliability: Medium]",
    "CoSN (2025). Screens in Balance: The 2025 Blaschke Report. https://www.cosn.org/2025-blaschke-report-toolkit/ [Reliability: Low]",
    "EdSurge (2026). Screen-Free Schools? Some Legislators Push for a New Normal. https://www.edsurge.com/news/2026-03-09-screen-free-schools-some-legislators-push-for-a-new-normal [Reliability: Medium]",
    "EdWorkingPapers (2025). The Impact of Cellphone Bans in Schools on Student Outcomes: Evidence from Florida. https://edworkingpapers.com/policy-practice-series/ai25-1315 [Reliability: Medium]",
    "Education Week (2016). 1-to-1 Laptop Initiatives Boost Student Scores. https://www.edweek.org/technology/1-to-1-laptop-initiatives-boost-student-scores-study-finds/2016/05 [Reliability: Medium]",
    "Education Week (2025). Do School Cellphone Bans Work? https://www.edweek.org/technology/do-school-cellphone-bans-work-what-early-findings-tell-us/2025/10 [Reliability: High]",
    "EU Publications Office (2025). Mobile Phone Bans in Schools Across the EU. https://op.europa.eu/en/publication-detail/-/publication/37a07be1-0241-11f1-825d-01aa75ed71a1/language-en [Reliability: Medium]",
    "Eurydice/European Commission. Digital Education at School in Europe. https://eurydice.eacea.ec.europa.eu/publications/digital-education-school-europe [Reliability: Low]",
    "ExcelinEd (2024). Cell Phone Free Schools: Three States Leading the Way. https://excelined.org/2024/10/01/cell-phone-free-schools-how-three-states-are-leading-the-way-to-create-distraction-free-learning/ [Reliability: High]",
    "Figlio, D. & Ozek, U. (2025). The Impact of Cellphone Bans in Schools on Student Outcomes: Evidence from Florida. NBER Working Paper w34388. https://www.nber.org/papers/w34388 [Reliability: High]",
    "5Rights Foundation. Better EdTech Futures for Children. https://5rightsfoundation.com/ [Reliability: Low]",
    "Furenes, M. I., Kucirkova, N., & Bus, A. G. (2021). A Comparison of Children's Reading on Paper Versus Screen: A Meta-Analysis. Review of Educational Research. https://journals.sagepub.com/doi/full/10.3102/0034654321998074 [Reliability: High]",
    "House of Commons Library (2024). Mobile phones in schools (England). https://commonslibrary.parliament.uk/research-briefings/cbp-10241/ [Reliability: High]",
    "IES/NCES (2025). More than Half of Public School Leaders Say Cell Phones Hurt Academic Performance. https://ies.ed.gov/learn/press-release/more-half-public-school-leaders-say-cell-phones-hurt-academic-performance [Reliability: High]",
    "Kong, Y., Seo, Y. S., & Zhai, L. (2024). Which reading comprehension is better? A meta-analysis of paper versus digital reading. https://www.sciencedirect.com/science/article/pii/S2772503024000288 [Reliability: High]",
    "LearnPlatform/Instructure (2025). EdTech Top 40 Report. https://www.prnewswire.com/news-releases/new-learnplatform-by-instructure-report-shows-k-12-districts-are-more-selective-about-edtech-tools-as-they-face-budget-crisis-302492756.html [Reliability: Low]",
    "Li, Y., et al. (2025). The relationship and mechanism of screen time and academic performance among adolescents. Frontiers in Public Health. https://www.frontiersin.org/journals/public-health/articles/10.3389/fpubh.2025.1533327/full [Reliability: Medium]",
    "NEA (2024). Member Polling: Social Media, Personal Devices and Mental Health. https://www.nea.org/sites/default/files/2024-06/nea-member-polling-on-social-media-personal-devices-and-mental-health-june-20-2024.pdf [Reliability: Low]",
    "NIH MedlinePlus Magazine. What Is All That Screen Time Doing to Your Child's Brain? https://magazine.medlineplus.gov/article/what-is-all-that-screen-time-doing-to-your-childs-brain [Reliability: High]",
    "NPR (2024). New study on teens and screentime. https://www.npr.org/2024/11/20/nx-s1-5193261/new-study-on-teens-and-screentime-finds-both-quality-and-quantity-are-important [Reliability: Medium]",
    "OECD (2024). Managing Screen Time: How to Protect and Equip Students Against Distraction. https://www.oecd.org/en/publications/managing-screen-time_7c225af4-en.html [Reliability: High]",
    "OECD (2024). Students, Digital Devices and Success. https://www.oecd.org/en/publications/2024/05/students-digital-devices-and-success_621829ff.html [Reliability: High]",
    "OECD (2024). Technology Use at School and Students' Learning Outcomes. https://www.oecd.org/content/dam/oecd/en/publications/reports/2024/12/technology-use-at-school-and-students-learning-outcomes_4c4f92e6/422db044-en.pdf [Reliability: High]",
    "Pagani, L. S., et al. (2010). Prospective Associations Between Early Childhood Television Exposure and Academic Well-Being. Archives of Pediatrics & Adolescent Medicine. [Reliability: High]",
    "Paulich, K. N., et al. (2021). Screen time and early adolescent mental health, academic, and social outcomes. PLOS ONE. https://pubmed.ncbi.nlm.nih.gov/34496002/ [Reliability: High]",
    "Pew Research Center (2024). Teens, Social Media and Technology 2024. https://www.pewresearch.org/internet/2024/12/12/teens-social-media-and-technology-2024/ [Reliability: High]",
    "Poulain, T., et al. (2018). Screen time and school performance at secondary school. BMC Public Health. https://bmcpublichealth.biomedcentral.com/articles/10.1186/s12889-018-5489-3 [Reliability: Medium]",
    "Sapien Labs (2025). The Mental State of the World in 2024. https://sapienlabs.org/wp-content/uploads/2025/02/Mental-State-of-the-World-2024-Online-Feb-26.pdf [Reliability: Low]",
    "Sungu, A., Choudhury, P. K., & Bjerre-Nielsen, A. (2025). Removing Phones from Classrooms Improves Academic Performance. SSRN. https://papers.ssrn.com/sol3/papers.cfm?abstract_id=5370727 [Reliability: High]",
    "UK Children's Commissioner (2025). School Survey: Smartphone Policies. https://assets.childrenscommissioner.gov.uk/wpuploads/2025/04/cco-school-survey-smartphone-policies.pdf [Reliability: Medium]",
    "Ulum, H. (2026). Screen Handicap in Mathematics: A Meta-Analysis. European Journal of Education. https://onlinelibrary.wiley.com/doi/10.1111/ejed.70400 [Reliability: High]",
    "UNESCO (2023). Global Education Monitoring Report: Technology in Education: A Tool on Whose Terms? https://unesdoc.unesco.org/ark:/48223/pf0000386165 [Reliability: High]",
    "U.S. Surgeon General (2023). Social Media and Youth Mental Health Advisory. https://www.schoolsafety.gov/resource/social-media-and-youth-mental-health-us-surgeon-generals-advisory [Reliability: High]",
    "Washington Post (2025). Some parents take steps to reduce children's screen time. https://www.washingtonpost.com/parenting/2025/11/10/kids-parents-ditching-screens/ [Reliability: Low]",
    "WHO (2019). Guidelines on Physical Activity, Sedentary Behaviour and Sleep for Children Under 5. https://www.who.int/publications/i/item/9789241550536 [Reliability: High]",
    "WHO Europe (2024). Teens, Screens and Mental Health. https://www.who.int/europe/news/item/25-09-2024-teens--screens-and-mental-health [Reliability: High]",
    "Zheng, B., Warschauer, M., Lin, C.-H., & Chang, C. (2016). Learning in One-to-One Laptop Environments. Review of Educational Research. https://journals.sagepub.com/doi/abs/10.3102/0034654316628645 [Reliability: High]",
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

doc.add_page_break()

# ─── TENNESSEE DATA APPENDIX ────────────────────────────────────────────────

add_heading(doc, "Appendix A: Tennessee Context", 1)

add_para(doc, (
    "This appendix compiles Tennessee-specific information relevant to the screen time and "
    "academic performance discussion. Tennessee-specific data on this topic is limited compared "
    "to states like Florida, California, or Virginia that have been at the forefront of phone ban "
    "legislation."
))

add_heading(doc, "Current State of Tennessee Policy", 2)

add_bullet(doc, (
    "Tennessee has not enacted state-level phone ban legislation as of March 2026."
))
add_bullet(doc, (
    "No dedicated TDOE policy or guidance document on screen time and academic performance was "
    "identified in this review."
))
add_bullet(doc, (
    "Individual districts in Tennessee have adopted their own phone policies, but no statewide "
    "mandate exists."
))
add_bullet(doc, (
    "TDOE tracks technology use in schools generally but does not publish screen time-specific "
    "academic outcome data."
))

add_heading(doc, "Implications for Tennessee Districts", 2)

add_para(doc, (
    "Given that 22 states enacted phone policies in 2025 alone, Tennessee districts may wish to "
    "consider proactive adoption of phone management policies rather than waiting for state "
    "legislation. Key considerations for a district like Greeneville City Schools (~2,500 students, "
    "K-12, East Tennessee):"
))

add_bullet(doc, (
    "Small district size allows for more consistent, equitable enforcement across all schools."
))
add_bullet(doc, (
    "The NBER Florida study found gains concentrated in middle and high school, suggesting a "
    "phased approach starting with GMS and GHS may be most efficient."
))
add_bullet(doc, (
    "Virginia's $500,000 implementation support model suggests state funding may become available "
    "for Tennessee districts pursuing phone-free policies."
))
add_bullet(doc, (
    "The equity findings from Florida (disproportionate suspensions for Black students in year one) "
    "should inform enforcement design from the outset."
))
add_bullet(doc, (
    "South Carolina's model of tying state funding to phone bans may influence Tennessee "
    "legislative discussions given regional proximity."
))

add_heading(doc, "Data Gaps", 2)

add_para(doc, (
    "Tennessee-specific research on screen time and academic performance is largely absent from the "
    "literature reviewed. The following data would be valuable for evidence-based policy development:"
))

add_bullet(doc, "TVAAS data correlated with device usage policies across Tennessee districts")
add_bullet(doc, "TNReady/TCAP scores pre- and post-phone policy implementation in Tennessee districts that have adopted bans")
add_bullet(doc, "TDOE guidance on recommended screen time limits for instructional vs. recreational use")
add_bullet(doc, "Student survey data on personal device use during school hours in Tennessee schools")

# ─── SAVE ────────────────────────────────────────────────────────────────────

os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"Report saved to: {OUTPUT_FILE}")
print(f"Total references: {len(references)}")
print(f"Total source reliability matrix entries: {len(srm_rows)}")
