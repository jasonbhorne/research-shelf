#!/usr/bin/env python3
"""Generate companion guide for In France Profound by T.D. Allman."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# -- Font defaults --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# -- Heading styles --
for level, size, color in [
    ("Heading 1", 22, RGBColor(0x1A, 0x3C, 0x5E)),
    ("Heading 2", 16, RGBColor(0x2B, 0x57, 0x7B)),
    ("Heading 3", 13, RGBColor(0x3A, 0x70, 0x9A)),
]:
    h = doc.styles[level]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == "Heading 1" else 14)
    h.paragraph_format.space_after = Pt(6)

# -- Block quote style --
bq_style = doc.styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
bq_style.font.name = "Calibri"
bq_style.font.size = Pt(10.5)
bq_style.font.italic = True
bq_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
bq_style.paragraph_format.left_indent = Inches(0.5)
bq_style.paragraph_format.right_indent = Inches(0.3)
bq_style.paragraph_format.space_before = Pt(8)
bq_style.paragraph_format.space_after = Pt(8)

# -- Helpers --
def add_quote(text, attribution=""):
    p = doc.add_paragraph(style="BlockQuote")
    p.add_run(f'"{text}"')
    if attribution:
        p.add_run(f"\n— {attribution}")

def add_body(text):
    doc.add_paragraph(text)

def add_rec_table(recs):
    """recs = list of (title, author, rationale)"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, label in enumerate(["Title", "Author", "Why Read It"]):
        hdr[i].text = label
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for title, author, rationale in recs:
        row = table.add_row().cells
        row[0].text = title
        row[1].text = author
        row[2].text = rationale
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()  # spacer
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Companion Guide")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = title_p2.add_run("In France Profound")
run2.font.size = Pt(28)
run2.font.bold = True
run2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = subtitle.add_run("The Long History of a House, a Mountain Town, and a People")
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = author_p.add_run("by T.D. Allman")
run4.font.size = Pt(14)
run4.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime("%B %d, %Y")
run5 = meta.add_run(f"Book | Full Spoilers | {today}")
run5.font.size = Pt(10)
run5.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# 1. QUICK TAKE
# ============================================================
doc.add_heading("Quick Take", level=1)

add_body(
    "In France Profound is the final work of T.D. Allman, the fearless American journalist who spent "
    "decades exposing the CIA's secret war in Laos, surviving kidnapping in Beirut, and witnessing "
    "Tiananmen Square. Published posthumously in August 2024, three months after Allman's death at 79, "
    "the book is his love letter to Lauzerte, a 1,500-person hilltop bastide in southwestern France "
    "where he lived for over 30 years. It is not a cozy expat memoir. It is a 480-page sweep from "
    "prehistoric cave paintings to the modern Yellow Vest movement, all viewed through the windows of an "
    "800-year-old house. Critics gave it starred reviews in both Kirkus and Publishers Weekly, praising "
    "its ambition and narrative energy while occasionally noting that Allman's baroque prose and "
    "digressive style demand patience. Goodreads readers land at 3.9/5, with most finding it engaging "
    "if sometimes sprawling. The consensus: a unique, deeply researched, and entertaining work of "
    "place-based history that rewards readers willing to follow Allman on his many detours."
)

# ============================================================
# 2. ABOUT THE AUTHOR
# ============================================================
doc.add_heading("About the Author", level=1)

add_body(
    "Timothy Damien Allman (1944-2024) was a native Floridian who became one of the great American "
    "foreign correspondents of the late twentieth century. After Harvard, he joined the Peace Corps, "
    "then launched a journalism career that would span nearly five decades and take him to the world's "
    "most dangerous places. In 1968, he broke the story of the CIA's covert war in Laos. He rescued "
    "massacre victims in Cambodia, was kidnapped in Beirut, and witnessed the Tiananmen Square crackdown "
    "firsthand. His byline appeared in The New York Times, Vanity Fair, The New Yorker, The Washington "
    "Post, Harper's, The Guardian, Le Monde, The Economist, and Rolling Stone."
)

add_body(
    "His books share a common method: take a place everyone thinks they understand, dig into its deep "
    "history, and show how the received story is wrong. Miami: City of the Future (1987) reframed the "
    "city as a global capital, not a retirement community. Rogue State: America at War with the World "
    "(2004) examined post-9/11 foreign policy. Finding Florida: The True History of the Sunshine State "
    "(2013) debunked sanitized narratives so effectively it was longlisted for the National Book Award. "
    "In France Profound applies the same contrarian lens to the village where Allman spent over three "
    "decades of his life."
)

add_body(
    "Allman's Harper's contributions alone spanned 46 years (1977-2023), covering everything from "
    "Jerry Brown to the Marcos regime to the Monroe Doctrine. He divided his time between New York and "
    "Lauzerte, living with his partner of 25 years, John Sui. He died of respiratory failure in a "
    "Manhattan hospital on May 12, 2024. In France Profound, his capstone work, was published three "
    "months later."
)

# ============================================================
# 3. WHAT THE CRITICS SAY
# ============================================================
doc.add_heading("What the Critics Say", level=1)

add_body(
    "The critical response was unanimously positive, if occasionally qualified. Book Marks aggregated "
    "five reviews and found zero pans or mixed notices: two raves and three positives."
)

doc.add_heading("The Raves", level=2)

add_quote(
    "An engaging, richly detailed tale... memoir, historical narrative, and travelogue, "
    "all delivered with zesty enthusiasm.",
    "Kirkus Reviews (Starred Review)"
)

add_quote(
    "A sumptuous account... the most penetrating aspect of Allman's narrative is his "
    "exploration of how his relationship with the town has altered his perception of what "
    "history is and how it moves. This enthralls.",
    "Publishers Weekly (Starred Review)"
)

add_body(
    "Two starred reviews from the major trade outlets is a strong showing for any nonfiction title, "
    "and particularly notable for a posthumous release without the usual author-driven publicity cycle."
)

doc.add_heading("The Praise", level=2)

add_quote(
    "His perspectives are grand, the history deep, the narrative conversational and enthusiastic.",
    "Dominic Green, Wall Street Journal"
)

add_quote(
    "His writing is often brilliant, warm, and clever. The stories are, by turns, gripping and amusing.",
    "Evan M. Anderson, Library Journal"
)

add_quote(
    "Intriguing, compelling, and full of novel insights.",
    "Booklist"
)

doc.add_heading("The Qualifications", level=2)

add_body(
    "The Christian Science Monitor offered the most balanced take, praising Allman's energy while "
    "identifying real weaknesses:"
)

add_quote(
    "At turns bombastic, entertaining, eccentric, and insightful... his writing, which verges on "
    "the baroque, sometimes spills over into overwrought.",
    "Bryn Stole, Christian Science Monitor"
)

add_body(
    "Stole also noted the book is \"thin on the texture of everyday life\" and contains few firsthand "
    "accounts from present-day villagers, a surprising gap for a 30-year resident. Some historical "
    "claims were flagged as lacking rigor, including an uncritical mention of the Children's Crusade "
    "and an overbroad claim about European innumeracy before Fibonacci."
)

doc.add_heading("Reader Response", level=2)

add_body(
    "On Goodreads, the book holds a 3.9/5 average from 82 ratings: 23% five-star, 50% four-star, "
    "23% three-star, and only 3% at two stars or below. The most common reader criticism was the "
    "absence of photographs and illustrations for such a place-centered book. Several readers noted "
    "the narrative structure is \"not always linked in the clearest of fashion.\" The book received "
    "no major award nominations and did not appear on bestseller lists, though the paperback release "
    "(March 2026) may generate renewed attention."
)

# ============================================================
# 4. THEMES AND WHAT MAKES IT WORK
# ============================================================
doc.add_heading("Themes and What Makes It Work", level=1)

doc.add_heading("La France Profonde as Idea and Reality", level=2)

add_body(
    "The title plays on la France profonde (\"Deep France\"), a term from a 1988 essay by Michel Dion "
    "describing the countryside where \"time stood still: where the customs, dialects, and feelings of "
    "belonging had held on defiantly through the turbulence of modernity.\" The concept has been claimed "
    "across the political spectrum, from Yellow Vest protesters to Marine Le Pen's National Rally to "
    "Eric Zemmour's campaigns. Allman's investigation asks what \"deep France\" actually means when you "
    "look at a specific place across eight centuries. His answer is ambivalent: Lauzerte is genuinely "
    "distinctive, shaped by forces different from Paris, but it is also subject to the same tides of "
    "centralization, modernization, and globalization that remake everywhere else."
)

doc.add_heading("The House as Protagonist", level=2)

add_body(
    "Allman's 800-year-old house on Lauzerte's principal square (dating from approximately 1175) is "
    "always capitalized as \"the House\" and functions as a character in its own right. It \"bestows "
    "gifts of magic and madness, joy, folly, good food and good wine.\" Previous occupants included "
    "French and English conquerors, Catholics and Protestants. Allman writes: \"Look out the windows of "
    "my House long enough, and you will witness the rise and decay of cultures, the formation and "
    "disintegration of economic systems.\" The House serves simultaneously as personal anchor, "
    "historical witness, and metaphor for the persistence of place against the transience of human "
    "affairs. That said, readers expecting an architectural or domestic history of the house itself "
    "will be disappointed. Allman uses it as a portal, not a subject."
)

doc.add_heading("La Longue Duree", level=2)

add_body(
    "Allman adopts what French Annales school historians call la longue duree, the \"long view\" that "
    "detects deep, enduring patterns beneath the surface of events. The narrative spans from the "
    "25,000-year-old cave paintings at Pech Merle through medieval bastide charters, the Hundred "
    "Years' War, the Wars of Religion, Napoleon, the Nazi occupation, and post-war modernization. "
    "This methodology is unusual in memoir or travel writing. It allows Allman to make arguments about "
    "continuity and change that most travel writers, limited to personal observation, simply cannot "
    "sustain. His central observation: \"France Profound doesn't make history. History periodically "
    "remakes it, and then rolls on its way.\""
)

doc.add_heading("The Forgotten Charter", level=2)

add_body(
    "Perhaps the book's most compelling historical recovery is the 1241 Lauzerte charter, written in "
    "Occitan (not Latin or French), which contained remarkably progressive provisions: due process, "
    "taxation with representation, abolition of serfdom, and recognition of women as legal persons. "
    "Allman argues this document deserves recognition alongside England's Magna Carta but has been "
    "systematically overlooked because it emerged from provincial France rather than the centers of "
    "power. This finding exemplifies his broader argument: the provinces have been marginalized in "
    "their own national narrative."
)

doc.add_heading("Religious Violence as Engine of Change", level=2)

add_body(
    "Writing from an agnostic perspective, Allman frames successive religious conflicts as the "
    "primary mechanism by which \"history periodically remakes\" the region. The Albigensian Crusade "
    "against the Cathars devastated Lauzerte's corner of France in the early 13th century, "
    "representing one of the first institutional religious genocides in Western Europe. In 1562, "
    "Protestant-Catholic violence killed approximately 567 people in the small town. The book's "
    "central square, visible from the House, was the site of massacres. Allman uses these episodes "
    "to challenge romanticized views of medieval provincial life."
)

doc.add_heading("Centralization vs. Provincial Autonomy", level=2)

add_body(
    "The region, once called Quercy, was sandwiched between the better-known Aquitaine and Languedoc, "
    "\"ever in thrall to the magnetic impulse of Paris.\" Allman traces how modernization eroded "
    "village autonomy: automobiles, chain stores, and fast food literally moved the town's center of "
    "gravity downhill from the medieval summit, as businesses relocated for parking convenience. "
    "Indoor plumbing did not arrive in Lauzerte until after World War II. This tension between Paris "
    "and the provinces is one of the oldest themes in French historiography. Allman gives it personal "
    "texture by living it."
)

# ============================================================
# 5. WHAT YOU MIGHT NOT KNOW
# ============================================================
doc.add_heading("What You Might Not Know", level=1)

add_body(
    "Allman was not an academic or historian by training. He was a gonzo-adjacent journalist whose "
    "career began by exposing covert CIA operations. His maximalist, digressive prose style connects "
    "directly to decades of writing for Harper's, Vanity Fair, and Rolling Stone. He compared "
    "Lauzerte's region to Faulkner's fictional Yoknapatawpha County, packed with peculiar characters "
    "and \"strange truths,\" signaling literary aspirations well beyond standard travel writing."
)

add_body(
    "No author interviews for In France Profound appear to exist. Allman died in May 2024, three "
    "months before the August publication date. The book was apparently complete or near-complete at "
    "the time of his death, already in the publication pipeline with Atlantic Monthly Press. His "
    "editor, George Gibson, praised Allman's ability to \"make connections nobody else would even "
    "imagine.\""
)

add_body(
    "The book contains no photographs or illustrations, a decision multiple readers found baffling "
    "for a work so rooted in a specific place and landscape. Whether this was Allman's choice or a "
    "publisher decision is unknown."
)

add_body(
    "Allman lived in Lauzerte for over 30 years, yet present-day residents appear in the book mainly "
    "in the acknowledgments. The critical divide maps neatly onto this gap: historians found the book "
    "imprecise, memoirists found it historically heavy, and travel readers found it personally thin. "
    "Readers who embraced the hybrid tended to rate it highly."
)

# ============================================================
# 6. IF YOU LIKED THIS
# ============================================================
doc.add_heading("If You Liked This", level=1)

add_body(
    "Five recommendations that approach similar territory from different angles:"
)

add_rec_table([
    (
        "Montaillou",
        "Emmanuel Le Roy Ladurie",
        "The closest intellectual companion. A landmark microhistory of a medieval Cathar village "
        "in the French Pyrenees, using Inquisition records to reconstruct daily life. Same region, "
        "same longue duree method, same revelatory detail."
    ),
    (
        "The Discovery of France",
        "Graham Robb",
        "Robb cycled 14,000 miles through France to reveal how much of the country remained terra "
        "incognita into the 19th century. Won the Ondaatje and Duff Cooper Prizes. Where Allman "
        "goes deep in one town, Robb goes wide across the whole country."
    ),
    (
        "Finding Florida",
        "T.D. Allman",
        "Allman's own previous book applies the same contrarian, myth-busting approach to the "
        "Sunshine State. Longlisted for the National Book Award. The natural companion read: "
        "same author, same method, different continent."
    ),
    (
        "Peasants into Frenchmen",
        "Eugen Weber",
        "The definitive academic treatment of how roads, railways, schools, and conscription "
        "transformed France's isolated rural populations into a unified citizenry (1870-1914). "
        "The scholarly backbone for everything Allman covers narratively."
    ),
    (
        "Something to Declare",
        "Julian Barnes",
        "Twenty years of essays on France by the Booker Prize winner: landscape, Flaubert, food, "
        "the Tour de France, and the English-French cultural relationship. More polished and "
        "essayistic than Allman, but a fine complement."
    ),
])

doc.add_paragraph()  # spacer

doc.add_heading("Watch and Listen", level=2)

add_body(
    "Etre et avoir (2002): Nicolas Philibert's documentary follows a teacher in a one-room "
    "schoolhouse in rural Auvergne. Captures the texture of everyday life in la France profonde "
    "that reviewers wished Allman had included more of."
)

add_body(
    "Visages Villages (Faces Places, 2017): Agnes Varda and street artist JR travel through rural "
    "France creating murals of ordinary people. Academy Award-nominated. A visual love letter to "
    "the people of deep France."
)

add_body(
    "Un Village Francais (2009-2017): A seven-season French TV drama set in a fictional village "
    "during the German occupation. Critically acclaimed for its nuanced portrayal of collaboration "
    "and resistance. Dramatizes one of the historical periods Allman covers."
)

add_body(
    "Cousin Jules (1972): A nearly wordless five-year film portrait of an elderly blacksmith-farmer "
    "in rural Burgundy. Won top prize at Locarno. Embodies the longue duree visually."
)

# ============================================================
# 7. SOURCES
# ============================================================
doc.add_heading("Sources", level=1)

sources = [
    "Anderson, Evan M. Library Journal (2024). Review of In France Profound.",
    "Book Marks (2024). In France Profound. https://bookmarks.reviews/reviews/in-france-profound-the-long-history-of-a-house-a-mountain-town-and-a-people/",
    "Booklist (2024). Review of In France Profound.",
    "Compact Magazine (2024). Searching for the Soul of France. https://www.compactmag.com/article/searching-for-the-soul-of-france/",
    "Goodreads (2024-2026). In France Profound. https://www.goodreads.com/book/show/197525435-in-france-profound",
    "Green, Dominic. Wall Street Journal (2024). Review of In France Profound.",
    "Grove Atlantic (2024). In France Profound. https://groveatlantic.com/book/in-france-profound/",
    "Grove Atlantic. T.D. Allman Author Page. https://groveatlantic.com/author/t-d-allman/",
    "Harper's Magazine. T.D. Allman Author Archive. https://harpers.org/author/tdallman/",
    "Kirkus Reviews (2024). In France Profound (Starred Review). https://www.kirkusreviews.com/book-reviews/td-allman/in-france-profound/",
    "Miami New Times (2024). T.D. Allman, Journalist and Miami Historian, Has Died at 79. https://www.miaminewtimes.com/uncategorized/td-allman-journalist-and-miami-historian-obituary-20212449/",
    "Parade Magazine (2024). Review blurb for In France Profound.",
    "Peebles, Thomas H. Washington Independent Review of Books (2024). In France Profound. https://www.washingtonindependentreviewofbooks.com/bookreview/in-france-profound",
    "Publishers Weekly (06/26/2024). In France Profound (Starred Review). https://www.publishersweekly.com/9780802127846",
    "Shelf Awareness (06/13/2024). Obituary Note: T.D. Allman. https://www.shelf-awareness.com/theshelf/2024-06-13/obituary_note:_t.d._allman.html",
    "Stole, Bryn. Christian Science Monitor (08/21/2024). T.D. Allman's 'In France Profound' marks history through the lens of village life. https://www.csmonitor.com/Books/Book-Reviews/2024/0821/t.d.-allman-in-france-profound",
    "Susancoventry.blogspot.com (10/2024). Book Review: In France Profound. https://susancoventry.blogspot.com/2024/10/book-review-in-france-profound-long.html",
]

for s in sources:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(s)
    run.font.size = Pt(9.5)

# -- Save --
output_path = "/Users/hornej/Documents/Research/in-france-profound/2026-03-18 In France Profound Companion Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
