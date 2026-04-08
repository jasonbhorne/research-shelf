#!/usr/bin/env python3
"""Generate companion guide for Radical Acceptance by Tara Brach."""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

today = date.today().strftime("%Y-%m-%d")
output_dir = os.path.expanduser("~/Documents/Research/radical-acceptance")
output_path = os.path.join(output_dir, f"{today} Radical Acceptance Companion Guide.docx")

doc = Document()

# --- Default style ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

# Heading styles
for level, size, color in [
    ("Heading 1", 22, RGBColor(0x1B, 0x3A, 0x4B)),
    ("Heading 2", 16, RGBColor(0x2C, 0x5F, 0x7C)),
    ("Heading 3", 13, RGBColor(0x3D, 0x7E, 0x9E)),
]:
    h = doc.styles[level]
    h.font.name = "Calibri"
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)

# --- Helper functions ---
def add_title_page():
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Radical Acceptance")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
    run.font.name = "Calibri"
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Embracing Your Life With the Heart of a Buddha")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
    run.italic = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("by Tara Brach, Ph.D.")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x3D, 0x3D, 0x3D)

    doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Companion Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x7C)
    run.bold = True

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Book  |  Full Spoilers  |  {today}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_blockquote(text, attribution=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if attribution:
        attr_p = doc.add_paragraph()
        attr_p.paragraph_format.left_indent = Inches(0.5)
        attr_p.paragraph_format.space_after = Pt(8)
        run = attr_p.add_run(f"-- {attribution}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def add_body(text):
    doc.add_paragraph(text)


def add_rec_table(recs):
    """Add a formatted recommendation table. Each rec is (title, author, rationale)."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, label in enumerate(["Title", "Author", "Why Read It"]):
        hdr[i].text = label
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for title, author, rationale in recs:
        row = table.add_row().cells
        row[0].text = title
        row[1].text = author
        row[2].text = rationale
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")


# --- Build the document ---

add_title_page()

# 1. Quick Take
doc.add_heading("Quick Take", level=2)
add_body(
    "Radical Acceptance is one of those books that earns its reputation honestly. "
    "First published in 2003, Tara Brach's debut blends Buddhist meditation, Western clinical "
    "psychology, and unflinching personal disclosure into a guide for people who carry a quiet "
    "conviction that something is fundamentally wrong with them. It is not a self-help book "
    "in the breezy, five-steps-to-happiness sense. It is slower, warmer, and more demanding "
    "than that. Brach asks you to sit with your pain rather than fix it, to meet your shame "
    "with tenderness rather than strategy. Over two decades and a 20th anniversary re-release "
    "later, it remains one of the most recommended books in the mindfulness space, with "
    "36,000+ Goodreads ratings (4.14/5) and a dedicated following among therapists, recovery "
    "communities, and everyday readers who found in it a permission they didn't know they needed."
)

# 2. About the Author
doc.add_heading("About the Author", level=2)
add_body(
    "Tara Brach (born 1953, East Orange, New Jersey) holds a rare dual credential in the "
    "mindfulness world: a Ph.D. in Clinical Psychology from Fielding Graduate University, where "
    "her dissertation focused on meditation as a therapeutic modality for addiction treatment, "
    "and formal ordination as a Buddhist lay priest. She completed a five-year teacher training "
    "program at Spirit Rock Meditation Center under Joseph Goldstein and Jack Kornfield, two "
    "of the founding figures of American insight meditation. She went on to co-found the "
    "Mindfulness Meditation Teacher Certification Program with Kornfield, which has trained "
    "over 7,000 meditation teachers across 74 countries."
)
add_body(
    "In 1998, she founded the Insight Meditation Community of Washington, D.C., which grew "
    "from a small Vipassana sitting group in Bethesda to over 200 weekly attendees. Her weekly "
    "podcast, launched in 2010, reaches approximately 2.5 million downloads per month, making "
    "it one of the most popular meditation podcasts in existence."
)
add_body(
    "Radical Acceptance was her debut. It was followed by True Refuge (2013), which debuted "
    "on The Washington Post's bestseller list; Radical Compassion (2019), which introduced the "
    "RAIN framework as a standalone practice; and Trusting the Gold (2021). Each subsequent book "
    "builds on the foundation Radical Acceptance laid, but none has matched its cultural reach."
)

doc.add_heading("The Personal Story Behind the Book", level=3)
add_body(
    "The book's credibility comes from more than credentials. After college, Brach spent "
    "roughly a decade living in a 3HO (Sikh Dharma) ashram near Boston, where organization "
    "leader Yogi Bhajan arranged her marriage to a man she barely knew. She eventually left "
    "after recognizing patterns of patriarchal control. The marriage ended five years later. "
    "During a pivotal moment in the desert after this upheaval, she confronted a fundamental "
    'question and chose to answer it: "I want to accept myself completely, just as I am."'
)
add_body(
    "She was later diagnosed with a connective tissue condition associated with Fragile X "
    "syndrome, adding chronic pain to her lived experience. These threads, the cult-like "
    "community, the failed arranged marriage, the chronic illness, run through the book's "
    'prologue, titled "Something Is Wrong with Me." Brach is not writing about suffering '
    "from a safe theoretical distance. She is writing from the middle of it."
)

# 3. What the Critics Say
doc.add_heading("What the Critics Say", level=2)
add_body(
    'Publishers Weekly called Radical Acceptance "a consoling and practical guide" with '
    '"gentle advice and guided meditation" garnished with "beautiful bits of poetry and '
    'well-loved dharma stories." Library Journal reviewer Mark Woodhouse noted that Brach\'s '
    '"tone is more logical and oriented toward psychology" than comparable Buddhist authors, '
    "an observation that gets at one of the book's defining qualities: it takes the spiritual "
    "seriously without abandoning the clinical."
)
add_blockquote(
    "A book, and a practice, that we all need.",
    "Stephen Cope, author of The Great Work of Your Life"
)
add_blockquote(
    "Gentle wisdom and tender healing... a most excellent medicine for our unworthiness and longing.",
    "Jack Kornfield"
)
add_body(
    "On Goodreads, the book carries a 4.14/5 average across 36,000+ ratings. Seventy-seven "
    "percent of readers give it four or five stars. Positive reviews consistently praise its "
    "practical guidance for self-acceptance and its accessibility as an introduction to Buddhist "
    "ideas. The most common criticism is that the client case studies can feel formulaic: "
    "clients arrive struggling, work with Brach, and experience breakthroughs. Some readers "
    "also note repetitiveness across chapters. These are fair observations, though they describe "
    "a feature of the therapeutic self-help genre as much as a flaw in this particular book."
)
add_body(
    "No New York Times or Kirkus reviews were located, suggesting the book may not have "
    "received major review attention at publication. Its reputation was built more through "
    "word of mouth, therapeutic recommendation, and Brach's growing podcast audience than "
    "through traditional literary gatekeeping."
)

# 4. Themes and What Makes It Work
doc.add_heading("Themes and What Makes It Work", level=2)

doc.add_heading("The Trance of Unworthiness", level=3)
add_body(
    "The book's central concept is what Brach calls \"the trance of unworthiness,\" the deeply "
    "held belief that something is fundamentally wrong with us. She argues this manifests in "
    '"crippling self-judgments and conflicts in our relationships, in addictions and perfectionism, '
    'in loneliness and overwork." The prologue is literally titled "Something Is Wrong with Me." '
    "Brach frames this not as individual pathology but as a pervasive cultural condition, "
    "a conditioned state of consciousness rather than truth, and radical acceptance is the "
    "practice of waking up from it."
)
add_body(
    "The framing matters. By calling it a \"trance,\" Brach simultaneously normalizes the "
    "experience (you are not uniquely broken; this is the water we all swim in) and reveals "
    "it as something that can be awakened from. It is one of the book's most distinctive "
    "rhetorical moves."
)

doc.add_heading("Two Pillars: Mindfulness and Compassion", level=3)
add_body(
    "Radical acceptance rests on two pillars. The first is clear seeing, or mindfulness: "
    "recognizing what is happening physically, mentally, and emotionally without being ruled "
    "by it. The second is compassion: responding to that experience with care and tenderness "
    "rather than harsh judgment. Brach draws on both Buddhist psychology and Western experiential "
    "therapy, arguing that experiencing and accepting the \"changing stream of sensations\" is "
    "central to transformation."
)
add_body(
    "This is not purely cognitive reframing. Brach insists on embodied experience. Emotions, "
    "she argues, are a combination of physical sensations and stories, and they \"continue to "
    'cause suffering until we experience them where they live in our body." The book\'s emphasis '
    "on somatic awareness, returning attention to the body when emotions arise, is one of its "
    "most therapeutically grounded features and anticipates the body-centered approaches that "
    "became mainstream in trauma therapy over the following decade."
)

doc.add_heading("The RAIN Framework", level=3)
add_body(
    "While RAIN (originally Recognize, Allow, Investigate, Non-identification) was first "
    "introduced by insight meditation teacher Michele McDonald, Brach popularized and modified "
    "it, changing the final step to \"Nurture.\" That single-word change is significant: it "
    "shifted RAIN from a purely observational practice to one that includes active self-compassion, "
    "consistent with the book's thesis. RAIN became central to Brach's later teaching and was "
    "given dedicated treatment in Radical Compassion (2019). The 20th anniversary edition of "
    "Radical Acceptance retroactively adds a bonus chapter, \"The RAIN of Compassion,\" connecting "
    "this framework back to the original book."
)

doc.add_heading("Structure as Practice", level=3)
add_body(
    "The book contains a prologue and twelve chapters, each ending with a guided meditation "
    "tied to the chapter's theme. The arc is deliberate: from diagnosing the problem (the trance "
    "of unworthiness), to individual healing practices (body awareness, desire, fear, self-compassion), "
    "to expanding outward (compassion for others, relationships), to ultimate realization. Each "
    "chapter functions as both essay and practice guide. Brach weaves three narrative modes "
    "throughout: personal autobiography, anonymized case histories from her psychotherapy "
    "practice, and fresh interpretations of traditional Buddhist stories and parables. Poetry "
    "punctuates the transitions."
)
add_body(
    "This tripartite structure gives Radical Acceptance a texture that purely instructional "
    "mindfulness books lack. The personal disclosure creates trust. The clinical stories provide "
    "universality. The Buddhist tales provide philosophical depth. And the guided meditations "
    "at the end of each chapter ensure the book is something you do, not just something you read."
)

# 5. What You Might Not Know
doc.add_heading("What You Might Not Know", level=2)
add_body(
    'The term "radical acceptance" has a parallel life in clinical psychology. Marsha Linehan '
    "coined it as a core concept in Dialectical Behavior Therapy (DBT), and it is also central "
    "to Acceptance and Commitment Therapy (ACT). Brach's use of the term comes from the Buddhist "
    "tradition rather than the clinical one, but the overlap is not coincidental. Meta-analyses "
    "show that self-compassion-related therapies produce significant improvements in anxiety "
    "and depressive symptoms, and a 2014 randomized controlled trial found ACT was significantly "
    "superior to control conditions in building self-compassion and reducing psychological "
    "distress. The science validates what Brach teaches through story and meditation."
)
add_body(
    "Therapists have taken notice. A continuing education course, \"Radical Acceptance with "
    "Tara Brach, Ph.D.: Integrating Meditation to Heal Shame and Fear in Clinical Practice,\" "
    "is available through CE Broker, allowing clinicians to earn professional credits while "
    "studying her approach. The book is frequently assigned alongside formal DBT and ACT treatment."
)
add_body(
    "The main audiobook is narrated by Cassandra Campbell (Audie Award winner, Audible Hall "
    "of Fame inductee), not by Brach herself. It runs 12 hours and 15 minutes. Listeners who "
    "want Brach's own voice should look to the companion 2-disc guided meditations set, the "
    "\"Radical Self-Acceptance\" audio program through Sounds True, or her podcast. The companion "
    "meditations include practices on radical acceptance of pain, tonglen (transforming suffering "
    "into compassion), and cultivating a forgiving heart."
)
add_body(
    "The 20th anniversary edition (November 2023) added a new introduction reflecting Brach's "
    "\"deepened understanding in an ever-changing world,\" a bonus chapter on \"The RAIN of "
    "Compassion,\" and additional guided meditations. It was the first hardback release of the book."
)
add_body(
    "Brach and Jack Kornfield are not just mutual admirers. They co-founded the Mindfulness "
    "Meditation Teacher Certification Program (serving participants from 74 countries), co-teach "
    "\"Mindfulness Daily\" (a 40-day online training), and share a joint teacher profile on "
    "Insight Timer. Reading Kornfield alongside Brach gives you both sides of a shared teaching "
    "lineage."
)

# 6. Beyond the Book
doc.add_heading("Beyond the Book", level=2)

doc.add_heading("Brach's Media Ecosystem", level=3)
add_body(
    "Radical Acceptance has an unusually rich companion media ecosystem for a nonfiction book. "
    "The Tara Brach Podcast (over 1,572 episodes, 2.5 million monthly downloads) is essentially "
    "a free, ongoing extension of the book's teachings. Each episode typically includes a talk "
    "and a guided meditation, and many directly revisit Radical Acceptance themes. A self-paced "
    "online course, \"Radical Acceptance: Freedom via Mindfulness and Meditation,\" is available "
    "through the Radical Compassion Institute (with 10% of proceeds donated to climate and racial "
    "justice nonprofits). The documentary \"A Joyful Mind\" (2016, 8.5/10 on IMDb) features "
    "Brach alongside neuroscientist Richard Davidson."
)

doc.add_heading("Brach's Other Books", level=3)
add_body(
    "Brach's four books form a natural progression. Radical Acceptance (2003) lays the "
    "groundwork. True Refuge (2013) deepens the spiritual dimension, exploring three refuges "
    "for times of difficulty. Radical Compassion (2019) provides the most actionable framework "
    "with a full treatment of the RAIN practice. Trusting the Gold (2021) distills the wisdom "
    "into shorter reflections. If you finish Radical Acceptance and want more, Radical Compassion "
    "is the natural next step."
)

doc.add_heading("Book Clubs and Recovery Communities", level=3)
add_body(
    "No official publisher's discussion guide exists, but the book is widely used in book clubs, "
    "recovery groups (notably addiction recovery forums like Talking Sober), and therapeutic "
    "reading groups. The chapter-ending guided meditations provide natural discussion breakpoints. "
    "Common discussion themes include the trance of unworthiness, using RAIN for difficult "
    "emotions, mindful parenting, compassionate communication, and how acceptance transforms "
    "relationships. The book's structure (personal stories + Buddhist tales + guided meditations) "
    "makes it unusually well-suited for group reading."
)

# 7. If You Liked This
doc.add_heading("If You Liked This", level=2)
add_body("Five books that share Radical Acceptance's territory from different angles:")
add_rec_table([
    (
        "When Things Fall Apart",
        "Pema Chodron",
        "The most frequently recommended companion. Chodron is more traditionally Buddhist where Brach integrates clinical psychology, but both offer practical wisdom for navigating pain."
    ),
    (
        "Self-Compassion",
        "Kristin Neff",
        "The empirical counterpart to Brach's contemplative approach. Neff's research on self-compassion provides the scientific grounding behind why Brach's practices work."
    ),
    (
        "A Path with Heart",
        "Jack Kornfield",
        "Kornfield and Brach are close collaborators and co-teachers. This foundational guide to Buddhist practice in daily life gives you the other side of a shared teaching lineage."
    ),
    (
        "The Miracle of Mindfulness",
        "Thich Nhat Hanh",
        "A slim, gentle classic. One of the primary Buddhist influences on Brach's work. Pairs naturally as an entry-level mindfulness text alongside Radical Acceptance's deeper psychological exploration."
    ),
    (
        "The Body Keeps the Score",
        "Bessel van der Kolk",
        "A different angle on the same territory. Brach approaches healing through Buddhist psychology; van der Kolk approaches it through trauma neuroscience. Both arrive at mindfulness and body awareness."
    ),
])

# 8. Sources
doc.add_heading("Sources", level=2)
sources = [
    "Amazon. Radical Acceptance: Embracing Your Life With the Heart of a Buddha. amazon.com",
    "Amazon. Radical Acceptance Guided Meditations (CD). amazon.com",
    "Buddhist Peace Fellowship. Tara Brach on Radical Acceptance and Spiritual Activism. bpf.org",
    "Calvin Rosser. Radical Acceptance: Summary & Notes. calvinrosser.com",
    "CE Broker. Radical Acceptance with Tara Brach, Ph.D. cebroker.com",
    "Chasing Brighter Podcast. March Book Club: Radical Acceptance. chasingbrighter.podbean.com",
    "Class Central. Udemy Course: Radical Acceptance. classcentral.com",
    "Encyclopedia.com. Brach, Tara. encyclopedia.com",
    "Goodreads. Radical Acceptance by Tara Brach. goodreads.com",
    "Goodreads. Similar Books to Radical Acceptance. goodreads.com",
    "GoodBooks. Radical Acceptance Recommendations. goodbooks.io",
    "Greater Good Science Center, UC Berkeley. Happiness Break with Tara Brach. greatergood.berkeley.edu",
    "IMDb. A Joyful Mind (2016). imdb.com",
    "Insight Timer. Jack Kornfield & Tara Brach. insighttimer.com",
    "Mindful.org. Investigate Anxiety with Tara Brach's RAIN Practice. mindful.org",
    "No Small Endeavor. Tara Brach Interview. nosmallendeavor.com",
    "Penguin Random House. Radical Acceptance. penguinrandomhouse.com",
    "Penguin Random House. Radical Compassion. penguinrandomhouse.com",
    "Penguin UK. Radical Acceptance (20th Anniversary Edition). penguin.co.uk",
    "PMC. Acceptance and Commitment Therapy and Self-Compassion RCT. pmc.ncbi.nlm.nih.gov",
    "Publishers Weekly. Tara Brach Author Page. publishersweekly.com",
    "Radical Compassion Institute. Radical Acceptance Course. courses.tarabrach.com",
    "Readings Books Australia. Radical Acceptance (2023 Edition). readings.com.au",
    "Schlow Library. Radical Acceptance Table of Contents. search.schlowlibrary.org",
    "Shepherd. Books Like Radical Acceptance. shepherd.com",
    "Shortform. Radical Acceptance Summary. shortform.com",
    "Sounds True. Radical Self-Acceptance (Audio Program). soundstrue.com",
    "Spirituality & Practice. Tara Brach Teacher Profile. spiritualityandpractice.com",
    "Springer. Self-Compassion Therapies Meta-Analysis. link.springer.com",
    "Talking Sober. Book Club: Radical Acceptance by Tara Brach. talkingsober.com",
    "Tantor Audio. Radical Acceptance Audiobook. tantor.com",
    "Tara Brach. About. tarabrach.com",
    "Tara Brach. Books. tarabrach.com",
    "Tara Brach. Free Guided Meditations for Radical Acceptance. tarabrach.com",
    "Tara Brach. RAIN Resources. tarabrach.com",
    "Washington Post (2013). Meditation Guru Tara Brach Is Calm Eye of Washington's Stress-Filled Storm. washingtonpost.com",
    "Wikipedia. Tara Brach. en.wikipedia.org",
]

for s in sources:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(s)
    run.font.size = Pt(9)

# --- Adjust section widths ---
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)

doc.save(output_path)
print(f"Saved: {output_path}")
