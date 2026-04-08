#!/usr/bin/env python3
"""Generate Project Hail Mary Companion Guide (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
import os

doc = Document()

# -- Style Setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# Quote style
quote_style = doc.styles.add_style('BlockQuote', 1)  # paragraph style
quote_style.font.name = 'Calibri'
quote_style.font.size = Pt(10.5)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
quote_style.paragraph_format.left_indent = Cm(1.5)
quote_style.paragraph_format.space_before = Pt(6)
quote_style.paragraph_format.space_after = Pt(6)

def add_quote(text, attribution=""):
    p = doc.add_paragraph(style='BlockQuote')
    p.add_run(f'"{text}"')
    if attribution:
        p.add_run(f'\n-- {attribution}')

def add_body(text):
    doc.add_paragraph(text)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Project Hail Mary')
run.font.size = Pt(36)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.bold = True

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run('A Companion Guide')
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run(f'Book by Andy Weir (2021)\n{date.today().strftime("%B %d, %Y")}\nFull spoilers included')
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================
# 1. QUICK TAKE
# ============================================================
doc.add_heading('Quick Take', level=1)

add_body(
    "Project Hail Mary is Andy Weir's triumphant third novel, the book that proved The Martian "
    "was no fluke. It follows Ryland Grace, an amnesiac middle-school science teacher who wakes "
    "up alone on a spaceship hurtling toward Tau Ceti, tasked with saving Earth from an extinction-level "
    "energy crisis caused by a sun-dimming microorganism called Astrophage. What begins as a "
    "familiar Weir survival puzzle, one man solving problems with science and gallows humor, "
    "transforms into something richer when Grace meets Rocky, a five-legged, rock-skinned alien "
    "from the 40 Eridani system who communicates through musical tones. Their friendship, built "
    "from first principles across every conceivable biological and linguistic barrier, is the "
    "emotional engine that elevates the book from clever to genuinely moving. Critics called it "
    '"nothing short of a science-fiction masterwork" (Kirkus, starred review), readers gave it a '
    "4.52 on Goodreads across 1.3 million ratings, and it won the 2021 Goodreads Choice Award "
    "for Best Science Fiction. Bill Gates finished it in one weekend. Barack Obama put it on his "
    "year-end list. If you like your sci-fi brainy, warm, and propulsive, this is the one."
)

# ============================================================
# 2. ABOUT THE AUTHOR
# ============================================================
doc.add_heading('About the Author', level=1)

add_body(
    "Andy Weir was born June 16, 1972, in Davis, California. His father was a physicist at "
    "Lawrence Livermore National Laboratory; his mother, an electrical engineer. He grew up steeped "
    "in science, started programming at Sandia National Laboratories at age 15, and spent two "
    "decades as a software engineer before fiction took over his life."
)

add_body(
    "His path to publication was unconventional. He ran a webcomic called Casey and Andy from 2001 "
    "to 2008, wrote the viral short story \"The Egg\" in 2009 (still making the rounds on Reddit "
    "to this day), and then self-published The Martian on his blog chapter by chapter. Readers "
    "asked him to put it on Kindle. It went viral. Crown Publishing picked it up in 2014. Ridley "
    "Scott turned it into a film starring Matt Damon in 2015. Weir won the 2016 John W. Campbell "
    "Award for Best New Writer."
)

add_body(
    "Then came the sophomore stumble. Artemis (2017), a heist novel set on a lunar colony, received "
    "mixed reviews and didn't capture the same magic. But Weir had already been working on bigger "
    "ideas. Before Artemis, he'd started a multi-volume space opera called Zhek about a substance "
    "that absorbs electromagnetic radiation and serves as fuel for interstellar travel. He wrote "
    "75,000 words before shelving it. Several core elements survived: the energy-absorbing substance "
    "became Astrophage, and a ruthless bureaucrat character became Eva Stratt. As Weir tells it, "
    "he realized that if a light-absorbing substance got into the Sun, \"that's where books come "
    "from!\" Project Hail Mary, published in 2021, was the result of that salvage operation, and "
    "it is widely regarded as his best work."
)

add_body(
    "Weir is refreshingly candid about his limitations. He has repeatedly called character depth "
    "\"my biggest weakness as a writer,\" adding: \"I've come up with interesting plots and "
    "interesting fake science, but my characters are kind of shallow and so I'm working on that.\" "
    "His compensating strategy is disarming: \"If you make exposition funny, the reader will forgive "
    "any amount of it.\" He identifies as a plot-driven author who doesn't always visualize his "
    "characters because he's \"just generally way more interested in the science involved.\" This "
    "self-awareness is part of his appeal. He's a nerd who writes for nerds, and he knows exactly "
    "what he's good at."
)

# ============================================================
# 3. WHAT THE CRITICS SAY
# ============================================================
doc.add_heading('What the Critics Say', level=1)

add_body(
    "The critical consensus on Project Hail Mary is overwhelmingly positive, with a consistent "
    "pattern: praise for the plotting and science-as-entertainment, occasional reservations about "
    "character depth, and near-universal love for Rocky."
)

doc.add_heading('The Reviews', level=2)

add_quote(
    "An unforgettable story of survival and the power of friendship, nothing short of a "
    "science-fiction masterwork.",
    "Kirkus Reviews (starred review)"
)

add_quote(
    "For readers who can forgive its shortcomings, the result is an engaging space odyssey.",
    "Alec Nevala-Lee, New York Times"
)

add_body(
    "The New York Times also named it an Editors' Choice, calling it \"a sensible course correction "
    "that supersizes the strategies of his most successful book.\" Mary Robinette Kowal, writing in "
    "the Washington Post, praised Grace's infectious enthusiasm for science but noted some problems "
    "\"could have been avoided with common sense and the use of checklists.\" Publishers Weekly "
    "called it \"a suspenseful portrait of human ingenuity and resilience\" and praised the toggle "
    "between present-day puzzle-solving and memory flashbacks."
)

doc.add_heading('Awards and Numbers', level=2)

add_body("The book's trophy case is substantial:")

awards = [
    "2021 Goodreads Choice Award, Best Science Fiction",
    "2021 Dragon Award, Best Science Fiction Novel",
    "2022 Audie Award, Audiobook of the Year AND Science Fiction",
    "2022 Seiun Award, Best Translated Long Work (Japan)",
    "2022 Hugo Award finalist, Best Novel",
    "4.52 average on Goodreads across 1.3 million ratings",
    "Debuted at #3 on the NYT Combined Print & E-Book Fiction list (May 2021)",
    "#1 on Locus Bestsellers for five consecutive months",
    "40 consecutive weeks on the NYT bestseller list (as of March 2026, driven by the film)",
    "Recommended by Bill Gates (finished it in one weekend) and Barack Obama (2021 reading list)",
]
for a in awards:
    doc.add_paragraph(a, style='List Bullet')

# ============================================================
# 4. THEMES AND WHAT MAKES IT WORK
# ============================================================
doc.add_heading('Themes and What Makes It Work', level=1)

doc.add_heading('Friendship Across Every Barrier', level=2)

add_body(
    "The heart of the book is the relationship between Ryland Grace and Rocky. They share no "
    "biology, no senses (Rocky is blind and navigates by sonar), no language, and no cultural "
    "frame of reference. Rocky breathes ammonia at 29 times Earth's atmospheric pressure and "
    "communicates through musical chords. Their friendship has to be built from absolute scratch, "
    "using the only thing they share: the scientific method. Numbers. Physical constants. "
    "Observation and inference. It is, in the best sense, a love letter to the idea that curiosity "
    "and goodwill are universal."
)

add_body(
    "What makes the friendship land emotionally is its asymmetry. Grace is funny, panicky, and "
    "self-deprecating. Rocky is methodical, brave, and earnest. They complement each other in "
    "ways that feel organic rather than engineered. And when Grace ultimately chooses to rescue "
    "Rocky instead of returning to Earth, the decision feels earned because Weir has spent 400 "
    "pages showing us why this alien matters more to Grace than abstract duty to a planet he "
    "barely remembers."
)

doc.add_heading('Scientific Problem-Solving as Narrative Engine', level=2)

add_body(
    "Weir's signature move is turning science into plot. Each chapter presents a problem, "
    "Grace works through it using real (or plausible) science, and the solution opens the next "
    "problem. This could be tedious, and in lesser hands it often is. But Weir has a genuine "
    "gift for making exposition entertaining. His trick, as he puts it: \"If you make exposition "
    "funny, the reader will forgive any amount of it.\" Grace's internal monologue, full of "
    "exclamation points, bad jokes, and genuine wonder, carries the technical passages the way "
    "a great teacher carries a lecture."
)

doc.add_heading('Sacrifice: Coerced vs. Chosen', level=2)

add_body(
    "The novel's central moral tension is the difference between forced and voluntary sacrifice. "
    "Grace didn't volunteer for a suicide mission. Eva Stratt, the ruthless head of the Hail Mary "
    "project, identified him as the best candidate and had his memory chemically suppressed so he "
    "couldn't refuse. He wakes up on the ship with no idea who he is or why he's there. The "
    "flashback timeline gradually reveals how he was conscripted, and the reader's sympathy shifts "
    "as we learn Grace initially tried to talk his way out of the mission."
)

add_body(
    "This makes the ending matter. When Grace chooses to save Rocky instead of returning to Earth, "
    "it's the first fully voluntary heroic act in the book. The coerced sacrifice imposed by Stratt "
    "is replaced by a genuine choice made from love and loyalty. Grace doesn't become a hero "
    "because the world needed one. He becomes a hero because he decided to be one."
)

doc.add_heading('Memory and Identity', level=2)

add_body(
    "Grace's amnesia is more than a plot device. It is the novel's structural principle. He must "
    "reconstruct his identity from scratch, and the book argues that identity is built through "
    "action, not memory. He experiments, he teaches Rocky, he solves problems, he cares. By the "
    "end, when he's teaching Eridian children on a planet 12 light-years from Earth, he has "
    "become his truest self: a teacher. The role he held before the mission, the role he holds "
    "after it, the role that defines him more than \"astronaut\" or \"savior\" ever could."
)

# ============================================================
# 5. THE DUAL TIMELINE AND NARRATIVE STRUCTURE
# ============================================================
doc.add_heading('How the Story Is Built', level=1)

add_body(
    "Project Hail Mary is written in first-person present tense, alternating between two timelines. "
    "In the present, Grace is aboard the Hail Mary, solving problems in real time. In the past, "
    "triggered by returning memories, we see how Earth discovered Astrophage, how Stratt assembled "
    "the mission, and how Grace was forced aboard."
)

add_body(
    "The amnesia framing solves hard sci-fi's perennial exposition problem: how do you deliver "
    "technical backstory without info-dumps? By making the backstory itself a mystery the "
    "protagonist is solving. Each flashback unlocks precisely when the present-day story needs it. "
    "The reader and Grace discover the story simultaneously, creating an unusually tight "
    "identification loop. You don't just read about Grace's confusion; you share it."
)

add_body(
    "This is Weir's most sophisticated narrative technique. The Martian was linear and "
    "journal-formatted. Artemis was a straightforward heist. Project Hail Mary uses structure "
    "as theme: the act of recovering memory mirrors the act of recovering identity, which mirrors "
    "the act of solving the scientific puzzle. All three threads converge at the climax."
)

# ============================================================
# 6. ROCKY AND THE ERIDIANS
# ============================================================
doc.add_heading('Rocky and the Eridians', level=1)

add_body(
    "Rocky is the book's secret weapon, and Weir built him from the ground up using rigorous "
    "worldbuilding. He started with a real exoplanet: 40 Eridani Ab. He calculated its conditions "
    "(high gravity, 29 times Earth's atmospheric pressure, extreme heat), then designed a biosphere "
    "that could plausibly evolve there. The result: five-legged crystalline beings with xenonite-based "
    "biology, no vision (they perceive through sonar-like echolocation), and communication via "
    "musical tones described in the text as quarter notes and eighth notes."
)

add_body(
    "Rocky is universally cited as the novel's most beloved element. \"I would die for Rocky\" is "
    "a common fan refrain, and it's not hard to see why. His alien-ness makes the friendship more "
    "meaningful, not less. Because Rocky shares nothing with Grace except intelligence and goodwill, "
    "every moment of connection feels hard-won. When Rocky learns to say \"fist my bump\" instead "
    "of \"fist bump,\" it is simultaneously hilarious and moving."
)

add_body(
    "For the 2026 film adaptation, the production team expanded the Eridian language to 250 words, "
    "using instruments including jug, ocarina, didgeridoo, contralto flute, and contralto clarinet. "
    "Rocky was brought to life with practical puppetry (six puppeteers on set) rather than CGI, "
    "a choice that audiences and critics have praised for giving the character tactile presence. "
    "James Ortiz performed Rocky's physical movements. Ryan Gosling improvised the now-iconic "
    "\"jazz hands\" first-contact scene."
)

# ============================================================
# 7. THE SCIENCE: REAL VS. SPECULATIVE
# ============================================================
doc.add_heading('The Science: Real vs. Speculative', level=1)

add_body(
    "Weir's approach to science fiction is what practitioners call the \"one big lie\" method. "
    "You make one or two speculative leaps, then rigorously follow the implications using real "
    "physics. The result feels trustworthy even when it isn't entirely real."
)

doc.add_heading('What Is Real', level=2)

real_items = [
    "Tau Ceti is a real star approximately 12 light-years from Earth. Tau Ceti e (Adrian in the novel) exists in exoplanet catalogs.",
    "40 Eridani is a real star system, and Weir used its actual properties to design the Eridian biosphere.",
    "Special relativity and time dilation are depicted accurately.",
    "Orbital mechanics and spacecraft engineering are, as astrophysicist Jacqueline McCleary put it, \"treated very fairly.\"",
    "The Hail Mary's spin-gravity design and centrifuge-based artificial gravity are internally consistent with known physics.",
]
for item in real_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('What Is Speculative', level=2)

add_body(
    "Astrophage is the biggest scientific stretch in the book, and it's a deliberate one. The "
    "organism violates the Second Law of Thermodynamics: it extracts energy from pure heat "
    "without a cold reservoir, which is physically impossible. The energy mismatch is enormous, "
    "as the Sun emits 10^26 joules per second, and a microorganism surviving five-million-degree "
    "temperatures to absorb that energy involves orders of magnitude of hand-waving. The Petrova "
    "line (a visible dimming line connecting affected stars) and Astrophage's mass-energy "
    "conversion mechanism are pure invention."
)

add_body(
    "But here is why it works: once you accept Astrophage exists and works the way Weir describes, "
    "everything downstream is rigorously worked out. The mission architecture, the fuel calculations, "
    "the biological experiments Grace runs, the interactions between Astrophage and Taumoeba, all "
    "of it follows logically. McCleary summarized it well: \"close enough to be enjoyable and, "
    "more importantly, self-consistent.\""
)

add_body(
    "For readers who want to go deeper, Scientific American, The Conversation, Britannica, "
    "the Royal Institution, and the Skeptical Inquirer have all published detailed analyses of "
    "the real science behind the fiction. The Skeptical Inquirer's piece on \"Astrophage and the "
    "Real Extremophiles\" is particularly good, connecting Weir's fiction to actual extremophile "
    "biology."
)

# ============================================================
# 8. WHAT YOU MIGHT NOT KNOW
# ============================================================
doc.add_heading('What You Might Not Know', level=1)

backstory_items = [
    (
        "The Zhek Connection",
        "The book's most distinctive elements, Astrophage and Eva Stratt, were salvaged from a "
        "75,000-word abandoned space opera called Zhek that Weir wrote before Artemis. He's described "
        "Project Hail Mary as a \"pastiche of ideas that had nothing to do with each other\" until "
        "the Sun connection clicked."
    ),
    (
        "Film Rights Sold Before Publication",
        "MGM bought the film rights in early 2020 for $3 million, with Ryan Gosling attached to "
        "star, before the book was even published. That's a remarkable bet on Weir's brand after "
        "Artemis underperformed."
    ),
    (
        "The Ending Was Always the Ending",
        "Weir has said Grace's choice to save Rocky instead of returning to Earth was always the "
        "planned ending. He intentionally left ambiguous whether Grace ever returns to Earth, telling "
        "Den of Geek: \"I didn't define that.\" This open thread has fueled sequel speculation."
    ),
    (
        "Gosling Added Depth Weir Couldn't",
        "Weir praised Ryan Gosling for adding \"so much depth and layers to Ryland that I never "
        "had in the book,\" a notably humble admission from the source material's author."
    ),
    (
        "The Four-Hour Cut",
        "Directors Phil Lord and Christopher Miller screened a nearly four-hour cut of the film to "
        "filmmaker friends before paring it to the 2.5-hour theatrical release."
    ),
    (
        "The Karaoke Scene Was Improvised",
        "One of the film's best scenes, a karaoke moment between Gosling and Sandra Huller (who plays "
        "Stratt), was not in the script."
    ),
]

for title, text in backstory_items:
    doc.add_heading(title, level=2)
    add_body(text)

# ============================================================
# 9. ON SCREEN
# ============================================================
doc.add_heading('On Screen', level=1)

add_body(
    "The film adaptation premiered in London on March 9, 2026, and opened in the US on March 20, "
    "2026, distributed by Amazon MGM Studios. Directed by Phil Lord and Christopher Miller (Into "
    "the Spider-Verse, The Lego Movie), with a screenplay by Drew Goddard (The Martian). The cast "
    "includes Ryan Gosling as Ryland Grace, Sandra Huller as Eva Stratt, Lionel Boyce, Milana "
    "Vayntrub, and Ken Leung. James Ortiz performed Rocky on set."
)

add_body(
    "The numbers tell the story: $80.6 million domestic opening weekend ($141 million global), "
    "the best opening of 2026 and a record for Amazon Studios. Total gross has passed $185 million "
    "and climbing. It sits at 95% on Rotten Tomatoes (critics) and 98% audience score, with a "
    "CinemaScore of \"A.\" It is only the second non-sequel in a decade to open that high, "
    "after Oppenheimer."
)

add_body(
    "Rocky was realized through practical puppetry, not primarily CGI, and the choice paid off. "
    "Six puppeteers brought the alien to life on set, giving Gosling something real to act against. "
    "The \"jazz hands\" first-contact scene, where Grace and Rocky attempt communication through "
    "improvised body language (hesitant squats, jazz hands, Usain Bolt's victory pose), has become "
    "a fan-favorite moment and a social media phenomenon. Rocky is being called the best movie "
    "character of 2026."
)

add_body(
    "The film has sparked lively debate on social media about book fidelity, humor frequency, "
    "and perceived political framing, but the controversy appears to have fueled rather than hurt "
    "the box office. The Hollywood Reporter noted the film \"ignited strong debate\" across platforms, "
    "covering everything from adaptation choices to America's role in the story."
)

doc.add_heading('The Audiobook: A Companion Experience', level=2)

add_body(
    "Ray Porter's narration of the audiobook won the 2022 Audie Award for Audiobook of the Year "
    "and is widely considered the definitive way to experience the story. Porter voices Rocky's "
    "musical-tone language over chord progressions, with the music gradually softening as Grace "
    "learns the language, a technique that translates the text's notation system into an auditory "
    "experience that print simply cannot replicate. AudioFile Magazine praised Porter for capturing "
    "\"the panic and semi-hysterical self-deprecation of an amnesiac\" and delivering \"a United "
    "Nations of accents\" where each one \"sounds like it's being voiced by a different actor.\" "
    "If you haven't experienced the audiobook, consider it essential listening, especially for "
    "Rocky's scenes."
)

# ============================================================
# 10. IF YOU LIKED THIS
# ============================================================
doc.add_heading('If You Liked This', level=1)

add_body(
    "Five recommendations depending on what you loved most about Project Hail Mary:"
)

# Recommendations table
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Book', 'Author', 'Why You\'ll Like It']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

recs = [
    ("Children of Time", "Adrian Tchaikovsky",
     "Speculative evolution meets first contact. Follows humanity encountering hyper-intelligent "
     "spiders on a terraformed world. Same sense of wonder at alien minds."),
    ("The Long Way to a Small, Angry Planet", "Becky Chambers",
     "Found-family aboard a diverse spaceship crew. Lighter on hard science but captures the same "
     "warmth and cross-species friendship."),
    ("Blindsight", "Peter Watts",
     "Hard sci-fi first contact that asks unsettling questions about consciousness and intelligence. "
     "The dark philosophical counterpoint to PHM's optimism."),
    ("Dragon's Egg", "Robert L. Forward",
     "Humans observe the million-times-faster evolution of tiny beings on a neutron star. Forward "
     "called it \"a textbook on neutron star physics disguised as a novel.\""),
    ("The Three-Body Problem", "Liu Cixin",
     "Hard science and impending alien contact on a civilizational scale. Shares PHM's love of "
     "physics-based problem-solving but with a darker geopolitical edge."),
]

for i, (title, author, why) in enumerate(recs):
    row = table.rows[i + 1]
    row.cells[0].text = title
    row.cells[1].text = author
    row.cells[2].text = why

add_body("")  # spacing

doc.add_heading('Companion Media', level=2)

companion_items = [
    "Andy Weir's other novels: The Martian (the obvious next read) and Artemis (less beloved, "
    "but still Weir). He has a new untitled standalone sci-fi novel sold to Del Rey, currently in progress.",
    "The Ray Porter audiobook (16 hours, Audie Award winner), essential for the Eridian language experience.",
    "Scientific American's \"Science Quickly\" podcast episode featuring Weir discussing the astrobiology of Astrophage.",
    "The Conversation's astrophysicist breakdown of the book's science.",
    "Skeptical Inquirer's \"Project Hail Mary, Astrophage, and the Real Extremophiles,\" connecting the "
    "fiction to real extremophile biology.",
    "The Project Hail Mary Fandom wiki (projecthailmary.fandom.com), where fans have calculated "
    "everything from Grace's travel time (~3.9 space years at 1.5g = ~13 Earth years) to the "
    "thermodynamic impossibility of Astrophage.",
]
for item in companion_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 11. BOOK CLUB DISCUSSION ANGLES
# ============================================================
doc.add_heading('Book Club Discussion Angles', level=1)

add_body("If you're reading this with a group, here are the conversations worth having:")

discussion_items = [
    "Stratt's utilitarianism: She operates on \"any sacrifice is justified to save the species.\" "
    "Is she right? Where's the line? Would you have made her choices?",
    "Grace as an unwilling hero: He didn't volunteer. He tried to talk his way out. Does his "
    "coerced participation change how we judge his eventual heroism?",
    "The teacher identity: Grace's superpower isn't being a genius scientist; it's being able "
    "to explain things simply. His ability to teach Rocky is what saves both civilizations. "
    "What does the book say about the value of teaching?",
    "The ending: Grace chooses Rocky over Earth. He spends 16 years teaching alien children. "
    "Is this a happy ending? A tragic one? Both?",
    "The science as storytelling: Does the hard science enhance or hinder your engagement? "
    "Did you skim the technical passages or lean in?",
    "Cross-species trust: Grace and Rocky build trust without shared language, culture, or biology. "
    "What does the book suggest about the foundations of trust?",
]
for item in discussion_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 12. SOURCES
# ============================================================
doc.add_heading('Sources', level=1)

sources = [
    "AudioFile Magazine. \"Narrator Ray Porter on Andy Weir's PROJECT HAIL MARY.\" audiofilemagazine.com.",
    "Big Think. \"Project Hail Mary: How Andy Weir Turns Science Into Drama.\" bigthink.com.",
    "Book Club Chat. \"Book Club Questions for Project Hail Mary by Andy Weir.\" bookclubchat.com.",
    "Britannica (2026). \"The Science Behind Project Hail Mary.\" britannica.com.",
    "Deadline (2020). \"MGM/Ryan Gosling Deal for The Hail Mary.\" deadline.com.",
    "Den of Geek (2026). \"Project Hail Mary Ending Explained.\" denofgeek.com.",
    "Gates Notes. \"Project Hail Mary Review.\" gatesnotes.com.",
    "Gizmodo (2026). \"Could There Be a Project Hail Mary Sequel?\" gizmodo.com.",
    "Gizmodo (2026). \"One of the Best Scenes in Project Hail Mary Wasn't in the Script.\" gizmodo.com.",
    "GradeSaver. \"Project Hail Mary Analysis.\" gradesaver.com.",
    "Hollywood Reporter (2026). \"Project Hail Mary Blasts Off to Huge $81M Box Office Opening.\" hollywoodreporter.com.",
    "Hollywood Reporter (2026). \"Why Ryan Gosling's Movie Ignited Strong Debate.\" hollywoodreporter.com.",
    "IndieWire (2026). \"Project Hail Mary Sound Design: Vocals for Rocky's Language.\" indiewire.com.",
    "Inverse (2026). \"Andy Weir Interview: Hard Sci-Fi Renaissance.\" inverse.com.",
    "Inverse (2026). \"Project Hail Mary Ending Explained.\" inverse.com.",
    "Kirkus Reviews. \"Project Hail Mary.\" kirkusreviews.com.",
    "LitCharts. \"Project Hail Mary Themes.\" litcharts.com.",
    "Looper. \"Project Hail Mary Writer Andy Weir on Ryan Gosling's Movie.\" looper.com.",
    "Medium (M. Hanafi). \"Project Hail Mary: A Study of Sacrifice, Betrayal, Solitude, Acceptance, and Friendship.\" medium.com.",
    "Motion Pictures Association (2026). \"Project Hail Mary Sound Designers on Creating Rocky's Alien Language.\" motionpictures.org.",
    "National Academy of Engineering. \"An Interview with Andy Weir.\" nae.edu.",
    "Northeastern University (2026). \"What Project Hail Mary Gets Right and Wrong About Astrophysics.\" news.northeastern.edu.",
    "NPR (2026). \"The Secret Weapon in Project Hail Mary Is Ryan Gosling's Star Power.\" npr.org.",
    "Penguin Random House. \"Science Fiction Books To Read if You Love Project Hail Mary.\" penguinrandomhouse.com.",
    "Project Hail Mary Wiki. projecthailmary.fandom.com.",
    "Rolling Stone. \"Andy Weir Profile.\" rollingstone.com.",
    "Rotten Tomatoes (2026). \"Project Hail Mary First Reviews.\" rottentomatoes.com.",
    "Science Friday (2026). \"Project Hail Mary Brings a New Kind of Alien to the Big Screen.\" sciencefriday.com.",
    "Scientific American (2026). \"How Accurate Is the Science in Project Hail Mary?\" scientificamerican.com.",
    "Scientific American Podcast. \"The Real Science (and the Fun Fiction) Behind Project Hail Mary.\" scientificamerican.com.",
    "Screen Rant (2026). \"Andy Weir Teases Next Sci-Fi Book & Potential Project Hail Mary Sequel.\" screenrant.com.",
    "Shelf Media Group. \"Interview: A Conversation with Andy Weir About Project Hail Mary.\" shelfmediagroup.com.",
    "Skeptical Inquirer. \"Project Hail Mary, Astrophage, and the Real Extremophiles.\" skepticalinquirer.org.",
    "Space.com (2026). \"Andy Weir Explains the Astrobiology Behind Project Hail Mary.\" space.com.",
    "SuperSummary. \"Project Hail Mary Study Guide.\" supersummary.com.",
    "The Bookseller. \"Del Rey Lands New Novel from Andy Weir.\" thebookseller.com.",
    "The Conversation (2026). \"Project Hail Mary Is Packed with Hard Science: An Astrophysicist Breaks It Down.\" theconversation.com.",
    "Variety (2026). \"Project Hail Mary Directors Screened Four-Hour Cut.\" variety.com.",
    "Wikipedia. \"Andy Weir.\" en.wikipedia.org.",
    "Wikipedia. \"Project Hail Mary.\" en.wikipedia.org.",
]

for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(2)

# ============================================================
# SAVE
# ============================================================
output_path = "/Users/hornej/Documents/Research/project-hail-mary/2026-03-28 Project Hail Mary Companion Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
