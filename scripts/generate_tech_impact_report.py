#!/usr/bin/env python3
"""
Generate research report: The Negative Impact of Technology in the K-12 Classroom
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('The Negative Impact of Technology in the K-12 Classroom', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Research Synthesis of Peer-Reviewed Studies, Meta-Analyses, and Seminal Works')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(89, 89, 89)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ============================================================
# SECTION 1: INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'Over the past two decades, K-12 school systems worldwide have invested billions of dollars in classroom technology\u2014'
    'laptops, tablets, interactive whiteboards, learning management systems, and one-to-one device programs. The implicit '
    'promise has been that these tools will modernize instruction, personalize learning, and boost student achievement. '
    'However, a growing and increasingly robust body of research suggests that this promise has gone largely unfulfilled '
    'and that, in many contexts, classroom technology is actively harming student learning, attention, social development, '
    'and mental health.'
)
doc.add_paragraph(
    'This report synthesizes peer-reviewed studies, systematic reviews, meta-analyses, and large-scale international '
    'assessment data on the negative impacts of technology in K-12 classrooms. It draws on the work of key researchers '
    'including Dr. Jared Cooney Horvath, Dr. Jean Twenge, Dr. Jonathan Haidt, Dr. Larry Rosen, and others. Where '
    'available, effect sizes and specific quantitative findings are reported. The evidence is organized thematically: '
    'overall effectiveness of educational technology, digital distraction and attention, reading comprehension on '
    'screens versus paper, one-to-one device programs, smartphones and social media in schools, mental health impacts, '
    'and international policy responses.'
)
doc.add_paragraph(
    'Note on the researcher named in the original request: The user asked specifically about "Jason Horvath." '
    'No researcher by that exact name was found publishing on this topic. However, Dr. Jared Cooney Horvath, a '
    'neuroscientist and educator who has conducted research at Harvard University and the University of Melbourne, '
    'is the most prominent Horvath publishing on technology and K-12 education. His 2025 book The Digital Delusion '
    'and his January 2026 testimony before the U.S. Senate Committee on Commerce, Science, and Transportation are '
    'central to the current discourse. This report treats his work extensively.'
)

# ============================================================
# SECTION 2: HORVATH
# ============================================================
doc.add_heading('2. The Work of Dr. Jared Cooney Horvath', level=1)

doc.add_heading('2.1 Background and Credentials', level=2)
doc.add_paragraph(
    'Dr. Jared Cooney Horvath (PhD, MEd) is a neuroscientist and educator who has conducted research and taught at '
    'Harvard University, Harvard Medical School, and the University of Melbourne. He has worked with more than 1,000 '
    'schools worldwide and is the author of six books and over fifty peer-reviewed articles. He currently serves as '
    'Director of LME Global, an organization dedicated to translating brain and behavioral science for educators, '
    'students, and communities. His work has appeared in The New Yorker, The Atlantic, The Economist, Harvard Business '
    'Review, and PBS\'s NOVA.'
)

doc.add_heading('2.2 The Digital Delusion (2025)', level=2)
doc.add_paragraph(
    'In his 2025 book The Digital Delusion: How Classroom Technology Harms Our Kids\' Learning\u2014And How to Help '
    'Them Thrive Again, Horvath draws on decades of neuroscience and education research to dismantle what he calls '
    '"the core myths fueling the edtech movement." His central argument is that digital tools fundamentally clash with '
    'the architecture of human learning\u2014how the brain encodes, consolidates, and retrieves information\u2014and that '
    'edtech platforms often degrade rather than enhance learning outcomes.'
)
doc.add_paragraph(
    'Horvath synthesized 398 meta-analyses covering more than 21,000 individual studies and found an overall weighted '
    'mean effect size of +0.29 standard deviations for educational technology interventions. He argues this falls '
    'below the threshold for meaningful educational impact. His subject-specific breakdowns are revealing:'
)

# Effect size table for Horvath
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Subject Area', 'Effect Size (d)', 'Meta-Analyses', 'Studies']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['Mathematics', '0.33', '22', '1,060'],
    ['Literacy', '0.25', '17', '736'],
    ['Sciences', '0.18', '6', '391'],
    ['Writing Quality', '0.32', '6', '75'],
    ['Specific Learning Needs', '0.61', '10', '216'],
    ['Overall Weighted Mean', '0.29', '398', '21,000+'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'For context, John Hattie\'s Visible Learning framework\u2014the largest synthesis of education research ever '
    'conducted, encompassing over 2,100 meta-analyses and 300 million students\u2014places the mean effect size of all '
    'educational interventions at 0.40 standard deviations, representing roughly one year of academic growth. Horvath '
    'argues that most edtech interventions fall well below this bar, particularly in science (d = 0.18) and literacy '
    '(d = 0.25), and that for disadvantaged populations the effect size drops to just 0.18.'
)

doc.add_heading('2.3 International Assessment Data', level=2)
doc.add_paragraph(
    'Horvath draws heavily on international assessment data to illustrate the negative correlation between heavy '
    'classroom computer use and student performance:'
)
bullets = [
    'PISA: Students who use computers more than six hours daily score 66 points lower than nonusers\u2014a gap '
    'equivalent to moving from the 50th to the 24th percentile, or roughly two letter grades.',
    'TIMSS 2019: Daily computer users scored 41 points lower in mathematics and 51 points lower in science '
    'compared to infrequent users.',
    'Collectively, Horvath asserts, "these assessments involve millions of students over decades and converge on '
    'the same conclusion: heavy classroom screen exposure is not improving learning outcomes at scale."'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('2.4 U.S. Senate Testimony (January 2026)', level=2)
doc.add_paragraph(
    'In January 2026, Horvath testified before the U.S. Senate Committee on Commerce, Science, and Transportation. '
    'He argued that if federal policy continues incentivizing large-scale digital adoption in schools without demanding '
    '"independent efficacy evidence, privacy protections, and developmental safeguards," there will be long-term '
    'educational harm. He told senators: "Our kids are less cognitively capable than we were at their age," describing '
    'what he termed a "decoupling" between time spent in school and cognitive development. He called for independent '
    'efficacy requirements before edtech products can be deployed in schools, drawing a parallel to FDA approval '
    'processes for pharmaceuticals.'
)

doc.add_heading('2.5 "The EdTech Revolution Has Failed" (After Babel, 2024)', level=2)
doc.add_paragraph(
    'In a widely cited 2024 article published on Jonathan Haidt\'s After Babel platform, Horvath presented his '
    'analysis of why schools across Europe and Southeast Asia are moving away from student-facing digital technology. '
    'He argued that education research systematically inflates positive effects of edtech because comparison conditions '
    'vary widely and often lack rigorous baselines. When educational interventions are benchmarked against established '
    'instructional methods rather than "business as usual" conditions, the apparent benefits of technology largely '
    'disappear.'
)

# ============================================================
# SECTION 3: DIGITAL DISTRACTION
# ============================================================
doc.add_heading('3. Digital Distraction and Attention in the Classroom', level=1)

doc.add_heading('3.1 The "Brain Drain" Effect', level=2)
doc.add_paragraph(
    'One of the most cited studies on technology and cognition is Ward, Duke, Gneezy, and Bos (2017), "Brain Drain: '
    'The Mere Presence of One\'s Own Smartphone Reduces Available Cognitive Capacity," published in the Journal of the '
    'Association for Consumer Research. In two experiments, participants were randomly assigned to place their smartphones '
    'on the desk (face down), in a bag or pocket, or in another room. Even when participants successfully maintained '
    'sustained attention and did not check their phones, those with phones on the desk or in their pocket performed '
    'significantly worse on tests of working memory capacity and fluid intelligence than those whose phones were in '
    'another room. The cognitive costs were highest for those most dependent on their smartphones. This finding has '
    'direct implications for classrooms where students have devices on their desks at all times.'
)

doc.add_heading('3.2 Laptop Multitasking and Off-Task Use', level=2)
doc.add_paragraph(
    'Ravizza, Uitvlugt, and Fenn (2017), in "Logged In and Zoned Out" (Psychological Science), monitored actual '
    'internet usage of students who brought laptops to an introductory psychology course via a proxy server. Key '
    'findings included: (1) nonacademic internet use was common and inversely related to class performance, even '
    'after controlling for motivation, interest, and intelligence; (2) the negative relationship between off-task '
    'laptop use and exam scores was twice as large for students with high ACT scores as for those with low ACT '
    'scores\u2014suggesting that even high-ability students are not immune to digital distraction.'
)

doc.add_heading('3.3 The Note-Taking Debate: Pen vs. Keyboard', level=2)
doc.add_paragraph(
    'Mueller and Oppenheimer (2014), in "The Pen Is Mightier Than the Keyboard" (Psychological Science), found that '
    'students who took notes on laptops performed worse on conceptual questions than students who took notes by hand. '
    'The proposed mechanism was that laptop users tended to transcribe lectures verbatim rather than processing and '
    'reframing information, leading to shallower encoding. However, subsequent replications and a meta-analysis have '
    'shown that the effect is smaller and less consistent than initially reported. While the specific pen-vs.-keyboard '
    'finding remains debated, the broader pattern\u2014that laptops invite distraction and shallower processing\u2014is '
    'supported by convergent evidence from multiple studies.'
)

doc.add_heading('3.4 UNESCO\'s Assessment of Distraction', level=2)
doc.add_paragraph(
    'The 2023 UNESCO Global Education Monitoring Report, covering 14 countries from pre-primary through higher '
    'education, found that smartphone notifications alone\u2014even without the student touching the device\u2014are '
    'sufficient to break attention. The report cited research showing it can take students up to 20 minutes to fully '
    'refocus on a learning task after a digital interruption.'
)

# ============================================================
# SECTION 4: SCREEN READING
# ============================================================
doc.add_heading('4. Screen Versus Paper: The Reading Comprehension Gap', level=1)
doc.add_paragraph(
    'Multiple meta-analyses have now established what researchers call the "screen inferiority effect"\u2014a consistent '
    'finding that reading comprehension is lower when text is presented on digital screens than on paper.'
)

bullets = [
    'Delgado, Vargas, Ackerman, and Salmer\u00f3n (2018) conducted a large meta-analysis published in an Elsevier journal, '
    'combining 54 studies (both between- and within-participant designs) involving 171,055 participants. Both designs '
    'yielded the same conclusion: paper reading produced better comprehension than digital reading.',
    'A 2024 meta-analysis of 49 studies specifically examining handheld digital devices (phones, tablets) found a '
    'small but statistically significant negative effect for digital reading compared to paper.',
    'The effect is moderated by text complexity and time pressure: the screen inferiority effect grows larger for '
    'longer, more complex texts and when readers are under time constraints.',
    'Students with low reading comprehension skills are disproportionately affected. While strong readers show '
    'comparable comprehension across media, struggling readers perform significantly worse on digital texts, '
    'particularly under time pressure.',
    'The "shallowing hypothesis" offers one explanation: constant exposure to fast-paced digital media trains '
    'the brain to process information more rapidly but less thoroughly, undermining the deep processing required '
    'for comprehension of complex texts.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# SECTION 5: 1:1 PROGRAMS
# ============================================================
doc.add_heading('5. One-to-One Device Programs', level=1)
doc.add_paragraph(
    'One-to-one (1:1) computing programs\u2014in which every student is issued a personal device such as a Chromebook '
    'or iPad\u2014have become standard in many U.S. school districts. Research on their effectiveness reveals a pattern '
    'of disappointing academic outcomes coupled with significant behavioral concerns.'
)

doc.add_heading('5.1 Academic Performance', level=2)
doc.add_paragraph(
    'Research conducted between 2011 and 2016 found that iPad use in schools can lead to decreased academic performance '
    'and attentional awareness because these devices act as a major distraction. A 2023 survey by Education Week found '
    'that 27% of educators reported a negative impact of 1:1 computing environments on classroom management\u2014up '
    'from 20% in 2019. Multiple studies have found that such programs "not only have little effect on academic '
    'achievement but also inappropriately prioritize screen-based learning over interpersonal communication and '
    'increase the digital divide in education."'
)

doc.add_heading('5.2 Behavioral and Social Effects', level=2)
doc.add_paragraph(
    'Research from Cal State found that Chromebooks are having negative effects on developing children\'s mental and '
    'physical health, ability to build peer relationships, and capacity to achieve interpersonal and emotional skills. '
    'Teachers report increased off-task behavior and students\' inability to sustain attention during class. '
    'Additionally, increased technological use in classrooms has been linked to negative effects on students\' ability '
    'to communicate face-to-face and interact personally with peers and teachers.'
)

doc.add_heading('5.3 Health Effects', level=2)
doc.add_paragraph(
    'Students exposed to large amounts of screen time have a higher risk of adverse health effects as well as '
    'learning deficits. Schools transitioning to 1:1 devices are increasing blue light exposure, which can disrupt '
    'circadian rhythms and interfere with sleep quality\u2014a critical factor for learning and memory consolidation '
    'in developing brains.'
)

# ============================================================
# SECTION 6: SMARTPHONES AND PHONE BANS
# ============================================================
doc.add_heading('6. Smartphones in Schools and the Impact of Phone Bans', level=1)

doc.add_heading('6.1 The Beland and Murphy Study (2016)', level=2)
doc.add_paragraph(
    'In a landmark study published in Labour Economics, Beland and Murphy (2016) examined the effect of smartphone '
    'bans across dozens of schools in Birmingham, Leicester, London, and Manchester. They combined survey data on '
    'mobile phone policies with administrative data on student achievement to create a longitudinal history of '
    'performance. Key findings:'
)
bullets = [
    'Student performance on high-stakes exams increased significantly post-ban, by approximately 0.07 standard '
    'deviations on average.',
    'The effect was concentrated among lower-achieving students: students in the bottom quintile of prior achievement '
    'gained 14.23% of a standard deviation, equivalent to adding roughly one hour of instruction per week.',
    'Students in the top quintile were neither positively nor negatively affected, suggesting that high achievers '
    'can self-regulate in the presence of phones while low achievers cannot.',
    'The authors concluded: "Banning mobile phones improves outcomes for the low-achieving students the most and '
    'does not significantly impact high achievers."'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.2 The Sungu, Choudhury, and Bjerre-Nielsen RCT (2025)', level=2)
doc.add_paragraph(
    'In the largest randomized controlled trial (RCT) on this topic to date, Sungu, Choudhury, and Bjerre-Nielsen '
    '(2025) studied nearly 17,000 students across ten higher education institutions in the Indian state of Odisha. '
    'Academic departments within each college were randomly assigned to either a "phone-ban" condition (students deposited '
    'phones in a wooden box at the start of class) or a control condition (unrestricted phone use). Results showed:'
)
bullets = [
    'Mandatory phone collection led to higher grades, particularly among lower-performing, first-year, and non-STEM '
    'students, with an average increase of 0.086 standard deviations.',
    'Random classroom spot checks revealed fewer instances of student chatter and disruptive behaviors, reduced '
    'phone usage, and increased teacher engagement in phone-ban classrooms.',
    'Students exposed to the ban became substantially more supportive of phone-use restrictions.',
    'Despite a mild increase in reported FOMO (fear of missing out), there were no significant changes in overall '
    'well-being, academic motivation, or experiences of online harassment.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('6.3 Florida Statewide Ban Evidence (2024-2025)', level=2)
doc.add_paragraph(
    'Research on Florida\'s statewide cellphone ban found that it increased scores on spring accountability tests '
    'by 1.1 percentiles overall in the second year. Cellphone bans also significantly reduced student unexcused '
    'absences, an effect that may explain a large fraction of the test score gains. However, implementation '
    'increased disciplinary incidents in the first year\u2014primarily among Black students, male students, and '
    'students in middle and high school\u2014though these increases disappeared by the second year.'
)

doc.add_heading('6.4 Global Policy Trends', level=2)
doc.add_paragraph(
    'As of the end of 2024, according to UNESCO, 79 education systems worldwide (roughly 40% of all systems with '
    'available data) have enacted bans on smartphone use in schools. This represents a rapid increase from 60 systems '
    '(30%) at the end of 2023. Countries that have removed smartphones from schools\u2014including Belgium, Spain, and '
    'the United Kingdom\u2014have reported improvements in learning outcomes, particularly for lower-performing students.'
)

# Table: Phone ban effects
doc.add_paragraph()
phone_table = doc.add_table(rows=5, cols=4)
phone_table.style = 'Light Shading Accent 1'
phone_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Study', 'Sample', 'Effect on Achievement', 'Key Finding']
for i, h in enumerate(headers):
    cell = phone_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

phone_data = [
    ['Beland & Murphy (2016)', '4 English cities', '+0.07 SD average', 'Low achievers gain 14.23% SD'],
    ['Sungu et al. (2025)', '~17,000 students (RCT)', '+0.086 SD', 'Strongest for low-performing, first-year'],
    ['Florida Ban (2024)', 'Statewide', '+1.1 percentile (Year 2)', 'Also reduced unexcused absences'],
    ['UNESCO Global Review', '79 education systems', 'Positive in Belgium, Spain, UK', 'Largest gains for disadvantaged students'],
]
for r, row_data in enumerate(phone_data):
    for c, val in enumerate(row_data):
        phone_table.rows[r+1].cells[c].text = val

doc.add_paragraph()

# ============================================================
# SECTION 7: TWENGE AND HAIDT
# ============================================================
doc.add_heading('7. Mental Health: The Work of Twenge, Haidt, and Rosen', level=1)

doc.add_heading('7.1 Jean Twenge and the iGen Thesis', level=2)
doc.add_paragraph(
    'Dr. Jean Twenge, a psychologist at San Diego State University, is widely recognized for her research linking '
    'the rise of smartphones and social media to a sharp decline in adolescent mental health beginning around 2012\u2014'
    'the year smartphone ownership first reached majority penetration among Americans. Her key findings include:'
)
bullets = [
    'Clinical-level depression doubled among U.S. 12- to 17-year-olds between 2011 and 2019.',
    'Heavy users of social media (five or more hours per day) are twice as likely to be depressed as nonusers.',
    'Depression rates begin to increase after just one hour of daily social media use, though the curve is shallower '
    'for other types of screen time, with higher depression rates appearing only after three to four hours per day.',
    'The increases in poor mental health indicators have been larger among girls and young women than among boys and '
    'young men.',
    'Twenge tracked PISA data showing that loneliness in school was stable from the early 2000s through 2012, then '
    'increased sharply across Europe and the U.S., though less so in East Asia.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('7.2 Jonathan Haidt and The Anxious Generation (2024)', level=2)
doc.add_paragraph(
    'Dr. Jonathan Haidt, a social psychologist at NYU Stern School of Business, extends Twenge\'s work in his 2024 '
    'book The Anxious Generation: How the Great Rewiring of Childhood Is Causing an Epidemic of Mental Illness. Haidt '
    'proposes four foundational reforms:'
)
bullets = [
    'No smartphones before high school (approximately age 14).',
    'No social media before age 16.',
    'Phone-free schools (including during breaks and lunch periods).',
    'Far more unsupervised play and childhood independence.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Haidt argues that phone access during school\u2014even outside of class, during lunch and breaks\u2014is '
    'detrimental to learning and social engagement. He identifies three of four "foundational harms of the phone-based '
    'childhood" that school phone bans can address: attention fragmentation, social deprivation, and addiction. He '
    'further notes that allowing phones in schools exacerbates disparities, with disadvantaged students suffering '
    'the greatest negative impact because they have less supervision and spend more time online.'
)

doc.add_heading('7.3 Larry Rosen and the Psychology of Technology', level=2)
doc.add_paragraph(
    'Dr. Larry Rosen, a research psychologist and professor emeritus at California State University, Dominguez Hills, '
    'has studied reactions to technology among more than 70,000 people in 23 countries over 30 years. His concept of '
    '"iDisorder"\u2014changes in brain processing and social relating caused by daily technology use that produce signs '
    'and symptoms of psychological disorders\u2014has influenced thinking about technology in schools. His research '
    'specifically examined: (1) the educational impact of text-message-induced task switching in the classroom; '
    '(2) the role of executive function and technological anxiety (FOMO) in course performance as mediated by '
    'technology usage and multitasking habits; and (3) the broader concept of how digital engagement restructures '
    'attention and self-regulation.'
)

doc.add_heading('7.4 Social Media, Cyberbullying, and Sleep in Schools', level=2)
doc.add_paragraph(
    'The CDC\'s 2023 Youth Risk Behavior Survey reported that 77% of high school students engage in frequent social '
    'media use, with documented associations between this use and experiences of bullying victimization, persistent '
    'feelings of sadness or hopelessness, and suicide risk. Forty-six percent of U.S. teens have experienced at least '
    'one form of cyberbullying (Pew Research Center, 2023). Poor sleep quality\u2014often driven by late-night social '
    'media use\u2014further exacerbates anxiety and depression symptoms and contributes to chronic emotional fatigue '
    'that undermines classroom performance.'
)

# ============================================================
# SECTION 8: ICT OVERUSE
# ============================================================
doc.add_heading('8. The Inverted U: ICT Overuse and Academic Performance', level=1)
doc.add_paragraph(
    'Gorj\u00f3n and Os\u00e9s (2023), in a study published in the Journal of Educational Computing Research using '
    'PISA 2018 data from 22 OECD countries, established a causal relationship between ICT overuse and declining student '
    'performance. Using Inverse Probability Weighting techniques, they found an inverted U-shaped relationship: low and '
    'medium ICT use at school improved mathematical performance compared to very low use, while very intensive use led '
    'to significant penalties. The most intensive users experienced underperformance equivalent to roughly half an '
    'academic year in Estonia, Finland, and Spain. This nonlinear finding is critical: it suggests that technology is '
    'not inherently harmful, but that the dosage and implementation matter enormously, and that current levels of use '
    'in many schools have crossed the point of diminishing\u2014and then negative\u2014returns.'
)

doc.add_paragraph(
    'PISA 2022 data reinforce this pattern. Students who spent up to one hour per day on digital devices for learning '
    'scored 14 points higher in mathematics than students who spent no time. But students who reported being distracted '
    'by other students\' digital devices in at least some mathematics lessons scored 15 points lower than students who '
    'reported this never or almost never happened, after controlling for socioeconomic factors.'
)

# ============================================================
# SECTION 9: SWEDEN
# ============================================================
doc.add_heading('9. Case Study: Sweden\'s Reversal on Digital Learning', level=1)
doc.add_paragraph(
    'Sweden provides a compelling real-world case study. After years of prioritizing digital learning\u2014replacing '
    'physical textbooks with screens beginning in 2009\u2014the Swedish government reversed course in May 2023. The '
    'catalyst was declining student performance on international assessments, including the PIRLS study, which showed '
    'that Swedish students\' reading comprehension scores declined between 2016 and 2021. Education Minister Lotta '
    'Edholm cited declining critical thinking skills and overreliance on screens as contributing factors.'
)
doc.add_paragraph(
    'The Swedish government committed \u20ac104 million (\u20ac60 million in 2023, with an additional \u20ac44 million '
    'for 2024-2025) to reintroduce physical textbooks. Teachers had reported that students struggled to sustain '
    'attention, and research cited by the government showed that reading on backlit screens demands more cognitive '
    'effort than paper, particularly for younger children. Several studies link heavy digital use to weaker memory '
    'retention and lower comprehension, especially among students who already struggle to focus.'
)

# ============================================================
# SECTION 10: SYNTHESIS TABLE
# ============================================================
doc.add_heading('10. Summary of Key Effect Sizes and Findings', level=1)
doc.add_paragraph(
    'The following table summarizes the quantitative findings from the major studies and meta-analyses reviewed:'
)

summary_table = doc.add_table(rows=13, cols=3)
summary_table.style = 'Light Shading Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Study/Source', 'Key Metric', 'Finding']
for i, h in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

summary_data = [
    ['Horvath (2025) meta-synthesis', 'Overall EdTech effect size', 'd = 0.29 (below Hattie\'s 0.40 threshold)'],
    ['Horvath: PISA data', 'Heavy computer users vs. nonusers', '-66 points (50th to 24th percentile)'],
    ['Horvath: TIMSS 2019', 'Daily vs. infrequent computer users', '-41 pts (math), -51 pts (science)'],
    ['Beland & Murphy (2016)', 'Phone ban effect on achievement', '+0.07 SD; low achievers +14.23% SD'],
    ['Sungu et al. (2025) RCT', 'Phone ban effect on grades', '+0.086 SD (n = 17,000)'],
    ['Ward et al. (2017)', 'Smartphone mere presence on cognition', 'Significant reduction in working memory'],
    ['Delgado et al. (2018)', 'Screen vs. paper reading', 'Paper advantage; n = 171,055'],
    ['Gorjon & Oses (2023)', 'ICT overuse penalty (OECD)', '~0.5 academic year loss at intensive use'],
    ['Twenge (2017-2022)', 'Depression in heavy SM users', '2x rate vs. nonusers (5+ hrs/day)'],
    ['PISA 2022', 'Device distraction in math class', '-15 points for distracted students'],
    ['Mueller & Oppenheimer (2014)', 'Laptop vs. handwritten notes', 'Worse conceptual recall on laptops'],
    ['Florida Ban Study (2024)', 'Statewide phone ban effect', '+1.1 percentile on state tests (Year 2)'],
]
for r, row_data in enumerate(summary_data):
    for c, val in enumerate(row_data):
        summary_table.rows[r+1].cells[c].text = val

doc.add_paragraph()

# ============================================================
# SECTION 11: EVIDENCE QUALITY
# ============================================================
doc.add_heading('11. Distinguishing Causal from Correlational Evidence', level=1)
doc.add_paragraph(
    'It is important to note the varying strength of the evidence reviewed in this report:'
)
bullets = [
    'Strong causal evidence: The Sungu et al. (2025) RCT and the Beland and Murphy (2016) quasi-experimental '
    'study provide the most robust causal evidence that removing phones from classrooms improves academic outcomes. '
    'The Gorjon and Oses (2023) study used Inverse Probability Weighting to establish causal claims about ICT overuse.',
    'Moderate causal evidence: The PISA and TIMSS correlations cited by Horvath are large-scale but cross-sectional '
    'and observational. They cannot rule out selection effects (e.g., schools that use more technology may differ '
    'in other ways from those that do not). However, the consistency of the pattern across countries and time periods '
    'strengthens the inference.',
    'Correlational evidence: Twenge\'s and Haidt\'s work on mental health relies primarily on time-series '
    'correlations and survey data. While the timing of the mental health decline aligns with smartphone adoption, '
    'and the evidence has been described as "reasonably robust" by Twenge herself, critics note that correlation '
    'does not establish causation and that other factors may contribute.',
    'Contested evidence: The Mueller and Oppenheimer pen-vs.-keyboard finding has been partially challenged by '
    'subsequent replications. The overall case against laptops rests more firmly on distraction and off-task use '
    'than on the specific mechanism of verbatim transcription.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
# SECTION 12: CONCLUSIONS
# ============================================================
doc.add_heading('12. Conclusions and Implications for K-12 Leaders', level=1)
doc.add_paragraph(
    'The weight of the evidence reviewed in this report supports several conclusions:'
)
bullets = [
    'Educational technology, as deployed at scale in K-12 settings, has not delivered on its promise. The overall '
    'effect size of edtech interventions (d = 0.29) falls below the threshold for meaningful impact, and the effects '
    'are weakest in the areas where they are needed most (literacy, science, disadvantaged populations).',
    'Heavy screen exposure in school is associated with lower academic performance across multiple international '
    'assessments and countries.',
    'Smartphones are a significant source of cognitive drain and distraction. Removing them from schools improves '
    'academic outcomes, particularly for lower-achieving and disadvantaged students.',
    'Reading comprehension is consistently lower on screens than on paper, particularly for complex texts, '
    'time-pressured conditions, and struggling readers.',
    'One-to-one device programs have not improved academic achievement and have introduced problems with '
    'distraction, off-task behavior, and diminished interpersonal communication.',
    'The mental health crisis among adolescents is temporally linked to the rise of smartphones and social '
    'media, with converging evidence from multiple research groups and datasets.',
    'International policy is moving rapidly toward restricting technology in schools, with 79 education systems '
    'now banning smartphones and countries like Sweden investing hundreds of millions of euros to restore '
    'physical textbooks.',
    'The relationship between technology and learning is nonlinear: moderate, purposeful, teacher-directed use '
    'can be beneficial, but the current default in many schools\u2014heavy, largely unstructured screen time\u2014'
    'has crossed into harmful territory.'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'For school district leaders, the implications are clear: technology decisions should be driven by evidence '
    'of efficacy, not vendor marketing; smartphone policies should be restrictive; device deployment should be '
    'purposeful and limited rather than blanket; and the return to physical textbooks and handwriting should be '
    'considered seriously, particularly for younger students and struggling readers.'
)

# ============================================================
# SOURCES
# ============================================================
doc.add_heading('Sources', level=1)
sources = [
    '1. Horvath, J. C. (2025). The Digital Delusion: How Classroom Technology Harms Our Kids\' Learning\u2014And How to Help Them Thrive Again. https://www.amazon.com/Digital-Delusion-Classroom-Technology-Learning/dp/B0G5622DQQ',
    '2. Horvath, J. C. (2024). "The EdTech Revolution Has Failed." After Babel. https://www.afterbabel.com/p/the-edtech-revolution-has-failed',
    '3. Horvath, J. C. (2026). Written Testimony before U.S. Senate Committee on Commerce, Science, and Transportation. https://www.commerce.senate.gov/services/files/A19DF2E8-3C69-4193-A676-430CF0C83DC2',
    '4. Ward, A. F., Duke, K., Gneezy, A., & Bos, M. W. (2017). "Brain Drain: The Mere Presence of One\'s Own Smartphone Reduces Available Cognitive Capacity." Journal of the Association for Consumer Research, 2(2). https://www.journals.uchicago.edu/doi/full/10.1086/691462',
    '5. Mueller, P. A., & Oppenheimer, D. M. (2014). "The Pen Is Mightier Than the Keyboard: Advantages of Longhand Over Laptop Note Taking." Psychological Science, 25(6). https://journals.sagepub.com/doi/abs/10.1177/0956797614524581',
    '6. Ravizza, S. M., Uitvlugt, M. G., & Fenn, K. M. (2017). "Logged In and Zoned Out: How Laptop Internet Use Relates to Classroom Learning." Psychological Science, 28(2). https://journals.sagepub.com/doi/10.1177/0956797616677314',
    '7. Beland, L.-P., & Murphy, R. (2016). "Ill Communication: Technology, Distraction & Student Performance." Labour Economics, 41, 61-76. https://www.sciencedirect.com/science/article/abs/pii/S0927537116300136',
    '8. Sungu, A., Choudhury, P. K., & Bjerre-Nielsen, A. (2025). "Removing Phones from Classrooms Improves Academic Performance." SSRN Working Paper. https://papers.ssrn.com/sol3/papers.cfm?abstract_id=5370727',
    '9. Gorj\u00f3n, L., & Os\u00e9s, A. (2023). "The Negative Impact of Information and Communication Technologies Overuse on Student Performance: Evidence From OECD Countries." Journal of Educational Computing Research. https://journals.sagepub.com/doi/abs/10.1177/07356331221133408',
    '10. Twenge, J. M. (2020). "Increases in Depression, Self-Harm, and Suicide Among U.S. Adolescents After 2012 and Links to Technology Use: Possible Mechanisms." Psychiatric Research and Clinical Practice. https://pmc.ncbi.nlm.nih.gov/articles/PMC9176070/',
    '11. Haidt, J. (2024). The Anxious Generation: How the Great Rewiring of Childhood Is Causing an Epidemic of Mental Illness. https://www.anxiousgeneration.com/',
    '12. Rosen, L. D. (2012). iDisorder: Understanding Our Obsession with Technology and Overcoming Its Hold on Us. Palgrave Macmillan. https://www.amazon.com/iDisorder-Understanding-Obsession-Technology-Overcoming/dp/1137278315',
    '13. Delgado, P., Vargas, C., Ackerman, R., & Salmer\u00f3n, L. (2018). "Don\'t Throw Away Your Printed Books: A Meta-Analysis on the Effects of Reading Media on Reading Comprehension." Educational Research Review, 25, 23-38. https://www.sciencedirect.com/science/article/pii/S1747938X18300101',
    '14. UNESCO (2023). Global Education Monitoring Report 2023: Technology in Education: A Tool on Whose Terms? https://gem-report-2023.unesco.org/',
    '15. OECD (2023). PISA 2022 Results (Volume I). https://www.oecd.org/en/publications/pisa-2022-results-volume-i_53f23881-en.html',
    '16. Hattie, J. (2009/updated). Visible Learning: A Synthesis of Over 800 Meta-Analyses Relating to Achievement. https://visible-learning.org/hattie-ranking-influences-effect-sizes-learning-achievement/',
    '17. Horvath, J. C., Lodge, J., & Hattie, J. (2016). From the Laboratory to the Classroom: Translating Science of Learning for Teachers. Routledge. https://www.amazon.com/Laboratory-Classroom-Jared-Horvath/dp/1138649643',
    '18. CDC (2023). Youth Risk Behavior Survey: Frequent Social Media Use and Mental Health Among High School Students. MMWR. https://www.cdc.gov/mmwr/volumes/73/su/su7304a3.htm',
    '19. U.S. Surgeon General (2023). Social Media and Youth Mental Health Advisory. https://www.hhs.gov/sites/default/files/sg-youth-mental-health-social-media-advisory.pdf',
    '20. Paragon Institute (2024). "Banning Smartphones in Schools: Review of the Literature Shows Positive Impact." https://paragoninstitute.org/public-health/banning-smartphones-in-schools/',
    '21. Rahali et al. (2024). "Smartphone Policies in Schools: What Does the Evidence Say?" London School of Economics. https://eprints.lse.ac.uk/125554/1/Smartphone_policies_in_schools_Rahali_et_al_2024_002_.pdf',
    '22. NBER (2025). "The Impact of Cellphone Bans in Schools on Student Outcomes: Evidence from Florida." Working Paper No. 34388. https://www.nber.org/system/files/working_papers/w34388/w34388.pdf',
    '23. Education Next (2025). "A Latter-Day Luddite Pulls the Plug on EdTech" (Review of The Digital Delusion). https://www.educationnext.org/latter-day-luddite-pulls-the-plug-on-edtech-book-review-digital-delusion-jared-cooney-horvath/',
    '24. MDPI Education Sciences (2024). "To Ban or Not to Ban? A Rapid Review on the Impact of Smartphone Bans in Schools." https://www.mdpi.com/2227-7102/14/8/906',
    '25. Frontiers in Psychology (2025). "A Meta-Analysis of the Impact of Technology Related Factors on Students\' Academic Performance." https://www.frontiersin.org/journals/psychology/articles/10.3389/fpsyg.2025.1524645/full',
    '26. EducationHQ (2025). "Tech Proliferation in Schools Is Harming Students\' Learning, Neuroscientist Warns." https://educationhq.com/news/tech-proliferation-in-schools-is-harming-students-learning-neuroscientist-warns-204990/',
    '27. Strom, A. (2020). "The Negative Effects of Technology for Students and Educators." Northwestern College Master\'s Thesis. https://nwcommons.nwciowa.edu/education_masters/326/',
    '28. PMC (2022). "Problem Technology Use, Academic Performance, and School Connectedness among Adolescents." https://pmc.ncbi.nlm.nih.gov/articles/PMC8871851/',
]

for s in sources:
    p = doc.add_paragraph(s)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)

# Save
output_path = '/Users/hornej/Documents/Research Reports/Negative Impact of Technology in K-12 Classrooms - Research Synthesis.docx'
doc.save(output_path)
print(f"Report saved to: {output_path}")
